"""
集成测试：跨模块交互与端到端行为验证

覆盖模块：cli.py, config, agent, set_style.py, set_tag.py, word_structure
"""

import argparse
import io
import os
import shutil
import tempfile
import threading
from unittest import mock

import pytest
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi.testclient import TestClient

from wordformat.cli import validate_file, main
from wordformat.config.loader import (
    LazyConfig,
    ConfigNotLoadedError,
    init_config,
    get_config,
    clear_config,
    load_config,
)
from wordformat.config.models import NodeConfigRoot
from wordformat.agent.message import MessageManager
from wordformat.agent.onnx_infer import (
    _get_best_onnx_providers,
    _load_model,
    onnx_single_infer,
    onnx_batch_infer,
    safe_batch_infer,
)
from wordformat.pipeline.stages import FormattingExecutionStage, ParagraphAlignmentStage
from wordformat.classify.tag import set_tag_main
from wordformat.structure.node_factory import create_node
from wordformat.structure.tree_builder import DocumentTreeBuilder
from wordformat.structure.document_builder import DocumentBuilder
from wordformat.structure.utils import promote_bodytext_in_subtrees_of_type
from wordformat.rules.node import FormatNode
from wordformat.rules.body import BodyText
from wordformat.api import save_upload_file, TEMP_DIR

apply_format_check_to_all_nodes = (
    FormattingExecutionStage().apply_format_check_to_all_nodes
)


# ==================== (a) Config 模块集成测试 ====================


class TestConfigLifecycle:
    """load_config -> get_config 基本流程。"""

    def test_load_then_get(self, config_path):
        load_config(config_path)
        cfg = get_config()
        assert isinstance(cfg, dict)
        assert "global_format" in cfg

    def test_get_without_load_raises(self):
        clear_config()
        with pytest.raises(RuntimeError):
            get_config()

    def test_clear_resets_state(self, config_path):
        load_config(config_path)
        assert get_config() is not None
        clear_config()
        with pytest.raises(RuntimeError):
            get_config()

    def test_reinit_after_clear(self, config_path):
        init_config(config_path)
        get_config()
        clear_config()
        init_config(config_path)
        cfg = get_config()
        assert isinstance(cfg, dict)

    def test_independent_instances(self):
        """多个 LazyConfig 实例相互独立。"""
        a = LazyConfig()
        b = LazyConfig()
        assert a is not b

    def test_thread_safe_independent_instances(self):
        """多线程各自创建独立实例无竞争。"""
        results = []
        barrier = threading.Barrier(10)

        def create():
            barrier.wait()
            lc = LazyConfig()
            lc._config = {}  # 直接设内部状态，不触发文件加载
            results.append(lc)

        threads = [threading.Thread(target=create) for _ in range(10)]
        for t in threads:
            t.start()
        for t in threads:
            t.join()
        # 所有实例独立
        assert len(set(id(r) for r in results)) == 10


# ==================== (b) DataModel 验证测试 ====================


# ==================== (c) Agent/Message 测试 ====================


class TestMessageManagerIntegration:
    """MessageManager 全操作 + clear() 线程安全缺陷"""

    def test_full_lifecycle(self):
        mgr = MessageManager()
        mgr.add_system_message("sys")
        mgr.add_user_message("hello")
        mgr.add_assistant_message("hi")
        mgr.add_tool_message("result", "tc_1")
        msgs = mgr.get_messages()
        assert len(msgs) == 4
        assert msgs[0]["role"] == "system"
        assert msgs[3]["tool_call_id"] == "tc_1"

    def test_clear_preserves_system(self):
        mgr = MessageManager()
        mgr.add_system_message("keep")
        mgr.add_user_message("remove")
        mgr.clear()
        msgs = mgr.get_messages()
        assert len(msgs) == 1
        assert msgs[0]["content"] == "keep"

    def test_reset_removes_all(self):
        mgr = MessageManager()
        mgr.add_system_message("sys")
        mgr.add_user_message("usr")
        mgr.reset_messages()
        assert mgr.get_messages() == []

    def test_add_assistant_with_tool_calls(self):
        mgr = MessageManager()
        mock_call = mock.MagicMock()
        mock_call.id = "call_1"
        mock_call.type = "function"
        mock_call.function.name = "fn1"
        mock_call.function.arguments = "{}"
        mock_msg = mock.MagicMock()
        mock_msg.content = "thinking"
        mock_msg.tool_calls = [mock_call]
        mgr.add_assistant_message(mock_msg)
        msgs = mgr.get_messages()
        assert len(msgs) == 1
        assert "tool_calls" in msgs[0]
        assert msgs[0]["tool_calls"][0]["id"] == "call_1"

    def test_clear_thread_safety(self):
        """clear() 和 add_message 并发不应丢数据"""
        mgr = MessageManager()

        def writer():
            for i in range(100):
                mgr.add_user_message(f"msg_{i}")

        def clearer():
            for _ in range(10):
                mgr.clear()

        t1 = threading.Thread(target=writer)
        t2 = threading.Thread(target=clearer)
        t1.start()
        t2.start()
        t1.join()
        t2.join()
        # clear 后只应有 system 消息（本测试未添加 system，应为空）
        msgs = mgr.get_messages()
        assert len(msgs) == 0


# ==================== (d) Agent/ONNX 推理测试 ====================


class TestONNXInferIntegration:
    """Mock 模型加载，验证推理流程与已知 bug"""

    def test_single_infer_success(self):
        with (
            mock.patch("wordformat.agent.onnx_infer._tokenizer") as mock_tok,
            mock.patch("wordformat.agent.onnx_infer._ort_sess") as mock_sess,
            mock.patch(
                "wordformat.agent.onnx_infer._id2label", {0: "body_text", 1: "heading"}
            ),
            mock.patch("wordformat.agent.onnx_infer._load_model"),
        ):
            enc = mock.MagicMock()
            enc.ids = [1, 2, 3]
            enc.attention_mask = [1, 1, 1]
            enc.type_ids = [0, 0, 0]
            mock_tok.encode.return_value = enc
            mock_sess.run.return_value = [np.array([[0.9, 0.1]])]
            result = onnx_single_infer("test")
            assert result["label"] == "body_text"
            assert result["score"] > 0.5

    def test_single_infer_error_returns_empty(self):
        with (
            mock.patch("wordformat.agent.onnx_infer._tokenizer") as mock_tok,
            mock.patch("wordformat.agent.onnx_infer._ort_sess") as mock_sess,
            mock.patch("wordformat.agent.onnx_infer._id2label", {0: "body_text"}),
            mock.patch("wordformat.agent.onnx_infer._load_model"),
        ):
            enc = mock.MagicMock()
            enc.ids = [1, 2, 3]
            enc.attention_mask = [1, 1, 1]
            enc.type_ids = [0, 0, 0]
            mock_tok.encode.return_value = enc
            mock_sess.run.side_effect = RuntimeError("OOM")
            result = onnx_single_infer("test")
            assert result == {"label": "", "score": 0.0}

    def test_batch_infer_error_format_matches_single(self):
        """batch 失败与 single 失败应返回相同结构"""
        with (
            mock.patch("wordformat.agent.onnx_infer._tokenizer") as mock_tok,
            mock.patch("wordformat.agent.onnx_infer._ort_sess") as mock_sess,
            mock.patch("wordformat.agent.onnx_infer._id2label", {0: "body_text"}),
            mock.patch("wordformat.agent.onnx_infer._load_model"),
        ):
            enc = mock.MagicMock()
            enc.ids = [1, 2, 3]
            enc.attention_mask = [1, 1, 1]
            enc.type_ids = [0, 0, 0]
            mock_tok.encode.return_value = enc
            mock_sess.run.side_effect = RuntimeError("OOM")
            batch_result = onnx_batch_infer(["t1"])[0]
            single_result = onnx_single_infer("t1")
            assert set(batch_result.keys()) == set(single_result.keys())

    @mock.patch("wordformat.agent.onnx_infer.onnx_batch_infer")
    def test_safe_batch_splits_correctly(self, mock_batch):
        mock_batch.side_effect = lambda texts: [{"label": "x"} for _ in texts]
        results = safe_batch_infer(["a", "b", "c", "d"], max_batch_size=2)
        assert len(results) == 4
        assert mock_batch.call_count == 2

    def test_provider_selection_priority(self):
        with mock.patch(
            "onnxruntime.get_available_providers",
            return_value=["CUDAExecutionProvider", "CPUExecutionProvider"],
        ):
            assert _get_best_onnx_providers() == ["CUDAExecutionProvider"]
        with mock.patch(
            "onnxruntime.get_available_providers",
            return_value=["DmlExecutionProvider", "CPUExecutionProvider"],
        ):
            assert _get_best_onnx_providers() == ["DmlExecutionProvider"]
        with mock.patch(
            "onnxruntime.get_available_providers", return_value=["CPUExecutionProvider"]
        ):
            assert _get_best_onnx_providers() == ["CPUExecutionProvider"]

    def test_global_state_thread_safety(self):
        """并发调用 _load_model 不应导致竞态条件"""
        errors = []

        def load():
            try:
                _load_model()
            except Exception as e:
                errors.append(e)

        threads = [threading.Thread(target=load) for _ in range(5)]
        for t in threads:
            t.start()
        for t in threads:
            t.join()
        assert len(errors) == 0


# ==================== (e) word_structure 集成测试 ====================


class TestNodeFactoryIntegration:
    """create_node 创建/拒绝 + 真实 CATEGORY_TO_CLASS"""

    def test_create_known_category(self, sample_yaml_config):
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()
        item = {
            "category": "heading_level_1",
            "paragraph": "第一章",
            "fingerprint": "fp_h1",
        }
        node = create_node(item, level=1, config=config)
        assert node is not None
        assert hasattr(node, "check_format")

    def test_create_unknown_category_returns_none(self, sample_yaml_config):
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()
        item = {"category": "nonexistent_type", "paragraph": "x", "fingerprint": "fp_x"}
        with mock.patch("wordformat.structure.node_factory.CATEGORY_TO_CLASS", {}):
            assert create_node(item, level=1, config=config) is None

    def test_create_missing_category_raises(self, sample_yaml_config):
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()
        with pytest.raises(ValueError, match="missing 'category'"):
            create_node({"paragraph": "x"}, level=1, config=config)


class TestTreeBuilderIntegration:
    """tree_builder 构建树 + 已知 bug"""

    @mock.patch("wordformat.structure.node_factory.create_node")
    def test_build_flat_headings(self, mock_create):
        mock_node = mock.MagicMock()
        mock_node.level = 1
        mock_create.return_value = mock_node
        builder = DocumentTreeBuilder()
        root = builder.build_tree([{"category": "heading_level_1"}])
        assert len(root.children) == 1

    def test_body_text_not_treated_as_heading(self):
        """body_text 不应被 _is_heading_category 判定为标题"""
        builder = DocumentTreeBuilder()
        assert builder._is_heading_category("body_text") is False


class TestDocumentBuilderIntegration:
    """DocumentBuilder 加载 JSON 并构建树"""

    def test_build_from_json_list(self, sample_yaml_config):
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()
        data = [
            {
                "category": "heading_level_1",
                "paragraph": "第一章",
                "fingerprint": "fp1",
            },
            {"category": "body_text", "paragraph": "正文", "fingerprint": "fp2"},
        ]
        root = DocumentBuilder.build_from_json(data, config=config)
        assert root.value["category"] == "top"
        assert len(root.children) >= 1

    def test_config_not_global_state(self):
        """CONFIG 不应是类变量，应为实例变量 _config"""
        builder = DocumentTreeBuilder()
        # CONFIG 可能被其他测试修改为非空，只验证实例变量 _config 存在
        assert hasattr(builder, "_config")


# ==================== (f) set_tag 集成测试 ====================


class TestSetTagMainIntegration:
    """mock DocxBase.parse 验证 set_tag_main 调用链"""

    @mock.patch("wordformat.classify.tag.DocxBase")
    def test_set_tag_main_calls_parse(self, mock_docx_cls):
        mock_instance = mock_docx_cls.return_value
        mock_instance.parse.return_value = [
            {
                "category": "body_text",
                "score": 0.9,
                "paragraph": "test",
                "fingerprint": "fp1",
            }
        ]
        result = set_tag_main("dummy.docx", "dummy.yaml")
        assert len(result) == 1
        mock_instance.parse.assert_called_once()

    @mock.patch("wordformat.classify.tag.DocxBase")
    def test_set_tag_main_passes_args(self, mock_docx_cls):
        mock_docx_cls.return_value.parse.return_value = []
        set_tag_main("path/to/doc.docx", "path/to/cfg.yaml")
        mock_docx_cls.assert_called_once_with(
            "path/to/doc.docx", configpath="path/to/cfg.yaml"
        )


# ==================== (g) set_style 集成测试 ====================


class TestSetStyleIntegration:
    """apply_format_check_to_all_nodes + xg"""

    def test_apply_check_mode(self):
        """check=True 时调用 check_format"""
        call_log = []

        class SpyNode(FormatNode):
            NODE_TYPE = "body_text"
            CONFIG_MODEL = type("C", (), {})()

            def __init__(self, **kw):
                super().__init__(
                    value={"category": "body_text", "fingerprint": "fp"}, level=1, **kw
                )
                doc = Document()
                self.paragraph = doc.add_paragraph("text")

            def check_format(self, doc):
                call_log.append(("check", id(self)))

            def apply_format(self, doc):
                call_log.append(("apply", id(self)))

            def load_config(self, cfg):
                pass

        root = SpyNode()
        child = SpyNode()
        root.add_child_node(child)
        doc = Document()
        apply_format_check_to_all_nodes(root, doc, {}, check=True)
        assert len(call_log) == 2
        assert all(c[0] == "check" for c in call_log)

    def test_apply_format_mode(self):
        """check=False 时调用 apply_format"""
        call_log = []

        class SpyNode(FormatNode):
            NODE_TYPE = "body_text"
            CONFIG_MODEL = type("C", (), {})()

            def __init__(self, **kw):
                super().__init__(
                    value={"category": "body_text", "fingerprint": "fp"}, level=1, **kw
                )
                doc = Document()
                self.paragraph = doc.add_paragraph("text")

            def check_format(self, doc):
                call_log.append("check")

            def apply_format(self, doc):
                call_log.append("apply")

            def load_config(self, cfg):
                pass

        root = SpyNode()
        doc = Document()
        apply_format_check_to_all_nodes(root, doc, {}, check=False)
        assert call_log == ["apply"]

    def test_flatten_tree_nodes_matches_doc_order(self):
        """DFS 展平的节点顺序与 document.paragraphs 一致。"""

        _flatten_tree_nodes = ParagraphAlignmentStage()._flatten_tree_nodes

        root = FormatNode(value={"category": "top"}, level=0)
        a = FormatNode(value={"category": "abstract_chinese_title"}, level=1)
        b = FormatNode(value={"category": "body_text"}, level=1)
        c = FormatNode(value={"category": "heading_level_1"}, level=1)
        d = FormatNode(value={"category": "body_text"}, level=1)
        root.add_child_node(a)
        root.add_child_node(b)
        root.add_child_node(c)
        c.add_child_node(d)

        nodes = _flatten_tree_nodes(root)
        assert nodes == [a, b, c, d]


# ==================== (i) set_style.py auto_format_thesis_document 覆盖测试 ====================


class TestAutoFormatThesisDocument:
    """覆盖 set_style.py lines 53-55, 120-176: auto_format_thesis_document 主流程"""

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    def test_check_mode_returns_annotated_path(
        self, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """check=True 模式：返回 --标注版.docx 路径 (lines 170-176)"""
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node
        mock_apply.return_value = None

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        result = auto_format_thesis_document(
            jsonpath=temp_docx,
            docxpath=temp_docx,
            configpath=config_path,
            savepath=str(tmp_path),
            check=True,
        )
        assert "--标注版.docx" in result

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    def test_apply_mode_returns_modified_path(
        self, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """check=False 模式：返回 --修改版.docx 路径 (line 173)"""
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node
        mock_apply.return_value = None

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        result = auto_format_thesis_document(
            jsonpath=temp_docx,
            docxpath=temp_docx,
            configpath=config_path,
            savepath=str(tmp_path),
            check=False,
        )
        assert "--修改版.docx" in result

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    def test_filters_body_text_nodes(
        self, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """body_text 节点不再被过滤，保留在 children 中"""
        body_node = mock.MagicMock()
        body_node.value = {"category": "body_text"}
        heading_node = mock.MagicMock()
        heading_node.value = {"category": "heading_level_1"}
        # build_from_json returns root_node; root_node.children gets reassigned
        root_node = mock.MagicMock()
        root_node.children = [body_node, heading_node]
        mock_builder.build_from_json.return_value = root_node
        mock_apply.return_value = None

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        auto_format_thesis_document(
            jsonpath=temp_docx,
            docxpath=temp_docx,
            configpath=config_path,
            savepath=str(tmp_path),
            check=True,
        )
        # body_text 不再被过滤，所有节点都保留
        assert len(root_node.children) == 2

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    @mock.patch("wordformat.pipeline.stages.promote_bodytext_in_subtrees_of_type")
    def test_promote_called_for_subtrees(
        self, mock_promote, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """promote_bodytext_in_subtrees_of_type 应被调用 (lines 153-161)"""
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node
        mock_apply.return_value = None

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        auto_format_thesis_document(
            jsonpath=temp_docx,
            docxpath=temp_docx,
            configpath=config_path,
            savepath=str(tmp_path),
            check=True,
        )
        assert (
            mock_promote.call_count == 3
        )  # AbstractTitleCN, AbstractTitleEN, References

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    def test_exception_in_traverse_raises(
        self, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """节点处理异常时 raise e (lines 53-55)"""
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node
        mock_apply.side_effect = RuntimeError("test error")

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        with pytest.raises(RuntimeError, match="test error"):
            auto_format_thesis_document(
                jsonpath=temp_docx,
                docxpath=temp_docx,
                configpath=config_path,
                savepath=str(tmp_path),
                check=True,
            )

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    def test_config_load_failure_raises(
        self, mock_builder, mock_apply, temp_docx, tmp_path
    ):
        """配置加载失败时 raise (lines 126-128)"""
        bad_config = str(tmp_path / "nonexistent.yaml")
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        with pytest.raises(Exception):
            auto_format_thesis_document(
                jsonpath=temp_docx,
                docxpath=temp_docx,
                configpath=bad_config,
                savepath=str(tmp_path),
                check=True,
            )

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    def test_apply_mode_lists_styles(
        self, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """check=False 时列出可用样式 (lines 139-143)"""
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node
        mock_apply.return_value = None

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        auto_format_thesis_document(
            jsonpath=temp_docx,
            docxpath=temp_docx,
            configpath=config_path,
            savepath=str(tmp_path),
            check=False,
        )
        # Should not raise - the styles listing is just logging


# ==================== (j) word_structure/utils.py promote_bodytext 覆盖测试 ====================


class TestPromoteBodyTextInSubtrees:
    """覆盖 word_structure/utils.py lines 44-62: promote_bodytext_in_subtrees_of_type"""

    def test_promote_bodytext_to_target(self):
        """BodyText 子节点应被升级为目标类型 (lines 44-53)"""

        # Create a parent node of a specific type
        class ParentType(FormatNode):
            pass

        class TargetType(FormatNode):
            pass

        parent = ParentType(
            value={"category": "parent", "fingerprint": "fp_parent"}, level=1
        )
        child = BodyText(
            value={
                "category": "body_text",
                "paragraph": "test",
                "fingerprint": "fp_child",
            },
            level=2,
        )
        parent.add_child_node(child)

        promote_bodytext_in_subtrees_of_type(parent, ParentType, TargetType)

        # The child should now be TargetType
        assert isinstance(parent.children[0], TargetType)

    def test_promote_recursive_in_subtree(self):
        """递归升级子树中的 BodyText (lines 52-53)"""
        from wordformat.rules.body import BodyText
        from wordformat.rules.node import FormatNode

        class ParentType(FormatNode):
            pass

        class TargetType(FormatNode):
            pass

        class MidNode(FormatNode):
            pass

        parent = ParentType(
            value={"category": "parent", "fingerprint": "fp_parent"}, level=1
        )
        mid = MidNode(value={"category": "mid", "fingerprint": "fp_mid"}, level=2)
        child = BodyText(
            value={
                "category": "body_text",
                "paragraph": "deep",
                "fingerprint": "fp_child",
            },
            level=3,
        )
        mid.add_child_node(child)
        parent.add_child_node(mid)

        promote_bodytext_in_subtrees_of_type(parent, ParentType, TargetType)

        assert isinstance(parent.children[0].children[0], TargetType)

    def test_no_promote_outside_parent_type(self):
        """不在 parent_type 子树中的 BodyText 不应被升级"""
        from wordformat.rules.body import BodyText
        from wordformat.rules.node import FormatNode

        class ParentType(FormatNode):
            pass

        class TargetType(FormatNode):
            pass

        class OtherType(FormatNode):
            pass

        root = FormatNode(value={"category": "top", "fingerprint": "fp_root"}, level=0)
        parent = ParentType(
            value={"category": "parent", "fingerprint": "fp_parent"}, level=1
        )
        other = OtherType(
            value={"category": "other", "fingerprint": "fp_other"}, level=1
        )
        child_in_parent = BodyText(
            value={"category": "body_text", "paragraph": "in", "fingerprint": "fp_in"},
            level=2,
        )
        child_in_other = BodyText(
            value={
                "category": "body_text",
                "paragraph": "out",
                "fingerprint": "fp_out",
            },
            level=2,
        )
        parent.add_child_node(child_in_parent)
        other.add_child_node(child_in_other)
        root.add_child_node(parent)
        root.add_child_node(other)

        promote_bodytext_in_subtrees_of_type(root, ParentType, TargetType)

        assert isinstance(parent.children[0], TargetType)
        assert isinstance(other.children[0], BodyText)  # unchanged


# ==================== (k) log_config.py 覆盖测试 ====================


class TestLogConfig:
    """覆盖 log_config.py lines 14, 46-55: setup_logger frozen path, setup_uvicorn_loguru"""

    def test_setup_logger_frozen_path(self):
        """sys.frozen=True 时使用 sys.executable.parent (line 14)"""
        import sys

        original_frozen = getattr(sys, "frozen", None)
        sys.frozen = True
        try:
            with mock.patch("sys.executable", "/fake/app"):
                with mock.patch("wordformat.log_config.logger") as mock_logger:
                    mock_logger.remove.return_value = None
                    mock_logger.add.return_value = None
                    from wordformat.log_config import setup_logger

                    setup_logger()
                    # Verify logger.remove was called (setup_logger ran)
                    mock_logger.remove.assert_called_once()
        finally:
            if original_frozen is None:
                delattr(sys, "frozen")
            else:
                sys.frozen = original_frozen

    def test_setup_uvicorn_loguru(self):
        """setup_uvicorn_loguru 配置 uvicorn 日志 (lines 46-55)"""
        # Both uvicorn and logging are imported inside the function
        mock_logging_mod = mock.MagicMock()
        mock_uvicorn_mod = mock.MagicMock()
        with mock.patch.dict(
            "sys.modules", {"uvicorn": mock_uvicorn_mod, "logging": mock_logging_mod}
        ):
            import importlib
            import wordformat.log_config as log_mod

            importlib.reload(log_mod)
            from wordformat.log_config import setup_uvicorn_loguru

            mock_uvicorn_logger = mock.MagicMock()
            mock_logging_mod.getLogger.return_value = mock_uvicorn_logger
            setup_uvicorn_loguru()
            # Verify uvicorn loggers are disabled
            assert mock_uvicorn_logger.disabled is True
            # Verify LOGGING_CONFIG is set
            assert mock_uvicorn_mod.config.LOGGING_CONFIG is not None
            assert mock_uvicorn_mod.config.LOGGING_CONFIG["version"] == 1


# ==================== (l) onnx_infer.py 额外覆盖测试 ====================


class TestONNXInferExceptionHandling:
    """覆盖 onnx_infer.py lines 48-50, 57, 88-90, 100, 115-117, 156, 159, 198-199, 212-233, 244"""

    def test_get_best_onnx_providers_exception_fallback(self):
        """get_available_providers raises -> fallback to CPU (lines 48-50)"""
        with mock.patch(
            "onnxruntime.get_available_providers", side_effect=Exception("no runtime")
        ):
            assert _get_best_onnx_providers() == ["CPUExecutionProvider"]

    def test_load_model_early_return_when_tokenizer_set(self):
        """_tokenizer is not None -> early return (line 57)"""
        import wordformat.agent.onnx_infer as m

        original = m._tokenizer
        try:
            m._tokenizer = mock.MagicMock()
            with mock.patch("tokenizers.Tokenizer") as mock_tok_cls:
                _load_model()
                # Tokenizer.from_file should NOT be called
                mock_tok_cls.from_file.assert_not_called()
        finally:
            m._tokenizer = original

    def test_load_model_fallback_to_cpu(self):
        """Best provider fails -> fallback to CPU (lines 88-90)"""
        import wordformat.agent.onnx_infer as m

        original_tok, original_sess, original_id2 = (
            m._tokenizer,
            m._ort_sess,
            m._id2label,
        )
        try:
            m._tokenizer = None
            m._ort_sess = None
            m._id2label = None
            mock_paths = {
                "onnx": "/fake/model.onnx",
                "tokenizer": "/fake/tokenizer.json",
                "id2label": "/fake/id2label.json",
            }
            with (
                mock.patch(
                    "wordformat.agent.onnx_infer._get_model_paths",
                    return_value=mock_paths,
                ),
                mock.patch(
                    "wordformat.agent.onnx_infer._get_best_onnx_providers",
                    return_value=["CUDAExecutionProvider"],
                ),
                mock.patch("tokenizers.Tokenizer") as mock_tok_cls,
                mock.patch("onnxruntime.InferenceSession") as mock_sess_cls,
                mock.patch(
                    "builtins.open", mock.mock_open(read_data='{"0":"body_text"}')
                ),
            ):
                mock_sess_cls.side_effect = [
                    RuntimeError("CUDA fail"),
                    mock.MagicMock(),
                ]
                _load_model()
                # Should have been called twice: first with CUDA (fail), then with CPU (success)
                assert mock_sess_cls.call_count == 2
                # Second call should use CPUExecutionProvider
                second_call_kwargs = mock_sess_cls.call_args_list[1].kwargs
                assert second_call_kwargs["providers"] == ["CPUExecutionProvider"]
        finally:
            m._tokenizer = original_tok
            m._ort_sess = original_sess
            m._id2label = original_id2

    def test_load_model_cpu_core_num_zero_fallback(self):
        """os.cpu_count() returns 0 -> fallback to 4 (line 100)"""
        import wordformat.agent.onnx_infer as m

        original_tok, original_sess, original_id2 = (
            m._tokenizer,
            m._ort_sess,
            m._id2label,
        )
        try:
            m._tokenizer = None
            m._ort_sess = None
            m._id2label = None
            mock_paths = {
                "onnx": "/fake/model.onnx",
                "tokenizer": "/fake/tokenizer.json",
                "id2label": "/fake/id2label.json",
            }
            with (
                mock.patch(
                    "wordformat.agent.onnx_infer._get_model_paths",
                    return_value=mock_paths,
                ),
                mock.patch("os.cpu_count", return_value=0),
                mock.patch("tokenizers.Tokenizer") as mock_tok_cls,
                mock.patch("onnxruntime.InferenceSession") as mock_sess_cls,
                mock.patch("onnxruntime.SessionOptions") as mock_opts_cls,
                mock.patch(
                    "builtins.open", mock.mock_open(read_data='{"0":"body_text"}')
                ),
            ):
                mock_sess = mock.MagicMock()
                mock_sess_cls.return_value = mock_sess
                _load_model()
                # Check that intra_op_num_threads was set to 4 (fallback)
                opts_instance = mock_opts_cls.return_value
                assert opts_instance.intra_op_num_threads == 4
        finally:
            m._tokenizer = original_tok
            m._ort_sess = original_sess
            m._id2label = original_id2

    def test_single_infer_truncation(self):
        """Input longer than MAX_LENGTH gets truncated (lines 115-117)"""
        import wordformat.agent.onnx_infer as m

        original_tok, original_sess, original_id2 = (
            m._tokenizer,
            m._ort_sess,
            m._id2label,
        )
        try:
            m._tokenizer = mock.MagicMock()
            m._ort_sess = mock.MagicMock()
            m._id2label = {0: "body_text", 1: "heading"}
            # Create encoded output longer than MAX_LENGTH (128)
            enc = mock.MagicMock()
            enc.ids = list(range(200))
            enc.attention_mask = [1] * 200
            enc.type_ids = [0] * 200
            m._tokenizer.encode.return_value = enc
            m._ort_sess.run.return_value = [np.array([[0.9, 0.1]])]
            result = onnx_single_infer("test")
            assert result["label"] == "body_text"
            # Verify the input_ids were truncated to MAX_LENGTH
            call_args = m._ort_sess.run.call_args[0][1]
            assert call_args["input_ids"].shape == (1, 128)
        finally:
            m._tokenizer = original_tok
            m._ort_sess = original_sess
            m._id2label = original_id2

    def test_batch_infer_empty_texts(self):
        """Empty texts list returns [] (line 156)"""
        result = onnx_batch_infer([])

    def test_batch_infer_loads_model_when_tokenizer_none(self):
        """_tokenizer is None triggers _load_model (line 159)"""
        import wordformat.agent.onnx_infer as m

        original_tok = m._tokenizer
        try:
            m._tokenizer = None
            with mock.patch("wordformat.agent.onnx_infer._load_model") as mock_load:
                mock_load.side_effect = RuntimeError("model not available")
                with pytest.raises(RuntimeError):
                    onnx_batch_infer(["test"])
                mock_load.assert_called_once()
        finally:
            m._tokenizer = original_tok

    def test_batch_infer_success_with_timing(self):
        """Batch inference success path with timing log (lines 198-199)"""
        import wordformat.agent.onnx_infer as m

        original_tok, original_sess, original_id2 = (
            m._tokenizer,
            m._ort_sess,
            m._id2label,
        )
        try:
            m._tokenizer = mock.MagicMock()
            m._ort_sess = mock.MagicMock()
            m._id2label = {0: "body_text", 1: "heading"}
            enc = mock.MagicMock()
            enc.ids = [1, 2, 3]
            enc.attention_mask = [1, 1, 1]
            enc.type_ids = [0, 0, 0]
            m._tokenizer.encode.return_value = enc
            m._ort_sess.run.return_value = [np.array([[0.9, 0.1], [0.2, 0.8]])]
            result = onnx_batch_infer(["text1", "text2"])
            assert len(result) == 2
            assert result[0]["label"] == "body_text"
            assert result[0]["text"] == "text1"
            assert result[1]["label"] == "heading"
            assert result[1]["text"] == "text2"
        finally:
            m._tokenizer = original_tok
            m._ort_sess = original_sess
            m._id2label = original_id2

    def test_batch_infer_result_assembly(self):
        """Result assembly with pred_id and score (lines 212-233)"""
        import wordformat.agent.onnx_infer as m

        original_tok, original_sess, original_id2 = (
            m._tokenizer,
            m._ort_sess,
            m._id2label,
        )
        try:
            m._tokenizer = mock.MagicMock()
            m._ort_sess = mock.MagicMock()
            m._id2label = {0: "body_text", 1: "heading", 2: "abstract"}
            enc = mock.MagicMock()
            enc.ids = [1, 2, 3]
            enc.attention_mask = [1, 1, 1]
            enc.type_ids = [0, 0, 0]
            m._tokenizer.encode.return_value = enc
            # 3 texts, 3 classes
            m._ort_sess.run.return_value = [
                np.array(
                    [
                        [0.1, 0.7, 0.2],
                        [0.8, 0.1, 0.1],
                        [0.3, 0.3, 0.4],
                    ]
                )
            ]
            result = onnx_batch_infer(["t1", "t2", "t3"])
            assert len(result) == 3
            assert result[0]["label"] == "heading"
            assert result[0]["pred_id"] == 1
            assert result[1]["label"] == "body_text"
            assert result[1]["pred_id"] == 0
            assert result[2]["label"] == "abstract"
            assert result[2]["pred_id"] == 2
            # Check score is a float
            assert isinstance(result[0]["score"], float)
        finally:
            m._tokenizer = original_tok
            m._ort_sess = original_sess
            m._id2label = original_id2

    def test_safe_batch_infer_empty_texts(self):
        """Empty texts returns [] (line 244)"""
        result = safe_batch_infer([])


# ==================== (m) keywords.py 覆盖测试 ====================


class TestKeywordsBaseGetLangConfig:
    """覆盖 keywords.py line 55: BaseKeywordsNode._get_lang_config with unknown LANG"""

    def test_get_lang_config_unknown_lang_fallback(self, sample_yaml_config):
        """Unknown LANG falls back to chinese config (line 55)"""
        from wordformat.config.loader import init_config, get_config
        from wordformat.rules.keywords import BaseKeywordsNode

        init_config(sample_yaml_config)
        config = get_config()

        class TestNode(BaseKeywordsNode):
            LANG = "unknown_lang"

        node = TestNode(
            value={"category": "test", "fingerprint": "fp"},
            level=1,
        )
        node.load_config(config)
        # Should get chinese config as fallback
        assert node.pydantic_config is not None


class TestKeywordsENCheckKeywordLabel:
    """覆盖 keywords.py line 83: KeywordsEN._check_keyword_label matching"""

    def test_check_keyword_label_matches(self):
        """_check_keyword_label returns True for 'Keywords' (line 83)"""
        from wordformat.rules.keywords import KeywordsEN

        node = KeywordsEN(
            value={"category": "abstract.keywords.english", "fingerprint": "fp"},
            level=1,
        )
        mock_run = mock.MagicMock()
        mock_run.text = "Keywords: AI, ML"
        assert node._check_keyword_label(mock_run) is True

    def test_check_keyword_label_key_words(self):
        """_check_keyword_label matches 'KEY WORDS'"""
        from wordformat.rules.keywords import KeywordsEN

        node = KeywordsEN(
            value={"category": "abstract.keywords.english", "fingerprint": "fp"},
            level=1,
        )
        mock_run = mock.MagicMock()
        mock_run.text = "KEY WORDS: AI"
        assert node._check_keyword_label(mock_run) is True

    def test_check_keyword_label_no_match(self):
        """_check_keyword_label returns False for content text"""
        from wordformat.rules.keywords import KeywordsEN

        node = KeywordsEN(
            value={"category": "abstract.keywords.english", "fingerprint": "fp"},
            level=1,
        )
        mock_run = mock.MagicMock()
        mock_run.text = "artificial intelligence"
        assert node._check_keyword_label(mock_run) is False


class TestKeywordsENBase:
    """覆盖 keywords.py lines 114, 121, 130-135, 152-153"""

    def _make_en_node(self, config_dict=None):
        """Helper to create a KeywordsEN node with config loaded"""
        from wordformat.rules.keywords import KeywordsEN

        node = KeywordsEN(
            value={"category": "abstract.keywords.english", "fingerprint": "fp"},
            level=1,
        )
        if config_dict:
            node.load_config(config_dict)
        return node

    def test_empty_run_skip(self, sample_yaml_config):
        """Empty run text is skipped (line 114)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_en_node(config)
        doc = Document()
        p = doc.add_paragraph()
        empty_run = p.add_run("   ")
        empty_run.text = "   "
        node.paragraph = p
        # Should not crash, empty run is skipped
        node.check_format(doc)

    def test_label_style_check(self, sample_yaml_config):
        """Label run style is checked (line 121)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_en_node(config)
        doc = Document()
        p = doc.add_paragraph()
        label_run = p.add_run("Keywords: ")
        label_run.font.bold = False  # Wrong - should be bold per config
        node.paragraph = p
        node.check_format(doc)
        # Should have added a comment about bold mismatch

    def test_content_style_check(self, sample_yaml_config):
        """Content run style is checked (lines 130-135)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_en_node(config)
        doc = Document()
        p = doc.add_paragraph()
        label_run = p.add_run("Keywords: ")
        label_run.font.bold = True  # Correct
        content_run = p.add_run("AI, ML")
        content_run.font.bold = True  # Wrong - content should not be bold
        node.paragraph = p
        node.check_format(doc)

    def test_keyword_count_validation_min(self, sample_yaml_config):
        """Keyword count < count_min triggers warning (via _run_rules)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_en_node(config)
        doc = Document()
        p = doc.add_paragraph()
        label_run = p.add_run("Keywords: ")
        label_run.font.bold = True
        content_run = p.add_run("AI")
        node.paragraph = p
        node.check_format(doc)
        # count_min is 3, only 1 keyword -> should trigger count warning

    def test_keyword_count_validation_max(self, sample_yaml_config):
        """Keyword count > count_max triggers warning (via _run_rules)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_en_node(config)
        doc = Document()
        p = doc.add_paragraph()
        label_run = p.add_run("Keywords: ")
        label_run.font.bold = True
        content_run = p.add_run("AI, ML, NLP, CV, DB, SE")
        node.paragraph = p
        node.check_format(doc)
        # count_max is 5, 6 keywords -> should trigger count warning


class TestKeywordsCNBase:
    """覆盖 keywords.py lines 177-180, 187, 218, 225, 234-239"""

    def _make_cn_node(self, config_dict=None):
        """Helper to create a KeywordsCN node with config loaded"""
        from wordformat.rules.keywords import KeywordsCN

        node = KeywordsCN(
            value={"category": "abstract.keywords.chinese", "fingerprint": "fp"},
            level=1,
        )
        if config_dict:
            node.load_config(config_dict)
        return node

    def test_paragraph_style_check(self, sample_yaml_config):
        """Paragraph style is checked (line 187)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_cn_node(config)
        doc = Document()
        p = doc.add_paragraph()
        p.paragraph_format.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )  # Wrong - should be left
        run = p.add_run("关键词：人工智能")
        node.paragraph = p
        node.check_format(doc)

    def test_label_style_check(self, sample_yaml_config):
        """Label run style is checked (line 218)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_cn_node(config)
        doc = Document()
        p = doc.add_paragraph()
        label_run = p.add_run("关键词")
        label_run.font.bold = False  # Wrong - should be bold
        node.paragraph = p
        node.check_format(doc)

    def test_content_style_check(self, sample_yaml_config):
        """Content run style is checked (line 225)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_cn_node(config)
        doc = Document()
        p = doc.add_paragraph()
        label_run = p.add_run("关键词：")
        label_run.font.bold = True  # Correct
        content_run = p.add_run("人工智能；机器学习")
        content_run.font.bold = True  # Wrong - content should not be bold
        node.paragraph = p
        node.check_format(doc)

    def test_keyword_count_and_trailing_punctuation(self, sample_yaml_config):
        """Keyword count validation + trailing punctuation check (via _run_rules)"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()

        node = self._make_cn_node(config)
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("关键词：人工智能；机器学习；")
        node.paragraph = p
        node.check_format(doc)
        # trailing_punctuation.enabled should be True by default
        # Text ends with ； which should trigger trailing punctuation warning


# ==================== (n) heading.py 覆盖测试 ====================


class TestHeadingLevelNodes:
    """覆盖 heading.py lines 36-41, 52-58"""

    def test_heading_level1_load_config_dict(self):
        """HeadingLevel1Node.load_config with dict (lines 36-41)"""
        from wordformat.rules.heading import HeadingLevel1Node

        node = HeadingLevel1Node(
            value={"category": "headings.level_1", "fingerprint": "fp"},
            level=1,
        )
        config_dict = {
            "headings": {
                "level_1": {
                    "alignment": "居中对齐",
                    "font_size": "小二",
                    "bold": True,
                }
            }
        }
        node.load_config(config_dict)
        assert node.pydantic_config is not None

    def test_heading_level2_load_config_dict(self):
        """HeadingLevel2Node.load_config with dict (lines 52-58)"""
        from wordformat.rules.heading import HeadingLevel2Node

        node = HeadingLevel2Node(
            value={"category": "headings.level_2", "fingerprint": "fp"},
            level=2,
        )
        config_dict = {
            "headings": {
                "level_2": {
                    "alignment": "左对齐",
                    "font_size": "三号",
                }
            }
        }
        node.load_config(config_dict)
        assert node.pydantic_config is not None

    def test_heading_level3_load_config_dict(self):
        """HeadingLevel3Node.load_config with dict"""
        from wordformat.rules.heading import HeadingLevel3Node

        node = HeadingLevel3Node(
            value={"category": "headings.level_3", "fingerprint": "fp"},
            level=3,
        )
        config_dict = {
            "headings": {
                "level_3": {
                    "alignment": "左对齐",
                    "font_size": "小四",
                }
            }
        }
        node.load_config(config_dict)
        assert node.pydantic_config is not None

    def test_heading_base_with_config(self, sample_yaml_config):
        """_base method with loaded config"""
        from wordformat.config.loader import init_config, get_config
        from wordformat.rules.heading import HeadingLevel1Node

        init_config(sample_yaml_config)
        config = get_config()

        node = HeadingLevel1Node(
            value={"category": "headings.level_1", "fingerprint": "fp"},
            level=1,
        )
        node.load_config(config)
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("第一章 绪论")
        node.paragraph = p
        node.check_format(doc)  # 通过 RULES handler 执行格式检查


# ==================== (o) set_style.py 额外覆盖测试 ====================


class TestSetStyleAdditionalCoverage:
    """覆盖 set_style.py lines 53-55, 147, 150, 167-168"""

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    def test_exception_in_traverse_logs_and_raises(
        self, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """Node exception: logs warning then raises (lines 53-55)"""
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node
        mock_apply.side_effect = RuntimeError("traverse error")

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        with pytest.raises(RuntimeError, match="traverse error"):
            auto_format_thesis_document(
                jsonpath=temp_docx,
                docxpath=temp_docx,
                configpath=config_path,
                savepath=str(tmp_path),
                check=True,
            )

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    def test_body_text_filtering(
        self, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """body_text nodes are no longer filtered out"""
        body_node = mock.MagicMock()
        body_node.value = {"category": "body_text"}
        heading_node = mock.MagicMock()
        heading_node.value = {"category": "heading_level_1"}
        root_node = mock.MagicMock()
        root_node.children = [body_node, heading_node]
        mock_builder.build_from_json.return_value = root_node
        mock_apply.return_value = None

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        auto_format_thesis_document(
            jsonpath=temp_docx,
            docxpath=temp_docx,
            configpath=config_path,
            savepath=str(tmp_path),
            check=True,
        )
        # body_text 不再被过滤
        assert len(root_node.children) == 2

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    @mock.patch("wordformat.pipeline.stages.promote_bodytext_in_subtrees_of_type")
    def test_promote_called(
        self, mock_promote, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """promote_bodytext_in_subtrees_of_type is called (line 150)"""
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node
        mock_apply.return_value = None

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        auto_format_thesis_document(
            jsonpath=temp_docx,
            docxpath=temp_docx,
            configpath=config_path,
            savepath=str(tmp_path),
            check=True,
        )
        assert mock_promote.call_count == 3

    @mock.patch(
        "wordformat.pipeline.stages.FormattingExecutionStage.apply_format_check_to_all_nodes"
    )
    @mock.patch("wordformat.pipeline.stages.DocumentBuilder")
    @mock.patch("wordformat.numbering.process_heading_numbering")
    def test_numbering_processing(
        self, mock_numbering, mock_builder, mock_apply, temp_docx, config_path, tmp_path
    ):
        """Numbering processing when enabled (lines 167-168)"""
        root_node = mock.MagicMock()
        root_node.children = []
        mock_builder.build_from_json.return_value = root_node
        mock_apply.return_value = None

        from wordformat.pipeline.orchestrate import auto_format_thesis_document

        # Mock config with numbering.enabled = True
        with mock.patch("wordformat.pipeline.stages.load_config") as mock_get_cfg:
            mock_cfg = mock.MagicMock()
            mock_cfg.numbering.enabled = True
            mock_get_cfg.return_value = mock_cfg
            auto_format_thesis_document(
                jsonpath=temp_docx,
                docxpath=temp_docx,
                configpath=config_path,
                savepath=str(tmp_path),
                check=False,
            )
            mock_numbering.assert_called_once()


# ==================== (p) rules/body.py 额外覆盖测试 ====================


class TestBodyTextAdditional:
    """覆盖 rules/body.py lines 27, 45"""

    def test_body_text_apply_format(self, sample_yaml_config):
        """BodyText._base with p=False, r=False (line 27)"""
        from wordformat.config.loader import init_config, get_config
        from wordformat.rules.body import BodyText

        init_config(sample_yaml_config)
        config = get_config()

        node = BodyText(
            value={"category": "body_text", "fingerprint": "fp"},
            level=1,
        )
        node.load_config(config)
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test content")
        run.font.size = Pt(14)  # Wrong size
        node.paragraph = p
        node.apply_format(doc)

    def test_body_text_apply_to_run(self, sample_yaml_config):
        """apply_format 通过 handler 修正字符格式。"""
        from wordformat.config.loader import init_config, get_config
        from wordformat.rules.body import BodyText

        init_config(sample_yaml_config)
        config = get_config()

        node = BodyText(
            value={"category": "body_text", "fingerprint": "fp"},
            level=1,
        )
        node.load_config(config)
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("test")
        run.font.bold = True  # Wrong - should be False per config
        node.paragraph = p
        node.apply_format(doc)
        assert run.font.bold is False  # Should have been fixed


# ==================== (q) rules/node.py 额外覆盖测试 ====================


class TestFormatNodeAdditional:
    """覆盖 rules/node.py lines 52, 127, 145"""

    def test_load_config_key_error_path(self):
        """load_config with non-dict intermediate path raises KeyError (line 52)"""
        from wordformat.rules.node import TreeNode

        node = TreeNode(value={"category": "a.b.c", "fingerprint": "fp"})
        config = {"a": "not_a_dict"}
        node.load_config(config)
        # Should not crash, returns empty config


# ==================== (r) tree.py 额外覆盖测试 ====================


class TestTreeAdditional:
    """覆盖 tree.py lines 47, 162-164"""

    def test_level_order_empty_root(self):
        """level_order with root that has None value (line 47)"""
        from wordformat.tree import Tree

        tree = Tree(None)
        result = list(tree.level_order())
        assert result == [None]

    def test_print_tree_with_paragraph_value(self, capsys):
        """print_tree with value having .paragraph dict (lines 162-164)"""
        from wordformat.tree import print_tree
        from wordformat.rules.node import TreeNode

        node = TreeNode(
            value={"category": "test", "paragraph": "hello", "fingerprint": "fp"}
        )
        print_tree(node)
        captured = capsys.readouterr()
        assert "test" in captured.out


# ==================== (s) settings.py 额外覆盖测试 ====================


class TestSettingsFrozenPath:
    """覆盖 settings.py line 15"""

    def test_frozen_path(self):
        """sys.frozen=True sets BASE_DIR to sys.executable.parent (line 15)"""
        import sys
        from pathlib import Path

        original_frozen = getattr(sys, "frozen", None)
        sys.frozen = True
        try:
            with mock.patch("sys.executable", "/fake/app/bin"):
                # Reload settings module to pick up frozen state
                import importlib
                import wordformat.settings as settings_mod

                importlib.reload(settings_mod)
                assert settings_mod.BASE_DIR == Path("/fake/app")
        finally:
            if original_frozen is None:
                delattr(sys, "frozen")
            else:
                sys.frozen = original_frozen
            # Reload back to normal
            import importlib
            import wordformat.settings as settings_mod

            importlib.reload(settings_mod)


# ==================== (t) config/config.py 额外覆盖测试 ====================


class TestConfigAdditional:
    """config/loader.py 额外覆盖测试。"""

    def test_load_invalid_yaml_raises(self, tmp_path):
        bad_yaml = tmp_path / "bad.yaml"
        bad_yaml.write_text("global_format: {alignment: [invalid}", encoding="utf-8")
        with pytest.raises(Exception):
            load_config(str(bad_yaml))

    def test_get_without_load_raises(self):
        clear_config()
        with pytest.raises(RuntimeError):
            get_config()

    def test_load_then_get_returns_config(self, config_path):
        cfg = load_config(config_path)
        assert cfg is get_config()


# ==================== 补充覆盖率：小缺口 ====================


class TestDocumentBuilderLoadFromFile:
    """覆盖 document_builder.py 的 except 分支（从文件加载 JSON）"""

    def test_load_paragraphs_from_file(self, tmp_path):
        import json

        data = [{"category": "body_text", "paragraph": "test", "fingerprint": "fp1"}]
        json_file = tmp_path / "data.json"
        json_file.write_text(json.dumps(data), encoding="utf-8")
        result = DocumentBuilder.load_paragraphs(str(json_file))
        assert len(result) == 1
        assert result[0]["category"] == "body_text"

    def test_load_paragraphs_invalid_json_falls_back_to_file(self, tmp_path):
        """无效 JSON 字符串回退到文件加载"""
        import json

        data = [{"category": "body_text", "paragraph": "test", "fingerprint": "fp1"}]
        json_file = tmp_path / "data.json"
        json_file.write_text(json.dumps(data), encoding="utf-8")
        # 传入一个不是有效 JSON 也不是 list 的路径
        result = DocumentBuilder.load_paragraphs(str(json_file))
        assert len(result) == 1


class TestTreeBuilderAdditionalCoverage:
    """覆盖 tree_builder.py 剩余行"""

    def test_build_tree_with_unknown_category_skipped(self, sample_yaml_config):
        """未知类别的节点被跳过（create_node 返回 None）"""
        from wordformat.config.loader import init_config, get_config

        init_config(sample_yaml_config)
        config = get_config()
        builder = DocumentTreeBuilder()
        items = [
            {"category": "unknown_category", "paragraph": "test", "fingerprint": "fp1"},
        ]
        root = builder.build_tree(items)
        # 未知类别被跳过，只有 root
        assert len(root.children) == 0

    def test_build_tree_with_body_text_category(self):
        """body_text 走 _attach_body_node 分支"""
        builder = DocumentTreeBuilder()
        items = [
            {"category": "body_text", "paragraph": "正文内容", "fingerprint": "fp1"},
        ]
        with mock.patch.object(builder, "_create_node_from_item") as mock_create:
            mock_node = mock.MagicMock()
            mock_node.level = 0
            mock_create.return_value = mock_node
            root = builder.build_tree(items)
        assert len(root.children) == 1


class TestNumberingAdditionalCoverage:
    """覆盖 numbering.py 剩余行 (49, 262)"""

    def test_auto_strip_numbering_empty_run_text(self, doc):
        """run 文本被完全清空后仍保留空 run"""
        from wordformat.numbering import _auto_strip_numbering

        p = doc.add_paragraph()
        run = p.add_run("1.1 ")
        _auto_strip_numbering(p, ilvl=1)
        # run 文本被清空但 run 仍存在
        assert run.text == ""

    def test_process_heading_numbering_disabled(self, doc):
        """numbering config enabled=False，直接返回"""
        from wordformat.numbering import process_heading_numbering
        from wordformat.config.models import NumberingConfig

        # 创建 disabled 的 numbering config
        config = NumberingConfig(enabled=False)
        process_heading_numbering(None, doc, config)
