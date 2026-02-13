#! /usr/bin/env python
# -*- coding: utf-8 -*-
import json
import os
import tempfile
from io import BytesIO

from docx import Document

from wordformat.set_style import auto_format_thesis_document


def test_auto_format_thesis_document_check_mode():
    """测试auto_format_thesis_document函数的检查模式"""
    # 创建内存中的docx文件
    doc = Document()
    doc.add_paragraph("测试段落")
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    # 创建临时json文件
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
        # 写入简单的JSON结构
        json_data = [
            {
                "fingerprint": "test_fingerprint_1",
                "category": "heading",
                "name": "标题1",
                "children": []
            }
        ]
        json.dump(json_data, f, ensure_ascii=False)
        temp_json_path = f.name

    # 创建临时输出目录
    with tempfile.TemporaryDirectory() as temp_output_dir:
        try:
            # 使用示例配置文件
            config_path = "example/undergrad_thesis.yaml"

            # 保存内存中的文档到临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                temp_docx_path = f.name
            doc.save(temp_docx_path)

            # 调用auto_format_thesis_document函数（检查模式）
            result_path = auto_format_thesis_document(
                jsonpath=temp_json_path,
                docxpath=temp_docx_path,
                configpath=config_path,
                savepath=temp_output_dir,
                check=True
            )

            # 验证结果文件被创建
            assert os.path.exists(result_path)
            assert "--标注版.docx" in result_path

        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_json_path):
                os.unlink(temp_json_path)


def test_auto_format_thesis_document_format_mode():
    """测试auto_format_thesis_document函数的格式化模式"""
    # 创建内存中的docx文件
    doc = Document()
    doc.add_paragraph("测试段落")
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    # 创建临时json文件
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
        # 写入简单的JSON结构
        json_data = [
            {
                "fingerprint": "test_fingerprint_1",
                "category": "heading",
                "name": "标题1",
                "children": []
            }
        ]
        json.dump(json_data, f, ensure_ascii=False)
        temp_json_path = f.name

    # 创建临时输出目录
    with tempfile.TemporaryDirectory() as temp_output_dir:
        try:
            # 使用示例配置文件
            config_path = "example/undergrad_thesis.yaml"

            # 保存内存中的文档到临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                temp_docx_path = f.name
            doc.save(temp_docx_path)

            # 调用auto_format_thesis_document函数（格式化模式）
            result_path = auto_format_thesis_document(
                jsonpath=temp_json_path,
                docxpath=temp_docx_path,
                configpath=config_path,
                savepath=temp_output_dir,
                check=False
            )

            # 验证结果文件被创建
            assert os.path.exists(result_path)
            assert "--修改版.docx" in result_path

        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_json_path):
                os.unlink(temp_json_path)


def test_auto_format_thesis_document_with_complex_structure():
    """测试auto_format_thesis_document函数处理复杂结构"""
    # 创建内存中的docx文件
    doc = Document()
    doc.add_paragraph("测试标题", style="Heading 1")
    doc.add_paragraph("测试正文")
    doc.add_paragraph("测试子标题", style="Heading 2")
    doc.add_paragraph("测试子正文")

    # 创建临时json文件
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
        # 写入复杂的JSON结构
        json_data = [
            {
                "fingerprint": "test_fingerprint_1",
                "category": "heading",
                "name": "测试标题",
                "children": [
                    {
                        "fingerprint": "test_fingerprint_2",
                        "category": "body",
                        "name": "测试正文",
                        "children": []
                    },
                    {
                        "fingerprint": "test_fingerprint_3",
                        "category": "heading",
                        "name": "测试子标题",
                        "children": [
                            {
                                "fingerprint": "test_fingerprint_4",
                                "category": "body",
                                "name": "测试子正文",
                                "children": []
                            }
                        ]
                    }
                ]
            }
        ]
        json.dump(json_data, f, ensure_ascii=False)
        temp_json_path = f.name

    # 创建临时输出目录
    with tempfile.TemporaryDirectory() as temp_output_dir:
        try:
            # 使用示例配置文件
            config_path = "example/undergrad_thesis.yaml"

            # 保存内存中的文档到临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                temp_docx_path = f.name
            doc.save(temp_docx_path)

            # 调用auto_format_thesis_document函数（检查模式）
            result_path = auto_format_thesis_document(
                jsonpath=temp_json_path,
                docxpath=temp_docx_path,
                configpath=config_path,
                savepath=temp_output_dir,
                check=True
            )

            # 验证结果文件被创建
            assert os.path.exists(result_path)
            assert "--标注版.docx" in result_path

        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_json_path):
                os.unlink(temp_json_path)


def test_xg_function():
    """测试xg函数"""
    from wordformat.set_style import xg
    from wordformat.rules.node import FormatNode

    # 创建一个简单的FormatNode树
    root_node = FormatNode(value={"category": "top"}, level=0)
    child_node = FormatNode(value={"category": "body", "fingerprint": "test_fingerprint"}, level=1)
    root_node.add_child(child_node)

    # 创建一个段落对象
    doc = Document()
    paragraph = doc.add_paragraph("测试段落")

    # 测试xg函数
    result = xg(root_node, paragraph)
    # 由于指纹不匹配，应该返回None
    assert result is None


def test_apply_format_check_to_all_nodes():
    """测试apply_format_check_to_all_nodes函数"""
    from wordformat.set_style import apply_format_check_to_all_nodes
    from wordformat.rules.node import FormatNode

    # 创建一个带有check_format方法的自定义节点类
    class TestFormatNode(FormatNode):
        def __init__(self, value, level):
            super().__init__(value, level)
            self.check_called = False
            self.apply_called = False
            # 设置paragraph属性，这样check_format和apply_format方法才会被调用
            doc = Document()
            self.paragraph = doc.add_paragraph("测试段落")
            # 确保children属性存在
            if not hasattr(self, 'children'):
                self.children = []

        def check_format(self, document):
            self.check_called = True

        def apply_format(self, document):
            self.apply_called = True

        def load_config(self, config):
            pass

        def add_child(self, child):
            """添加子节点"""
            if not hasattr(self, 'children'):
                self.children = []
            self.children.append(child)

    # 创建一个节点树
    root_node = TestFormatNode(value={"category": "body", "fingerprint": "root_fingerprint"}, level=0)
    child_node = TestFormatNode(value={"category": "body", "fingerprint": "child_fingerprint"}, level=1)
    root_node.add_child(child_node)

    # 验证子节点是否被正确添加
    assert len(root_node.children) == 1
    assert root_node.children[0] == child_node

    # 创建一个文档对象
    doc = Document()

    # 测试检查模式
    apply_format_check_to_all_nodes(root_node, doc, {}, check=True)
    assert root_node.check_called
    assert child_node.check_called

    # 重置标志
    root_node.check_called = False
    child_node.check_called = False

    # 测试应用模式
    apply_format_check_to_all_nodes(root_node, doc, {}, check=False)
    assert root_node.apply_called
    assert child_node.apply_called


def test_auto_format_thesis_document_with_body_text_category():
    """测试auto_format_thesis_document函数处理body_text类别"""
    # 创建内存中的docx文件
    doc = Document()
    doc.add_paragraph("测试标题", style="Heading 1")
    doc.add_paragraph("测试正文")

    # 创建临时json文件，包含body_text类别
    with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
        json_data = [
            {
                "fingerprint": "test_fingerprint_1",
                "category": "heading",
                "name": "测试标题",
                "children": []
            },
            {
                "fingerprint": "test_fingerprint_2",
                "category": "body_text",
                "name": "测试正文",
                "children": []
            }
        ]
        json.dump(json_data, f, ensure_ascii=False)
        temp_json_path = f.name

    # 创建临时输出目录
    with tempfile.TemporaryDirectory() as temp_output_dir:
        try:
            # 使用示例配置文件
            config_path = "example/undergrad_thesis.yaml"

            # 保存内存中的文档到临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
                temp_docx_path = f.name
            doc.save(temp_docx_path)

            # 调用auto_format_thesis_document函数
            result_path = auto_format_thesis_document(
                jsonpath=temp_json_path,
                docxpath=temp_docx_path,
                configpath=config_path,
                savepath=temp_output_dir,
                check=True
            )

            # 验证结果文件被创建
            assert os.path.exists(result_path)

        finally:
            # 清理临时文件
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            if os.path.exists(temp_json_path):
                os.unlink(temp_json_path)
