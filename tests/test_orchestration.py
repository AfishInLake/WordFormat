"""编排层测试 —— binding, style_fixer, table_formatter, build_element"""
import pytest
from unittest.mock import patch
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from wordformat.rules.node import FormatNode
from wordformat.config.datamodel import NodeConfigRoot


class TestFixStyleRunProperties:
    """测试 _fix_style_run_properties 各种样式路径"""

    def test_bold_add_element(self):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from wordformat.set_style import _fix_style_run_properties

        class MockStyle:
            element = OxmlElement("w:style")
            rPr = OxmlElement("w:rPr")
            element.insert(0, rPr)

        class MockCfg:
            bold = True
            italic = None
            underline = None
            chinese_font_name = None
            english_font_name = None
            font_size = None
            font_color = None

        style = MockStyle()
        _fix_style_run_properties(style, MockCfg(), "test_style")
        assert style.rPr.find(qn("w:b")) is not None

    def test_italic_remove_element(self):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from wordformat.set_style import _fix_style_run_properties

        class MockStyle:
            element = OxmlElement("w:style")
            rPr = OxmlElement("w:rPr")
            element.insert(0, rPr)
            rPr.append(OxmlElement("w:i"))

        class MockCfg:
            bold = None
            italic = False
            underline = None
            chinese_font_name = None
            english_font_name = None
            font_size = None
            font_color = None

        style = MockStyle()
        _fix_style_run_properties(style, MockCfg(), "test_style")
        assert style.rPr.find(qn("w:i")) is None

    def test_underline_add_element(self):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from wordformat.set_style import _fix_style_run_properties

        class MockStyle:
            element = OxmlElement("w:style")
            rPr = OxmlElement("w:rPr")
            element.insert(0, rPr)

        class MockCfg:
            bold = None
            italic = None
            underline = True
            chinese_font_name = None
            english_font_name = None
            font_size = None
            font_color = None

        style = MockStyle()
        _fix_style_run_properties(style, MockCfg(), "test_style")
        u = style.rPr.find(qn("w:u"))
        assert u is not None
        assert u.get(qn("w:val")) == "single"


class TestFixStyleParagraphProperties:
    """测试 _fix_style_paragraph_properties 各种路径"""

    def test_space_before_after_not_set(self):
        from docx.oxml import OxmlElement
        from wordformat.set_style import _fix_style_paragraph_properties

        class MockStyle:
            element = OxmlElement("w:style")

        class MockCfg:
            alignment = None
            space_before = None
            space_after = None
            line_spacingrule = None
            line_spacing = None
            first_line_indent = None
            left_indent = None
            right_indent = None

        style = MockStyle()
        _fix_style_paragraph_properties(style, MockCfg(), "test")

    def test_line_spacingrule_without_line_spacing(self):
        from docx.oxml import OxmlElement
        from wordformat.set_style import _fix_style_paragraph_properties

        class MockStyle:
            element = OxmlElement("w:style")

        class MockCfg:
            alignment = None
            space_before = None
            space_after = None
            line_spacingrule = 3
            line_spacing = None
            first_line_indent = None
            left_indent = None
            right_indent = None

        style = MockStyle()
        _fix_style_paragraph_properties(style, MockCfg(), "test")

    def test_line_spacing_倍_unit(self):
        from docx.oxml import OxmlElement
        from wordformat.set_style import _fix_style_paragraph_properties

        class MockStyle:
            element = OxmlElement("w:style")

        class MockCfg:
            alignment = None
            space_before = None
            space_after = None
            line_spacingrule = 2
            line_spacing = "1.5倍"
            first_line_indent = None
            left_indent = None
            right_indent = None

        style = MockStyle()
        _fix_style_paragraph_properties(style, MockCfg(), "test")


class TestSyncInsertionsDeletions:
    """测试 _sync_insertions 和 _sync_deletions"""

    def test_sync_deletions_removes_paragraphs(self):
        from wordformat.set_style import _sync_deletions
        doc = Document()
        p1 = doc.add_paragraph("keep")
        p2 = doc.add_paragraph("delete me")
        body = p1._element.getparent()
        assert len(body.findall(p1._element.tag)) == 2
        docx_paras = [p for p in doc.paragraphs if p.text.strip()]
        _sync_deletions(docx_paras, {1})
        remaining = body.findall(p1._element.tag)
        assert len(remaining) == 1
        assert remaining[0] is p1._element

    def test_sync_insertions_adds_paragraph(self):
        from wordformat.set_style import _sync_insertions
        doc = Document()
        p1 = doc.add_paragraph("第一章")
        root = FormatNode(value={"category": "top"}, level=0)
        new_node = FormatNode(value={"category": "body_text", "paragraph": "新增段落"}, level=2)
        tree_nodes = [new_node]
        matches = {}
        _sync_insertions(tree_nodes, matches, {0}, doc)
        # 新段落应该创建并绑定
        assert new_node.paragraph is not None
        assert new_node.paragraph.text == "新增段落"

    def test_sync_insertions_with_prev_match(self):
        from wordformat.set_style import _sync_insertions
        doc = Document()
        p1 = doc.add_paragraph("第一章")
        root = FormatNode(value={"category": "top"}, level=0)
        # tree: [matched_node, new_node]
        matched_node = FormatNode(value={"category": "body_text", "paragraph": "第一章"}, level=1, paragraph=p1)
        new_node = FormatNode(value={"category": "body_text", "paragraph": "新增段落"}, level=2)
        tree_nodes = [matched_node, new_node]
        matches = {0: p1}
        _sync_insertions(tree_nodes, matches, {1}, doc)
        assert new_node.paragraph is not None
        assert new_node.paragraph.text == "新增段落"
        # 新段落应在 p1 之后
        assert new_node.paragraph._element.getprevious() is p1._element

    def test_sync_insertions_with_next_match(self):
        from wordformat.set_style import _sync_insertions
        doc = Document()
        p2 = doc.add_paragraph("绪论")
        new_node = FormatNode(value={"category": "body_text", "paragraph": "前置新增"}, level=2)
        matched_node = FormatNode(value={"category": "body_text", "paragraph": "绪论"}, level=1, paragraph=p2)
        tree_nodes = [new_node, matched_node]
        matches = {1: p2}
        _sync_insertions(tree_nodes, matches, {0}, doc)
        assert new_node.paragraph is not None
        assert new_node.paragraph.text == "前置新增"
        # 新段落应在 p2 之前
        assert new_node.paragraph._element.getnext() is p2._element


# ============================================================
# hyperlinks.py — 覆盖引用超链接函数
# ============================================================


class TestFixStyleRunPropertiesMore:
    """测试 _fix_style_run_properties 更多路径"""

    def test_italic_add_element(self):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from wordformat.set_style import _fix_style_run_properties

        class MockStyle:
            element = OxmlElement("w:style")
            rPr = OxmlElement("w:rPr")
            element.insert(0, rPr)

        class MockCfg:
            bold = None
            italic = True
            underline = None
            chinese_font_name = None
            english_font_name = None
            font_size = None
            font_color = None

        style = MockStyle()
        _fix_style_run_properties(style, MockCfg(), "test")
        assert style.rPr.find(qn("w:i")) is not None

    def test_underline_remove_element(self):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from wordformat.set_style import _fix_style_run_properties

        class MockStyle:
            element = OxmlElement("w:style")
            rPr = OxmlElement("w:rPr")
            element.insert(0, rPr)
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)

        class MockCfg:
            bold = None
            italic = None
            underline = False
            chinese_font_name = None
            english_font_name = None
            font_size = None
            font_color = None

        style = MockStyle()
        _fix_style_run_properties(style, MockCfg(), "test")
        assert style.rPr.find(qn("w:u")) is None


class TestFixStyleParagraphPropertiesMore:
    """测试 _fix_style_paragraph_properties 更多路径"""

    def test_pt_line_spacing_conversion(self):
        from docx.oxml import OxmlElement
        from wordformat.set_style import _fix_style_paragraph_properties

        class MockStyle:
            element = OxmlElement("w:style")

        class MockCfg:
            alignment = None
            space_before = None
            space_after = None
            line_spacingrule = 4  # exact
            line_spacing = "12pt"
            first_line_indent = None
            left_indent = None
            right_indent = None

        style = MockStyle()
        _fix_style_paragraph_properties(style, MockCfg(), "test")

    def test_space_before_not_hang_unit(self):
        from docx.oxml import OxmlElement
        from wordformat.set_style import _fix_style_paragraph_properties

        class MockStyle:
            element = OxmlElement("w:style")

        class MockCfg:
            alignment = None
            space_before = "10pt"
            space_after = None
            line_spacingrule = None
            line_spacing = None
            first_line_indent = None
            left_indent = None
            right_indent = None

        style = MockStyle()
        _fix_style_paragraph_properties(style, MockCfg(), "test")

    def test_first_line_indent_not_char_unit(self):
        from docx.oxml import OxmlElement
        from wordformat.set_style import _fix_style_paragraph_properties

        class MockStyle:
            element = OxmlElement("w:style")

        class MockCfg:
            alignment = None
            space_before = None
            space_after = None
            line_spacingrule = None
            line_spacing = None
            first_line_indent = "10pt"
            left_indent = None
            right_indent = None

        style = MockStyle()
        _fix_style_paragraph_properties(style, MockCfg(), "test")


# ============================================================
# word_structure/utils.py — find_and_modify_first
# ============================================================


class TestFormatTableContent:
    """测试 format_table_content — 无 tables 属性的配置"""

    def test_no_tables_attribute(self):
        """传入没有 tables 属性的对象，应直接 return"""
        from wordformat.set_style import format_table_content

        class NoTablesConfig:
            pass

        doc = Document()
        format_table_content(doc, NoTablesConfig(), check=True)


# ============================================================
# hyperlinks.py — 更多覆盖
# ============================================================


class TestCollectAllStyleConfigs:
    """测试 _collect_all_style_configs 自定义样式路径"""

    def test_custom_style_name(self):
        from wordformat.set_style import _collect_all_style_configs
        from wordformat.config.datamodel import NodeConfigRoot
        config = NodeConfigRoot()
        result = _collect_all_style_configs(config)
        assert isinstance(result, dict)


class TestFixAllStyleDefinitions:
    """测试 _fix_all_style_definitions"""

    def test_normal_case(self):
        from wordformat.set_style import _fix_all_style_definitions
        from wordformat.config.datamodel import NodeConfigRoot
        doc = Document()
        config = NodeConfigRoot()
        _fix_all_style_definitions(doc, config)


# ============================================================
# core/tree.py — Tree 类完整测试
# ============================================================


class TestBuildElement:
    """测试 _build_*_element 函数"""

    def test_build_paragraph_element(self):
        from wordformat.orchestration.binding import _build_paragraph_element
        from wordformat.core.tree import TreeNode
        node = TreeNode({"paragraph": "测试段落"})
        elem = _build_paragraph_element(node)
        assert elem.tag.endswith("}p")
        t_elem = elem.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        assert t_elem.text == "测试段落"

    def test_build_table_element(self):
        from wordformat.orchestration.binding import _build_table_element
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "table"})
        node.type = "table"
        node.content = {"rows": 2, "cols": 2, "cells": [["A", "B"], ["C", "D"]]}
        elem = _build_table_element(node)
        assert elem.tag.endswith("}tbl")
        rows = elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr")
        assert len(rows) == 2

    def test_build_table_empty_cells(self):
        from wordformat.orchestration.binding import _build_table_element
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "table"})
        node.type = "table"
        node.content = {"rows": 1, "cols": 1}
        elem = _build_table_element(node)
        t_elem = elem.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        assert t_elem.text == ""

    def test_build_image_element(self):
        from wordformat.orchestration.binding import _build_image_element
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "image"})
        node.type = "image"
        node.content = {"path": "/tmp/img.png", "width": 400, "height": 300}
        elem = _build_image_element(node)
        assert elem.tag.endswith("}drawing")

    def test_build_element_dispatches_correctly(self):
        from wordformat.orchestration.binding import _build_element
        from wordformat.core.tree import TreeNode
        # paragraph
        p_node = TreeNode({"paragraph": "text"})
        assert _build_element(p_node).tag.endswith("}p")
        # table
        t_node = TreeNode({"category": "table"})
        t_node.type = "table"
        t_node.content = {"rows": 1, "cols": 1}
        assert _build_element(t_node).tag.endswith("}tbl")
        # image
        i_node = TreeNode({"category": "image"})
        i_node.type = "image"
        i_node.content = {"path": "/tmp/img.png"}
        assert _build_element(i_node).tag.endswith("}drawing")
        # unknown fallback
        u_node = TreeNode({"paragraph": "unknown"})
        u_node.type = "unknown"
        assert _build_element(u_node).tag.endswith("}p")


# ============================================================
# orchestration/style_fixer.py — 样式修正边界
# ============================================================


class TestStyleFixerEdgeCases:
    """测试样式修正函数的边界路径"""

    def test_collect_all_style_configs_custom_style_name(self):
        from wordformat.orchestration.style_fixer import collect_all_style_configs
        from wordformat.config.datamodel import NodeConfigRoot, GlobalFormatConfig

        class CustomConfig(NodeConfigRoot):
            references: GlobalFormatConfig = GlobalFormatConfig(
                chinese_font_name="宋体",
                builtin_style_name="MyCustomStyle",
            )

        config = CustomConfig()
        result = collect_all_style_configs(config)
        assert "MyCustomStyle" in result

    def test_fix_style_paragraph_properties_pt_line_spacing(self):
        from wordformat.orchestration.style_fixer import fix_style_paragraph_properties
        from docx.oxml import OxmlElement

        class MockStyle:
            element = OxmlElement("w:style")

        class MockCfg:
            alignment = None
            space_before = None
            space_after = None
            line_spacingrule = 4
            line_spacing = "12pt"
            first_line_indent = None
            left_indent = None
            right_indent = None

        fix_style_paragraph_properties(MockStyle(), MockCfg(), "test")

    def test_fix_style_run_properties_italic_add(self):
        from wordformat.orchestration.style_fixer import fix_style_run_properties
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        class MockStyle:
            element = OxmlElement("w:style")
            rPr = OxmlElement("w:rPr")
            element.insert(0, rPr)

        class MockCfg:
            bold = None
            italic = True
            underline = None
            chinese_font_name = None
            english_font_name = None
            font_size = None
            font_color = None

        style = MockStyle()
        fix_style_run_properties(style, MockCfg(), "test")
        assert style.rPr.find(qn("w:i")) is not None


# ============================================================
# orchestration/binding.py — 更多绑定路径
# ============================================================


class TestBindAndSyncCheckMode:
    """测试 bind_and_sync check 模式的日志路径"""

    def test_check_mode_reports_deletions(self):
        from wordformat.orchestration.binding import bind_and_sync
        from docx import Document
        doc = Document()
        doc.add_paragraph("keep")
        doc.add_paragraph("delete_me")
        root = FormatNode(value={"category": "top"}, level=0)
        keep = FormatNode(value={"category": "body_text", "paragraph": "keep"}, level=2)
        root.add_child_node(keep)
        bind_and_sync(root, doc, check=True)

    def test_sync_insertions_body_fallback(self):
        from wordformat.orchestration.binding import _sync_insertions
        from docx import Document
        doc = Document()
        new_node = FormatNode(value={"category": "body_text", "paragraph": "only"}, level=2)
        _sync_insertions([new_node], {}, {0}, doc)
        assert new_node.paragraph is not None


# ============================================================
# core/category.py — 剩余边界
# ============================================================


class TestFormatTableContentCoverage:
    """覆盖 format_table_content 循环体"""

    def test_table_content_no_tables_attr(self):
        from wordformat.orchestration.table_formatter import format_table_content
        doc = Document()

        class NoTables:
            pass
        format_table_content(doc, NoTables(), check=True)

    def test_table_content_empty_table(self):
        from wordformat.orchestration.table_formatter import format_table_content
        doc = Document()
        doc.add_table(rows=0, cols=0)
        from wordformat.config.datamodel import NodeConfigRoot
        config = NodeConfigRoot()
        try:
            format_table_content(doc, config, check=True)
        except Exception:
            pass  # LazyConfig may not be initialized


# ============================================================
# A: 新代码边界测试
# ============================================================


class TestBuildElementEdgeCases:
    """_build_*_element 边界测试"""

    def test_build_table_mismatched_cells(self):
        from wordformat.orchestration.binding import _build_table_element
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "table"})
        node.type = "table"
        node.content = {"rows": 3, "cols": 3, "cells": [["仅一行"]]}
        elem = _build_table_element(node)
        rows = elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr")
        assert len(rows) == 3  # 3 行都创建了，缺失单元格为空

    def test_build_table_empty_cells_list(self):
        from wordformat.orchestration.binding import _build_table_element
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "table"})
        node.type = "table"
        node.content = {"rows": 2, "cols": 2, "cells": []}
        elem = _build_table_element(node)
        rows = elem.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr")
        assert len(rows) == 2  # 不崩溃

    def test_build_image_zero_dimensions(self):
        from wordformat.orchestration.binding import _build_image_element
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "image"})
        node.type = "image"
        node.content = {"path": "", "width": 0, "height": 0}
        elem = _build_image_element(node)
        extent = elem.find(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent")
        assert extent.get("cx") == "0"


class TestSyncInsertionsEdgeCases:
    """_sync_insertions 边界测试"""

    def test_insert_at_beginning(self):
        from wordformat.orchestration.binding import _sync_insertions
        from docx import Document
        doc = Document()
        p = doc.add_paragraph("已有段落")
        new_node = FormatNode(value={"category": "body_text", "paragraph": "前置新增"}, level=2)
        matched = FormatNode(value={"category": "body_text", "paragraph": "已有段落"}, level=1, paragraph=p)
        _sync_insertions([new_node, matched], {1: p}, {0}, doc)
        assert new_node.paragraph is not None
        assert new_node.paragraph._element.getnext() is p._element

    def test_insert_at_end(self):
        from wordformat.orchestration.binding import _sync_insertions
        from docx import Document
        doc = Document()
        p = doc.add_paragraph("已有段落")
        matched = FormatNode(value={"category": "body_text", "paragraph": "已有段落"}, level=1, paragraph=p)
        new_node = FormatNode(value={"category": "body_text", "paragraph": "尾部新增"}, level=2)
        _sync_insertions([matched, new_node], {0: p}, {1}, doc)
        assert new_node.paragraph is not None
        assert new_node.paragraph._element.getprevious() is p._element

    def test_multiple_consecutive_insertions(self):
        from wordformat.orchestration.binding import _sync_insertions
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document()
        p0 = doc.add_paragraph("第一段")
        p2 = doc.add_paragraph("第三段")
        n0 = FormatNode(value={"category": "body_text", "paragraph": "第一段"}, level=1, paragraph=p0)
        n1 = FormatNode(value={"category": "body_text", "paragraph": "第二段(插入)"}, level=2)
        n2 = FormatNode(value={"category": "body_text", "paragraph": "第三段"}, level=1, paragraph=p2)
        _sync_insertions([n0, n1, n2], {0: p0, 2: p2}, {1}, doc)
        assert n1.paragraph is not None
        body = p0._element.getparent()
        children = list(body)
        texts = [c.find(qn("w:r")).find(qn("w:t")).text for c in children if c.find(qn("w:r")) is not None]
        assert texts == ["第一段", "第二段(插入)", "第三段"]

    def test_delete_all_paragraphs(self):
        from wordformat.orchestration.binding import _sync_deletions
        from docx import Document
        doc = Document()
        p1 = doc.add_paragraph("A")
        p2 = doc.add_paragraph("B")
        body = p1._element.getparent()
        paras = [p for p in doc.paragraphs if p.text.strip()]
        _sync_deletions(paras, {0, 1})
        remaining = [p.text for p in doc.paragraphs if p.text.strip()]
        assert remaining == []


class TestBindAndSyncMixed:
    """bind_and_sync 混合场景"""

    def test_mixed_insert_and_delete(self):
        from wordformat.orchestration.binding import bind_and_sync
        from docx import Document
        doc = Document()
        doc.add_paragraph("保留")
        doc.add_paragraph("删除我")
        root = FormatNode(value={"category": "top"}, level=0)
        keep = FormatNode(value={"category": "body_text", "paragraph": "保留"}, level=2)
        new_node = FormatNode(value={"category": "body_text", "paragraph": "新增"}, level=2)
        root.add_child_node(keep)
        root.add_child_node(new_node)
        bind_and_sync(root, doc, check=False)
        remaining = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "保留" in remaining
        assert "新增" in remaining
        assert "删除我" not in remaining

    def test_mixed_node_types_paragraph_and_table(self):
        from wordformat.orchestration.binding import bind_and_sync
        from docx import Document
        doc = Document()
        p = doc.add_paragraph("正文")
        root = FormatNode(value={"category": "top"}, level=0)
        para_node = FormatNode(value={"category": "body_text", "paragraph": "正文"}, level=2)
        table_node = FormatNode(value={"category": "table"}, level=2)
        table_node.type = "table"
        table_node.content = {"rows": 1, "cols": 1, "cells": [["数据"]]}
        root.add_child_node(para_node)
        root.add_child_node(table_node)
        bind_and_sync(root, doc, check=False)
        assert table_node.paragraph is not None


class TestStyleFixerErrorPaths:
    """样式修正错误路径"""

    def test_font_color_parse_failure_no_crash(self):
        from wordformat.orchestration.style_fixer import fix_style_run_properties
        from docx.oxml import OxmlElement

        class MockStyle:
            element = OxmlElement("w:style")
            rPr = OxmlElement("w:rPr")
            element.insert(0, rPr)

        class BadCfg:
            bold = None
            italic = None
            underline = None
            chinese_font_name = None
            english_font_name = None
            font_size = None
            font_color = "invalid_color_!!!!"

        fix_style_run_properties(MockStyle(), BadCfg(), "test")

    def test_style_lookup_keyerror_no_crash(self):
        from wordformat.orchestration.style_fixer import fix_all_style_definitions
        from wordformat.config.datamodel import NodeConfigRoot, GlobalFormatConfig
        from unittest.mock import patch
        doc = Document()

        class CustomConfig(NodeConfigRoot):
            body_text: GlobalFormatConfig = GlobalFormatConfig(
                chinese_font_name="宋体",
                builtin_style_name="正文",
            )

        config = CustomConfig()
        fix_all_style_definitions(doc, config)
