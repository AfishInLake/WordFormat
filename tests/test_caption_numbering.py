"""题注编号处理模块的单元测试。"""

from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from wordformat.config.datamodel import CaptionNumberingConfig
from wordformat.utils import _from_chinese_num, _from_roman, parse_caption_text

# ======================== _from_roman ========================


class TestFromRoman:
    def test_basic_singles(self):
        assert _from_roman("I") == 1
        assert _from_roman("V") == 5
        assert _from_roman("X") == 10

    def test_subtractive(self):
        assert _from_roman("IV") == 4
        assert _from_roman("IX") == 9
        assert _from_roman("XL") == 40
        assert _from_roman("CM") == 900

    def test_composite(self):
        assert _from_roman("XIV") == 14
        assert _from_roman("XXVII") == 27

    def test_case_insensitive(self):
        assert _from_roman("iv") == 4

    def test_empty_raises(self):
        with pytest.raises(ValueError):
            _from_roman("")

    def test_invalid_raises(self):
        with pytest.raises(ValueError):
            _from_roman("ABC")


# ======================== _from_chinese_num ========================


class TestFromChineseNum:
    def test_single_digit(self):
        assert _from_chinese_num("一") == 1
        assert _from_chinese_num("九") == 9

    def test_ten(self):
        assert _from_chinese_num("十") == 10

    def test_teens(self):
        assert _from_chinese_num("十一") == 11
        assert _from_chinese_num("十五") == 15

    def test_tens(self):
        assert _from_chinese_num("二十") == 20
        assert _from_chinese_num("九十九") == 99

    def test_hundred(self):
        assert _from_chinese_num("一百") == 100

    def test_hundreds_complex(self):
        assert _from_chinese_num("一百二十三") == 123

    def test_financial_digits(self):
        assert _from_chinese_num("壹") == 1
        assert _from_chinese_num("叁") == 3

    def test_empty_raises(self):
        with pytest.raises(ValueError):
            _from_chinese_num("")

    def test_invalid_raises(self):
        with pytest.raises(ValueError):
            _from_chinese_num("abc")


# ======================== parse_caption_text ========================


class TestParseCaptionText:
    def test_basic_figure_arabic(self):
        result = parse_caption_text("图1.1 系统架构图")
        assert result is not None
        assert result["label"] == "图"
        assert result["chapter_num"] == 1
        assert result["separator"] == "."
        assert result["number_num"] == 1
        assert result["name"] == "系统架构图"

    def test_hyphen_separator(self):
        result = parse_caption_text("图1-1 数据流程图")
        assert result is not None
        assert result["separator"] == "-"

    def test_colon_separator(self):
        result = parse_caption_text("表1:1 测试数据")
        assert result["separator"] == ":"

    def test_em_dash_separator(self):
        result = parse_caption_text("图1—1 架构设计")
        assert result is not None
        assert result["separator"] == "—"

    def test_en_dash_separator(self):
        result = parse_caption_text("图1–1 测试图")
        assert result["separator"] == "–"

    def test_chinese_chapter_number(self):
        result = parse_caption_text("图一.1 系统架构图")
        assert result is not None
        assert result["chapter_text"] == "一"
        assert result["chapter_num"] == 1

    def test_roman_chapter_number(self):
        result = parse_caption_text("图I.1 系统架构图")
        assert result["chapter_text"] == "I"
        assert result["chapter_num"] == 1

    def test_roman_chapter_lowercase(self):
        result = parse_caption_text("图ii.1 数据图")
        assert result["chapter_num"] == 2

    def test_fullwidth_space(self):
        result = parse_caption_text("图1.1　全角空格名称")
        assert result is not None
        assert result["name"] == "全角空格名称"

    def test_empty_returns_none(self):
        assert parse_caption_text("") is None

    def test_plain_text_returns_none(self):
        assert parse_caption_text("这是一段普通正文") is None

    def test_no_space_before_name_returns_none(self):
        """编号后无空格无法可靠提取名称，返回 None。"""
        result = parse_caption_text("图1.1测试")
        assert result is None

    def test_space_after_label(self):
        """标签后有空格：图 1.1 测试。"""
        result = parse_caption_text("图 1.1 系统架构图")
        assert result is not None
        assert result["label"] == "图"
        assert result["chapter_num"] == 1
        assert result["number_num"] == 1
        assert result["name"] == "系统架构图"

    def test_continued_table_caption(self):
        """续表 5.3 API接口测试结果 → 正确解析"""
        result = parse_caption_text("续表5.3 API接口测试结果")
        assert result is not None
        assert result["label"] == "表"
        assert result["chapter_num"] == 5
        assert result["separator"] == "."
        assert result["number_num"] == 3
        assert result["name"] == "API接口测试结果"
        assert result["is_continued"] is True

    def test_continued_figure_caption(self):
        """续图 2.1 系统架构图 → 正确解析"""
        result = parse_caption_text("续图2.1 系统架构图")
        assert result is not None
        assert result["label"] == "图"
        assert result["chapter_num"] == 2
        assert result["number_num"] == 1
        assert result["name"] == "系统架构图"
        assert result["is_continued"] is True

    def test_continued_caption_with_space_after_label(self):
        """续 表 5.3 xxx → 去掉续后正常匹配。"""
        result = parse_caption_text("续 表 5.3 测试表格")
        assert result is not None
        assert result["label"] == "表"
        assert result["is_continued"] is True
        assert result["name"] == "测试表格"

    def test_continued_table_with_hyphen(self):
        """续表5-3 测试 → 连字符分隔符"""
        result = parse_caption_text("续表5-3 测试")
        assert result is not None
        assert result["label"] == "表"
        assert result["chapter_num"] == 5
        assert result["separator"] == "-"
        assert result["number_num"] == 3
        assert result["is_continued"] is True

    def test_regular_caption_not_continued(self):
        """普通题注 is_continued 为 False"""
        result = parse_caption_text("表5.3 测试")
        assert result is not None
        assert result["is_continued"] is False


# ======================== _replace_paragraph_text ========================


class TestReplaceParagraphText:
    def test_single_run(self):
        from wordformat.rules.caption import _replace_paragraph_text

        doc = Document()
        p = doc.add_paragraph()
        p.add_run("旧文本")
        _replace_paragraph_text(p, "新文本")
        assert p.text == "新文本"

    def test_multiple_runs(self):
        from wordformat.rules.caption import _replace_paragraph_text

        doc = Document()
        p = doc.add_paragraph()
        p.add_run("第一部分")
        p.add_run("第二部分")
        _replace_paragraph_text(p, "全新文本")
        assert p.text == "全新文本"
        assert p.runs[0].text == "全新文本"
        assert p.runs[1].text == ""

    def test_empty_runs(self):
        from wordformat.rules.caption import _replace_paragraph_text

        doc = Document()
        p = doc.add_paragraph()
        _replace_paragraph_text(p, "新文本")
        assert p.text == ""


# ======================== _check_caption_numbering ========================


class TestCheckCaptionNumbering:
    """直接测试 _check_caption_numbering 函数。"""

    def _make_caption_figure(self, paragraph, chapter=1, seq=1):
        from wordformat.rules.caption import CaptionFigure

        node = CaptionFigure(
            value={
                "category": "caption_figure",
                "fingerprint": "fig001",
                "chapter_number": chapter,
                "sequence_number": seq,
            },
            level=0,
            paragraph=paragraph,
        )
        return node

    def _make_paragraph(self, text):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run(text)
        return p

    def test_correct_no_comment(self):
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("图1.1 系统架构图")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_not_called()

    def test_wrong_chapter_adds_comment(self):
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("图2.1 测试图")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_called_once()
        assert "章节号错误" in mock_comment.call_args[0][2]

    def test_wrong_separator_adds_comment(self):
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("图1-1 测试图")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_called_once()
        assert "分隔符错误" in mock_comment.call_args[0][2]

    def test_wrong_label_adds_comment(self):
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("表1.1 测试")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_called_once()
        assert "标签错误" in mock_comment.call_args[0][2]

    def test_wrong_sequence_adds_comment(self):
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("图1.3 第三张图")
        node = self._make_caption_figure(p, chapter=1, seq=2)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_called_once()
        assert "编号错误" in mock_comment.call_args[0][2]

    def test_unparseable_adds_comment(self):
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("这是正文内容")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_called_once()
        assert "格式错误" in mock_comment.call_args[0][2]

    def test_label_space_enabled_no_space_in_text(self):
        """label_number_space=true，题注无空格 → 添加批注。"""
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("图1.1 测试")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".", label_number_space=True)

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_called_once()
        assert "间距错误" in mock_comment.call_args[0][2]

    def test_label_space_disabled_with_space_in_text(self):
        """label_number_space=false，题注有空格 → 添加批注。"""
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("图 1.1 测试")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".", label_number_space=False)

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_called_once()
        assert "间距错误" in mock_comment.call_args[0][2]

    def test_disabled_does_nothing(self):
        """配置禁用时不检查。此逻辑在 _base() 中判断，这里验证 disabled 的 cfg 不触发。"""
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("图2.1 测试")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=False, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        # 即使 disabled，_check_caption_numbering 仍会执行检查
        # 开关判断在 _base() 的调用方
        mock_comment.assert_called_once()  # 章节号错误仍会报

    def test_empty_paragraph_skipped(self):
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "图", cfg)

        mock_comment.assert_not_called()

    def test_continued_table_correct_no_comment(self):
        """续表格式正确不添加批注。"""
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("续表5.3 测试")
        node = self._make_caption_figure(p, chapter=5, seq=3)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "表", cfg)

        mock_comment.assert_not_called()

    def test_continued_table_label_space_wrong(self):
        """续表 label_number_space=True 但无空格 → 添加批注。"""
        from wordformat.rules.caption import _check_caption_numbering

        doc = Document()
        p = self._make_paragraph("续表5.3 测试")
        node = self._make_caption_figure(p, chapter=5, seq=3)
        cfg = CaptionNumberingConfig(enabled=True, separator=".", label_number_space=True)

        with patch.object(node, "add_comment") as mock_comment:
            _check_caption_numbering(node, doc, "表", cfg)

        mock_comment.assert_called_once()
        assert "间距错误" in mock_comment.call_args[0][2]


# ======================== _apply_caption_numbering ========================


class TestApplyCaptionNumbering:
    """直接测试 _apply_caption_numbering 函数。"""

    def _make_caption_figure(self, paragraph, chapter=1, seq=1):
        from wordformat.rules.caption import CaptionFigure

        node = CaptionFigure(
            value={
                "category": "caption_figure",
                "fingerprint": "fig001",
                "chapter_number": chapter,
                "sequence_number": seq,
            },
            level=0,
            paragraph=paragraph,
        )
        return node

    def _make_paragraph(self, text):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run(text)
        return p

    def test_rewrites_caption(self):
        from wordformat.rules.caption import _apply_caption_numbering

        p = self._make_paragraph("图2-1 旧名称")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        _apply_caption_numbering(node, "图", cfg)

        assert p.text == "图1.1 旧名称"

    def test_preserves_caption_name(self):
        from wordformat.rules.caption import _apply_caption_numbering

        p = self._make_paragraph("图一.9 基于深度学习的图像识别算法")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        _apply_caption_numbering(node, "图", cfg)

        assert p.text == "图1.1 基于深度学习的图像识别算法"

    def test_apply_with_label_space(self):
        """label_number_space=true → 标签后加空格。"""
        from wordformat.rules.caption import _apply_caption_numbering

        p = self._make_paragraph("图2-1 旧名称")
        node = self._make_caption_figure(p, chapter=1, seq=1)
        cfg = CaptionNumberingConfig(enabled=True, separator=".", label_number_space=True)

        _apply_caption_numbering(node, "图", cfg)

        assert p.text == "图 1.1 旧名称"

    def test_continued_table_preserves_numbering(self):
        """续表保留原标题注编号并保留续前缀。"""
        from wordformat.rules.caption import _apply_caption_numbering

        p = self._make_paragraph("续表5.3 API接口测试结果")
        node = self._make_caption_figure(p, chapter=5, seq=3)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        _apply_caption_numbering(node, "表", cfg)

        assert p.text == "续表5.3 API接口测试结果"

    def test_continued_table_preserves_numbering_with_label_space(self):
        """续表 label_number_space=true → 标签后加空格。"""
        from wordformat.rules.caption import _apply_caption_numbering

        p = self._make_paragraph("续表5.3 API接口测试结果")
        node = self._make_caption_figure(p, chapter=5, seq=3)
        cfg = CaptionNumberingConfig(enabled=True, separator=".", label_number_space=True)

        _apply_caption_numbering(node, "表", cfg)

        assert p.text == "续表 5.3 API接口测试结果"

    def test_continued_table_corrects_separator(self):
        """续表修正分隔符。"""
        from wordformat.rules.caption import _apply_caption_numbering

        p = self._make_paragraph("续表5-3 测试")
        node = self._make_caption_figure(p, chapter=5, seq=3)
        cfg = CaptionNumberingConfig(enabled=True, separator=".")

        _apply_caption_numbering(node, "表", cfg)

        assert p.text == "续表5.3 测试"


# ======================== apply_format_check_to_all_nodes 集成 ========================


class TestCaptionNumberingIntegration:
    """测试 caption numbering 在 apply_format_check_to_all_nodes 中的集成。"""

    @pytest.fixture
    def caption_yaml(self, tmp_path):
        """创建带题注编号配置的临时 YAML 文件。"""
        yaml_content = """
global_format:
  alignment: '两端对齐'
  first_line_indent: '2字符'
  chinese_font_name: '宋体'
  english_font_name: 'Times New Roman'
  font_size: '小四'
  font_color: '黑色'
  builtin_style_name: '正文'
figures:
  caption_position: 'below'
  caption_prefix: '图'
  rules:
    caption_numbering:
      enabled: true
      separator: '.'
tables:
  caption_position: 'above'
  caption_prefix: '表'
  content:
    font_size: '五号'
  rules:
    caption_numbering:
      enabled: true
      separator: '.'
numbering:
  enabled: false
  captions:
    enabled: true
    separator: '.'
"""
        path = tmp_path / "caption_test.yaml"
        path.write_text(yaml_content, encoding="utf-8")
        return str(path)

    def _init_config(self, caption_yaml):
        from wordformat.config.config import get_config, init_config

        init_config(caption_yaml)
        return get_config()

    def _make_heading_node(self, children=None):
        node = MagicMock()
        node.value = {"category": "heading_level_1"}
        node.children = children or []
        node.paragraph = None
        return node

    def _make_caption_figure(self, paragraph):
        from wordformat.rules.caption import CaptionFigure

        node = CaptionFigure(
            value={"category": "caption_figure", "fingerprint": "fig001"},
            level=0,
            paragraph=paragraph,
        )
        return node

    def _make_caption_table(self, paragraph):
        from wordformat.rules.caption import CaptionTable

        node = CaptionTable(
            value={"category": "caption_table", "fingerprint": "tab001"},
            level=0,
            paragraph=paragraph,
        )
        return node

    def _make_paragraph(self, text):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run(text)
        return p

    def test_check_mode_injects_chapter_and_seq(self, caption_yaml):
        """验证遍历时章节号和顺序号被注入到 node.value 中。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p = self._make_paragraph("图1.1 系统架构图")
        fig = self._make_caption_figure(p)
        heading = self._make_heading_node(children=[fig])
        config = self._init_config(caption_yaml)

        with patch.object(fig, "add_comment"):
            apply_format_check_to_all_nodes(heading, doc, config, check=True)

        assert fig.value["chapter_number"] == 1
        assert fig.value["sequence_number"] == 1

    # 抑制格式检查产生的 comment，仅关注编号逻辑
    @pytest.fixture
    def _suppress_format_comments(self):
        """Mock ParagraphStyle 和 CharacterStyle 使其不产生格式差异。"""
        mock_ps = MagicMock()
        mock_ps.diff_from_paragraph.return_value = {}
        mock_ps.apply_to_paragraph.return_value = {}
        mock_cs = MagicMock()
        mock_cs.diff_from_run.return_value = {}
        mock_cs.apply_to_run.return_value = {}
        with patch(
            "wordformat.style.check_format.ParagraphStyle.from_config",
            return_value=mock_ps,
        ), patch(
            "wordformat.style.check_format.CharacterStyle",
            return_value=mock_cs,
        ):
            yield

    @pytest.mark.usefixtures("_suppress_format_comments")
    def test_check_mode_correct_no_comment(self, caption_yaml):
        """check 模式：格式正确，不添加批注。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p = self._make_paragraph("图1.1 系统架构图")
        fig = self._make_caption_figure(p)
        heading = self._make_heading_node(children=[fig])
        config = self._init_config(caption_yaml)

        with patch.object(fig, "add_comment") as mock_comment:
            apply_format_check_to_all_nodes(heading, doc, config, check=True)

        mock_comment.assert_not_called()

    @pytest.mark.usefixtures("_suppress_format_comments")
    def test_check_mode_wrong_chapter_adds_comment(self, caption_yaml):
        """check 模式：章节号错误应添加批注。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p = self._make_paragraph("图2.1 测试图")
        fig = self._make_caption_figure(p)
        heading = self._make_heading_node(children=[fig])
        config = self._init_config(caption_yaml)

        with patch.object(fig, "add_comment") as mock_comment:
            apply_format_check_to_all_nodes(heading, doc, config, check=True)

        mock_comment.assert_called_once()
        assert "章节号错误" in mock_comment.call_args[0][2]

    @pytest.mark.usefixtures("_suppress_format_comments")
    def test_apply_mode_rewrites_caption(self, caption_yaml):
        """apply 模式：重写题注文本。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p = self._make_paragraph("图2-1 旧名称")
        fig = self._make_caption_figure(p)
        heading = self._make_heading_node(children=[fig])
        config = self._init_config(caption_yaml)

        apply_format_check_to_all_nodes(heading, doc, config, check=False)

        assert p.text == "图1.1 旧名称"

    @pytest.mark.usefixtures("_suppress_format_comments")
    def test_per_type_counters(self, caption_yaml):
        """图和表使用独立计数器。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p_fig = self._make_paragraph("图1.1 图")
        p_tab = self._make_paragraph("表1.1 表")

        fig = self._make_caption_figure(p_fig)
        tab = self._make_caption_table(p_tab)
        heading = self._make_heading_node(children=[fig, tab])
        config = self._init_config(caption_yaml)

        with patch.object(fig, "add_comment") as mc1, \
             patch.object(tab, "add_comment") as mc2:
            apply_format_check_to_all_nodes(heading, doc, config, check=True)

        mc1.assert_not_called()
        mc2.assert_not_called()
        assert fig.value["sequence_number"] == 1
        assert tab.value["sequence_number"] == 1

    @pytest.mark.usefixtures("_suppress_format_comments")
    def test_multi_chapter_counters_reset(self, caption_yaml):
        """不同章节的题注计数器独立重置。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p1 = self._make_paragraph("图1.1 第一章图")
        p2 = self._make_paragraph("图2.1 第二章图")

        fig1 = self._make_caption_figure(p1)
        fig2 = self._make_caption_figure(p2)
        h1 = self._make_heading_node(children=[fig1])
        h2 = self._make_heading_node(children=[fig2])

        root = MagicMock()
        root.value = {"category": "top"}
        root.children = [h1, h2]
        root.paragraph = None

        config = self._init_config(caption_yaml)

        with patch.object(fig1, "add_comment"), \
             patch.object(fig2, "add_comment"):
            apply_format_check_to_all_nodes(root, doc, config, check=True)

        assert fig1.value["chapter_number"] == 1
        assert fig1.value["sequence_number"] == 1
        assert fig2.value["chapter_number"] == 2
        assert fig2.value["sequence_number"] == 1

    @pytest.mark.usefixtures("_suppress_format_comments")
    def test_disabled_skips_numbering_check(self, caption_yaml):
        """rules.caption_numbering.enabled=False 时不检查编号（但仍注入 chapter/seq）。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p = self._make_paragraph("图2.1 测试")
        fig = self._make_caption_figure(p)
        heading = self._make_heading_node(children=[fig])
        config = self._init_config(caption_yaml)
        config.figures.rules.caption_numbering.enabled = False

        with patch.object(fig, "add_comment") as mock_comment:
            apply_format_check_to_all_nodes(heading, doc, config, check=True)

        mock_comment.assert_not_called()
        assert fig.value["chapter_number"] == 1
        assert fig.value["sequence_number"] == 1

    @pytest.mark.usefixtures("_suppress_format_comments")
    def test_continued_table_does_not_increment_counter(self, caption_yaml):
        """续表不消耗题注编号，后续普通题注编号正确递增。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p1 = self._make_paragraph("表5.1 岗位信息表")
        p2 = self._make_paragraph("续表5.1 岗位信息表")
        p3 = self._make_paragraph("表5.2 薪资表")

        tab1 = self._make_caption_table(p1)
        tab_continued = self._make_caption_table(p2)
        tab2 = self._make_caption_table(p3)
        heading = self._make_heading_node(children=[tab1, tab_continued, tab2])
        config = self._init_config(caption_yaml)

        with patch.object(tab1, "add_comment"), \
             patch.object(tab_continued, "add_comment"), \
             patch.object(tab2, "add_comment"):
            apply_format_check_to_all_nodes(heading, doc, config, check=True)

        # 表5.1: 章节号5，序号1
        assert tab1.value["chapter_number"] == 1
        assert tab1.value["sequence_number"] == 1
        # 续表5.1: 保留原编号，不递增
        assert tab_continued.value["chapter_number"] == 5
        assert tab_continued.value["sequence_number"] == 1
        # 表5.2: 章节号1，序号2（续表未消耗编号）
        assert tab2.value["chapter_number"] == 1
        assert tab2.value["sequence_number"] == 2

    @pytest.mark.usefixtures("_suppress_format_comments")
    def test_continued_table_apply_mode_preserves_prefix(self, caption_yaml):
        """apply 模式：续表文本保留续前缀。"""
        from wordformat.set_style import apply_format_check_to_all_nodes

        doc = Document()
        p = self._make_paragraph("续表5.3 API接口测试结果")
        tab = self._make_caption_table(p)
        heading = self._make_heading_node(children=[tab])
        config = self._init_config(caption_yaml)

        apply_format_check_to_all_nodes(heading, doc, config, check=False)

        assert p.text == "续表5.3 API接口测试结果"
