#!/usr/bin/env python
"""
comprehensive tests for style modules:
  check_format, get_some, set_some, style_enum, utils
"""

import pytest
from types import SimpleNamespace
from unittest.mock import MagicMock, PropertyMock, patch as mock_patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from wordformat.style.diff import DIFFResult, CharacterStyle, ParagraphStyle
from wordformat.style.reader import (
    paragraph_get_alignment, paragraph_get_space_before, paragraph_get_space_after,
    paragraph_get_line_spacing, paragraph_get_first_line_indent,
    paragraph_get_builtin_style_name, run_get_font_name, run_get_font_size_pt,
    run_get_font_color, run_get_font_bold, run_get_font_italic,
    run_get_font_underline, GetIndent, _get_style_spacing,
)
from wordformat.style.writer import (
    run_set_font_name, set_paragraph_space_before_by_lines,
    set_paragraph_space_after_by_lines, _paragraph_space_by_lines,
    SetSpacing, SetLineSpacing, SetIndent, SetFirstLineIndent,
)
from wordformat.style.defs import (
    FontName, FontSize, FontColor, Alignment, LineSpacingRule, LineSpacing,
    FirstLineIndent, LeftIndent, RightIndent, BuiltInStyle, SpaceBefore, SpaceAfter,
    Spacing, UnitLabelEnum,
)
from wordformat.style.units import extract_unit_from_string, UnitResult


# ---------------------------------------------------------------------------
# fixtures & helpers
# ---------------------------------------------------------------------------

@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def mock_warning():
    w = MagicMock()
    for attr in ("bold", "italic", "underline", "font_size", "font_color",
                 "font_name", "alignment", "space_before", "space_after",
                 "line_spacing", "line_spacingrule", "first_line_indent",
                 "left_indent", "right_indent", "builtin_style_name"):
        setattr(w, attr, True)
    return w


@pytest.fixture
def mock_warning_off():
    w = MagicMock()
    for attr in ("bold", "italic", "underline", "font_size", "font_color",
                 "font_name", "alignment", "space_before", "space_after",
                 "line_spacing", "line_spacingrule", "first_line_indent",
                 "left_indent", "right_indent", "builtin_style_name"):
        setattr(w, attr, False)
    return w


def _set_warning(w):
    import wordformat.style.diff as m
    m.style_checks_warning = w


def _clear_warning():
    import wordformat.style.diff as m
    m.__dict__.pop("style_checks_warning", None)


# ===========================================================================
# set_some
# ===========================================================================

class TestSetSomeFontName:
    def test_set_and_verify_xml(self, doc):
        run = doc.add_paragraph().add_run("x")
        run_set_font_name(run, "黑体")
        assert run._element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"



class TestSetSomeSpaceByLines:
    def test_set_before_and_after(self, doc):
        p = doc.add_paragraph()
        set_paragraph_space_before_by_lines(p, 0.5)
        set_paragraph_space_after_by_lines(p, 1.0)
        assert paragraph_get_space_before(p) == 0.5
        assert paragraph_get_space_after(p) == 1.0

    def test_set_both(self, doc):
        p = doc.add_paragraph()
        _paragraph_space_by_lines(p, before_lines=0.3, after_lines=0.7)
        assert paragraph_get_space_before(p) == 0.3
        assert paragraph_get_space_after(p) == 0.7

    def test_zero_preserves_existing(self, doc):
        p = doc.add_paragraph()
        set_paragraph_space_after_by_lines(p, 1.0)
        set_paragraph_space_before_by_lines(p, 0.0)
        assert paragraph_get_space_after(p) == 1.0



class TestSetSpacingHang:
    def test_set_and_clamp(self, doc):
        p = doc.add_paragraph()
        SetSpacing.set_hang(p, "before", 0.5)
        assert paragraph_get_space_before(p) == 0.5
        SetSpacing.set_hang(p, "before", 50.0)
        assert paragraph_get_space_before(p) == 10.0



class TestSetLineSpacing:
    @pytest.mark.parametrize("method,val", [("set_pt", 20), ("set_cm", 1.0)])
    def test_sets_exactly_rule(self, doc, method, val):
        p = doc.add_paragraph()
        getattr(SetLineSpacing, method)(p, val)
        assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY



class TestSetIndent:
    def test_set_char_left(self, doc):
        p = doc.add_paragraph()
        SetIndent.set_char(p, "R", 2.5)
        assert GetIndent.left_indent(p) == 2.5

    def test_set_char_returns_true(self, doc):
        assert SetIndent.set_char(doc.add_paragraph(), "R", 2) is True

    def test_set_char_invalid_returns_false(self, doc):
        assert SetIndent.set_char(doc.add_paragraph(), "Z", 1) is False

    @pytest.mark.parametrize("method,indent_type", [
        ("set_pt", "R"), ("set_cm", "X"), ("set_inch", "R"), ("set_mm", "R"),
    ])
    def test_physical_units_set_indent(self, doc, method, indent_type):
        p = doc.add_paragraph()
        getattr(SetIndent, method)(p, indent_type, 1.0)
        attr = "left_indent" if indent_type == "R" else "right_indent"
        assert getattr(p.paragraph_format, attr) is not None

    def test_apply_indent_invalid_raises(self, doc):
        with pytest.raises(ValueError, match="无效的缩进类型"):
            SetIndent._apply_indent(doc.add_paragraph(), "Z", 10)



class TestSetFirstLineIndent:
    def test_set_and_clear(self, doc):
        p = doc.add_paragraph()
        SetFirstLineIndent.set_char(p, 2)
        assert paragraph_get_first_line_indent(p) == 2.0
        SetFirstLineIndent.clear(p)
        assert paragraph_get_first_line_indent(p) is None

    def test_clear_preserves_left_right(self, doc):
        p = doc.add_paragraph()
        SetIndent.set_pt(p, "R", 12)
        SetIndent.set_pt(p, "X", 12)
        SetFirstLineIndent.set_char(p, 2)
        SetFirstLineIndent.clear(p)
        ind = p._element.pPr.find(qn("w:ind"))
        assert ind is not None
        assert ind.get(qn("w:left")) is not None and ind.get(qn("w:right")) is not None
        assert ind.get(qn("w:firstLineChars")) is None

    def test_clear_no_indent_ok(self, doc):
        SetFirstLineIndent.clear(doc.add_paragraph())  # no crash

    @pytest.mark.parametrize("val,expected", [(0, 0.0), (-1, -1.0)])
    def test_first_line_and_hanging_indent(self, doc, val, expected):
        p = doc.add_paragraph()
        SetFirstLineIndent.set_char(p, val)
        assert paragraph_get_first_line_indent(p) == expected

    @pytest.mark.parametrize("method", ["set_pt", "set_cm", "set_inch", "set_mm"])
    def test_physical_units_no_firstLineChars(self, doc, method):
        p = doc.add_paragraph()
        getattr(SetFirstLineIndent, method)(p, 1.0)
        assert paragraph_get_first_line_indent(p) is None



# ===========================================================================
# Additional coverage tests for set_some.py
# ===========================================================================


class TestSetSpacingUnits:
    """Cover lines 137-140, 151-154, 165-168, 179-182, 193-196: _SetSpacing unit methods"""

    def test_set_pt_before(self, doc):
        """_SetSpacing.set_pt with spacing_type='before' (lines 151-152)"""
        p = doc.add_paragraph()
        SetSpacing.set_pt(p, "before", 12)
        assert p.paragraph_format.space_before is not None

    def test_set_pt_after(self, doc):
        """_SetSpacing.set_pt with spacing_type='after' (lines 153-154)"""
        p = doc.add_paragraph()
        SetSpacing.set_pt(p, "after", 12)
        assert p.paragraph_format.space_after is not None

    def test_set_cm_before(self, doc):
        """_SetSpacing.set_cm with spacing_type='before' (lines 165-166)"""
        p = doc.add_paragraph()
        SetSpacing.set_cm(p, "before", 1.0)
        assert p.paragraph_format.space_before is not None

    def test_set_cm_after(self, doc):
        """_SetSpacing.set_cm with spacing_type='after' (lines 167-168)"""
        p = doc.add_paragraph()
        SetSpacing.set_cm(p, "after", 1.0)
        assert p.paragraph_format.space_after is not None

    def test_set_inch_before(self, doc):
        """_SetSpacing.set_inch with spacing_type='before' (lines 179-180)"""
        p = doc.add_paragraph()
        SetSpacing.set_inch(p, "before", 0.5)
        assert p.paragraph_format.space_before is not None

    def test_set_inch_after(self, doc):
        """_SetSpacing.set_inch with spacing_type='after' (lines 181-182)"""
        p = doc.add_paragraph()
        SetSpacing.set_inch(p, "after", 0.5)
        assert p.paragraph_format.space_after is not None

    def test_set_mm_before(self, doc):
        """_SetSpacing.set_mm with spacing_type='before' (lines 193-194)"""
        p = doc.add_paragraph()
        SetSpacing.set_mm(p, "before", 5.0)
        assert p.paragraph_format.space_before is not None

    def test_set_mm_after(self, doc):
        """_SetSpacing.set_mm with spacing_type='after' (lines 195-196)"""
        p = doc.add_paragraph()
        SetSpacing.set_mm(p, "after", 5.0)
        assert p.paragraph_format.space_after is not None

    def test_set_hang_sets_twips_to_zero(self, doc):
        """set_hang 写入 w:before="0" 覆盖样式级 pt 间距"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "200")
        pPr.append(spacing)
        SetSpacing.set_hang(p, "before", 0.5)
        spacing_after = pPr.find(qn("w:spacing"))
        assert spacing_after.get(qn("w:before")) == "0"



class TestSetLineSpacingUnits:
    """Cover lines 234-235, 245-246: _SetLineSpacing additional unit methods"""

    def test_set_inch(self, doc):
        """_SetLineSpacing.set_inch (lines 234-235)"""
        p = doc.add_paragraph()
        SetLineSpacing.set_inch(p, 0.5)
        assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY

    def test_set_mm(self, doc):
        """_SetLineSpacing.set_mm (lines 245-246)"""
        p = doc.add_paragraph()
        SetLineSpacing.set_mm(p, 10.0)
        assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.EXACTLY



class TestSetIndentUnits:
    """Cover lines 288-291, 300-302, 461-463: _SetIndent and _SetFirstLineIndent unit methods"""

    def test_set_char_zero_clears_attrs(self, doc):
        """_SetIndent.set_char with value=0 clears attributes (lines 286-291)"""
        p = doc.add_paragraph()
        SetIndent.set_char(p, "R", 2)
        SetIndent.set_char(p, "R", 0)
        # After setting to 0, the char attr should be removed
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                assert ind.get(qn("w:leftChars")) is None

    def test_set_char_exception_returns_false(self, doc):
        """_SetIndent.set_char with exception -> returns False (lines 300-302)"""
        p = doc.add_paragraph()
        # Force an exception by removing _element
        original = p._element
        p._element = None
        assert SetIndent.set_char(p, "R", 2) is False
        p._element = original

    def test_set_pt(self, doc):
        """_SetIndent.set_pt (line 313-314)"""
        p = doc.add_paragraph()
        SetIndent.set_pt(p, "R", 12)
        assert p.paragraph_format.left_indent is not None

    def test_set_cm(self, doc):
        """_SetIndent.set_cm (line 325-326)"""
        p = doc.add_paragraph()
        SetIndent.set_cm(p, "X", 1.0)
        assert p.paragraph_format.right_indent is not None

    def test_set_inch(self, doc):
        """_SetIndent.set_inch (line 337-338)"""
        p = doc.add_paragraph()
        SetIndent.set_inch(p, "R", 0.5)
        assert p.paragraph_format.left_indent is not None

    def test_set_mm(self, doc):
        """_SetIndent.set_mm (line 349-350)"""
        p = doc.add_paragraph()
        SetIndent.set_mm(p, "X", 5.0)
        assert p.paragraph_format.right_indent is not None



class TestSetFirstLineIndentUnits:
    """Cover lines 478, 494, 510, 526: _SetFirstLineIndent physical unit methods"""

    def test_set_inch(self, doc):
        """_SetFirstLineIndent.set_inch (line 479)"""
        p = doc.add_paragraph()
        SetFirstLineIndent.set_inch(p, 0.5)
        assert p.paragraph_format.first_line_indent is not None

    def test_set_mm(self, doc):
        """_SetFirstLineIndent.set_mm (line 495)"""
        p = doc.add_paragraph()
        SetFirstLineIndent.set_mm(p, 5.0)
        assert p.paragraph_format.first_line_indent is not None

    def test_set_pt(self, doc):
        """_SetFirstLineIndent.set_pt (line 511)"""
        p = doc.add_paragraph()
        SetFirstLineIndent.set_pt(p, 24)
        assert p.paragraph_format.first_line_indent is not None

    def test_set_cm(self, doc):
        """_SetFirstLineIndent.set_cm (line 527)"""
        p = doc.add_paragraph()
        SetFirstLineIndent.set_cm(p, 0.5)
        assert p.paragraph_format.first_line_indent is not None

    def test_set_inch_clears_firstLineChars(self, doc):
        """set_inch clears firstLineChars to avoid priority conflict (line 478)"""
        p = doc.add_paragraph()
        SetFirstLineIndent.set_char(p, 2)
        SetFirstLineIndent.set_inch(p, 0.5)
        # firstLineChars should be cleared
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                assert ind.get(qn("w:firstLineChars")) is None

    def test_set_char_exception_returns_false(self, doc):
        """set_char with exception -> returns False (lines 461-463)"""
        p = doc.add_paragraph()
        # The try block accesses paragraph._element.get_or_add_pPr()
        # We need to make it fail inside the try, after clear() succeeds.
        # clear() calls get_or_add_pPr first, then we make it fail.
        call_count = [0]
        original_get_or_add = p._element.get_or_add_pPr
        def bad_get_or_add():
            call_count[0] += 1
            if call_count[0] > 1:
                # First call is from clear(), second is from the try block
                raise RuntimeError("test error")
            return original_get_or_add()
        p._element.get_or_add_pPr = bad_get_or_add
        assert SetFirstLineIndent.set_char(p, 2) is False
        p._element.get_or_add_pPr = original_get_or_add

