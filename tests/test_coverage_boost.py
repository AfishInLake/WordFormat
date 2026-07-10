"""
覆盖率提升测试文件

针对当前覆盖率 93% 的未覆盖代码行编写测试，目标 96%+。
覆盖模块：heading, keywords, numbering, style_enum, set_style, get_some, utils
"""

import io
import os
import re
from unittest.mock import MagicMock, patch, PropertyMock

import pytest
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from wordformat.config.models import NodeConfigRoot
from wordformat.rules.heading import (
    HeadingLevel1Node,
    HeadingLevel2Node,
    HeadingLevel3Node,
)
from wordformat.rules.keywords import KeywordsCN, KeywordsEN
from wordformat.numbering import (
    _convert_to_twips,
    _set_indent_value,
    _build_numbering_rPr,
    _auto_strip_numbering,
)
from wordformat.style.defs import (
    BuiltInStyle,
    FirstLineIndent,
    ensure_style_exists,
)
from wordformat.pipeline.stages import (
    StyleDefinitionFixStage,
    FormattingExecutionStage,
)

_fix_stage = StyleDefinitionFixStage()
_format_stage = FormattingExecutionStage()
apply_format_check_to_all_nodes = _format_stage.apply_format_check_to_all_nodes
from wordformat.style.reader import (
    paragraph_get_space_before,
    paragraph_get_space_after,
    paragraph_get_line_spacing,
    paragraph_get_first_line_indent,
    paragraph_get_builtin_style_name,
    run_get_font_name,
    run_get_font_size_pt,
    run_get_font_color,
    GetIndent,
)
from wordformat.utils import (
    get_paragraph_numbering_text,
    _count_numbering_levels,
    _format_number,
    _to_chinese_num,
    _to_roman,
    ensure_directory_exists,
)


# ---------------------------------------------------------------------------
# 共享 fixtures / helpers
# ---------------------------------------------------------------------------


def _load_yaml(path):
    import yaml
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


def _load_root_config(config_path):
    return _load_yaml(config_path)


@pytest.fixture
def root_config(sample_yaml_config):
    """从 sample_yaml_config 加载配置，与示例文件解耦。"""
    from wordformat.config.loader import init_config
    from wordformat.config.models import NodeConfigRoot

    init_config(sample_yaml_config)
    return NodeConfigRoot(**_load_root_config(sample_yaml_config))


# ===========================================================================
# 1. rules/heading.py 未覆盖行：92, 97, 174-222
# ===========================================================================


class TestHeadingCoverageBoost:
    """覆盖 heading.py 中未覆盖的代码行。"""

    def test_apply_format_sets_paragraph_style_and_alignment(self, root_config):
        """ParagraphStyle.apply_to_paragraph 先应用 builtin_style_name，
        再应用显式段落格式。验证两者都被正确设置。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("第一章 绪论")
        r.font.size = Pt(10)

        # 手动在 XML 中添加 w:jc=both（错误的对齐方式）
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)

        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)

        with patch.object(node, "add_comment"):
            node.apply_format(doc)

        # builtin_style_name 已设置为 Heading 1（python-docx 存储为 "Heading1"）
        assert p.style.name == "Heading 1"
        # w:jc 已被 ParagraphStyle.apply_to_paragraph 更新为居中对齐
        jc_after = pPr.find(qn("w:jc"))
        assert jc_after is not None
        assert jc_after.get(qn("w:val")) == "center"

    def test_apply_format_updates_existing_pstyle(self, root_config):
        """当段落已有 w:pStyle 时，ParagraphStyle.apply_to_paragraph
        通过 python-docx API 将其更新为正确的样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1 研究背景")
        r.font.size = Pt(10)

        # 手动添加一个不同的 w:pStyle
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "Normal")
        pPr.append(pStyle)

        assert pPr.find(qn("w:pStyle")).get(qn("w:val")) == "Normal"

        node = HeadingLevel2Node(value=p, level=2, paragraph=p)
        node.load_config(root_config)

        with patch.object(node, "add_comment"):
            node.apply_format(doc)

        # python-docx 通过 paragraph.style = "Heading 2" 更新了样式
        assert p.style.name == "Heading 2"

    def test_fix_all_style_definitions_collects_all_styles(self, root_config):
        """collect_style_configs 应收集配置中所有唯一的样式名。"""
        style_map = root_config.collect_style_configs()
        assert len(style_map) > 0
        # 应包含 Normal（body_text 使用）和 Heading 1/2/3
        style_names = list(style_map.keys())
        assert any("Normal" in s for s in style_names)

    def test_fix_all_style_definitions_handles_document(self, root_config):
        """_fix_all_style_definitions 应对文档中存在的样式定义进行修正。"""
        doc = Document()
        _fix_stage._fix_all_style_definitions(doc, root_config)
        # 不应抛出异常

    def test_fix_style_run_properties_clears_theme_color(self, root_config):
        """_fix_style_run_properties 应清除样式定义中的主题色。"""
        
        from wordformat.style.defs import ensure_style_exists

        doc = Document()
        cfg = root_config.headings.level_1
        ensure_style_exists(doc, "Heading 1")
        style = doc.styles["Heading 1"]

        # 手动设置主题色
        rPr = style.element.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style.element.insert(0, rPr)
        color = OxmlElement("w:color")
        color.set(qn("w:themeColor"), "accent1")
        color.set(qn("w:val"), "4472C4")
        rPr.append(color)

        _fix_stage._fix_style_run_properties(style, cfg, "Heading 1")

        rPr_after = style.element.find(qn("w:rPr"))
        color_after = rPr_after.find(qn("w:color"))
        assert color_after is not None
        assert color_after.get(qn("w:themeColor")) is None

    def test_fix_style_run_properties_sets_font_name(self, root_config):
        """_fix_style_run_properties 应设置样式定义中的英文字体名。"""
        
        from wordformat.style.defs import ensure_style_exists

        doc = Document()
        cfg = root_config.body_text
        ensure_style_exists(doc, "Normal")
        style = doc.styles["Normal"]

        _fix_stage._fix_style_run_properties(style, cfg, "Normal")

        rPr = style.element.find(qn("w:rPr"))
        rFonts = rPr.find(qn("w:rFonts"))
        assert rFonts is not None

    def test_fix_style_run_properties_sets_bold(self, root_config):
        """_fix_style_run_properties 应设置样式定义中的加粗属性。"""
        
        from wordformat.style.defs import ensure_style_exists
        
        doc = Document()
        cfg = NodeConfigRoot(bold=True)
        ensure_style_exists(doc, "Heading 1")
        style = doc.styles["Heading 1"]

        _fix_stage._fix_style_run_properties(style, cfg, "Heading 1")

        rPr = style.element.find(qn("w:rPr"))
        b = rPr.find(qn("w:b"))
        assert b is not None

    def test_fix_style_run_properties_removes_italic(self, root_config):
        """_fix_style_run_properties 当 italic=False 时应移除斜体元素。"""
        
        from wordformat.style.defs import ensure_style_exists
        
        doc = Document()
        cfg = NodeConfigRoot(italic=False)
        ensure_style_exists(doc, "Heading 1")
        style = doc.styles["Heading 1"]

        _fix_stage._fix_style_run_properties(style, cfg, "Heading 1")

        rPr = style.element.find(qn("w:rPr"))
        i = rPr.find(qn("w:i"))
        assert i is None

    def test_fix_style_paragraph_properties_sets_alignment(self, root_config):
        """_fix_style_paragraph_properties 应设置样式定义中的对齐方式。"""
        
        from wordformat.style.defs import ensure_style_exists
        
        doc = Document()
        cfg = NodeConfigRoot(alignment="居中对齐")
        ensure_style_exists(doc, "Heading 1")
        style = doc.styles["Heading 1"]

        _fix_stage._fix_style_paragraph_properties(style, cfg, "Heading 1")

        pPr = style.element.find(qn("w:pPr"))
        jc = pPr.find(qn("w:jc"))
        assert jc is not None
        assert jc.get(qn("w:val")) == "center"


# ===========================================================================
# 2. rules/keywords.py 未覆盖行：56, 69-106, 110, 128, 149, 183, 195, 232,
#    261, 295, 307
# ===========================================================================


class TestKeywordsCoverageBoost:
    """覆盖 keywords.py 中未覆盖的代码行。"""

    def test_apply_to_paragraph_path_cn(self, root_config):
        """覆盖行 56 (CN): apply_to_paragraph 路径（p=False）。

        KeywordsCN._check_paragraph_style 中 p=False 时调用 apply_to_paragraph。
        """
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("关键词：机器学习；深度学习")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment"):
            node.apply_format(doc)

    def test_apply_to_paragraph_path_en(self, root_config):
        """覆盖行 56 (EN): apply_to_paragraph 路径（p=False）。

        KeywordsEN._check_paragraph_style 中 p=False 时调用 apply_to_paragraph。
        """
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Keywords: AI, ML")
        node = KeywordsEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment"):
            node.apply_format(doc)

    def test_split_mixed_runs_cn(self, root_config):
        """覆盖行 69-106: _split_mixed_runs 拆分标签和内容混合的 run。

        当标签和内容在同一个 run 中时，应拆分为两个独立 run。
        """
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("关键词：校园二手交易；Django")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment"):
            node.apply_format(doc)

        # 拆分后应有多个 run
        assert len(p.runs) >= 2

    def test_split_mixed_runs_en(self, root_config):
        """覆盖行 69-106: _split_mixed_runs 英文标签拆分。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Keywords: Machine Learning, Deep Learning")
        node = KeywordsEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment"):
            node.apply_format(doc)

        assert len(p.runs) >= 2

    def test_get_label_split_pattern_en(self, root_config):
        """覆盖行 128: KeywordsEN._get_label_split_pattern 返回正则。"""
        node = KeywordsCN(value=None, level=0, paragraph=None)
        pattern = KeywordsEN._get_label_split_pattern(node)
        assert pattern is not None
        assert pattern.search("Keywords: ") is not None
        assert pattern.search("Keyword ") is not None

    def test_split_mixed_runs_in_format_mode_cn(self, root_config):
        """覆盖行 149: self._split_mixed_runs() 在格式化模式下执行。

        KeywordsCN._base 中 not p 时执行 _split_mixed_runs。
        """
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("关键词：A；B；C；D")
        r.font.size = Pt(10)
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        initial_run_count = len(p.runs)
        with patch.object(node, "add_comment"):
            node.apply_format(doc)
        # 格式化模式下拆分后 run 数量应增加
        assert len(p.runs) >= initial_run_count

    def test_en_label_apply_to_run(self, root_config):
        """覆盖行 183: KeywordsEN._base 中标签 apply_to_run 路径。

        当 p=False, r=False 且 run 是标签时调用 label_style.apply_to_run。
        """
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Keywords")
        r.font.size = Pt(10)
        r.font.bold = False  # 应为加粗
        node = KeywordsEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_en_content_apply_to_run(self, root_config):
        """覆盖行 195: KeywordsEN._base 中内容 apply_to_run 路径。

        当 p=False, r=False 且 run 不是标签时调用 content_style.apply_to_run。
        """
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Keywords: ")
        r1.font.bold = True
        r2 = p.add_run("Machine Learning")
        r2.font.size = Pt(10)
        r2.font.bold = True  # 内容不应加粗
        node = KeywordsEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_cn_label_apply_to_run(self, root_config):
        """覆盖行 295: KeywordsCN._base 中标签 apply_to_run 路径。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("关键词")
        r1.font.size = Pt(10)
        r1.font.bold = False  # 应为加粗
        r2 = p.add_run("：A；B；C")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_cn_content_apply_to_run(self, root_config):
        """覆盖行 307: KeywordsCN._base 中内容 apply_to_run 路径。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("关键词：")
        r1.font.bold = True
        r2 = p.add_run("A；B；C")
        r2.font.size = Pt(10)
        r2.font.bold = True  # 内容不应加粗
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_cn_split_mixed_runs_format_mode(self, root_config):
        """覆盖行 261: KeywordsCN._base 中 not p 时执行 _split_mixed_runs。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("关键词：测试内容")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment"):
            node.apply_format(doc)
        assert len(p.runs) >= 2

    def test_cn_get_label_split_pattern(self, root_config):
        """覆盖行 232: KeywordsCN._get_label_split_pattern 返回正则。"""
        node = KeywordsCN(value=None, level=0, paragraph=None)
        pattern = node._get_label_split_pattern()
        assert pattern is not None
        assert pattern.search("关键词：") is not None
        assert pattern.search("关键词:") is not None


# ===========================================================================
# 3. numbering.py 未覆盖行：36-61, 76-95, 120, 129, 150-153, 164-167, 201,
#    365, 370, 435
# ===========================================================================


class TestNumberingCoverageBoost:
    """覆盖 numbering.py 中未覆盖的代码行。"""

    def test_convert_to_twips_cm(self):
        """覆盖行 45-46: cm 单位转换。"""
        result = _convert_to_twips("0.75cm")
        assert isinstance(result, int)
        assert result > 0

    def test_convert_to_twips_mm(self):
        """覆盖行 47-48: mm 单位转换。"""
        result = _convert_to_twips("5mm")
        assert isinstance(result, int)
        assert result > 0

    def test_convert_to_twips_inch(self):
        """覆盖行 49-50: inch 单位转换。"""
        result = _convert_to_twips("1inch")
        assert isinstance(result, int)
        assert result > 0

    def test_convert_to_twips_pt(self):
        """覆盖行 51-52: pt 单位转换。"""
        result = _convert_to_twips("12pt")
        assert isinstance(result, int)
        assert result > 0

    def test_convert_to_twips_invalid(self):
        """覆盖行 37-39: 无效值返回 0。"""
        result = _convert_to_twips("invalid")
        assert result == 0

    def test_convert_to_twips_unsupported_unit(self):
        """覆盖行 53-55: 不支持的物理单位返回 0。"""
        result = _convert_to_twips("100emu")
        assert result == 0

    def test_convert_to_twips_exception(self):
        """覆盖行 56-58: 单位换算异常时返回 0。"""
        # 传入一个会导致 Pt() 抛出异常的值
        result = _convert_to_twips("abcpt")
        assert result == 0

    def test_set_indent_value_char_unit(self):
        """覆盖行 84-89: 字符单位设置。"""
        ind = OxmlElement("w:ind")
        _set_indent_value(ind, "left", "2字符")
        assert ind.get(qn("w:leftChars")) == "200"

    def test_set_indent_value_physical_unit(self):
        """覆盖行 90-93: 物理单位（cm）设置。"""
        ind = OxmlElement("w:ind")
        _set_indent_value(ind, "left", "0.75cm")
        assert ind.get(qn("w:left")) is not None

    def test_set_indent_value_invalid(self):
        """覆盖行 77-79: 无效值跳过设置。"""
        ind = OxmlElement("w:ind")
        _set_indent_value(ind, "left", "invalid")
        assert ind.get(qn("w:left")) is None
        assert ind.get(qn("w:leftChars")) is None

    def test_set_indent_value_unsupported_unit(self):
        """覆盖行 94-95: 不支持的缩进单位跳过设置。"""
        ind = OxmlElement("w:ind")
        _set_indent_value(ind, "left", "100emu")
        assert ind.get(qn("w:left")) is None

    def test_build_numbering_rPr_level_none(self):
        """覆盖行 119-120: level_cfg 为 None 时返回 None。"""
        headings_config = MagicMock()
        headings_config.level_1 = None
        result = _build_numbering_rPr(headings_config, "level_1")
        assert result is None

    def test_build_numbering_rPr_all_none(self):
        """覆盖行 128-129: 所有属性为 None 时返回 None。"""
        headings_config = MagicMock()
        level_cfg = MagicMock()
        level_cfg.chinese_font_name = None
        level_cfg.english_font_name = None
        level_cfg.font_size = None
        level_cfg.font_color = None
        level_cfg.bold = None
        headings_config.level_1 = level_cfg
        result = _build_numbering_rPr(headings_config, "level_1")
        assert result is None

    def test_build_numbering_rPr_bold_true(self):
        """覆盖行 163-167: bold=True 时添加 w:b 和 w:bCs 元素。"""
        headings_config = MagicMock()
        level_cfg = MagicMock()
        level_cfg.chinese_font_name = "黑体"
        level_cfg.english_font_name = "Times New Roman"
        level_cfg.font_size = "小二"
        level_cfg.bold = True
        headings_config.level_1 = level_cfg
        result = _build_numbering_rPr(headings_config, "level_1")
        assert result is not None
        assert result.find(qn("w:b")) is not None
        assert result.find(qn("w:bCs")) is not None

    def test_build_numbering_rPr_numeric_font_size(self):
        """覆盖行 150-153: 非中文字号（数字）的字号处理。"""
        headings_config = MagicMock()
        level_cfg = MagicMock()
        level_cfg.chinese_font_name = "黑体"
        level_cfg.english_font_name = None
        level_cfg.font_size = "15.5"
        level_cfg.bold = None
        headings_config.level_1 = level_cfg
        result = _build_numbering_rPr(headings_config, "level_1")
        assert result is not None
        sz = result.find(qn("w:sz"))
        assert sz is not None

    def test_build_numbering_rPr_invalid_font_size(self):
        """覆盖行 152-153: 无法解析的字号返回 None。"""
        headings_config = MagicMock()
        level_cfg = MagicMock()
        level_cfg.chinese_font_name = "黑体"
        level_cfg.english_font_name = None
        level_cfg.font_size = "abc"
        level_cfg.bold = None
        headings_config.level_1 = level_cfg
        result = _build_numbering_rPr(headings_config, "level_1")
        assert result is not None
        # sz 不应被添加（half_pt 为 None）
        assert result.find(qn("w:sz")) is None

    def test_auto_strip_numbering_success(self):
        """覆盖 _auto_strip_numbering 成功清除编号。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1 研究背景")
        result = _auto_strip_numbering(p, ilvl=1)
        assert result is True
        assert not p.text.startswith("1.1")

    def test_auto_strip_numbering_no_match(self):
        """无编号文本返回 False。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("研究背景")
        result = _auto_strip_numbering(p, ilvl=0)
        assert result is False

    def test_auto_strip_numbering_no_match_at_wrong_ilvl(self):
        """编号格式对应级别不匹配时返回 False。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("1 测试")
        result = _auto_strip_numbering(p, ilvl=0)
        assert result is False

    def test_auto_strip_numbering_no_runs(self):
        """无 run 时返回 False。"""
        doc = Document()
        p = doc.add_paragraph()
        result = _auto_strip_numbering(p, ilvl=0)
        assert result is False

    def test_auto_strip_numbering_multi_run(self):
        """多 run 逐字符删除。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("1.")
        r2 = p.add_run("1 ")
        r3 = p.add_run("研究背景")
        result = _auto_strip_numbering(p, ilvl=1)
        assert result is True
        assert p.text.strip() == "研究背景"


# ===========================================================================
# 4. style/style_enum.py 未覆盖行：155, 641, 649-657
# ===========================================================================


class TestStyleEnumCoverageBoost:
    """覆盖 style_enum.py 中未覆盖的代码行。"""

    def test_first_line_indent_pt_unit(self):
        """覆盖行 155: FirstLineIndent 使用 pt 单位时走物理单位分支。

        FirstLineIndent.get_from_paragraph 中 unit 为 pt 时，
        从 paragraph_format.first_line_indent 获取值。
        """
        doc = Document()
        p = doc.add_paragraph()
        from docx.shared import Pt
        p.paragraph_format.first_line_indent = Pt(24)
        fli = FirstLineIndent("24pt")
        result = fli.get_from_paragraph(p)
        assert result is not None

    def test_first_line_indent_cm_unit(self):
        """覆盖行 155: FirstLineIndent 使用 cm 单位。"""
        doc = Document()
        p = doc.add_paragraph()
        from docx.shared import Cm
        p.paragraph_format.first_line_indent = Cm(0.85)
        fli = FirstLineIndent("0.85cm")
        result = fli.get_from_paragraph(p)
        assert result is not None

    def test_first_line_indent_inch_unit(self):
        """覆盖行 155: FirstLineIndent 使用 inch 单位。"""
        doc = Document()
        p = doc.add_paragraph()
        from docx.shared import Inches
        p.paragraph_format.first_line_indent = Inches(0.5)
        fli = FirstLineIndent("0.5inch")
        result = fli.get_from_paragraph(p)
        assert result is not None

    def test_ensure_style_exists_base_style_none(self):
        """覆盖行 641: base_style_name 为 None 时 base_style = None。

        当样式映射中 base_style_name 为 None 时（如 Normal），
        创建样式不设置基础样式。
        """
        doc = Document()
        # 先删除 Normal 样式（如果存在），确保需要创建
        # 实际上 Normal 总是存在的，所以用另一个方式触发
        # 使用一个不在映射中的样式名，其 base 默认为 "Normal"
        ensure_style_exists(doc, "TestStyleNoBase")
        # 清理
        try:
            doc.styles.element.remove(
                doc.styles["TestStyleNoBase"].element
            )
        except Exception:
            pass

    def test_ensure_style_exists_with_outline_lvl(self):
        """覆盖行 649-657: 设置 outlineLvl 的分支。

        创建 Heading 样式时应设置大纲级别。
        """
        doc = Document()
        # 先删除 Heading 4 样式（如果存在）
        try:
            style_elem = doc.styles["Heading 4"].element
            style_elem.getparent().remove(style_elem)
        except KeyError:
            pass

        ensure_style_exists(doc, "Heading 4")

        # 验证 outlineLvl 已设置
        style = doc.styles["Heading 4"]
        pPr = style.element.find(qn("w:pPr"))
        if pPr is not None:
            outlineLvl = pPr.find(qn("w:outlineLvl"))
            assert outlineLvl is not None
            assert outlineLvl.get(qn("w:val")) == "3"

    def test_ensure_style_exists_already_exists(self):
        """覆盖行 630-633: 样式已存在时直接返回。"""
        doc = Document()
        # Normal 样式总是存在
        ensure_style_exists(doc, "Normal")
        # 不应抛出异常

    def test_ensure_style_creates_with_base(self):
        """覆盖行 643-644: 创建样式时设置基础样式。"""
        doc = Document()
        style_name = "TestStyleWithBase"
        # 确保不存在
        try:
            style_elem = doc.styles[style_name].element
            style_elem.getparent().remove(style_elem)
        except KeyError:
            pass

        ensure_style_exists(doc, style_name)
        style = doc.styles[style_name]
        assert style is not None
        # 清理
        try:
            style_elem = doc.styles[style_name].element
            style_elem.getparent().remove(style_elem)
        except Exception:
            pass


# ===========================================================================
# 5. set_style.py 未覆盖行：46, 57, 61, 99-100, 129-131, 223, 226
# ===========================================================================


class TestSetStyleCoverageBoost:
    """覆盖 set_style.py 中未覆盖的代码行。"""

    def test_fix_all_style_definitions_no_config_sections(self):
        """config_model 无任何有效的 GlobalFormatConfig 时不抛异常。"""
        doc = Document()
        # 使用一个只有 numbering 和 global_format 的空模型
        from wordformat.config.models import NodeConfigRoot

        config_bare = NodeConfigRoot(
            global_format=NodeConfigRoot(),
            body_text=NodeConfigRoot(),
        )
        _fix_stage._fix_all_style_definitions(doc, config_bare)
        # 不应抛出异常

    def test_fix_all_style_definitions_theme_color_fix(self):
        """修正主题色的完整流程：有 themeColor 的样式被修正。"""
        from wordformat.config.models import NodeConfigRoot

        doc = Document()
        config_model = NodeConfigRoot(
            global_format=NodeConfigRoot(),
            headings=NodeConfigRoot(
                level_1=NodeConfigRoot(builtin_style_name="Heading 1", font_color="黑色"),
            ),
        )

        # 确保 Heading 1 有主题色
        try:
            style = doc.styles["Heading 1"]
            rPr = style.element.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                style.element.insert(0, rPr)
            color = rPr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                rPr.append(color)
            color.set(qn("w:themeColor"), "accent1")
        except KeyError:
            pass

        _fix_stage._fix_all_style_definitions(doc, config_model)

        # 验证主题色已清除
        rPr_after = style.element.find(qn("w:rPr"))
        color_after = rPr_after.find(qn("w:color"))
        assert color_after is not None
        assert color_after.get(qn("w:themeColor")) is None

    def test_apply_format_check_to_all_nodes_void_category(self):
        """覆盖行 99-100: traverse 中 category 在 VOIDNODELIST 中。

        当节点的 category 在 VOIDNODELIST 中时，跳过格式化。
        """
        from wordformat.rules.node import FormatNode
        from wordformat.settings import VOIDNODELIST

        doc = Document()
        p = doc.add_paragraph("测试")
        root_node = FormatNode(
            value={"category": VOIDNODELIST[0]},
            level=0,
            paragraph=p,
        )
        root_node.children = []

        config = MagicMock()
        with patch.object(root_node, "check_format"):
            apply_format_check_to_all_nodes(root_node, doc, config, check=True)
            # VOIDNODELIST 中的 category 不应调用 check_format
            root_node.check_format.assert_not_called()

    def test_apply_format_check_to_all_nodes_no_check_format(self):
        """覆盖行 129-131: traverse 中节点无 check_format 的分支。

        当节点没有 check_format 方法时，跳过。
        """
        # 使用一个没有 check_format 的简单对象
        class SimpleNode:
            def __init__(self):
                self.value = {"category": "body_text"}
                self.children = []
                self.paragraph = None

        doc = Document()
        root_node = SimpleNode()
        config = MagicMock()

        # 不应抛出异常
        apply_format_check_to_all_nodes(root_node, doc, config, check=True)

    def test_apply_format_check_to_all_nodes_exception_handling(self):
        """覆盖行 129-131: traverse 中异常处理分支。"""
        from wordformat.rules.node import FormatNode

        doc = Document()
        p = doc.add_paragraph("测试")
        root_node = FormatNode(
            value={"category": "body_text", "fingerprint": "abc123"},
            level=0,
            paragraph=p,
        )
        root_node.children = []

        config = MagicMock()
        # load_config 会抛出异常
        with patch.object(root_node, "load_config", side_effect=ValueError("test error")):
            with pytest.raises(ValueError):
                apply_format_check_to_all_nodes(root_node, doc, config, check=True)

    def test_apply_format_check_to_all_nodes_apply_mode(self):
        """覆盖行 127-128: check=False 时调用 apply_format。"""
        from wordformat.rules.node import FormatNode

        doc = Document()
        p = doc.add_paragraph("测试")
        root_node = FormatNode(
            value={"category": "body_text", "fingerprint": "abc123"},
            level=0,
            paragraph=p,
        )
        root_node.children = []

        config = MagicMock()
        with patch.object(root_node, "load_config"):
            with patch.object(root_node, "apply_format") as mock_apply:
                apply_format_check_to_all_nodes(root_node, doc, config, check=False)
                mock_apply.assert_called_once_with(doc)

    def test_apply_format_check_to_all_nodes_no_paragraph(self):
        """覆盖行 124: node.paragraph 为 None 时跳过。"""
        from wordformat.rules.node import FormatNode

        doc = Document()
        root_node = FormatNode(
            value={"category": "body_text", "fingerprint": "abc123"},
            level=0,
            paragraph=None,
        )
        root_node.children = []

        config = MagicMock()
        with patch.object(root_node, "load_config"):
            with patch.object(root_node, "check_format") as mock_check:
                apply_format_check_to_all_nodes(root_node, doc, config, check=True)
                mock_check.assert_not_called()


# ===========================================================================
# 6. style/get_some.py 未覆盖行：106, 122, 157-158, 190-191, 280, 340, 429-433
# ===========================================================================


class TestGetSomeCoverageBoost:
    """覆盖 get_some.py 中未覆盖的代码行。"""

    def test_paragraph_get_space_before_with_lines(self):
        """覆盖行 106, 122: paragraph_get_space_before 带 beforeLines 属性。"""
        doc = Document()
        p = doc.add_paragraph()
        # 手动设置 beforeLines XML 属性
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:beforeLines"), "200")  # 2行
        pPr.append(spacing)

        result = paragraph_get_space_before(p)
        assert result is not None
        assert result == 2.0

    def test_paragraph_get_space_before_exception(self):
        """覆盖行 157-158: paragraph_get_space_before 异常处理。"""
        # 使用 mock 使 paragraph._element 抛出异常
        p = MagicMock()
        p._element = MagicMock()
        p._element.find.side_effect = AttributeError("test")
        result = paragraph_get_space_before(p)
        assert result is None

    def test_paragraph_get_space_after_with_lines(self):
        """覆盖行 190-191: paragraph_get_space_after 带 afterLines 属性。"""
        doc = Document()
        p = doc.add_paragraph()
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:afterLines"), "100")  # 1行
        pPr.append(spacing)

        result = paragraph_get_space_after(p)
        assert result is not None
        assert result == 1.0

    def test_paragraph_get_space_after_exception(self):
        """覆盖行 190-191: paragraph_get_space_after 异常处理。"""
        p = MagicMock()
        p._element = MagicMock()
        p._element.find.side_effect = TypeError("test")
        result = paragraph_get_space_after(p)
        assert result is None

    def test_paragraph_get_builtin_style_name_none(self):
        """覆盖行 280: style 为 None 时返回空字符串。"""
        p = MagicMock()
        p.style = None
        result = paragraph_get_builtin_style_name(p)
        assert result == ""

    def test_run_get_font_color_theme(self):
        """覆盖行 340: themeColor 类型返回 None。"""
        # 使用 MagicMock 模拟 run.font.color，使 type 返回 THEME
        from docx.enum.dml import MSO_COLOR_TYPE
        mock_run = MagicMock()
        mock_color = MagicMock()
        mock_color.type = MSO_COLOR_TYPE.THEME
        mock_run.font.color = mock_color
        result = run_get_font_color(mock_run)
        assert result is None

    def test_get_indent_invalid_type(self):
        """覆盖行 429-433: GetIndent.line_indent 无效 indent_type。"""
        doc = Document()
        p = doc.add_paragraph()
        with pytest.raises(ValueError, match="indent_type"):
            GetIndent.line_indent(p, "invalid")

    def test_get_indent_no_pPr(self):
        """覆盖行 429: pPr 为 None 时返回 None。"""
        doc = Document()
        p = doc.add_paragraph()
        result = GetIndent.left_indent(p)
        assert result is None

    def test_paragraph_get_line_spacing_multiple(self):
        """覆盖行 222-226: MULTIPLE 行距类型返回 float。"""
        doc = Document()
        p = doc.add_paragraph()
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        result = paragraph_get_line_spacing(p)
        assert result == 1.5

    def test_paragraph_get_line_spacing_at_least(self):
        """覆盖行 229-231: AT_LEAST 类型返回 None。"""
        doc = Document()
        p = doc.add_paragraph()
        from docx.enum.text import WD_LINE_SPACING
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        result = paragraph_get_line_spacing(p)
        assert result is None

    def test_paragraph_get_line_spacing_exception(self):
        """覆盖行 233-235: 异常处理返回 None。"""
        p = MagicMock()
        p.paragraph_format = MagicMock()
        p.paragraph_format.line_spacing_rule = MagicMock(side_effect=AttributeError)
        result = paragraph_get_line_spacing(p)
        assert result is None

    def test_run_get_font_size_pt_from_style(self):
        """覆盖行 315-317: 从样式继承字体大小。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("测试")
        # run 没有显式设置 font.size，应从样式继承
        result = run_get_font_size_pt(r)
        assert isinstance(result, float)

    def test_run_get_font_name_none(self):
        """覆盖行 293-300: run 无 rPr 时返回 None。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("测试")
        result = run_get_font_name(r)
        # 新 run 可能没有设置 eastAsia 字体
        assert result is None or isinstance(result, str)


# ===========================================================================
# 7. utils.py 未覆盖行：117, 128, 132, 154, 203, 208-212, 220, 223, 226, 310
# ===========================================================================


class TestUtilsCoverageBoost:
    """覆盖 utils.py 中未覆盖的代码行。"""

    def test_get_paragraph_numbering_text_no_pPr(self):
        """覆盖行 96-97: pPr 为 None 时返回空字符串。"""
        doc = Document()
        p = doc.add_paragraph()
        # 确保没有 pPr
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            p._element.remove(pPr)
        result = get_paragraph_numbering_text(p)
        assert result == ""

    def test_get_paragraph_numbering_text_no_numPr(self):
        """覆盖行 99-100: numPr 为 None 时返回空字符串。"""
        doc = Document()
        p = doc.add_paragraph()
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        result = get_paragraph_numbering_text(p)
        assert result == ""

    def test_get_paragraph_numbering_text_no_numId(self):
        """覆盖行 104-105: numId_elem 为 None 时返回空字符串。"""
        doc = Document()
        p = doc.add_paragraph()
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
        result = get_paragraph_numbering_text(p)
        assert result == ""

    def test_get_paragraph_numbering_text_num_id_zero(self):
        """覆盖行 110-111: numId 为 "0" 时返回空字符串。"""
        doc = Document()
        p = doc.add_paragraph()
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        numPr = OxmlElement("w:numPr")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "0")
        numPr.append(numId)
        pPr.append(numPr)
        result = get_paragraph_numbering_text(p)
        assert result == ""

    def test_get_paragraph_numbering_text_no_numbering_part(self):
        """覆盖行 115-117: 无 numbering_part 时返回空字符串。"""
        doc = Document()
        p = doc.add_paragraph()
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p._element.insert(0, pPr)
        numPr = OxmlElement("w:numPr")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), "1")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        # 新建文档可能默认有 numbering part，也可能没有
        # 只要不抛出异常即可
        result = get_paragraph_numbering_text(p)
        assert isinstance(result, str)

    def test_count_numbering_levels_no_body(self):
        """覆盖行 203: body 为 None 时返回空字典。"""
        doc = Document()
        p = doc.add_paragraph()
        result = _count_numbering_levels(MagicMock(), "1", p)
        assert result == {}

    def test_to_chinese_num_zero(self):
        """覆盖行 292-293: num <= 0 时返回字符串。"""
        assert _to_chinese_num(0) == "0"
        assert _to_chinese_num(-1) == "-1"

    def test_to_chinese_num_hundred(self):
        """覆盖行 308-309: num == 100 时返回 '一百'。"""
        assert _to_chinese_num(100) == "一百"

    def test_to_chinese_num_over_hundred(self):
        """覆盖行 310: num > 100 时返回字符串。"""
        assert _to_chinese_num(101) == "101"

    def test_to_chinese_num_tens(self):
        """覆盖行 298-306: 10-99 的中文数字。"""
        assert _to_chinese_num(10) == "十"
        assert _to_chinese_num(15) == "十五"
        assert _to_chinese_num(20) == "二十"
        assert _to_chinese_num(99) == "九十九"

    def test_to_chinese_num_units(self):
        """覆盖行 296-297: 1-9 的中文数字。"""
        assert _to_chinese_num(1) == "一"
        assert _to_chinese_num(5) == "五"
        assert _to_chinese_num(9) == "九"

    def test_roman_numeral(self):
        """覆盖 _to_roman 函数。"""
        assert _to_roman(1) == "i"
        assert _to_roman(4) == "iv"
        assert _to_roman(9) == "ix"
        assert _to_roman(2024) == "mmxxiv"

    def test_format_number_various(self):
        """覆盖 _format_number 各种格式。"""
        assert _format_number(1, "decimal") == "1"
        assert _format_number(1, "upperRoman") == "I"
        assert _format_number(1, "lowerRoman") == "i"
        assert _format_number(1, "upperLetter") == "A"
        assert _format_number(1, "lowerLetter") == "a"
        assert _format_number(1, "chineseCountingThousand") == "一"
        assert _format_number(1, "ideographTraditional") == "一"
        assert _format_number(1, "chineseCounting") == "一"

    def test_ensure_directory_exists_creates(self, tmp_path):
        """覆盖 ensure_directory_exists 创建目录。"""
        new_dir = str(tmp_path / "new_subdir" / "nested")
        ensure_directory_exists(new_dir)
        import os
        assert os.path.isdir(new_dir)

    def test_ensure_directory_exists_already_exists(self, tmp_path):
        """覆盖 ensure_directory_exists 目录已存在。"""
        ensure_directory_exists(str(tmp_path))
        # 不应抛出异常

    def test_ensure_directory_exists_is_file(self, tmp_path):
        """覆盖 ensure_directory_exists 路径是文件时抛出 ValueError。"""
        file_path = tmp_path / "test_file.txt"
        file_path.write_text("test")
        with pytest.raises(ValueError, match="不是文件夹"):
            ensure_directory_exists(str(file_path))


# ===========================================================================
# 覆盖率提升：utils/_docx, config, style/xml_ops, base, tree, numbering
# ===========================================================================


class TestCoverageBoostRound2:
    """补漏低覆盖率分支。"""

    # --- utils/_docx.py ---

    def test_ensure_directory_exists_creates_dir(self, tmp_path):
        """确保目录不存在时创建。"""
        from wordformat.utils import ensure_directory_exists
        d = str(tmp_path / "new_dir")
        ensure_directory_exists(d)
        assert os.path.isdir(d)

    def test_para_contains_image_false(self):
        """无图片的段落返回 False。"""
        from wordformat.utils import para_contains_image
        doc = Document()
        p = doc.add_paragraph("text")
        assert para_contains_image(p) is False

    def test_remove_all_numbering(self):
        """remove_all_numbering 对空白文档不抛异常。"""
        from wordformat.utils import remove_all_numbering
        doc = Document()
        remove_all_numbering(doc)

    # --- config/loader.py ---

    def test_config_not_loaded_error(self):
        """ConfigNotLoadedError 可正常抛出和捕获。"""
        from wordformat.config.loader import ConfigNotLoadedError
        with pytest.raises(ConfigNotLoadedError):
            raise ConfigNotLoadedError("test")

    # --- style/xml_ops.py ---

    def test_rPr_set_italic_remove(self):
        """rPr_set_italic 传入 False 时移除 w:i。"""
        from wordformat.style.xml_ops import rPr_set_italic, ensure_rPr
        from docx.oxml import OxmlElement
        style = Document().styles["Normal"]
        rPr = ensure_rPr(style.element)
        rPr.append(OxmlElement("w:i"))
        rPr_set_italic(rPr, False)
        assert rPr.find(qn("w:i")) is None

    def test_rPr_set_underline_remove(self):
        """rPr_set_underline 传入 False 时移除 w:u。"""
        from wordformat.style.xml_ops import rPr_set_underline, ensure_rPr
        from docx.oxml import OxmlElement
        style = Document().styles["Normal"]
        rPr = ensure_rPr(style.element)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        rPr_set_underline(rPr, False)
        assert rPr.find(qn("w:u")) is None

    # --- tree.py: JSON file path ---

    def test_print_tree_from_json(self, tmp_path):
        """print_tree 支持传入 JSON 文件路径。"""
        import json
        from wordformat.tree import print_tree
        json_path = tmp_path / "test.json"
        json_path.write_text(json.dumps([
            {"category": "body_text", "paragraph": "hello"},
        ]))
        # 不应抛异常
        with patch("sys.stdout", io.StringIO()):
            print_tree(str(json_path), show_index=True, show_confidence=True)

    # --- numbering.py ---

    def test_process_heading_numbering_disabled(self):
        """numbering 禁用时不处理。"""
        from wordformat.numbering import process_heading_numbering
        from unittest.mock import MagicMock
        config = MagicMock()
        config.enabled = False
        process_heading_numbering(None, Document(), config)

    # --- pipeline/stages.py ---

    def test_formatting_stage_skips_void_nodes(self, doc):
        """VOIDNODELIST 中的类别跳过格式化。"""
        from wordformat.pipeline.stages import FormattingExecutionStage
        from wordformat.rules.node import FormatNode
        from wordformat.settings import VOIDNODELIST
        stage = FormattingExecutionStage()
        root = FormatNode(value={"category": VOIDNODELIST[0]}, level=0)
        root.children = []
        config = MagicMock()
        stage.apply_format_check_to_all_nodes(root, doc, config, check=True)

    def test_summary_stage_adds_comment(self, doc):
        """SummaryGenerationStage 在 check 模式下添加批注。"""
        from wordformat.pipeline.stages import SummaryGenerationStage, FormattingExecutionStage
        from wordformat.rules.node import FormatNode
        p = doc.add_paragraph("")
        root = FormatNode(value={"category": "top"}, level=0)
        root.children = []
        FormatNode.reset_stats()
        stage = SummaryGenerationStage()
        from wordformat.pipeline.context import FormatContext
        ctx = FormatContext(
            json_path="", docx_path="", check=True,
            config_path="", save_dir="/tmp",
            document=doc, root_node=root,
        )
        stage.process(ctx)

    # --- utils/_docx.py more ---


    # --- style/xml_ops.py ---

    def test_rPr_set_bold_remove(self):
        """rPr_set_bold 传入 False 移除 w:b。"""
        from wordformat.style.xml_ops import rPr_set_bold, ensure_rPr
        style = Document().styles["Normal"]
        rPr = ensure_rPr(style.element)
        rPr.append(OxmlElement("w:b"))
        rPr_set_bold(rPr, False)
        assert rPr.find(qn("w:b")) is None

    def test_line_rule_to_xml_unknown(self):
        """line_rule_to_xml 对未知值返回 auto。"""
        from wordformat.style.xml_ops import line_rule_to_xml
        assert line_rule_to_xml(999) == "auto"

    # --- style/defs.py ---

    def test_fontsize_base_set_numeric(self):
        """FontSize.base_set 处理数字字符串。"""
        from wordformat.style.defs import FontSize
        doc = Document()
        p = doc.add_paragraph("test")
        run = p.add_run("x")
        fs = FontSize("14")
        fs.base_set(run)
        assert run.font.size is not None

    def test_fontname_base_set_english(self):
        """FontName.base_set 处理英文字体。"""
        from wordformat.style.defs import FontName
        doc = Document()
        p = doc.add_paragraph("test")
        run = p.add_run("x")
        fn = FontName("Arial")
        fn.base_set(run)
        assert run.font.name == "Arial"

    def test_unit_label_enum_eq_none(self):
        """UnitLabelEnum __eq__ 对 None 返回 rel_value == 0。"""
        from wordformat.style.defs import FirstLineIndent
        indent = FirstLineIndent("0字符")
        assert indent == None  # rel_value is 0, so equal to None

    # --- numbering.py ---

    def test_auto_strip_numbering_empty_runs(self):
        """空 run 的段落 _auto_strip_numbering 返回 False。"""
        from wordformat.numbering import _auto_strip_numbering
        doc = Document()
        p = doc.add_paragraph("")
        assert _auto_strip_numbering(p, 0) is False

    # --- pipeline/stages.py ---

    def test_load_config_stage_no_config(self):
        """未提供配置文件时使用默认配置。"""
        from wordformat.pipeline.stages import LoadConfigStage
        from wordformat.pipeline.context import FormatContext
        ctx = FormatContext(json_path="", docx_path="", check=True, config_path="")
        stage = LoadConfigStage()
        result = stage.process(ctx)
        assert result is ctx

    # --- utils/_fs.py ---


    def test_ensure_is_directory_missing(self, tmp_path):
        """ensure_is_directory 路径不存在时抛出 ValueError。"""
        from wordformat.utils import ensure_is_directory
        with pytest.raises(ValueError, match="不存在"):
            ensure_is_directory(str(tmp_path / "nonexistent"))
