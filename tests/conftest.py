"""
共享 fixtures 和测试工具

将所有测试文件中重复定义的 fixture 集中管理。
"""
import os
import json
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# 项目根目录
PROJECT_ROOT = Path(__file__).parent.parent


# ==================== Document Fixtures ====================

@pytest.fixture
def doc():
    """创建一个空的 Document 对象"""
    return Document()


@pytest.fixture
def doc_with_paragraphs(doc):
    """创建包含多个段落的 Document"""
    for text in ["第一章 绪论", "1.1 研究背景", "这是正文内容", "参考文献"]:
        doc.add_paragraph(text)
    return doc


@pytest.fixture
def doc_with_runs(doc):
    """创建包含带格式 Run 的 Document"""
    p = doc.add_paragraph()
    run = p.add_run("测试文本")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.bold = True
    return doc


@pytest.fixture
def temp_docx(tmp_path):
    """创建一个临时 docx 文件并返回路径"""
    doc = Document()
    doc.add_paragraph("测试段落")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def temp_json(tmp_path):
    """创建一个临时 JSON 文件（模拟 set_tag_main 输出）"""
    data = [
        {
            "category": "heading_level_1",
            "score": 0.95,
            "paragraph": "第一章 绪论",
            "fingerprint": "abc123",
        },
        {
            "category": "body_text",
            "score": 0.88,
            "paragraph": "这是正文内容",
            "fingerprint": "def456",
        },
    ]
    path = tmp_path / "test.json"
    path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    return str(path)


# ==================== Config Fixtures ====================

_INLINE_YAML = """\
global_format:
  alignment: '两端对齐'
  space_before: "0.5行"
  space_after: "0.5行"
  line_spacingrule: "1.5倍行距"
  line_spacing: '1.5倍'
  left_indent: "0字符"
  right_indent: "0字符"
  first_line_indent: '2字符'
  builtin_style_name: '正文'
  chinese_font_name: '宋体'
  english_font_name: 'Times New Roman'
  font_size: '小四'
  font_color: '黑色'
  bold: false
  italic: false
  underline: false
abstract:
  chinese:
    chinese_title:
      alignment: '居中对齐'
      first_line_indent: '0字符'
      chinese_font_name: '黑体'
      font_size: '小二'
      bold: true
    chinese_content:
      alignment: '两端对齐'
  english:
    english_title:
      alignment: '居中对齐'
      first_line_indent: '0字符'
      font_size: '四号'
      bold: true
    english_content:
      alignment: '两端对齐'
  keywords:
    chinese:
      label:
        chinese_font_name: '黑体'
        font_size: '三号'
        bold: true
      alignment: '左对齐'
      first_line_indent: '0字符'
      chinese_font_name: '宋体'
      font_size: '小四'
      bold: false
      rules:
        keyword_count:
          enabled: true
          count_min: 3
          count_max: 5
        trailing_punctuation:
          enabled: true
    english:
      label:
        font_size: '三号'
        bold: true
      alignment: '左对齐'
      first_line_indent: '0字符'
      font_size: '小四'
      bold: false
      rules:
        keyword_count:
          enabled: true
          count_min: 3
          count_max: 5
headings:
  level_1:
    alignment: '居中对齐'
    first_line_indent: '0字符'
    chinese_font_name: '黑体'
    english_font_name: 'Times New Roman'
    font_size: '小二'
    font_color: '黑色'
    bold: true
    builtin_style_name: 'Heading 1'
  level_2:
    alignment: '左对齐'
    first_line_indent: '0字符'
    chinese_font_name: '黑体'
    font_size: '三号'
    builtin_style_name: 'Heading 2'
  level_3:
    alignment: '左对齐'
    first_line_indent: '0字符'
    chinese_font_name: '黑体'
    font_size: '小四'
    builtin_style_name: 'Heading 3'
body_text:
  alignment: '两端对齐'
  chinese_font_name: '宋体'
  english_font_name: 'Times New Roman'
  font_size: '小四'
figures:
  caption_prefix: '图'
tables:
  caption_prefix: '表'
  content:
    font_size: '五号'
references:
  title:
    alignment: '居中对齐'
    first_line_indent: '0字符'
    chinese_font_name: '黑体'
    font_size: '三号'
    bold: true
  content:
    alignment: '左对齐'
    first_line_indent: '0字符'
    font_size: '五号'
acknowledgements:
  title:
    alignment: '居中对齐'
    first_line_indent: '0字符'
    chinese_font_name: '黑体'
    font_size: '小二'
    bold: true
  content:
    alignment: '两端对齐'
    first_line_indent: '2字符'
numbering:
  enabled: false
"""

@pytest.fixture
def config_path(tmp_path):
    """用内联 YAML 生成临时配置文件，不依赖外部文件。"""
    config_file = tmp_path / "test_config.yaml"
    config_file.write_text(_INLINE_YAML, encoding="utf-8")
    return str(config_file)


@pytest.fixture
def mock_config():
    """创建一个 mock 配置对象，模拟 NodeConfigRoot"""
    config = MagicMock()
    config.style_checks_warning = MagicMock(
        bold=True, italic=True, underline=True,
        font_size=True, font_name=True, font_color=True,
        alignment=True, space_before=True, space_after=True,
        line_spacing=True, line_spacingrule=True,
        left_indent=True, right_indent=True,
        first_line_indent=True, builtin_style_name=True,
    )
    config.numbering = MagicMock(enabled=False)
    return config




@pytest.fixture
def sample_yaml_config(tmp_path):
    """用内联 YAML 生成临时配置文件。"""
    path = tmp_path / "test_config.yaml"
    path.write_text(_INLINE_YAML, encoding="utf-8")
    return str(path)


# ==================== Mock Fixtures ====================

@pytest.fixture
def mock_onnx_infer():
    """Mock ONNX 推理函数，返回可控结果"""
    def _infer(texts):
        if isinstance(texts, str):
            texts = [texts]
        return [
            {"text": t, "label": "body_text", "pred_id": 0, "score": 0.9}
            for t in texts
        ]
    return _infer


@pytest.fixture
def mock_onnx_single():
    """Mock 单条 ONNX 推理"""
    return lambda text: {"label": "body_text", "score": 0.9}


@pytest.fixture(autouse=True)
def reset_config():
    """每个测试前后自动清理配置状态"""
    from wordformat.config.loader import clear_config
    clear_config()
    yield
    clear_config()


@pytest.fixture(autouse=True)
def reset_style_warning():
    """每个测试前后重置警告缓存。"""
    from wordformat.style import diff
    diff._warnings = None
    yield
    diff._warnings = None
