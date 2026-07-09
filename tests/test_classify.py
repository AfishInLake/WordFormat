"""
Core 模块综合测试

覆盖 tree.py, utils.py, rules/node.py, numbering.py, settings.py
"""
import os
import pytest
from io import StringIO
from unittest.mock import MagicMock, patch

from docx import Document
from docx.oxml.ns import qn

from wordformat.tree import Tree, Stack, print_tree
from wordformat.rules.node import TreeNode, FormatNode
from wordformat.numbering import (
    _auto_strip_numbering,
    _strip_reference_numbering,
    apply_auto_numbering,
    create_numbering_definition,
    process_heading_numbering,
)
from wordformat.utils import (
    get_file_name,
    ensure_is_directory,
    ensure_directory_exists,
    _to_roman,
    _to_chinese_num,
    load_yaml_with_merge,
    get_paragraph_numbering_text,
    remove_all_numbering,
    _format_number,
    _get_level_fmt,
    _count_numbering_levels,
)
from wordformat.style.reader import _get_style_spacing
from wordformat.base import DocxBase
from wordformat import settings


# ============================================================
# tree.py — Tree
# ============================================================


# ============================================================
# base.py — DocxBase
# ============================================================


class TestDocxBase:
    """测试 DocxBase 类的初始化和 parse 方法"""

    def test_init(self, temp_docx):
        """测试 DocxBase 初始化"""
        base = DocxBase(temp_docx, "/fake/config.yaml")
        assert base.docx_file == temp_docx
        assert base.document is not None
        assert base.re_dict == {}

    def _create_multi_para_docx(self, tmp_path, texts):
        """辅助方法：创建包含多个段落的 docx 文件"""
        doc = Document()
        for text in texts:
            doc.add_paragraph(text)
        path = str(tmp_path / "multi.docx")
        doc.save(path)
        return path

    def test_parse_with_mocked_batch_infer(self, tmp_path):
        """测试 parse 方法使用 mock 的批量推理"""
        path = self._create_multi_para_docx(tmp_path, ["绪论", "研究背景", "正文内容"])
        mock_batch_results = [
            {"label": "heading_level_1", "score": 0.95},
            {"label": "body_text", "score": 0.88},
            {"label": "body_text", "score": 0.75},
        ]

        with patch("wordformat.base.onnx_batch_infer", return_value=mock_batch_results):
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert len(result) == 3
        assert result[0]["category"] == "heading_level_1"
        assert result[1]["category"] == "body_text"
        assert "paragraph" in result[0]
        assert "score" in result[0]

    def test_parse_low_score_forced_to_body_text(self, tmp_path):
        """测试低置信度结果被强制设为 body_text"""
        path = self._create_multi_para_docx(tmp_path, ["绪论", "研究背景"])
        mock_batch_results = [
            {"label": "heading_level_1", "score": 0.95},
            {"label": "heading_level_2", "score": 0.3},
        ]

        with patch("wordformat.base.onnx_batch_infer", return_value=mock_batch_results):
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert result[0]["category"] == "heading_level_1"
        assert result[1]["category"] == "body_text"
        assert "强制设为" in result[1]["comment"]


    def test_parse_batch_failure_fallback_to_single(self, temp_docx):
        """测试批量推理失败时降级到单条推理"""
        mock_single_result = {"label": "body_text", "score": 0.9}

        with patch("wordformat.base.onnx_batch_infer", side_effect=RuntimeError("ONNX error")):
            with patch("wordformat.base.onnx_single_infer", return_value=mock_single_result):
                base = DocxBase(temp_docx, "/fake/config.yaml")
                result = base.parse()

        # temp_docx 只有一个段落
        assert len(result) >= 1
        assert all(item["category"] == "body_text" for item in result)

    def test_parse_empty_document(self, tmp_path):
        """测试空文档的解析"""
        doc = Document()
        path = str(tmp_path / "empty.docx")
        doc.save(path)

        with patch("wordformat.base.onnx_batch_infer", return_value=[]) as mock_infer:
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert result == []
        mock_infer.assert_not_called()

    def test_parse_batches_correctly(self, tmp_path):
        """测试按 BATCH_SIZE 分批推理"""
        path = self._create_multi_para_docx(tmp_path, [f"段落 {i}" for i in range(5)])

        call_count = 0

        def mock_batch(texts):
            nonlocal call_count
            call_count += 1
            return [{"label": "body_text", "score": 0.9}] * len(texts)

        with patch("wordformat.base.onnx_batch_infer", side_effect=mock_batch):
            with patch("wordformat.base.BATCH_SIZE", 2):
                base = DocxBase(path, "/fake/config.yaml")
                result = base.parse()

        assert len(result) == 5
        # BATCH_SIZE=2, 5 个段落 → 3 次调用 (2+2+1)
        assert call_count == 3

    def test_parse_includes_numbering_text(self, tmp_path):
        """测试解析时包含自动编号文字"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        doc = Document()

        # 先移除已有的 numbering 关系
        rels = doc.part.rels
        to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
        for k in to_remove:
            del rels[k]

        p = doc.add_paragraph("绪论")

        # 添加 numbering
        numbering_elm = OxmlElement("w:numbering")
        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), "0")
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "decimal")
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "%1.")
        lvl.append(lvlText)
        abstract_num.append(lvl)
        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), "0")
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), "1")
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), "0")
        numPr.append(ilvl_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)

        path = str(tmp_path / "numbered.docx")
        doc.save(path)

        with patch("wordformat.base.onnx_batch_infer", return_value=[{"text": "1. 绪论", "label": "body_text", "pred_id": 0, "score": 0.9}]):
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert len(result) == 1
        # 段落文本应包含编号 "1. 绪论"
        assert result[0]["paragraph"] == "1. 绪论"


# ============================================================
# utils.py — _format_number 额外覆盖测试
# ============================================================



