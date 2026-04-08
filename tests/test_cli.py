#! /usr/bin/env python
# -*- coding: utf-8 -*-
import os
import tempfile
import pytest
import time
from unittest import mock
from pathlib import Path

from wordformat.cli import validate_file, main


def test_validate_file():
    """测试 validate_file 函数"""
    # 测试有效的文件
    with tempfile.NamedTemporaryFile(delete=False) as f:
        temp_file = f.name
    try:
        result = validate_file(temp_file, "测试文件")
        assert result == os.path.abspath(temp_file)
    finally:
        os.unlink(temp_file)

    # 测试不存在的文件
    with pytest.raises(Exception):
        validate_file("nonexistent_file_12345.txt", "测试文件")

    # 测试文件夹路径
    with tempfile.TemporaryDirectory() as temp_dir:
        with pytest.raises(Exception):
            validate_file(temp_dir, "测试文件")


@mock.patch('sys.argv')
def test_main_gj_generate_json(mock_argv):
    """测试 main 函数 gj(generate-json) 模式（自动生成JSON）"""
    from docx import Document

    # 临时 docx
    doc = Document()
    doc.add_paragraph("测试正文")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx = f.name
    doc.save(temp_docx)

    # 临时配置
    with tempfile.NamedTemporaryFile(suffix='.yaml', delete=False, mode='w', encoding='utf-8') as f:
        cfg = """
style_checks_warning:
  bold: true
global_format:
  alignment: "两端对齐"
abstract:
  chinese:
    chinese_title:
      alignment: "居中对齐"
  english:
    english_title:
      alignment: "居中对齐"
  keywords:
    chinese:
      count_min: 3
      count_max: 6
headings:
  level_1:
    alignment: "居中对齐"
  level_2:
    alignment: "左对齐"
  level_3:
    alignment: "左对齐"
body_text:
  <<: *global_format
figures:
  caption_position: below
tables:
  caption_position: above
references:
  title:
    alignment: "居中对齐"
acknowledgements:
  title:
    alignment: "居中对齐"
"""
        f.write(cfg)
        temp_config = f.name

    # 输出目录
    output_dir = tempfile.mkdtemp()

    try:
        # 模拟命令：wf gj -d xx.docx -c config.yaml -o output
        mock_argv.__getitem__.side_effect = lambda i: [
            "wordformat",
            "gj",
            "-d", temp_docx,
            "-c", temp_config,
            "-o", output_dir
        ][i]
        mock_argv.__len__.return_value = 7

        main()

        # 验证生成 json
        json_files = list(Path(output_dir).glob("*.json"))
        assert len(json_files) >= 1
        assert json_files[0].exists()

    finally:
        if os.path.exists(temp_docx):
            os.unlink(temp_docx)
        if os.path.exists(temp_config):
            os.unlink(temp_config)
        import shutil
        if os.path.exists(output_dir):
            shutil.rmtree(output_dir)


@mock.patch('sys.argv')
def test_main_cf_check_format(mock_argv):
    """测试 main 函数 cf(check-format) 模式"""
    from docx import Document

    # 临时 docx
    doc = Document()
    doc.add_paragraph("测试段落")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx = f.name
    doc.save(temp_docx)

    # 临时配置
    with tempfile.NamedTemporaryFile(suffix='.yaml', delete=False, mode='w', encoding='utf-8') as f:
        f.write("""
style_checks_warning:
  bold: true
global_format:
  alignment: "左对齐"
headings:
  level_1:
    alignment: "居中对齐"
""")
        temp_config = f.name

    # 临时 JSON
    with tempfile.NamedTemporaryFile(suffix='.json', delete=False, mode='w', encoding='utf-8') as f:
        f.write('[{"fingerprint": "test", "category": "body_text", "content": "测试"}]')
        temp_json = f.name

    output_dir = tempfile.mkdtemp()

    try:
        # 模拟命令：wf cf -d xx.docx -c xx.yaml -f xx.json -o output
        mock_argv.__getitem__.side_effect = lambda i: [
            "wordformat",
            "cf",
            "-d", temp_docx,
            "-c", temp_config,
            "-f", temp_json,
            "-o", output_dir
        ][i]
        mock_argv.__len__.return_value = 9

        main()

    finally:
        for p in [temp_docx, temp_config, temp_json]:
            if os.path.exists(p):
                os.unlink(p)
        import shutil
        if os.path.exists(output_dir):
            shutil.rmtree(output_dir)


@mock.patch('sys.argv')
def test_main_af_apply_format(mock_argv):
    """测试 main 函数 af(apply-format) 模式"""
    from docx import Document

    # 临时 docx
    doc = Document()
    doc.add_paragraph("测试段落")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx = f.name
    doc.save(temp_docx)

    # 临时配置
    with tempfile.NamedTemporaryFile(suffix='.yaml', delete=False, mode='w', encoding='utf-8') as f:
        f.write("""
style_checks_warning:
  bold: true
global_format:
  alignment: "两端对齐"
""")
        temp_config = f.name

    # 临时 JSON
    with tempfile.NamedTemporaryFile(suffix='.json', delete=False, mode='w', encoding='utf-8') as f:
        f.write('[{"fingerprint": "test", "category": "body_text", "content": "测试"}]')
        temp_json = f.name

    output_dir = tempfile.mkdtemp()

    try:
        # 模拟命令：wf af -d xx.docx -c xx.yaml -f xx.json -o output
        mock_argv.__getitem__.side_effect = lambda i: [
            "wordformat",
            "af",
            "-d", temp_docx,
            "-c", temp_config,
            "-f", temp_json,
            "-o", output_dir
        ][i]
        mock_argv.__len__.return_value = 9

        main()

    finally:
        for p in [temp_docx, temp_config, temp_json]:
            if os.path.exists(p):
                os.unlink(p)
        import shutil
        if os.path.exists(output_dir):
            shutil.rmtree(output_dir)


@mock.patch('sys.argv')
def test_main_no_args_show_help(mock_argv):
    """无参数直接打印帮助"""
    mock_argv.__getitem__.side_effect = lambda i: ["wordformat"][i]
    mock_argv.__len__.return_value = 1
    main()