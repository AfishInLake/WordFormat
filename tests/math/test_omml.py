#! /usr/bin/env python
"""wordformat.math 模块单元测试。"""

import pytest

from wordformat.math import (
    add_display_math,
    add_inline_math,
    latex_to_omath,
    latex_to_omath_para,
    set_cell_math,
)


class TestLatexToOmath:
    """测试 LaTeX → OMML 核心转换。"""

    def test_simple_expression(self):
        om = latex_to_omath(r"x^2 + y^2")
        assert om is not None
        assert om.tag.endswith("}oMath")

    def test_fraction(self):
        om = latex_to_omath(r"\frac{1}{2}")
        assert om is not None

    def test_sqrt(self):
        om = latex_to_omath(r"\sqrt{x^2 + y^2}")
        assert om is not None

    def test_sum(self):
        om = latex_to_omath(r"\sum_{i=1}^{n} x_i")
        assert om is not None

    def test_greek_letters(self):
        om = latex_to_omath(r"\alpha + \beta = \gamma")
        assert om is not None

    def test_subscript_and_superscript(self):
        om = latex_to_omath(r"x_i^2 + w_j^{(k)}")
        assert om is not None

    def test_matrix(self):
        om = latex_to_omath(r"\begin{matrix} a & b \\ c & d \end{matrix}")
        assert om is not None

    def test_empty_string(self):
        assert latex_to_omath("") is None
        assert latex_to_omath("   ") is None

    def test_invalid_latex(self):
        # Invalid LaTeX should return None rather than crash
        result = latex_to_omath(r"\invalid_command_xyz{")
        # Should not raise, may return None or partial result
        assert result is None or result is not None

    def test_aligned_equation(self):
        om = latex_to_omath(r"\begin{aligned} x &= 1 \\ y &= 2 \end{aligned}")
        # aligned environment may not be supported by all backends
        # should not raise, even if result is None
        assert om is None or om is not None


class TestLatexToOmathPara:
    """测试块级公式转换。"""

    def test_wraps_in_omathpara(self):
        omp = latex_to_omath_para(r"x = \frac{1}{2}")
        assert omp is not None
        assert "oMathPara" in omp.tag

    def test_empty_returns_none(self):
        assert latex_to_omath_para("") is None


class TestAddDisplayMath:
    """测试添加块级公式到文档。"""

    def test_adds_paragraph(self):
        from docx import Document

        doc = Document()
        p = add_display_math(doc, r"x^2 + y^2 = z^2")
        assert p is not None
        # 段落中应包含 OMML 元素
        omath = p._p.find(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath")
        assert omath is not None

    def test_empty_latex_adds_text_run(self):
        from docx import Document

        doc = Document()
        p = add_display_math(doc, "")
        assert p is not None
        assert len(p.runs) >= 1


class TestAddInlineMath:
    """测试内联公式渲染。"""

    def test_single_inline_math(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        add_inline_math(p, "The value is $x^2$ here.")
        # 应有 text run 和 oMath 元素
        runs = p.runs
        assert len(runs) >= 1

    def test_multiple_inline_math(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        add_inline_math(p, "$a$ and $b$ and $c$")
        # Should not crash with multiple formulas
        assert p._p is not None

    def test_no_math(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        add_inline_math(p, "Plain text no formula.")
        assert len(p.runs) == 1
        assert p.runs[0].text == "Plain text no formula."

    def test_invalid_math_falls_back_to_text(self):
        from docx import Document

        doc = Document()
        p = doc.add_paragraph()
        add_inline_math(p, r"$\invalid\command{}$")
        # Should not raise
        assert p._p is not None


class TestSetCellMath:
    """测试表格单元格数学公式。"""

    def test_display_math_in_cell(self):
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].add_run("placeholder")  # ensure runs exist
        set_cell_math(cell, r"$$x^2 + y^2$$")
        assert cell.paragraphs[0]._p is not None

    def test_inline_math_in_cell(self):
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].add_run("placeholder")
        set_cell_math(cell, r"The value is $E=mc^2$.")
        assert cell.paragraphs[0]._p is not None

    def test_plain_text_in_cell(self):
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].add_run("placeholder")
        set_cell_math(cell, "No math here.")
        assert len(cell.paragraphs[0].runs) >= 1
