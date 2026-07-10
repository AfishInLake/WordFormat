"""
rules 模块测试 —— 聚焦真实行为验证，无填充。
"""
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from wordformat.config.models import NodeConfigRoot
from wordformat.rules import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleContentCN,
    AbstractTitleContentEN,
    AbstractTitleEN,
    Acknowledgements,
    AcknowledgementsCN,
    BodyText,
    CaptionFigure,
    CaptionTable,
    FormatNode,
    HeadingLevel1Node,
    HeadingLevel2Node,
    HeadingLevel3Node,
    KeywordsCN,
    KeywordsEN,
    ReferenceEntry,
    References,
)
from wordformat.rules.node import FormatNode as FormatNodeBase

# ---------------------------------------------------------------------------
# 共享 fixtures / helpers
# ---------------------------------------------------------------------------


def _make_node(cls, text="测试文本"):
    """创建一个带 paragraph 的 FormatNode 实例。"""
    doc = Document()
    p = doc.add_paragraph(text)
    return cls(value=p, level=0, paragraph=p)


def _load_root_config(config_path):
    """从 YAML 路径加载配置 dict。"""
    return _load_yaml(config_path)


def _load_yaml(path):
    import yaml
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


@pytest.fixture
def root_config(sample_yaml_config):
    """从 sample_yaml_config 加载 NodeConfigRoot，与示例文件解耦。"""
    from wordformat.config.loader import init_config
    init_config(sample_yaml_config)
    return _load_root_config(sample_yaml_config)


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def para(doc):
    return doc.add_paragraph("测试段落")


@pytest.fixture
def run_with_text(para):
    r = para.add_run("关键词")
    r.font.size = Pt(12)
    r.font.bold = True
    return r


# ---------------------------------------------------------------------------
# 1. FormatNode 基类行为
# ---------------------------------------------------------------------------



class TestKeywordsLogic:
    """关键词节点的标签识别、数量校验、标点校验。"""

    def test_cn_label_detection(self):
        """中文关键词标签识别（使用真实 paragraph runs）。"""
        doc = Document()
        p = doc.add_paragraph("关键词")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        assert node._check_keyword_label(p.runs[0]) is True

        p2 = doc.add_paragraph("关 键 词")
        node.paragraph = p2
        assert node._check_keyword_label(p2.runs[0]) is True

        p3 = doc.add_paragraph("机器学习")
        node.paragraph = p3
        assert node._check_keyword_label(p3.runs[0]) is False

    def test_en_label_detection(self):
        """英文关键词标签识别（使用真实 paragraph runs）。"""
        doc = Document()
        p = doc.add_paragraph("Keywords")
        node = KeywordsEN(value=p, level=0, paragraph=p)
        assert node._check_keyword_label(p.runs[0]) is True

        p2 = doc.add_paragraph("Keyword")
        node.paragraph = p2
        assert node._check_keyword_label(p2.runs[0]) is True

        p3 = doc.add_paragraph("KEY WORDS")
        node.paragraph = p3
        assert node._check_keyword_label(p3.runs[0]) is True

        p4 = doc.add_paragraph("machine learning")
        node.paragraph = p4
        assert node._check_keyword_label(p4.runs[0]) is False

    def test_cn_count_validation_too_few(self, root_config):
        """中文关键词数量不足时应触发 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("关键词：机器学习")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        # 应至少有一条数量不足的 comment
        texts = [c.kwargs["text"] for c in mock_comment.call_args_list]
        assert any("数量过少" in t for t in texts)

    def test_cn_count_validation_too_many(self, root_config):
        """中文关键词数量超限时应触发 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("关键词：A；B；C；D；E；F")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        texts = [c.kwargs["text"] for c in mock_comment.call_args_list]
        assert any("数量过多" in t for t in texts)

    def test_cn_trailing_punct_detection(self, root_config):
        """中文关键词末尾标点校验。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("关键词：机器学习；深度学习；")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        texts = [c.kwargs["text"] for c in mock_comment.call_args_list]
        assert any("标点错误" in t for t in texts)

    def test_en_count_validation_too_few(self, root_config):
        """英文关键词数量不足。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Keywords: AI")
        node = KeywordsEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        texts = [c.kwargs["text"] for c in mock_comment.call_args_list]
        assert any("数量过少" in t for t in texts)

# ---------------------------------------------------------------------------
# 6. _base 实现实际调用 diff/apply 逻辑
# ---------------------------------------------------------------------------
