#! /usr/bin/env python
# -*- coding: utf-8 -*-
import tempfile
import os
from unittest import mock
from docx import Document

from wordformat.base import DocxBase


def test_docx_base_init():
    """测试DocxBase类的初始化功能"""
    # 创建临时docx文件
    doc = Document()
    doc.add_paragraph("测试段落")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx_path = f.name
    doc.save(temp_docx_path)
    
    try:
        # 使用示例配置文件
        config_path = "example/undergrad_thesis.yaml"
        
        # 初始化DocxBase实例
        dox = DocxBase(temp_docx_path, configpath=config_path)
        
        # 验证实例属性
        assert dox.docx_file == temp_docx_path
        assert dox.document is not None
        
    finally:
        # 清理临时文件
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)


def test_docx_base_parse():
    """测试DocxBase类的parse方法"""
    # 创建临时docx文件
    doc = Document()
    doc.add_paragraph("测试段落")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx_path = f.name
    doc.save(temp_docx_path)
    
    try:
        # 使用示例配置文件
        config_path = "example/undergrad_thesis.yaml"
        
        # 初始化DocxBase实例并调用parse方法
        dox = DocxBase(temp_docx_path, configpath=config_path)
        result = dox.parse()
        
        # 验证解析结果
        assert isinstance(result, list)
        
    finally:
        # 清理临时文件
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)


def test_docx_base_get_tag_by_regex():
    """测试DocxBase类的get_tag_by_regex方法"""
    # 创建临时docx文件
    doc = Document()
    doc.add_paragraph("测试段落")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx_path = f.name
    doc.save(temp_docx_path)
    
    try:
        # 使用示例配置文件
        config_path = "example/undergrad_thesis.yaml"
        
        # 初始化DocxBase实例
        dox = DocxBase(temp_docx_path, configpath=config_path)
        
        # 测试空输入
        assert dox.get_tag_by_regex("") == ("", "")
        assert dox.get_tag_by_regex(None) == ("", "")
        
        # 测试正常输入
        # 注意：这里需要根据实际配置文件中的正则表达式来测试
        # 由于配置可能不同，这里只测试方法能正常执行
        result = dox.get_tag_by_regex("测试段落")
        assert isinstance(result, tuple)
        assert len(result) == 2
        
    finally:
        # 清理临时文件
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)

@mock.patch('wordformat.base.onnx_batch_infer')
def test_docx_base_parse_with_low_confidence(mock_onnx_batch_infer):
    """测试DocxBase类的parse方法（低置信度情况）"""
    # 创建临时docx文件
    doc = Document()
    doc.add_paragraph("测试段落")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx_path = f.name
    doc.save(temp_docx_path)
    
    try:
        # 模拟低置信度的推理结果
        mock_onnx_batch_infer.return_value = [{"预测标签": "test", "预测概率": 0.5}]
        
        # 使用示例配置文件
        config_path = "example/undergrad_thesis.yaml"
        
        # 初始化DocxBase实例并调用parse方法
        dox = DocxBase(temp_docx_path, configpath=config_path)
        result = dox.parse()
        
        # 验证解析结果
        assert isinstance(result, list)
        assert len(result) == 1
        assert result[0]["category"] == "body_text"  # 低置信度应被设为body_text
        
    finally:
        # 清理临时文件
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)

@mock.patch('wordformat.base.onnx_batch_infer')
def test_docx_base_parse_with_heading_fulu(mock_onnx_batch_infer):
    """测试DocxBase类的parse方法（遇到heading_fulu提前终止）"""
    # 创建临时docx文件
    doc = Document()
    doc.add_paragraph("测试段落1")
    doc.add_paragraph("测试段落2")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx_path = f.name
    doc.save(temp_docx_path)
    
    try:
        # 模拟包含heading_fulu的推理结果
        mock_onnx_batch_infer.return_value = [
            {"预测标签": "heading_fulu", "预测概率": 0.9},
            {"预测标签": "body_text", "预测概率": 0.9}
        ]
        
        # 使用示例配置文件
        config_path = "example/undergrad_thesis.yaml"
        
        # 初始化DocxBase实例并调用parse方法
        dox = DocxBase(temp_docx_path, configpath=config_path)
        result = dox.parse()
        
        # 验证解析结果（应该只返回第一个结果，因为遇到heading_fulu会提前终止）
        assert isinstance(result, list)
        assert len(result) == 1
        assert result[0]["category"] == "heading_fulu"
        
    finally:
        # 清理临时文件
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)

@mock.patch('wordformat.base.onnx_batch_infer')
@mock.patch('wordformat.base.onnx_single_infer')
def test_docx_base_parse_with_batch_error(mock_onnx_single_infer, mock_onnx_batch_infer):
    """测试DocxBase类的parse方法（批量推理失败，降级到单条处理）"""
    # 创建临时docx文件
    doc = Document()
    doc.add_paragraph("测试段落")
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        temp_docx_path = f.name
    doc.save(temp_docx_path)
    
    try:
        # 模拟批量推理失败
        mock_onnx_batch_infer.side_effect = Exception("Batch inference error")
        # 模拟单条推理成功，返回与批量推理一致的格式
        mock_onnx_single_infer.return_value = {"预测标签": "body_text", "预测概率": 0.9}
        
        # 使用示例配置文件
        config_path = "example/undergrad_thesis.yaml"
        
        # 初始化DocxBase实例并调用parse方法
        dox = DocxBase(temp_docx_path, configpath=config_path)
        result = dox.parse()
        
        # 验证解析结果
        assert isinstance(result, list)
        
    finally:
        # 清理临时文件
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
