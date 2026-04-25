"""
Core 模块综合测试

覆盖 tree.py, utils.py, rules/node.py, numbering.py, settings.py
"""
import os
import pytest
from io import StringIO
from unittest.mock import MagicMock, patch

from docx import Document
from docx.oxml.ns import qn

from wordformat.tree import Tree, Stack, print_tree
from wordformat.rules.node import TreeNode, FormatNode
from wordformat.numbering import (
    strip_manual_numbering,
    apply_auto_numbering,
    create_numbering_definition,
    process_heading_numbering,
)
from wordformat.utils import (
    check_duplicate_fingerprints,
    get_file_name,
    ensure_is_directory,
    ensure_directory_exists,
    _to_roman,
    _to_chinese_num,
    load_yaml_with_merge,
    get_paragraph_numbering_text,
    remove_all_numbering,
    _format_number,
    _get_level_fmt,
    _count_numbering_levels,
)
from wordformat.style.get_some import _get_style_spacing
from wordformat.base import DocxBase
from wordformat import settings


# ============================================================
# tree.py — Tree
# ============================================================


class TestTreeCreation:
    """Tree 创建与基本属性"""

    def test_create_tree_with_root_value(self):
        tree = Tree("root")
        assert tree.root.value == "root"

    def test_create_tree_with_dict_value(self):
        tree = Tree({"category": "top"})
        assert tree.root.value == {"category": "top"}

    def test_tree_repr(self):
        tree = Tree("my_root")
        assert repr(tree) == "Tree(root=my_root)"

    def test_is_empty_always_false(self):
        """BUG: is_empty() 始终返回 False，因为 root 在 __init__ 中总是被赋值"""
        tree = Tree("root")
        assert tree.is_empty() is False


class TestTreeTraversals:
    """三种遍历：前序、后序、层序"""

    def setup_method(self):
        tree = Tree("A")
        b = tree.root.add_child("B")
        c = tree.root.add_child("C")
        b.add_child("D")
        b.add_child("E")
        c.add_child("F")
        self.tree = tree

    def test_preorder(self):
        assert list(self.tree.preorder()) == ["A", "B", "D", "E", "C", "F"]

    def test_postorder(self):
        assert list(self.tree.postorder()) == ["D", "E", "B", "F", "C", "A"]

    def test_level_order(self):
        assert list(self.tree.level_order()) == ["A", "B", "C", "D", "E", "F"]

    def test_preorder_single_node(self):
        tree = Tree("only")
        assert list(tree.preorder()) == ["only"]

    def test_postorder_single_node(self):
        tree = Tree("only")
        assert list(tree.postorder()) == ["only"]

    def test_level_order_single_node(self):
        tree = Tree("only")
        assert list(tree.level_order()) == ["only"]


class TestTreeFindAndMetrics:
    """find_by_condition, height, size"""

    def setup_method(self):
        tree = Tree("A")
        b = tree.root.add_child("B")
        c = tree.root.add_child("C")
        b.add_child("D")
        b.add_child("E")
        c.add_child("F")
        self.tree = tree

    def test_find_by_condition_exists(self):
        node = self.tree.find_by_condition(lambda v: v == "E")
        assert node is not None
        assert node.value == "E"

    def test_find_by_condition_not_exists(self):
        node = self.tree.find_by_condition(lambda v: v == "Z")
        assert node is None

    def test_find_by_condition_first_match(self):
        node = self.tree.find_by_condition(lambda v: isinstance(v, str) and len(v) == 1)
        assert node.value == "A"  # DFS 先遇到根

    def test_height_single_node(self):
        assert Tree("x").height() == 0

    def test_height_deep_tree(self):
        tree = Tree("1")
        tree.root.add_child("2").add_child("3").add_child("4")
        assert tree.height() == 3

    def test_height_balanced_tree(self):
        assert self.tree.height() == 2

    def test_size_single_node(self):
        assert Tree("x").size() == 1

    def test_size(self):
        assert self.tree.size() == 6


# ============================================================
# tree.py — Stack
# ============================================================


class TestStack:
    """Stack 的全部操作"""

    def test_push_and_pop(self):
        s = Stack()
        s.push(10)
        s.push(20)
        assert s.pop() == 20
        assert s.pop() == 10

    def test_peek(self):
        s = Stack()
        s.push("hello")
        assert s.peek() == "hello"
        assert s.size() == 1  # peek 不弹出

    def test_peek_safe_on_empty(self):
        s = Stack()
        assert s.peek_safe() is None

    def test_peek_safe_returns_top(self):
        s = Stack()
        s.push(42)
        assert s.peek_safe() == 42

    def test_is_empty(self):
        s = Stack()
        assert s.is_empty() is True
        s.push(1)
        assert s.is_empty() is False

    def test_size(self):
        s = Stack()
        assert s.size() == 0
        s.push("a")
        s.push("b")
        assert s.size() == 2

    def test_clear(self):
        s = Stack()
        s.push(1)
        s.push(2)
        s.clear()
        assert s.is_empty() is True
        assert s.size() == 0

    def test_pop_empty_raises(self):
        s = Stack()
        with pytest.raises(IndexError, match="pop from empty stack"):
            s.pop()

    def test_peek_empty_raises(self):
        s = Stack()
        with pytest.raises(IndexError, match="peek from empty stack"):
            s.peek()

    def test_bool_truthy(self):
        s = Stack()
        assert bool(s) is False
        s.push(1)
        assert bool(s) is True

    def test_repr(self):
        s = Stack()
        s.push(1)
        s.push(2)
        assert "1" in repr(s)
        assert "2" in repr(s)


# ============================================================
# tree.py — print_tree
# ============================================================


class TestPrintTree:
    """print_tree 输出捕获"""

    def test_print_single_node(self):
        node = TreeNode("hello")
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(node)
        output = buf.getvalue()
        assert "hello" in output
        assert "└──" in output

    def test_print_dict_value_node(self):
        node = TreeNode({
            "category": "body_text",
            "paragraph": "这是一段正文内容",
            "fingerprint": "fp001",
        })
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(node)
        output = buf.getvalue()
        assert "body_text" in output

    def test_print_tree_with_children(self):
        root = TreeNode("root")
        root.add_child("child1")
        root.add_child("child2")
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(root)
        output = buf.getvalue()
        assert "root" in output
        assert "child1" in output
        assert "child2" in output


# ============================================================
# utils.py — check_duplicate_fingerprints
# ============================================================


class TestCheckDuplicateFingerprints:
    def test_no_duplicates(self):
        data = [
            {"fingerprint": "aaa"},
            {"fingerprint": "bbb"},
        ]
        # 不应抛出异常
        check_duplicate_fingerprints(data)

    def test_has_duplicates(self):
        data = [
            {"fingerprint": "aaa"},
            {"fingerprint": "aaa"},
        ]
        # 不抛异常，仅打日志
        check_duplicate_fingerprints(data)

    def test_missing_fingerprint_raises(self):
        data = [{"category": "body_text"}]
        with pytest.raises(ValueError, match="fingerprint"):
            check_duplicate_fingerprints(data)

    def test_empty_list(self):
        check_duplicate_fingerprints([])


# ============================================================
# utils.py — get_file_name
# ============================================================


class TestGetFileName:
    @pytest.mark.parametrize(
        "path, expected",
        [
            ("/home/user/doc.docx", "doc"),
            ("simple.txt", "simple"),
            ("/path/to/my.file.name.pdf", "my.file.name"),
        ],
    )
    def test_extracts_name(self, path, expected):
        assert get_file_name(path) == expected


# ============================================================
# utils.py — ensure_is_directory
# ============================================================


class TestEnsureIsDirectory:
    def test_valid_directory(self, tmp_path):
        # tmp_path 本身就是目录
        ensure_is_directory(str(tmp_path))

    def test_nonexistent_path_raises(self):
        with pytest.raises(ValueError, match="路径不存在"):
            ensure_is_directory("/nonexistent/path/that/does/not/exist")

    def test_file_instead_of_directory_raises(self, tmp_path):
        file_path = tmp_path / "afile.txt"
        file_path.write_text("hello")
        with pytest.raises(ValueError, match="不是一个文件夹"):
            ensure_is_directory(str(file_path))


# ============================================================
# utils.py — ensure_directory_exists
# ============================================================


class TestEnsureDirectoryExists:
    def test_create_new_directory(self, tmp_path):
        new_dir = str(tmp_path / "sub" / "dir")
        ensure_directory_exists(new_dir)
        assert os.path.isdir(new_dir)

    def test_existing_directory_no_error(self, tmp_path):
        ensure_directory_exists(str(tmp_path))

    def test_existing_file_raises(self, tmp_path):
        file_path = tmp_path / "blocked.txt"
        file_path.write_text("data")
        with pytest.raises(ValueError, match="不是文件夹"):
            ensure_directory_exists(str(file_path))


# ============================================================
# utils.py — _to_roman
# ============================================================


class TestToRoman:
    @pytest.mark.parametrize(
        "num, expected",
        [
            (1, "i"),
            (4, "iv"),
            (5, "v"),
            (9, "ix"),
            (10, "x"),
            (40, "xl"),
            (50, "l"),
            (90, "xc"),
            (100, "c"),
            (400, "cd"),
            (500, "d"),
            (900, "cm"),
            (1000, "m"),
            (1984, "mcmlxxxiv"),
            (3999, "mmmcmxcix"),
        ],
    )
    def test_valid_numbers(self, num, expected):
        assert _to_roman(num) == expected

    def test_zero_returns_empty_string(self):
        """_to_roman(0) 静默返回空字符串而非抛出异常"""
        assert _to_roman(0) == ""

    def test_negative_returns_empty_string(self):
        """_to_roman(-5) 静默返回空字符串而非抛出异常"""
        assert _to_roman(-5) == ""


# ============================================================
# utils.py — _to_chinese_num
# ============================================================


class TestToChineseNum:
    @pytest.mark.parametrize(
        "num, expected",
        [
            (1, "一"),
            (5, "五"),
            (9, "九"),
            (10, "十"),
            (11, "十一"),
            (20, "二十"),
            (21, "二十一"),
            (99, "九十九"),
        ],
    )
    def test_valid_numbers(self, num, expected):
        assert _to_chinese_num(num) == expected

    def test_zero(self):
        assert _to_chinese_num(0) == "0"

    def test_negative(self):
        assert _to_chinese_num(-3) == "-3"

    def test_hundred_should_convert(self):
        assert _to_chinese_num(100) == "一百"

    def test_hundred_fallback(self):
        # _to_chinese_num(100) 返回 "一百"
        assert _to_chinese_num(100) == "一百"


# ============================================================
# utils.py — load_yaml_with_merge
# ============================================================


class TestLoadYamlWithMerge:
    def test_load_valid_yaml(self, tmp_path):
        yaml_file = tmp_path / "config.yaml"
        yaml_file.write_text("key: value\nnested:\n  a: 1\n", encoding="utf-8")
        result = load_yaml_with_merge(str(yaml_file))
        assert result["key"] == "value"
        assert result["nested"]["a"] == 1

    def test_nonexistent_file_raises(self):
        with pytest.raises(FileNotFoundError):
            load_yaml_with_merge("/nonexistent/file.yaml")


# ============================================================
# rules/node.py — TreeNode
# ============================================================


class TestTreeNode:
    def test_init_with_simple_value(self):
        node = TreeNode("hello")
        assert node.value == "hello"
        assert node.children == []

    def test_init_with_dict_value_and_fingerprint(self):
        node = TreeNode({"category": "body_text", "fingerprint": "abc123"})
        assert node.fingerprint == "abc123"

    def test_init_with_top_category_skips_fingerprint(self):
        node = TreeNode({"category": "top"})
        assert node.fingerprint is None

    def test_init_missing_fingerprint_raises(self):
        with pytest.raises(ValueError, match="fingerprint"):
            TreeNode({"category": "body_text"})

    def test_config_default_empty(self):
        node = TreeNode("x")
        assert node.config == {}

    def test_load_config_nested_path(self):
        node = TreeNode("x")
        node.NODE_TYPE = "a.b.c"
        full = {"a": {"b": {"c": {"key": "val"}}}}
        node.load_config(full)
        assert node.config == {"key": "val"}

    def test_load_config_missing_path(self):
        node = TreeNode("x")
        node.NODE_TYPE = "x.y.z"
        node.load_config({"a": 1})
        assert node.config == {}

    def test_load_config_non_dict_input(self):
        node = TreeNode("x")
        node.load_config("not a dict")
        assert node.config == {}

    def test_add_child_returns_child(self):
        node = TreeNode("parent")
        child = node.add_child("child_val")
        assert child.value == "child_val"
        assert len(node.children) == 1

    def test_add_child_node(self):
        parent = TreeNode("parent")
        child = TreeNode("child")
        parent.add_child_node(child)
        assert parent.children[0] is child

    def test_repr(self):
        node = TreeNode("test_val")
        assert repr(node) == "TreeNode(test_val)"

    def test_fingerprint_attribute_exists_for_non_dict(self):
        node = TreeNode("simple_string")
        assert hasattr(node, "fingerprint")


# ============================================================
# rules/node.py — FormatNode
# ============================================================


class TestFormatNode:
    def test_init_defaults(self):
        node = FormatNode(value="test", level=1)
        assert node.level == 1
        assert node.paragraph is None
        assert node.expected_rule is None
        assert node._pydantic_config is None

    def test_pydantic_config_raises_before_load(self):
        node = FormatNode(value="test", level=1)
        with pytest.raises(ValueError, match="尚未加载Pydantic配置"):
            _ = node.pydantic_config

    def test_load_yaml_config_file_not_found(self):
        with pytest.raises(FileNotFoundError, match="配置文件"):
            FormatNode.load_yaml_config("/nonexistent/config.yaml")

    def test_load_yaml_config_invalid_yaml(self, tmp_path):
        bad_yaml = tmp_path / "bad.yaml"
        bad_yaml.write_text(":\n  - [invalid", encoding="utf-8")
        with pytest.raises((ValueError, RuntimeError)):
            FormatNode.load_yaml_config(str(bad_yaml))

    def test_update_paragraph(self, doc):
        node = FormatNode(value="test", level=1)
        p = doc.add_paragraph("hello")
        node.update_paragraph(p)
        assert node.paragraph is p

    def test_base_raises_not_implemented(self, doc):
        node = FormatNode(value="test", level=1)
        with pytest.raises(NotImplementedError, match="Subclasses should implement"):
            node._base(doc, p=True, r=True)

    def test_check_format_raises(self, doc):
        node = FormatNode(value="test", level=1)
        with pytest.raises(NotImplementedError):
            node.check_format(doc)

    def test_apply_format_raises(self, doc):
        node = FormatNode(value="test", level=1)
        with pytest.raises(NotImplementedError):
            node.apply_format(doc)

    def test_add_comment_with_text(self, doc):
        node = FormatNode(value="test", level=1)
        p = doc.add_paragraph("hello")
        run = p.runs[0]
        with patch.object(doc, "add_comment") as mock_add:
            node.add_comment(doc, run, "格式错误")
            mock_add.assert_called_once_with(
                runs=run, text="格式错误", author="论文解析器", initials="afish"
            )

    def test_add_comment_empty_text_skipped(self, doc):
        node = FormatNode(value="test", level=1)
        p = doc.add_paragraph("hello")
        run = p.runs[0]
        # 空文本不应调用 add_comment
        node.add_comment(doc, run, "   ")

    def test_load_config_heading_level_bug(self):
        """HeadingLevelConfig 映射到 full_config.headings 而非具体级别"""
        from wordformat.config.datamodel import HeadingLevelConfig, NodeConfigRoot

        class TestFormatNode(FormatNode[HeadingLevelConfig]):
            NODE_TYPE = "headings.level_1"
            CONFIG_MODEL = HeadingLevelConfig

        node = TestFormatNode(value="test", level=1)
        root_config = NodeConfigRoot()
        # load_config 应为 level_1 分配配置，但 BUG 将整个 headings 赋值
        node.load_config(root_config)
        # BUG: pydantic_config 是 HeadingsConfig 而非 HeadingLevelConfig
        assert node.pydantic_config is not None


# ============================================================
# numbering.py — strip_manual_numbering
# ============================================================


class TestStripManualNumbering:
    def test_matching_pattern_strips(self, doc):
        p = doc.add_paragraph()
        run = p.add_run("第一章 绪论")
        result = strip_manual_numbering(p, r"^第[一二三四五六七八九十百千零]+章\s*")
        assert result is True
        assert p.text == "绪论"

    def test_non_matching_pattern_returns_false(self, doc):
        p = doc.add_paragraph()
        p.add_run("绪论")
        result = strip_manual_numbering(p, r"^\d+\.\d+\s+")
        assert result is False
        assert p.text == "绪论"

    def test_empty_pattern_returns_false(self, doc):
        p = doc.add_paragraph()
        p.add_run("1.1 背景")
        assert strip_manual_numbering(p, "") is False

    def test_no_runs_returns_false(self, doc):
        p = doc.add_paragraph()
        assert strip_manual_numbering(p, r"^第.+章\s*") is False

    def test_multi_run_strips_correctly(self, doc):
        p = doc.add_paragraph()
        r1 = p.add_run("第一章")
        r2 = p.add_run(" 绪论")
        result = strip_manual_numbering(p, r"^第[一二三四五六七八九十百千零]+章\s*")
        assert result is True
        assert p.text == "绪论"


# ============================================================
# numbering.py — apply_auto_numbering
# ============================================================


class TestApplyAutoNumbering:
    def test_adds_numPr_to_paragraph(self, doc):
        p = doc.add_paragraph("test")
        apply_auto_numbering(p, num_id="100", ilvl="0")
        from docx.oxml.ns import qn
        pPr = p._element.find(qn("w:pPr"))
        assert pPr is not None
        numPr = pPr.find(qn("w:numPr"))
        assert numPr is not None
        numId_elem = numPr.find(qn("w:numId"))
        assert numId_elem.get(qn("w:val")) == "100"
        ilvl_elem = numPr.find(qn("w:ilvl"))
        assert ilvl_elem.get(qn("w:val")) == "0"

    def test_replaces_existing_numPr(self, doc):
        p = doc.add_paragraph("test")
        apply_auto_numbering(p, num_id="50", ilvl="1")
        apply_auto_numbering(p, num_id="99", ilvl="2")
        from docx.oxml.ns import qn
        pPr = p._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr"))
        assert numPr.find(qn("w:numId")).get(qn("w:val")) == "99"
        assert numPr.find(qn("w:ilvl")).get(qn("w:val")) == "2"


# ============================================================
# numbering.py — create_numbering_definition
# ============================================================


class TestCreateNumberingDefinition:
    def test_disabled_config_returns_empty(self, doc):
        config = MagicMock()
        config.enabled = False
        result = create_numbering_definition(doc, config)
        assert result == {}

    def test_enabled_config_creates_definitions(self, doc):
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        result = create_numbering_definition(doc, config)
        assert "level_1" in result
        assert "level_2" not in result
        assert "level_3" not in result

    def test_multiple_levels(self, doc):
        config = MagicMock()
        config.enabled = True
        for key in ("level_1", "level_2", "level_3"):
            lvl = MagicMock()
            lvl.enabled = True
            lvl.template = f"template_{key}"
            lvl.suffix = "space"
            lvl.numbering_indent = None
            lvl.text_indent = None
            setattr(config, key, lvl)

        result = create_numbering_definition(doc, config)
        assert len(result) == 3
        # num_id 应该递增
        ids = [int(v) for v in result.values()]
        assert ids == sorted(ids)


# ============================================================
# numbering.py — process_heading_numbering
# ============================================================


class TestProcessHeadingNumbering:
    def test_disabled_config_noop(self, doc):
        root = TreeNode({"category": "top"})
        config = MagicMock()
        config.enabled = False
        process_heading_numbering(root, doc, config)
        # 无异常即通过


# ============================================================
# settings.py
# ============================================================


class TestSettings:
    def test_batch_size_is_string(self):
        assert isinstance(settings.BATCH_SIZE, str)

    def test_voidnodelist_contains_top(self):
        assert "top" in settings.VOIDNODELIST

    def test_voidnodelist_is_list(self):
        assert isinstance(settings.VOIDNODELIST, list)

    def test_host_is_string(self):
        assert isinstance(settings.HOST, str)

    def test_port_is_int(self):
        assert isinstance(settings.PORT, int)

    def test_server_host_format(self):
        assert settings.SERVER_HOST.startswith("http://")
        assert str(settings.PORT) in settings.SERVER_HOST


# ============================================================
# utils.py — get_paragraph_numbering_text
# ============================================================


class TestGetParagraphNumberingText:
    """测试从段落 XML 中提取自动编号文字"""

    def test_no_pPr_returns_empty(self, doc):
        """段落没有 pPr 时返回空字符串"""
        p = doc.add_paragraph("hello")
        # 确保没有 pPr
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            p._element.remove(pPr)
        assert get_paragraph_numbering_text(p) == ""

    def test_no_numPr_returns_empty(self, doc):
        """段落有 pPr 但没有 numPr 时返回空字符串"""
        p = doc.add_paragraph("hello")
        from docx.oxml import OxmlElement
        pPr = OxmlElement("w:pPr")
        p._element.insert(0, pPr)
        assert get_paragraph_numbering_text(p) == ""

    def test_no_numId_returns_empty(self, doc):
        """numPr 中没有 numId 时返回空字符串"""
        p = doc.add_paragraph("hello")
        from docx.oxml import OxmlElement
        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
        p._element.insert(0, pPr)
        assert get_paragraph_numbering_text(p) == ""

    def test_numId_zero_returns_empty(self, doc):
        """numId 为 0 时返回空字符串"""
        p = doc.add_paragraph("hello")
        from docx.oxml import OxmlElement
        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), "0")
        numPr.append(numId_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)
        assert get_paragraph_numbering_text(p) == ""

    def test_no_numbering_part_returns_empty(self, doc):
        """文档没有 numbering part 时返回空字符串（NotImplementedError 未被源码捕获，属于已知 bug）"""
        p = doc.add_paragraph("hello")
        from docx.oxml import OxmlElement
        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), "1")
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), "0")
        numPr.append(ilvl_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)
        # 移除 numbering 关系使 numbering_part 访问抛出异常
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        rels = doc.part.rels
        to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
        for k in to_remove:
            del rels[k]
        # 源码只捕获 AttributeError/KeyError，但 python-docx 可能抛 NotImplementedError
        with pytest.raises(NotImplementedError):
            get_paragraph_numbering_text(p)

    def _setup_numbering(self, doc, num_fmt="decimal", lvl_text="%1.", num_id="1", abstract_num_id="0"):
        """辅助方法：为文档创建 numbering part 和编号定义"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        # 先移除已有的 numbering 关系（Document() 默认模板有 numbering part）
        try:
            rels = doc.part.rels
            to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
            for k in to_remove:
                del rels[k]
        except Exception:
            pass

        numbering_elm = OxmlElement("w:numbering")

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), abstract_num_id)

        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), num_fmt)
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), lvl_text)
        lvl.append(lvlText)
        abstract_num.append(lvl)
        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), abstract_num_id)
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        return numbering_elm

    def _add_numPr_to_paragraph(self, p, num_id="1", ilvl="0"):
        """辅助方法：为段落添加 numPr"""
        from docx.oxml import OxmlElement
        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), num_id)
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), ilvl)
        numPr.append(ilvl_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)

    def test_with_numbering_definition(self, doc):
        """测试有完整 numbering 定义时能提取编号文字"""
        # BUG: qn() 无法处理 XPath 表达式（如 w:num[@w:numId='1']），
        # 导致 get_paragraph_numbering_text 在查找 num 元素时抛出 ValueError
        self._setup_numbering(doc, num_fmt="decimal", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "1."

    def test_chinese_counting_format(self, doc):
        """测试中文计数格式（chineseCounting）"""
        self._setup_numbering(doc, num_fmt="chineseCounting", lvl_text="第%1章")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "第一章"

    def test_upper_roman_format(self, doc):
        """测试大写罗马数字格式"""
        self._setup_numbering(doc, num_fmt="upperRoman", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "I."

    def test_lower_roman_format(self, doc):
        """测试小写罗马数字格式"""
        self._setup_numbering(doc, num_fmt="lowerRoman", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "i."

    def test_upper_letter_format(self, doc):
        """测试大写字母格式"""
        self._setup_numbering(doc, num_fmt="upperLetter", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "A."

    def test_lower_letter_format(self, doc):
        """测试小写字母格式"""
        self._setup_numbering(doc, num_fmt="lowerLetter", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "a."

    def test_missing_abstract_num_returns_empty(self, doc):
        """abstractNumId 对应的 abstractNum 不存在时返回空字符串"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        # 先移除已有的 numbering 关系
        rels = doc.part.rels
        to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
        for k in to_remove:
            del rels[k]

        numbering_elm = OxmlElement("w:numbering")

        # 只创建 num 引用，不创建 abstractNum
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), "999")
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        p = doc.add_paragraph("test")
        self._add_numPr_to_paragraph(p)

        assert get_paragraph_numbering_text(p) == ""

    def test_missing_lvl_returns_empty(self, doc):
        """abstractNum 中没有对应 ilvl 的 lvl 时返回空字符串"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        # 先移除已有的 numbering 关系
        rels = doc.part.rels
        to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
        for k in to_remove:
            del rels[k]

        numbering_elm = OxmlElement("w:numbering")

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), "0")
        # 不创建任何 lvl
        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), "0")
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        p = doc.add_paragraph("test")
        self._add_numPr_to_paragraph(p)

        assert get_paragraph_numbering_text(p) == ""


# ============================================================
# utils.py — _format_number
# ============================================================


class TestFormatNumber:
    """测试 _format_number 各种格式化类型"""

    def test_decimal(self):
        assert _format_number(5, "decimal") == "5"

    def test_upper_roman(self):
        assert _format_number(3, "upperRoman") == "III"

    def test_lower_roman(self):
        assert _format_number(4, "lowerRoman") == "iv"

    def test_upper_letter(self):
        assert _format_number(1, "upperLetter") == "A"
        assert _format_number(3, "upperLetter") == "C"

    def test_lower_letter(self):
        assert _format_number(1, "lowerLetter") == "a"
        assert _format_number(3, "lowerLetter") == "c"

    def test_chinese_counting_thousand(self):
        assert _format_number(5, "chineseCountingThousand") == "五"

    def test_ideograph_traditional(self):
        assert _format_number(3, "ideographTraditional") == "三"

    def test_chinese_counting(self):
        assert _format_number(7, "chineseCounting") == "七"

    def test_unknown_format_fallback_to_str(self):
        """未知格式回退到 str"""
        assert _format_number(42, "unknownFormat") == "42"


# ============================================================
# utils.py — _get_level_fmt
# ============================================================


class TestGetLevelFmt:
    """测试 _get_level_fmt 获取指定级别的 numFmt"""

    def test_existing_level_returns_fmt(self):
        from lxml import etree
        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:numFmt w:val="upperRoman"/>'
            '  </w:lvl>'
            '</w:abstractNum>'
        )
        assert _get_level_fmt(abstract_num, 0) == "upperRoman"

    def test_missing_level_returns_decimal(self):
        from lxml import etree
        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '</w:abstractNum>'
        )
        assert _get_level_fmt(abstract_num, 0) == "decimal"

    def test_no_numFmt_returns_decimal(self):
        from lxml import etree
        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:start w:val="1"/>'
            '  </w:lvl>'
            '</w:abstractNum>'
        )
        assert _get_level_fmt(abstract_num, 0) == "decimal"


# ============================================================
# utils.py — remove_all_numbering
# ============================================================


class TestRemoveAllNumbering:
    """测试 remove_all_numbering 移除标题样式的编号绑定"""

    def test_removes_numPr_from_heading_styles(self, doc):
        """从 Heading 1/2/3 样式中移除 numPr"""
        from docx.oxml import OxmlElement

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = doc.styles[style_name]
            style_element = style._element

            # 确保 pPr 存在
            pPr = style_element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                style_element.insert(0, pPr)

            # 添加 numPr
            numPr = OxmlElement("w:numPr")
            pPr.append(numPr)

        remove_all_numbering(doc)

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = doc.styles[style_name]
            pPr = style._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                assert numPr is None, f"{style_name} 的 numPr 未被移除"

    def test_removes_outlineLvl_from_heading_styles(self, doc):
        """从 Heading 1/2/3 样式中移除 outlineLvl"""
        from docx.oxml import OxmlElement

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = doc.styles[style_name]
            style_element = style._element

            pPr = style_element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                style_element.insert(0, pPr)

            # 先移除已有的 outlineLvl
            existing = pPr.find(qn("w:outlineLvl"))
            if existing is not None:
                pPr.remove(existing)

            outlineLvl = OxmlElement("w:outlineLvl")
            outlineLvl.set(qn("w:val"), "0")
            pPr.append(outlineLvl)

        remove_all_numbering(doc)

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = doc.styles[style_name]
            pPr = style._element.find(qn("w:pPr"))
            if pPr is not None:
                outlineLvl = pPr.find(qn("w:outlineLvl"))
                assert outlineLvl is None, f"{style_name} 的 outlineLvl 未被移除"

    def test_no_pPr_in_style_no_error(self, doc):
        """样式中没有 pPr 时不报错"""
        # Heading 1 的 pPr 可能存在也可能不存在，直接调用不应报错
        remove_all_numbering(doc)

    def test_style_not_in_doc_no_error(self, doc):
        """如果样式不存在（不太可能），不报错"""
        # Document() 默认包含 Heading 1/2/3，所以这个测试主要验证不会抛异常
        remove_all_numbering(doc)


# ============================================================
# style/get_some.py — _get_style_spacing
# ============================================================


class TestGetStyleSpacing:
    """测试 _get_style_spacing 递归查找样式间距"""

    def test_none_style_returns_none(self):
        assert _get_style_spacing(None) is None

    def test_style_with_direct_spacing(self):
        """样式自身有 beforeLines 时直接返回"""
        mock_style = MagicMock()
        mock_elem = MagicMock()
        mock_pPr = MagicMock()
        mock_spacing = MagicMock()

        mock_style.element = mock_elem
        mock_elem.find.return_value = mock_pPr
        mock_pPr.find.return_value = mock_spacing
        mock_spacing.get.return_value = "200"  # 2.0 行

        result = _get_style_spacing(mock_style, "before")
        assert result == 2.0

    def test_style_with_zero_spacing_falls_to_base(self):
        """样式自身 spacing 为 0 时递归查基样式"""
        mock_style = MagicMock()
        mock_base = MagicMock()
        mock_elem = MagicMock()
        mock_pPr = MagicMock()
        mock_spacing = MagicMock()
        mock_base_elem = MagicMock()
        mock_base_pPr = MagicMock()
        mock_base_spacing = MagicMock()

        mock_style.element = mock_elem
        mock_style.base_style = mock_base
        mock_elem.find.return_value = mock_pPr
        mock_pPr.find.return_value = mock_spacing
        mock_spacing.get.return_value = "0"  # 0 行

        mock_base.element = mock_base_elem
        mock_base_elem.find.return_value = mock_base_pPr
        mock_base_pPr.find.return_value = mock_base_spacing
        mock_base_spacing.get.return_value = "150"  # 1.5 行

        result = _get_style_spacing(mock_style, "before")
        assert result == 1.5

    def test_no_pPr_falls_to_base(self):
        """样式没有 pPr 时递归查基样式"""
        mock_style = MagicMock()
        mock_base = MagicMock()
        mock_elem = MagicMock()
        mock_base_elem = MagicMock()
        mock_base_pPr = MagicMock()
        mock_base_spacing = MagicMock()

        mock_style.element = mock_elem
        mock_style.base_style = mock_base
        mock_elem.find.return_value = None  # pPr 为 None

        mock_base.element = mock_base_elem
        mock_base_elem.find.return_value = mock_base_pPr
        mock_base_pPr.find.return_value = mock_base_spacing
        mock_base_spacing.get.return_value = "100"  # 1.0 行

        result = _get_style_spacing(mock_style, "before")
        assert result == 1.0

    def test_no_spacing_falls_to_base(self):
        """样式有 pPr 但没有 spacing 时递归查基样式"""
        mock_style = MagicMock()
        mock_base = MagicMock()
        mock_elem = MagicMock()
        mock_pPr = MagicMock()
        mock_base_elem = MagicMock()
        mock_base_pPr = MagicMock()
        mock_base_spacing = MagicMock()

        mock_style.element = mock_elem
        mock_style.base_style = mock_base
        mock_elem.find.return_value = mock_pPr
        mock_pPr.find.return_value = None  # spacing 为 None

        mock_base.element = mock_base_elem
        mock_base_elem.find.return_value = mock_base_pPr
        mock_base_pPr.find.return_value = mock_base_spacing
        mock_base_spacing.get.return_value = "300"  # 3.0 行

        result = _get_style_spacing(mock_style, "before")
        assert result == 3.0

    def test_after_spacing_type(self):
        """测试 after 类型的间距查找"""
        mock_style = MagicMock()
        mock_elem = MagicMock()
        mock_pPr = MagicMock()
        mock_spacing = MagicMock()

        mock_style.element = mock_elem
        mock_elem.find.return_value = mock_pPr
        mock_pPr.find.return_value = mock_spacing
        mock_spacing.get.return_value = "500"  # 5.0 行

        result = _get_style_spacing(mock_style, "after")
        assert result == 5.0

    def test_no_base_style_returns_none(self):
        """没有基样式且自身无 spacing 时返回 None"""
        mock_style = MagicMock()
        mock_elem = MagicMock()
        mock_pPr = MagicMock()

        mock_style.element = mock_elem
        mock_style.base_style = None
        mock_elem.find.return_value = mock_pPr
        mock_pPr.find.return_value = None

        # base_style 为 None 时 AttributeError 会被捕获
        result = _get_style_spacing(mock_style, "before")
        assert result is None

    def test_style_element_none_falls_to_base(self):
        """style.element 为 None 时递归查基样式"""
        mock_style = MagicMock()
        mock_base = MagicMock()
        mock_base_elem = MagicMock()
        mock_base_pPr = MagicMock()
        mock_base_spacing = MagicMock()

        mock_style.element = None
        mock_style.base_style = mock_base

        mock_base.element = mock_base_elem
        mock_base_elem.find.return_value = mock_base_pPr
        mock_base_pPr.find.return_value = mock_base_spacing
        mock_base_spacing.get.return_value = "100"

        result = _get_style_spacing(mock_style, "before")
        assert result == 1.0


# ============================================================
# numbering.py — strip_manual_numbering (empty result)
# ============================================================


class TestStripManualNumberingEmptyResult:
    """测试 strip_manual_numbering 在文本变空时的行为"""

    def test_text_becomes_empty_after_strip(self, doc):
        """匹配整个文本后，段落文本变为空"""
        p = doc.add_paragraph()
        run = p.add_run("1.1 ")
        result = strip_manual_numbering(p, r"^1\.1\s*")
        assert result is True
        # 整个文本被清除后，段落文本为空
        assert p.text == ""

    def test_single_run_fully_consumed(self, doc):
        """单个 run 的文本完全被匹配清除"""
        p = doc.add_paragraph()
        run = p.add_run("第一章")
        result = strip_manual_numbering(p, r"^第[一二三四五六七八九十百千零]+章")
        assert result is True
        assert p.text == ""


# ============================================================
# numbering.py — create_numbering_definition (XML creation)
# ============================================================


class TestCreateNumberingDefinitionXML:
    """测试 create_numbering_definition 的 XML 创建逻辑"""

    def test_creates_abstractNum_and_num_elements(self, doc):
        """验证创建的 XML 中包含 abstractNum 和 num 元素"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "%1."
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        result = create_numbering_definition(doc, config)

        assert "level_1" in result

        # 验证 numbering part 中有 abstractNum 和 num
        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element

        abstract_nums = numbering_elm.findall(qn("w:abstractNum"))
        nums = numbering_elm.findall(qn("w:num"))
        assert len(abstract_nums) >= 1
        assert len(nums) >= 1

    def test_chinese_template_creates_chinese_numFmt(self, doc):
        """中文模板（含'第'和'章'）使用 chineseCountingThousand 格式"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        create_numbering_definition(doc, config)

        numbering_elm = doc.part.numbering_part._element
        # 找到最后一个 abstractNum（create_numbering_definition 追加的）
        abstract_nums = numbering_elm.findall(qn("w:abstractNum"))
        abstract_num = abstract_nums[-1]
        lvl = abstract_num.find(qn("w:lvl"))
        numFmt = lvl.find(qn("w:numFmt"))
        assert numFmt.get(qn("w:val")) == "chineseCountingThousand"

    def test_decimal_template_creates_decimal_numFmt(self, doc):
        """非中文模板使用 decimal 格式"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "%1."
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        create_numbering_definition(doc, config)

        numbering_elm = doc.part.numbering_part._element
        abstract_nums = numbering_elm.findall(qn("w:abstractNum"))
        abstract_num = abstract_nums[-1]
        lvl = abstract_num.find(qn("w:lvl"))
        numFmt = lvl.find(qn("w:numFmt"))
        assert numFmt.get(qn("w:val")) == "decimal"

    def test_no_template_skips_level(self, doc):
        """template 为空时跳过该级别"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = ""
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        result = create_numbering_definition(doc, config)
        assert "level_1" not in result

    def test_creates_numbering_part_when_missing(self, tmp_path):
        """文档没有 numbering part 时自动创建"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        # 创建一个没有 numbering 关系的文档
        doc = Document()
        # 移除已有的 numbering 关系（如果有）
        try:
            rels = doc.part.rels
            to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
            for k in to_remove:
                del rels[k]
        except Exception:
            pass

        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "%1."
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        result = create_numbering_definition(doc, config)
        assert "level_1" in result

        # 现在应该有了
        numbering_part = doc.part.numbering_part
        assert numbering_part is not None


# ============================================================
# numbering.py — process_heading_numbering (enabled)
# ============================================================


class TestProcessHeadingNumberingEnabled:
    """测试 process_heading_numbering 在启用时的完整流程"""

    def test_enabled_config_processes_headings(self, doc):
        """启用配置时处理标题节点：清除手动编号 + 应用自动编号"""
        config = MagicMock()
        config.enabled = True

        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.strip_pattern = r"^第[一二三四五六七八九十百千零]+章\s*"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None

        level_2 = MagicMock()
        level_2.enabled = True
        level_2.template = "%1.%2"
        level_2.strip_pattern = r"^\d+\.\d+\s+"
        level_2.suffix = "space"
        level_2.numbering_indent = None
        level_2.text_indent = None

        level_3 = MagicMock()
        level_3.enabled = False
        level_3.template = ""
        level_3.strip_pattern = ""

        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        # 创建标题节点（heading 类别需要 fingerprint）
        root = TreeNode({"category": "top"})

        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        p1 = doc.add_paragraph("第一章 绪论")
        h1_node.paragraph = p1
        root.add_child_node(h1_node)

        h2_node = TreeNode({"category": "heading_level_2", "fingerprint": "fp2"})
        p2 = doc.add_paragraph("1.1 研究背景")
        h2_node.paragraph = p2
        h1_node.add_child_node(h2_node)

        body_node = TreeNode({"category": "body_text", "fingerprint": "fp3"})
        p3 = doc.add_paragraph("这是正文内容")
        body_node.paragraph = p3
        h2_node.add_child_node(body_node)

        process_heading_numbering(root, doc, config)

        # 验证手动编号被清除
        assert p1.text == "绪论"
        assert p2.text == "研究背景"

        # 验证自动编号被应用
        pPr1 = p1._element.find(qn("w:pPr"))
        numPr1 = pPr1.find(qn("w:numPr"))
        assert numPr1 is not None

        pPr2 = p2._element.find(qn("w:pPr"))
        numPr2 = pPr2.find(qn("w:numPr"))
        assert numPr2 is not None

    def test_enabled_but_no_paragraph_skips(self, doc):
        """启用配置但节点没有 paragraph 属性时跳过"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.strip_pattern = r"^第.+章\s*"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        root = TreeNode({"category": "top"})
        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        # 不设置 paragraph 属性
        root.add_child_node(h1_node)

        # 不应抛异常
        process_heading_numbering(root, doc, config)

    def test_enabled_level_disabled_skips_strip(self, doc):
        """级别未启用时跳过该级别的处理"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = False
        level_1.template = ""
        level_1.strip_pattern = ""
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        root = TreeNode({"category": "top"})
        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        p1 = doc.add_paragraph("第一章 绪论")
        h1_node.paragraph = p1
        root.add_child_node(h1_node)

        process_heading_numbering(root, doc, config)

        # 文本不应被修改
        assert p1.text == "第一章 绪论"

    def test_no_strip_pattern_skips_strip(self, doc):
        """没有 strip_pattern 时跳过手动编号清除"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.strip_pattern = ""  # 空 strip_pattern
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        root = TreeNode({"category": "top"})
        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        p1 = doc.add_paragraph("第一章 绪论")
        h1_node.paragraph = p1
        root.add_child_node(h1_node)

        process_heading_numbering(root, doc, config)

        # 文本不应被修改（因为没有 strip_pattern）
        assert p1.text == "第一章 绪论"
        # 但自动编号应该被应用
        pPr = p1._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr"))
        assert numPr is not None

    def test_empty_num_id_map_returns_early(self, doc):
        """num_id_map 为空时提前返回"""
        config = MagicMock()
        config.enabled = True
        # 所有级别都禁用或没有 template → create_numbering_definition 返回空
        level_1 = MagicMock()
        level_1.enabled = False
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        root = TreeNode({"category": "top"})
        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        p1 = doc.add_paragraph("第一章 绪论")
        h1_node.paragraph = p1
        root.add_child_node(h1_node)

        # 不应抛异常，且不应修改段落
        process_heading_numbering(root, doc, config)
        assert p1.text == "第一章 绪论"


# ============================================================
# base.py — DocxBase
# ============================================================


class TestDocxBase:
    """测试 DocxBase 类的初始化和 parse 方法"""

    def test_init(self, temp_docx):
        """测试 DocxBase 初始化"""
        base = DocxBase(temp_docx, "/fake/config.yaml")
        assert base.docx_file == temp_docx
        assert base.document is not None
        assert base.re_dict == {}

    def _create_multi_para_docx(self, tmp_path, texts):
        """辅助方法：创建包含多个段落的 docx 文件"""
        doc = Document()
        for text in texts:
            doc.add_paragraph(text)
        path = str(tmp_path / "multi.docx")
        doc.save(path)
        return path

    def test_parse_with_mocked_batch_infer(self, tmp_path):
        """测试 parse 方法使用 mock 的批量推理"""
        path = self._create_multi_para_docx(tmp_path, ["绪论", "研究背景", "正文内容"])
        mock_batch_results = [
            {"label": "heading_level_1", "score": 0.95},
            {"label": "body_text", "score": 0.88},
            {"label": "body_text", "score": 0.75},
        ]

        with patch("wordformat.base.onnx_batch_infer", return_value=mock_batch_results):
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert len(result) == 3
        assert result[0]["category"] == "heading_level_1"
        assert result[1]["category"] == "body_text"
        assert "fingerprint" in result[0]
        assert "paragraph" in result[0]
        assert "score" in result[0]

    def test_parse_low_score_forced_to_body_text(self, tmp_path):
        """测试低置信度结果被强制设为 body_text"""
        path = self._create_multi_para_docx(tmp_path, ["绪论", "研究背景"])
        mock_batch_results = [
            {"label": "heading_level_1", "score": 0.95},
            {"label": "heading_level_2", "score": 0.3},
        ]

        with patch("wordformat.base.onnx_batch_infer", return_value=mock_batch_results):
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert result[0]["category"] == "heading_level_1"
        assert result[1]["category"] == "body_text"
        assert "强制设为" in result[1]["comment"]

    def test_parse_heading_fulu_early_return(self, tmp_path):
        """测试遇到 heading_fulu 时提前终止"""
        path = self._create_multi_para_docx(tmp_path, ["正文", "附录", "更多内容"])
        mock_batch_results = [
            {"label": "body_text", "score": 0.9},
            {"label": "heading_fulu", "score": 0.85},
            {"label": "body_text", "score": 0.9},  # 不应被处理
        ]

        with patch("wordformat.base.onnx_batch_infer", return_value=mock_batch_results):
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert len(result) == 2
        assert result[0]["category"] == "body_text"
        assert result[1]["category"] == "heading_fulu"

    def test_parse_batch_failure_fallback_to_single(self, temp_docx):
        """测试批量推理失败时降级到单条推理"""
        mock_single_result = {"label": "body_text", "score": 0.9}

        with patch("wordformat.base.onnx_batch_infer", side_effect=RuntimeError("ONNX error")):
            with patch("wordformat.base.onnx_single_infer", return_value=mock_single_result):
                base = DocxBase(temp_docx, "/fake/config.yaml")
                result = base.parse()

        # temp_docx 只有一个段落
        assert len(result) >= 1
        assert all(item["category"] == "body_text" for item in result)

    def test_parse_empty_document(self, tmp_path):
        """测试空文档的解析"""
        doc = Document()
        path = str(tmp_path / "empty.docx")
        doc.save(path)

        with patch("wordformat.base.onnx_batch_infer", return_value=[]) as mock_infer:
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert result == []
        mock_infer.assert_not_called()

    def test_parse_batches_correctly(self, tmp_path):
        """测试按 BATCH_SIZE 分批推理"""
        path = self._create_multi_para_docx(tmp_path, [f"段落 {i}" for i in range(5)])

        call_count = 0

        def mock_batch(texts):
            nonlocal call_count
            call_count += 1
            return [{"label": "body_text", "score": 0.9}] * len(texts)

        with patch("wordformat.base.onnx_batch_infer", side_effect=mock_batch):
            with patch("wordformat.base.BATCH_SIZE", "2"):
                base = DocxBase(path, "/fake/config.yaml")
                result = base.parse()

        assert len(result) == 5
        # BATCH_SIZE=2, 5 个段落 → 3 次调用 (2+2+1)
        assert call_count == 3

    def test_parse_includes_numbering_text(self, tmp_path):
        """测试解析时包含自动编号文字"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        doc = Document()

        # 先移除已有的 numbering 关系
        rels = doc.part.rels
        to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
        for k in to_remove:
            del rels[k]

        p = doc.add_paragraph("绪论")

        # 添加 numbering
        numbering_elm = OxmlElement("w:numbering")
        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), "0")
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "decimal")
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), "%1.")
        lvl.append(lvlText)
        abstract_num.append(lvl)
        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), "0")
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), "1")
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), "0")
        numPr.append(ilvl_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)

        path = str(tmp_path / "numbered.docx")
        doc.save(path)

        with patch("wordformat.base.onnx_batch_infer", return_value=[{"text": "1. 绪论", "label": "body_text", "pred_id": 0, "score": 0.9}]):
            base = DocxBase(path, "/fake/config.yaml")
            result = base.parse()

        assert len(result) == 1
        # 段落文本应包含编号 "1. 绪论"
        assert result[0]["paragraph"] == "1. 绪论"


# ============================================================
# utils.py — _format_number 额外覆盖测试
# ============================================================


class TestFormatNumberAdditional:
    """补充 _format_number 的边界情况和完整格式覆盖"""

    def test_decimal_zero(self):
        """decimal 格式处理 0"""
        assert _format_number(0, "decimal") == "0"

    def test_decimal_negative(self):
        """decimal 格式处理负数"""
        assert _format_number(-5, "decimal") == "-5"

    def test_upper_roman_zero(self):
        """upperRoman 格式处理 0（_to_roman 返回空字符串后 upper()）"""
        assert _format_number(0, "upperRoman") == ""

    def test_upper_roman_negative(self):
        """upperRoman 格式处理负数"""
        assert _format_number(-3, "upperRoman") == ""

    def test_lower_roman_zero(self):
        """lowerRoman 格式处理 0"""
        assert _format_number(0, "lowerRoman") == ""

    def test_upper_letter_boundary_26(self):
        """upperLetter 格式 n=26 返回 Z（1 <= n <= 26 范围内）"""
        assert _format_number(26, "upperLetter") == "Z"

    def test_upper_letter_boundary_1(self):
        """upperLetter 格式 n=1 返回 A"""
        assert _format_number(1, "upperLetter") == "A"

    def test_upper_letter_boundary_27(self):
        """upperLetter 格式 n=27 超出 A-Z 范围回退到 str"""
        assert _format_number(27, "upperLetter") == "27"

    def test_lower_letter_boundary_26(self):
        """lowerLetter 格式 n=26 返回 z（1 <= n <= 26 范围内）"""
        assert _format_number(26, "lowerLetter") == "z"

    def test_lower_letter_boundary_1(self):
        """lowerLetter 格式 n=1 返回 a"""
        assert _format_number(1, "lowerLetter") == "a"

    def test_chinese_counting_zero(self):
        """chineseCounting 格式处理 0"""
        assert _format_number(0, "chineseCounting") == "0"

    def test_chinese_counting_negative(self):
        """chineseCounting 格式处理负数"""
        assert _format_number(-3, "chineseCounting") == "-3"

    def test_chinese_counting_hundred(self):
        """chineseCounting 格式处理 100（_to_chinese_num 返回 "一百"）"""
        assert _format_number(100, "chineseCounting") == "一百"

    def test_chinese_counting_thousand_zero(self):
        """chineseCountingThousand 格式处理 0"""
        assert _format_number(0, "chineseCountingThousand") == "0"

    def test_ideograph_traditional_zero(self):
        """ideographTraditional 格式处理 0"""
        assert _format_number(0, "ideographTraditional") == "0"

    def test_unknown_format_with_zero(self):
        """未知格式回退到 str，处理 0"""
        assert _format_number(0, "totallyUnknown") == "0"

    def test_unknown_format_with_large_number(self):
        """未知格式回退到 str，处理大数"""
        assert _format_number(99999, "unknownFmt") == "99999"


# ============================================================
# utils.py — _get_level_fmt 额外覆盖测试
# ============================================================


class TestGetLevelFmtAdditional:
    """补充 _get_level_fmt 的边界情况测试"""

    def test_level_missing_numFmt_returns_decimal(self):
        """lvl 存在但缺少 numFmt 元素时返回 decimal"""
        from lxml import etree
        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:start w:val="1"/>'
            '  </w:lvl>'
            '</w:abstractNum>'
        )
        assert _get_level_fmt(abstract_num, 0) == "decimal"

    def test_level_missing_numFmt_val_returns_decimal(self):
        """numFmt 元素存在但缺少 w:val 属性时返回 decimal"""
        from lxml import etree
        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:numFmt/>'
            '  </w:lvl>'
            '</w:abstractNum>'
        )
        assert _get_level_fmt(abstract_num, 0) == "decimal"

    def test_nonexistent_level_returns_decimal(self):
        """请求不存在的级别返回 decimal"""
        from lxml import etree
        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:numFmt w:val="decimal"/>'
            '  </w:lvl>'
            '</w:abstractNum>'
        )
        assert _get_level_fmt(abstract_num, 5) == "decimal"


# ============================================================
# utils.py — _count_numbering_levels 覆盖测试
# ============================================================


class TestCountNumberingLevels:
    """测试 _count_numbering_levels 编号级别计数逻辑"""

    def _setup_numbering_doc(self, doc, num_fmt="decimal", lvl_text="%1.", num_id="1", abstract_num_id="0"):
        """辅助方法：为文档创建 numbering part"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        # 移除已有的 numbering 关系
        try:
            rels = doc.part.rels
            to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
            for k in to_remove:
                del rels[k]
        except Exception:
            pass

        numbering_elm = OxmlElement("w:numbering")

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), abstract_num_id)

        for lvl_idx in range(3):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(lvl_idx))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            numFmt = OxmlElement("w:numFmt")
            numFmt.set(qn("w:val"), num_fmt)
            lvl.append(numFmt)
            lvlText = OxmlElement("w:lvlText")
            lvlText.set(qn("w:val"), lvl_text)
            lvl.append(lvlText)
            abstract_num.append(lvl)

        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), abstract_num_id)
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        return numbering_elm

    def _add_numPr_to_paragraph(self, p, num_id="1", ilvl="0"):
        """辅助方法：为段落添加 numPr"""
        from docx.oxml import OxmlElement
        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), num_id)
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), ilvl)
        numPr.append(ilvl_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)

    def test_single_numbered_paragraph(self, doc):
        """单个编号段落返回 {0: 1}"""
        self._setup_numbering_doc(doc)
        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p, num_id="1", ilvl="0")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p)
        assert result == {0: 1}

    def test_multiple_level1_paragraphs(self, doc):
        """多个一级编号段落正确计数"""
        self._setup_numbering_doc(doc)
        p1 = doc.add_paragraph("第一章")
        self._add_numPr_to_paragraph(p1, num_id="1", ilvl="0")
        p2 = doc.add_paragraph("正文")
        doc.add_paragraph("无编号段落")
        p3 = doc.add_paragraph("第二章")
        self._add_numPr_to_paragraph(p3, num_id="1", ilvl="0")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p3)
        assert result == {0: 2}

    def test_mixed_levels_reset_counters(self, doc):
        """混合级别时下级计数器在上级重启"""
        self._setup_numbering_doc(doc)
        p1 = doc.add_paragraph("第一章")
        self._add_numPr_to_paragraph(p1, num_id="1", ilvl="0")
        p2 = doc.add_paragraph("1.1 背景")
        self._add_numPr_to_paragraph(p2, num_id="1", ilvl="1")
        p3 = doc.add_paragraph("1.2 方法")
        self._add_numPr_to_paragraph(p3, num_id="1", ilvl="1")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p3)
        assert result == {0: 1, 1: 2}

    def test_level_reset_after_new_parent(self, doc):
        """新的上级段落重置下级计数"""
        self._setup_numbering_doc(doc)
        p1 = doc.add_paragraph("第一章")
        self._add_numPr_to_paragraph(p1, num_id="1", ilvl="0")
        p2 = doc.add_paragraph("1.1 背景")
        self._add_numPr_to_paragraph(p2, num_id="1", ilvl="1")
        p3 = doc.add_paragraph("第二章")
        self._add_numPr_to_paragraph(p3, num_id="1", ilvl="0")
        p4 = doc.add_paragraph("2.1 方法")
        self._add_numPr_to_paragraph(p4, num_id="1", ilvl="1")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p4)
        assert result == {0: 2, 1: 1}

    def test_no_matching_num_ids_returns_empty(self, doc):
        """没有匹配的 numId 时返回空字典"""
        self._setup_numbering_doc(doc, num_id="1", abstract_num_id="0")
        p = doc.add_paragraph("test")
        self._add_numPr_to_paragraph(p, num_id="1", ilvl="0")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        # 使用不存在的 abstract_num_id
        result = _count_numbering_levels(numbering_elm, "999", p)
        assert result == {}

    def test_no_numbered_paragraphs_before_target(self, doc):
        """目标段落之前没有编号段落时返回 {ilvl: 1}"""
        self._setup_numbering_doc(doc)
        doc.add_paragraph("无编号段落")
        p = doc.add_paragraph("编号段落")
        self._add_numPr_to_paragraph(p, num_id="1", ilvl="0")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p)
        assert result == {0: 1}
