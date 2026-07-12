"""
Core 模块综合测试

覆盖 tree.py, utils.py, rules/node.py, numbering.py, settings.py
"""

import os
import pytest
from io import StringIO
from unittest.mock import MagicMock, patch

from docx import Document
from docx.oxml.ns import qn

from wordformat.tree import Tree, Stack, print_tree
from wordformat.config.dotdict import DotDict
from wordformat.rules.node import TreeNode, FormatNode
from wordformat.numbering import (
    _auto_strip_numbering,
    _strip_reference_numbering,
    apply_auto_numbering,
    create_numbering_definition,
    process_heading_numbering,
)
from wordformat.utils import (
    get_file_name,
    ensure_is_directory,
    ensure_directory_exists,
    _to_roman,
    _to_chinese_num,
    load_yaml_with_merge,
    get_paragraph_numbering_text,
    remove_all_numbering,
    _format_number,
    _get_level_fmt,
    _count_numbering_levels,
)
from wordformat.base import DocxBase
from wordformat import settings


# ============================================================
# tree.py — Tree
# ============================================================


# ============================================================
# tree.py — Tree
# ============================================================


class TestTreeCreation:
    """Tree 创建与基本属性"""

    def test_create_tree_with_root_value(self):
        tree = Tree("root")
        assert tree.root.value == "root"

    def test_create_tree_with_dict_value(self):
        tree = Tree({"category": "top"})
        assert tree.root.value == {"category": "top"}

    def test_tree_repr(self):
        tree = Tree("my_root")
        assert repr(tree) == "Tree(root=my_root)"

    def test_is_empty(self):
        """空树（仅 root 无子节点）返回 True，有子节点返回 False"""
        tree = Tree("root")
        assert tree.is_empty() is True
        tree.root.add_child("child")
        assert tree.is_empty() is False


class TestTreeTraversals:
    """三种遍历：前序、后序、层序"""

    def setup_method(self):
        tree = Tree("A")
        b = tree.root.add_child("B")
        c = tree.root.add_child("C")
        b.add_child("D")
        b.add_child("E")
        c.add_child("F")
        self.tree = tree

    def test_preorder(self):
        assert list(self.tree.preorder()) == ["A", "B", "D", "E", "C", "F"]

    def test_postorder(self):
        assert list(self.tree.postorder()) == ["D", "E", "B", "F", "C", "A"]

    def test_level_order(self):
        assert list(self.tree.level_order()) == ["A", "B", "C", "D", "E", "F"]

    def test_preorder_single_node(self):
        tree = Tree("only")
        assert list(tree.preorder()) == ["only"]

    def test_postorder_single_node(self):
        tree = Tree("only")
        assert list(tree.postorder()) == ["only"]

    def test_level_order_single_node(self):
        tree = Tree("only")
        assert list(tree.level_order()) == ["only"]


class TestTreeFindAndMetrics:
    """find_by_condition, height, size"""

    def setup_method(self):
        tree = Tree("A")
        b = tree.root.add_child("B")
        c = tree.root.add_child("C")
        b.add_child("D")
        b.add_child("E")
        c.add_child("F")
        self.tree = tree

    def test_find_by_condition_exists(self):
        node = self.tree.find_by_condition(lambda v: v == "E")
        assert node is not None
        assert node.value == "E"

    def test_find_by_condition_not_exists(self):
        node = self.tree.find_by_condition(lambda v: v == "Z")
        assert node is None

    def test_find_by_condition_first_match(self):
        node = self.tree.find_by_condition(lambda v: isinstance(v, str) and len(v) == 1)
        assert node.value == "A"  # DFS 先遇到根

    def test_height_single_node(self):
        assert Tree("x").height() == 0

    def test_height_deep_tree(self):
        tree = Tree("1")
        tree.root.add_child("2").add_child("3").add_child("4")
        assert tree.height() == 3

    def test_height_balanced_tree(self):
        assert self.tree.height() == 2

    def test_size_single_node(self):
        assert Tree("x").size() == 1

    def test_size(self):
        assert self.tree.size() == 6


# ============================================================
# tree.py — Stack
# ============================================================


class TestStack:
    """Stack 的全部操作"""

    def test_push_and_pop(self):
        s = Stack()
        s.push(10)
        s.push(20)
        assert s.pop() == 20
        assert s.pop() == 10

    def test_peek(self):
        s = Stack()
        s.push("hello")
        assert s.peek() == "hello"
        assert s.size() == 1  # peek 不弹出

    def test_peek_safe_on_empty(self):
        s = Stack()
        assert s.peek_safe() is None

    def test_peek_safe_returns_top(self):
        s = Stack()
        s.push(42)
        assert s.peek_safe() == 42

    def test_is_empty(self):
        s = Stack()
        assert s.is_empty() is True
        s.push(1)
        assert s.is_empty() is False

    def test_size(self):
        s = Stack()
        assert s.size() == 0
        s.push("a")
        s.push("b")
        assert s.size() == 2

    def test_clear(self):
        s = Stack()
        s.push(1)
        s.push(2)
        s.clear()
        assert s.is_empty() is True
        assert s.size() == 0

    def test_pop_empty_raises(self):
        s = Stack()
        with pytest.raises(IndexError, match="pop from empty stack"):
            s.pop()

    def test_peek_empty_raises(self):
        s = Stack()
        with pytest.raises(IndexError, match="peek from empty stack"):
            s.peek()

    def test_bool_truthy(self):
        s = Stack()
        assert bool(s) is False
        s.push(1)
        assert bool(s) is True

    def test_repr(self):
        s = Stack()
        s.push(1)
        s.push(2)
        assert "1" in repr(s)
        assert "2" in repr(s)


# ============================================================
# tree.py — print_tree
# ============================================================


class TestPrintTree:
    """print_tree 输出捕获"""

    def test_print_single_node(self):
        node = TreeNode("hello")
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(node)
        output = buf.getvalue()
        assert "hello" in output
        # rich 渲染单根节点无线条前缀

    def test_print_dict_value_node(self):
        node = TreeNode(
            {
                "category": "body_text",
                "paragraph": "这是一段正文内容",
                "fingerprint": "fp001",
            }
        )
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(node)
        output = buf.getvalue()
        assert "body_text" in output

    def test_print_tree_with_children(self):
        root = TreeNode("root")
        root.add_child("child1")
        root.add_child("child2")
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(root)
        output = buf.getvalue()
        assert "root" in output
        assert "child1" in output
        assert "child2" in output


# ============================================================
# rules/node.py — TreeNode
# ============================================================


class TestTreeNode:
    def test_init_with_simple_value(self):
        node = TreeNode("hello")
        assert node.value == "hello"
        assert node.children == []

    def test_init_fingerprint_is_none(self):
        """fingerprint 已废弃，始终为 None。"""
        node = TreeNode({"category": "body_text"})
        assert node.fingerprint is None

        node2 = TreeNode({"category": "body_text", "fingerprint": "abc123"})
        assert node2.fingerprint is None

    def test_config_default_empty(self):
        node = TreeNode("x")
        assert node.config == {}

    def test_load_config_nested_path(self):
        node = TreeNode("x")
        node.NODE_TYPE = "a.b.c"
        full = {"a": {"b": {"c": {"key": "val"}}}}
        node.load_config(full)
        assert node.config == {"key": "val"}

    def test_load_config_missing_path(self):
        node = TreeNode("x")
        node.NODE_TYPE = "x.y.z"
        node.load_config({"a": 1})
        assert node.config == {}

    def test_load_config_non_dict_input(self):
        node = TreeNode("x")
        node.load_config("not a dict")
        assert node.config == {}

    def test_add_child_returns_child(self):
        node = TreeNode("parent")
        child = node.add_child("child_val")
        assert child.value == "child_val"
        assert len(node.children) == 1

    def test_add_child_node(self):
        parent = TreeNode("parent")
        child = TreeNode("child")
        parent.add_child_node(child)
        assert parent.children[0] is child

    def test_repr(self):
        node = TreeNode("test_val")
        assert repr(node) == "TreeNode(test_val)"

    def test_fingerprint_attribute_is_none_for_non_dict(self):
        node = TreeNode("simple_string")
        assert node.fingerprint is None


# ============================================================
# rules/node.py — FormatNode
# ============================================================


class TestFormatNode:
    def test_init_defaults(self):
        node = FormatNode(value="test", level=1)
        assert node.level == 1
        assert node.paragraph is None
        assert node.expected_rule is None
        assert isinstance(node.pydantic_config, DotDict)
        assert len(node.pydantic_config) == 0

    def test_pydantic_config_raises_before_load(self):
        node = FormatNode(value="test", level=1)
        # 未加载配置时 pydantic_config 返回空 DotDict（不再抛异常）
        cfg = node.pydantic_config
        assert isinstance(cfg, DotDict)

    def test_update_paragraph(self, doc):
        node = FormatNode(value="test", level=1)
        p = doc.add_paragraph("hello")
        node.update_paragraph(p)
        assert node.paragraph is p

    def test_base_is_noop(self, doc):
        """_base() 默认为空操作。"""
        node = FormatNode(value="test", level=1)
        node._base(doc, p=True, r=True)
        node._base(doc, p=False, r=False)

    def test_check_format_no_config(self, doc):
        """未加载配置时 check_format 不抛异常（handler 跳过）。"""
        p = doc.add_paragraph("test")
        node = FormatNode(value="test", level=1, paragraph=p)
        node.check_format(doc)  # 不应抛异常

    def test_apply_format_no_config(self, doc):
        """未加载配置时 apply_format 不抛异常（handler 跳过）。"""
        p = doc.add_paragraph("test")
        node = FormatNode(value="test", level=1, paragraph=p)
        node.apply_format(doc)  # 不应抛异常

    def test_add_comment_buffers(self, doc):
        """add_comment 缓冲文本，_flush_comments 合并写入为一条批注。"""
        node = FormatNode(value="test", level=1)
        p = doc.add_paragraph("hello")
        node.paragraph = p
        node.add_comment(doc, p.runs[0], "标题-字号错误：小四，规范：五号")
        node.add_comment(doc, p.runs[0], "标题-加粗错误：加粗，规范：不加粗")
        node._flush_comments(doc)
        comments = list(doc.comments)
        assert len(comments) == 1
        text = comments[0].text
        assert "字号错误" in text
        assert "加粗错误" in text
        # 每条问题各占一段
        assert len(comments[0].paragraphs) == 2

    def test_add_comment_empty_text_skipped(self, doc):
        node = FormatNode(value="test", level=1)
        p = doc.add_paragraph("hello")
        run = p.runs[0]
        # 空文本不应调用 add_comment
        node.add_comment(doc, run, "   ")

    def test_load_config_with_node_type(self):
        """load_config 沿 NODE_TYPE 路径提取 dict 并与 DEFAULTS 合并。"""

        class TestNode(FormatNode):
            NODE_TYPE = "headings.level_1"
            DEFAULTS = {"font_size": "小二", "alignment": "居中对齐"}

        node = TestNode(value="test", level=1)
        node.load_config(
            {
                "headings": {
                    "level_1": {"font_size": "三号", "chinese_font_name": "黑体"}
                }
            }
        )
        assert node.pydantic_config.font_size == "三号"  # YAML 覆盖
        assert node.pydantic_config.alignment == "居中对齐"  # DEFAULTS
        assert node.pydantic_config.chinese_font_name == "黑体"  # YAML


class TestDotDict:
    """覆盖 config/dotdict.py。"""

    def test_setattr(self):
        from wordformat.config.dotdict import DotDict

        d = DotDict()
        d.alignment = "居中"
        assert d["alignment"] == "居中"

    def test_delattr_existing_key(self):
        from wordformat.config.dotdict import DotDict

        d = DotDict(a=1)
        del d.a
        assert "a" not in d

    def test_delattr_missing_key_raises(self):
        from wordformat.config.dotdict import DotDict
        import pytest

        d = DotDict()
        with pytest.raises(AttributeError):
            del d.nonexistent

    def test_getattr_nested_dict_returns_dotdict(self):
        from wordformat.config.dotdict import DotDict

        d = DotDict({"outer": {"inner": 1}})
        assert isinstance(d.outer, DotDict)
        assert d.outer.inner == 1

    def test_deep_merge_nested(self):
        from wordformat.config.dotdict import deep_merge

        base = {"a": {"b": 1, "c": 2}}
        override = {"a": {"b": 99}}
        result = deep_merge(base, override)
        assert result["a"]["b"] == 99  # overridden
        assert result["a"]["c"] == 2  # kept


class TestExportDefaults:
    """覆盖 structure/registry.py export_defaults()。"""

    def test_export_returns_dict_with_key_sections(self):
        from wordformat.structure.registry import export_defaults

        # 需要先触发 @register（已通过 test 导入完成）
        from wordformat.rules import (  # noqa: F401
            AbstractContentCN,
            AbstractTitleCN,
            BodyText,
            HeadingLevel1Node,
        )

        data = export_defaults()
        assert isinstance(data, dict)
        assert "template_name" in data
        assert "style_checks_warning" in data
        assert "abstract" in data
        assert "headings" in data
        assert "body" in data
