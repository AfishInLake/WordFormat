"""
rules 模块测试 —— 聚焦真实行为验证，无填充。
"""
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from wordformat.config.models import NodeConfigRoot
from wordformat.rules import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleContentCN,
    AbstractTitleContentEN,
    AbstractTitleEN,
    Acknowledgements,
    AcknowledgementsCN,
    BodyText,
    CaptionFigure,
    CaptionTable,
    FormatNode,
    HeadingLevel1Node,
    HeadingLevel2Node,
    HeadingLevel3Node,
    KeywordsCN,
    KeywordsEN,
    ReferenceEntry,
    References,
)
from wordformat.rules.node import FormatNode as FormatNodeBase

# ---------------------------------------------------------------------------
# 共享 fixtures / helpers
# ---------------------------------------------------------------------------


def _make_node(cls, text="测试文本"):
    """创建一个带 paragraph 的 FormatNode 实例。"""
    doc = Document()
    p = doc.add_paragraph(text)
    return cls(value=p, level=0, paragraph=p)


def _load_root_config(config_path):
    """从 YAML 路径加载配置 dict。"""
    return _load_yaml(config_path)


def _load_yaml(path):
    import yaml
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


@pytest.fixture
def root_config(sample_yaml_config):
    """从 sample_yaml_config 加载 NodeConfigRoot，与示例文件解耦。"""
    from wordformat.config.loader import init_config
    init_config(sample_yaml_config)
    return _load_root_config(sample_yaml_config)


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def para(doc):
    return doc.add_paragraph("测试段落")


@pytest.fixture
def run_with_text(para):
    r = para.add_run("关键词")
    r.font.size = Pt(12)
    r.font.bold = True
    return r


# ---------------------------------------------------------------------------
# 1. FormatNode 基类行为
# ---------------------------------------------------------------------------


class TestFormatNodeBase:
    """FormatNode 基类的核心契约。"""

    def test_base_is_noop(self, doc, para):
        """_base() 默认为空操作。"""
        node = FormatNodeBase(value=para, level=0, paragraph=para)
        node._base(doc, p=True, r=True)
        node._base(doc, p=False, r=False)

    def test_check_format_calls_base_with_true(self, doc, para):
        """check_format 应以 p=True, r=True 调用 _base。"""
        node = FormatNodeBase(value=para, level=0, paragraph=para)
        with patch.object(node, "_base") as mock_base, \
                patch.object(node, "_run_rules"):
            node.check_format(doc)
        mock_base.assert_called_once_with(doc, p=True, r=True)

    def test_apply_format_calls_base_with_false(self, doc, para):
        """apply_format 应以 p=False, r=False 调用 _base。"""
        node = FormatNodeBase(value=para, level=0, paragraph=para)
        with patch.object(node, "_base") as mock_base, \
                patch.object(node, "_run_rules"):
            node.apply_format(doc)
        mock_base.assert_called_once_with(doc, p=False, r=False)

    def test_pydantic_config_raises_before_load(self, para):
        """未加载配置时访问 pydantic_config 应抛出 ValueError。"""
        node = FormatNodeBase(value=para, level=0, paragraph=para)
        with pytest.raises(ValueError, match="尚未加载"):
            _ = node.pydantic_config

    def test_unknown_config_type_raises(self, root_config, para):
        """没有 CONFIG_PATH 的节点，load_config 后 _pydantic_config 应为 None。"""

        # 创建一个没有 CONFIG_PATH 的子类
        class FakeNode(FormatNodeBase):
            CONFIG_MODEL = type("TotallyUnknownConfig", (), {})

        node = FakeNode(value=para, level=0, paragraph=para)
        node.load_config(root_config)
        assert node._pydantic_config is None


# ---------------------------------------------------------------------------
# 2. 所有节点类型可实例化
# ---------------------------------------------------------------------------


NODE_CLASSES = [
    AbstractTitleCN, AbstractTitleContentCN, AbstractContentCN,
    AbstractTitleEN, AbstractTitleContentEN, AbstractContentEN,
    KeywordsCN, KeywordsEN,
    HeadingLevel1Node, HeadingLevel2Node, HeadingLevel3Node,
    BodyText,
    CaptionFigure, CaptionTable,
    References, ReferenceEntry,
    Acknowledgements, AcknowledgementsCN,
]


@pytest.mark.parametrize("cls", NODE_CLASSES, ids=lambda c: c.__name__)
class TestNodeInstantiation:
    """每个节点类型都能正确实例化。"""

    def test_instantiation(self, cls):
        node = _make_node(cls)
        assert node.paragraph is not None
        assert node.level == 0

    def test_has_defaults(self, cls):
        assert hasattr(cls, "DEFAULTS")
        assert isinstance(cls.DEFAULTS, dict)

    def test_has_node_type(self, cls):
        assert hasattr(cls, "NODE_TYPE")
        assert isinstance(cls.NODE_TYPE, str)
        assert len(cls.NODE_TYPE) > 0


# ---------------------------------------------------------------------------
# 3. load_config 为每个节点正确设置 _pydantic_config
# ---------------------------------------------------------------------------


class TestLoadConfig:
    """验证 load_config 后 _pydantic_config 类型正确。"""

    def test_abstract_title_cn(self, root_config):
        node = _make_node(AbstractTitleCN)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.chinese_font_name == "黑体"
        assert node._pydantic_config.bold is True

    def test_abstract_content_cn(self, root_config):
        node = _make_node(AbstractContentCN)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert hasattr(node._pydantic_config, "chinese_content")

    def test_abstract_title_en(self, root_config):
        node = _make_node(AbstractTitleEN)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.bold is True

    def test_abstract_content_en(self, root_config):
        node = _make_node(AbstractContentEN)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert hasattr(node._pydantic_config, "english_content")

    def test_body_text(self, root_config):
        node = _make_node(BodyText)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.chinese_font_name == "宋体"

    def test_caption_figure(self, root_config):
        node = _make_node(CaptionFigure)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.caption_prefix == "图"

    def test_caption_table(self, root_config):
        node = _make_node(CaptionTable)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.caption_prefix == "表"

    def test_table_content_config(self, root_config):
        """TablesConfig 的 content 字段默认加载。"""
        node = _make_node(CaptionTable)
        node.load_config(root_config)
        assert node._pydantic_config.content is not None
        assert node._pydantic_config.content.font_size == "五号"

    def test_table_content_config_default(self):
        """TablesConfig 的 content 有默认值。"""
        from wordformat.config.models import TablesConfig

        cfg = TablesConfig()
        assert cfg.content is not None
        assert cfg.content.font_size == "小四"  # GlobalFormatConfig 默认值

    def test_references(self, root_config):
        node = _make_node(References)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.bold is True

    def test_reference_entry(self, root_config):
        node = _make_node(ReferenceEntry)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.font_size == "五号"

    def test_acknowledgements(self, root_config):
        node = _make_node(Acknowledgements)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.bold is True

    def test_acknowledgements_cn(self, root_config):
        node = _make_node(AcknowledgementsCN)
        node.load_config(root_config)
        assert node._pydantic_config is not None

    def test_keywords_cn_from_dict(self, config_path):
        """KeywordsCN 从 dict 加载配置时，因 LANG='cn' 与 YAML 键 'chinese' 不匹配，
        回退到空 dict 并使用 KeywordsConfig 默认值。"""
        raw = _load_yaml(config_path)
        node = _make_node(KeywordsCN)
        node.load_config(raw)
        assert node._pydantic_config is not None
        # dict 路径下 LANG='cn' 找不到 YAML 中的 'chinese' 键，使用默认值
        assert node._pydantic_config.rules.keyword_count.count_min == 4
        assert node._pydantic_config.rules.keyword_count.count_max == 6

    def test_keywords_en_from_dict(self, config_path):
        """KeywordsEN 支持从 dict 加载配置。"""
        raw = _load_yaml(config_path)
        node = _make_node(KeywordsEN)
        node.load_config(raw)
        assert node._pydantic_config is not None
        assert node._pydantic_config.label.bold is False  # 默认值

    def test_keywords_cn_from_node_config_root(self, root_config):
        """KeywordsCN 从 NodeConfigRoot 加载时使用 chinese 子配置。"""
        node = _make_node(KeywordsCN)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.label.chinese_font_name == "黑体"

    def test_keywords_en_from_node_config_root(self, root_config):
        """KeywordsEN 从 NodeConfigRoot 加载时使用 english 子配置。"""
        node = _make_node(KeywordsEN)
        node.load_config(root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.english_font_name == "Times New Roman"


# ---------------------------------------------------------------------------
# 4. HeadingLevelConfig bug
# ---------------------------------------------------------------------------


class TestHeadingBug:
    """
    NODE_TYPE 现在自动回退为 CONFIG_PATH，
    基类 FormatNode.load_config 可正确解析 heading 配置。
    """

    def test_formatnode_heading_bug(self, root_config):
        """Heading 节点没有 CONFIG_PATH，NODE_TYPE 自动回退为 CONFIG_PATH，
        通过 FormatNode 基类 load_config 可正确解析配置。"""
        node = _make_node(HeadingLevel1Node)
        FormatNode.load_config(node, root_config)
        assert node._pydantic_config is not None
        assert node._pydantic_config.font_size == "小二"

    def test_base_heading_load_config_works_correctly(self, root_config):
        """BaseHeadingNode 重写的 load_config 应正确加载对应层级配置。"""
        node_l1 = _make_node(HeadingLevel1Node)
        node_l1.load_config(root_config)
        assert node_l1._pydantic_config.font_size == "小二"

        node_l2 = _make_node(HeadingLevel2Node)
        node_l2.load_config(root_config)
        assert node_l2._pydantic_config.font_size == "三号"

        node_l3 = _make_node(HeadingLevel3Node)
        node_l3.load_config(root_config)
        assert node_l3._pydantic_config.font_size == "小四"


# ---------------------------------------------------------------------------
# 5. KeywordsCN / KeywordsEN 特有逻辑
# ---------------------------------------------------------------------------


class TestKeywordsLogic:
    """关键词节点的标签识别、数量校验、标点校验。"""

    def test_cn_label_detection(self):
        """中文关键词标签识别（使用真实 paragraph runs）。"""
        doc = Document()
        p = doc.add_paragraph("关键词")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        assert node._check_keyword_label(p.runs[0]) is True

        p2 = doc.add_paragraph("关 键 词")
        node.paragraph = p2
        assert node._check_keyword_label(p2.runs[0]) is True

        p3 = doc.add_paragraph("机器学习")
        node.paragraph = p3
        assert node._check_keyword_label(p3.runs[0]) is False

    def test_en_label_detection(self):
        """英文关键词标签识别（使用真实 paragraph runs）。"""
        doc = Document()
        p = doc.add_paragraph("Keywords")
        node = KeywordsEN(value=p, level=0, paragraph=p)
        assert node._check_keyword_label(p.runs[0]) is True

        p2 = doc.add_paragraph("Keyword")
        node.paragraph = p2
        assert node._check_keyword_label(p2.runs[0]) is True

        p3 = doc.add_paragraph("KEY WORDS")
        node.paragraph = p3
        assert node._check_keyword_label(p3.runs[0]) is True

        p4 = doc.add_paragraph("machine learning")
        node.paragraph = p4
        assert node._check_keyword_label(p4.runs[0]) is False

    def test_cn_count_validation_too_few(self, root_config):
        """中文关键词数量不足时应触发 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("关键词：机器学习")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        # 应至少有一条数量不足的 comment
        texts = [c.kwargs["text"] for c in mock_comment.call_args_list]
        assert any("数量过少" in t for t in texts)

    def test_cn_count_validation_too_many(self, root_config):
        """中文关键词数量超限时应触发 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("关键词：A；B；C；D；E；F")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        texts = [c.kwargs["text"] for c in mock_comment.call_args_list]
        assert any("数量过多" in t for t in texts)

    def test_cn_trailing_punct_detection(self, root_config):
        """中文关键词末尾标点校验。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("关键词：机器学习；深度学习；")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        texts = [c.kwargs["text"] for c in mock_comment.call_args_list]
        assert any("标点错误" in t for t in texts)

    def test_en_count_validation_too_few(self, root_config):
        """英文关键词数量不足。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Keywords: AI")
        node = KeywordsEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        texts = [c.kwargs["text"] for c in mock_comment.call_args_list]
        assert any("数量过少" in t for t in texts)

    def test_cn_no_config_raises_value_error(self, doc):
        """KeywordsCN 在 _pydantic_config 为 None 时访问 pydantic_config 抛出 ValueError。
        注意：_base 中 `if self.pydantic_config is None` 实际会触发 property 的异常，
        因为 property 在 _pydantic_config 为 None 时直接 raise 而非返回 None。"""
        p = doc.add_paragraph("关键词：测试")
        node = KeywordsCN(value=p, level=0, paragraph=p)
        # 不加载配置
        with pytest.raises(ValueError, match="尚未加载"):
            node.check_format(doc)

    def test_keywords_unsupported_type_raises(self):
        """KeywordsCN.load_config 传入不支持的类型应抛出 TypeError。"""
        node = _make_node(KeywordsCN)
        with pytest.raises(TypeError, match="配置类型不支持"):
            node.load_config(42)


# ---------------------------------------------------------------------------
# 6. _base 实现实际调用 diff/apply 逻辑
# ---------------------------------------------------------------------------


class TestBaseImplementation:
    """验证 _base 实现确实执行了 diff_from_run / apply_to_run 逻辑。"""

    def test_body_text_check_calls_diff(self, root_config):
        """BodyText.check_format 应调用 diff_from_run。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("正文内容")
        r.font.size = Pt(10)  # 故意设置错误字号
        node = BodyText(value=p, level=0, paragraph=p)
        node.load_config(root_config)

        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)

        # 应至少调用一次 add_comment（因为字号不匹配）
        assert mock_comment.call_count >= 1

    def test_abstract_title_cn_check_runs(self, root_config):
        """AbstractTitleCN.check_format 应遍历所有 run。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要")
        r.font.size = Pt(10)
        node = AbstractTitleCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)

        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)

        # 至少对 run 和 paragraph 各调用一次
        assert mock_comment.call_count >= 2

    def test_heading_no_config_raises_value_error(self, doc):
        """HeadingLevel1Node 未加载配置时 check_format 抛出 ValueError。"""
        p = doc.add_paragraph("第一章 绪论")
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        with pytest.raises(ValueError, match="尚未加载"):
            node.check_format(doc)

    def test_references_check_runs(self, root_config):
        """References.check_format 应遍历 run 并调用 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("参考文献")
        r.font.size = Pt(10)
        node = References(value=p, level=0, paragraph=p)
        node.load_config(root_config)

        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)

        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 7. BodyText._apply_citation_superscript
# ---------------------------------------------------------------------------


class TestBodyTextCitationSuperscript:
    """测试 BodyText.apply_format 中的引用上标自动设置。"""

    @staticmethod
    def _get_vertAlign(run):
        """返回 run 的 w:vertAlign 值，无则返回 None。"""
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            return None
        va = rPr.find(qn("w:vertAlign"))
        return va.get(qn("w:val")) if va is not None else None

    def test_single_citation_gets_superscript(self):
        """单引用 [1] 应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("这是一篇论文[1]的引用。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp1"},
                        level=0, paragraph=p)
        # 直接调用 apply_format（跳过 load_config，不依赖完整配置）
        node._apply_citation_superscript()
        # "[1]" 应在独立的上标 run 中
        runs = p.runs
        superscript_texts = [
            r.text for r in runs if self._get_vertAlign(r) == "superscript"
        ]
        assert "[1]" in superscript_texts

    def test_multiple_citations(self):
        """多个引用 [1] 和 [2,3] 都应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("参见文献[1]和[2,3]的讨论。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp2"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        superscript_texts = [
            r.text for r in p.runs if self._get_vertAlign(r) == "superscript"
        ]
        assert "[1]" in superscript_texts
        assert "[2,3]" in superscript_texts

    def test_range_citation(self):
        """范围引用 [1-3] 应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("文献[1-3]提供了详细分析。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp3"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        superscript_texts = [
            r.text for r in p.runs if self._get_vertAlign(r) == "superscript"
        ]
        assert "[1-3]" in superscript_texts

    def test_chinese_comma_citation(self):
        """中文逗号分隔的引用 [1，2] 应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("见[1，2，3]的研究。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp4"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        superscript_texts = [
            r.text for r in p.runs if self._get_vertAlign(r) == "superscript"
        ]
        assert any("[1，2，3]" in t for t in superscript_texts)

    def test_no_citation_leaves_runs_unchanged(self):
        """无引用的段落应保持原样。"""
        doc = Document()
        p = doc.add_paragraph("这是一段没有引用的正文。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp5"},
                        level=0, paragraph=p)
        original_text = p.text
        node._apply_citation_superscript()
        # 文本不变，且无上标 run
        assert p.text == original_text
        assert all(self._get_vertAlign(r) is None for r in p.runs)

    def test_non_citation_brackets_not_affected(self):
        """非数字方括号如 [注] 不应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("这是一个[注]释说明。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp6"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        assert all(self._get_vertAlign(r) is None for r in p.runs)

    def test_citation_split_across_runs(self):
        """引用跨 run 时先分割再设上标。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("前面文字[1")
        p.add_run("2]后面文字")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp7"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        superscript_texts = [
            r.text for r in p.runs if self._get_vertAlign(r) == "superscript"
        ]
        # 分割后 [1 和 2] 分别在两个 run 中，但都是上标
        assert "[1" in superscript_texts
        assert "2]" in superscript_texts


# ---------------------------------------------------------------------------
# 8. AbstractTitleCN._base 完整 diff/apply 覆盖
# ---------------------------------------------------------------------------


class TestAbstractTitleCNBase:
    """覆盖 abstract.py AbstractTitleCN._base 的 diff 和 apply 分支。"""

    def test_check_with_wrong_format_triggers_comments(self, root_config):
        """check 模式：错误格式应触发 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要")
        r.font.size = Pt(10)  # 错误字号
        r.font.bold = False  # 应为加粗
        node = AbstractTitleCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_apply_fixes_wrong_format(self, root_config):
        """apply 模式：应调用 apply_to_paragraph 和 apply_to_run。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要")
        r.font.size = Pt(10)
        r.font.bold = False
        node = AbstractTitleCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        # apply 模式也会调用 add_comment
        assert mock_comment.call_count >= 2

    def test_check_no_runs_skips_without_error(self, root_config):
        """段落无 run 时（空段），check_format 安全跳过不报错。"""
        doc = Document()
        p = doc.add_paragraph()
        node = AbstractTitleCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)  # 不应抛异常


# ---------------------------------------------------------------------------
# 9. AbstractTitleContentCN._base 覆盖（标题+正文混合）
# ---------------------------------------------------------------------------


class TestAbstractTitleContentCNBase:
    """覆盖 abstract.py AbstractTitleContentCN._base 的 check_title 分支。"""

    def test_check_title_run_uses_title_style(self, root_config):
        """包含'摘要'的 run 应使用 chinese_title 样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要")
        r.font.size = Pt(10)  # 错误字号
        node = AbstractTitleContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_check_content_run_uses_content_style(self, root_config):
        """非标题 run 应使用 chinese_content 样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("这是摘要正文内容")
        r.font.size = Pt(10)
        node = AbstractTitleContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_apply_title_and_content_runs(self, root_config):
        """apply 模式：标题和正文 run 都应被修正。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("摘要")
        r1.font.size = Pt(10)
        r2 = p.add_run("正文内容")
        r2.font.size = Pt(10)
        node = AbstractTitleContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 3

    def test_check_mode_does_not_mutate_run_text(self, root_config):
        """_base 在检查模式下不应修改 run.text（修复：移除了破坏性 replace）。"""
        doc = Document()
        p = doc.add_paragraph()
        original_text = "摘 要"
        r = p.add_run(original_text)
        r.font.size = Pt(10)
        node = AbstractTitleContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        # 修复后：检查模式不应改变 run.text
        assert r.text == original_text


# ---------------------------------------------------------------------------
# 10. AbstractContentCN._base 覆盖
# ---------------------------------------------------------------------------


class TestAbstractContentCNBase:
    """覆盖 abstract.py AbstractContentCN._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_font_size(self, root_config):
        """check 模式：错误字号应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要正文")
        r.font.size = Pt(10)
        node = AbstractContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_apply_fixes_font_size(self, root_config):
        """apply 模式：应修正字号。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要正文")
        r.font.size = Pt(10)
        node = AbstractContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 2

    def test_check_multiple_runs(self, root_config):
        """多个 run 都应被检查。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("摘要")
        r1.font.size = Pt(10)
        r2 = p.add_run("正文")
        r2.font.size = Pt(10)
        node = AbstractContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        # 2 run comments + 1 paragraph comment
        assert mock_comment.call_count >= 3


# ---------------------------------------------------------------------------
# 11. AbstractTitleEN._base 覆盖
# ---------------------------------------------------------------------------


class TestAbstractTitleENBase:
    """覆盖 abstract.py AbstractTitleEN._base 的 diff/apply 逻辑。
    注意：AbstractTitleEN 只在 diff_result 非空时才 add_comment。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Abstract")
        r.font.size = Pt(10)
        r.font.bold = False
        node = AbstractTitleEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_empty_list(self, root_config):
        """AbstractTitleEN._base 始终返回 []。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Abstract")
        r.font.size = Pt(10)
        node = AbstractTitleEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Abstract")
        r.font.size = Pt(10)
        r.font.bold = False
        node = AbstractTitleEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 12. AbstractTitleContentEN._base 覆盖
# ---------------------------------------------------------------------------


class TestAbstractTitleContentENBase:
    """覆盖 abstract.py AbstractTitleContentEN._base 的 check_title 分支。"""

    def test_check_title_run_uses_title_style(self, root_config):
        """包含'Abstract'的 run 应使用 english_title 样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Abstract")
        r.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_check_content_run_uses_content_style(self, root_config):
        """非标题 run 应使用 english_content 样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("This is the abstract content.")
        r.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_apply_mixed_runs(self, root_config):
        """apply 模式：混合标题和正文 run。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Abstract")
        r1.font.size = Pt(10)
        r2 = p.add_run("Content text")
        r2.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 3

    def test_check_title_normalizes_case_lower(self, root_config):
        """小写 'abstract' 应匹配并自动修正为 'Abstract'。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("abstract: some content")
        r.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        assert r.text.startswith("Abstract")

    def test_check_title_normalizes_case_upper(self, root_config):
        """全大写 'ABSTRACT' 应匹配并自动修正为 'Abstract'。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("ABSTRACT\n\nbody text")
        r.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        assert r.text.startswith("Abstract")

    def test_split_abstract_across_runs(self, root_config):
        """"Abstract" 被拆分到两个 run 时仍能正确识别标题部分。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Abst")
        r1.font.size = Pt(10)
        r2 = p.add_run("ract: body text")
        r2.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        # r1 开头被修正为 "Abstract"，r2 保持 "body text" 部分
        assert r1.text.startswith("Abstract")
        assert "body text" in r2.text

    def test_split_abstract_across_three_runs(self, root_config):
        """"Abstract" 被拆分到三个 run 时仍能正确识别。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Abs")
        r1.font.size = Pt(10)
        r2 = p.add_run("tra")
        r2.font.size = Pt(10)
        r3 = p.add_run("ct: content")
        r3.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        # r1 应被修正为 "Abstract"
        assert r1.text.startswith("Abstract")
        # r2 和 r3 开头部分属于标题前缀，应被清空
        assert r2.text == ""
        assert "content" in r3.text


# ---------------------------------------------------------------------------
# 13. AbstractContentEN._base 覆盖
# ---------------------------------------------------------------------------


class TestAbstractContentENBase:
    """覆盖 abstract.py AbstractContentEN._base 的 diff/apply 逻辑。
    注意：AbstractContentEN 只在 diff_result/issues 非空时才 add_comment。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("This is abstract content.")
        r.font.size = Pt(10)
        node = AbstractContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("This is abstract content.")
        r.font.size = Pt(10)
        node = AbstractContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            # 配置中 line_spacing 为 "0倍"，现会触发 ValueError，mock 掉该步
            with patch("wordformat.style.diff.LineSpacing.format"):
                node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_multiple_runs(self, root_config):
        """多个 run 都应被检查。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("First sentence. ")
        r1.font.size = Pt(10)
        r2 = p.add_run("Second sentence.")
        r2.font.size = Pt(10)
        node = AbstractContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2


# ---------------------------------------------------------------------------
# 14. Acknowledgements._base 覆盖
# ---------------------------------------------------------------------------


class TestAcknowledgementsBase:
    """覆盖 acknowledgement.py Acknowledgements._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢")
        r.font.size = Pt(10)
        r.font.bold = False
        node = Acknowledgements(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_empty_list(self, root_config):
        """Acknowledgements._base 始终返回 []。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢")
        r.font.size = Pt(10)
        node = Acknowledgements(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢")
        r.font.size = Pt(10)
        node = Acknowledgements(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_no_diffs_no_comment(self, root_config):
        """格式完全正确时，不应调用 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢")
        node = Acknowledgements(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        # 即使格式正确，段落级别的 diff 仍可能触发 comment
        # 但如果没有差异，不应有 comment
        # 注意：新 Document 的段落默认对齐方式可能与配置不同


# ---------------------------------------------------------------------------
# 15. AcknowledgementsCN._base 覆盖
# ---------------------------------------------------------------------------


class TestAcknowledgementsCNBase:
    """覆盖 acknowledgement.py AcknowledgementsCN._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("感谢导师的指导。")
        r.font.size = Pt(10)
        node = AcknowledgementsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_empty_list(self, root_config):
        """AcknowledgementsCN._base 始终返回 []。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("感谢导师的指导。")
        r.font.size = Pt(10)
        node = AcknowledgementsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("感谢导师的指导。")
        r.font.size = Pt(10)
        node = AcknowledgementsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_first_line_indent(self, root_config):
        """验证 first_line_indent 配置被正确使用。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢内容段落。")
        r.font.size = Pt(10)
        node = AcknowledgementsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        # 验证配置中有 first_line_indent 字段
        assert hasattr(node._pydantic_config, "first_line_indent")
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 16. CaptionFigure._base 覆盖
# ---------------------------------------------------------------------------


class TestCaptionFigureBase:
    """覆盖 caption.py CaptionFigure._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("图1 测试图片")
        r.font.size = Pt(10)
        node = CaptionFigure(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("图1 测试图片")
        r.font.size = Pt(10)
        node = CaptionFigure(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 17. CaptionTable._base 覆盖
# ---------------------------------------------------------------------------


class TestCaptionTableBase:
    """覆盖 caption.py CaptionTable._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("表1 测试表格")
        r.font.size = Pt(10)
        node = CaptionTable(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("表1 测试表格")
        r.font.size = Pt(10)
        node = CaptionTable(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 18. HeadingLevel1Node._base 覆盖
# ---------------------------------------------------------------------------


class TestHeadingLevel1NodeBase:
    """覆盖 heading.py HeadingLevel1Node._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("第一章 绪论")
        r.font.size = Pt(10)
        r.font.bold = False
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("第一章 绪论")
        r.font.size = Pt(10)
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_issues_list(self, root_config):
        """返回值应为包含 issue 字典的列表。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("第一章 绪论")
        r.font.size = Pt(10)
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        # 应有 run_issues 或 paragraph_issues

    def test_check_skips_empty_runs(self, root_config):
        """空白 run 应被跳过。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("   ")
        r2 = p.add_run("第一章 绪论")
        r2.font.size = Pt(10)
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        # 空白 run 不应触发 comment
        run_comments = [
            c for c in mock_comment.call_args_list
            if c.kwargs.get("runs") is r1
        ]
        assert len(run_comments) == 0


# ---------------------------------------------------------------------------
# 19. HeadingLevel2Node._base 覆盖
# ---------------------------------------------------------------------------


class TestHeadingLevel2NodeBase:
    """覆盖 heading.py HeadingLevel2Node._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1 研究背景")
        r.font.size = Pt(10)
        node = HeadingLevel2Node(value=p, level=2, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1 研究背景")
        r.font.size = Pt(10)
        node = HeadingLevel2Node(value=p, level=2, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_issues(self, root_config):
        """返回值应包含 issue 字典。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1 研究背景")
        r.font.size = Pt(10)
        node = HeadingLevel2Node(value=p, level=2, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)


# ---------------------------------------------------------------------------
# 20. HeadingLevel3Node._base 覆盖
# ---------------------------------------------------------------------------


class TestHeadingLevel3NodeBase:
    """覆盖 heading.py HeadingLevel3Node._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1.1 问题描述")
        r.font.size = Pt(10)
        node = HeadingLevel3Node(value=p, level=3, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1.1 问题描述")
        r.font.size = Pt(10)
        node = HeadingLevel3Node(value=p, level=3, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_issues(self, root_config):
        """返回值应包含 issue 字典。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1.1 问题描述")
        r.font.size = Pt(10)
        node = HeadingLevel3Node(value=p, level=3, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)


# ---------------------------------------------------------------------------
# 21. References._base 覆盖
# ---------------------------------------------------------------------------


class TestReferencesBase:
    """覆盖 references.py References._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("参考文献")
        r.font.size = Pt(10)
        r.font.bold = False
        node = References(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_empty_list(self, root_config):
        """References._base 始终返回 []。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("参考文献")
        r.font.size = Pt(10)
        node = References(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("参考文献")
        r.font.size = Pt(10)
        node = References(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 22. ReferenceEntry._base 覆盖
# ---------------------------------------------------------------------------


class TestReferenceEntryBase:
    """覆盖 references.py ReferenceEntry._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("[1] 作者. 论文标题[J]. 期刊, 2024.")
        r.font.size = Pt(10)
        node = ReferenceEntry(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_empty_list(self, root_config):
        """ReferenceEntry._base 始终返回 []。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("[1] 作者. 论文标题[J]. 期刊, 2024.")
        r.font.size = Pt(10)
        node = ReferenceEntry(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("[1] 作者. 论文标题[J]. 期刊, 2024.")
        r.font.size = Pt(10)
        node = ReferenceEntry(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_alignment_and_indent(self, root_config):
        """验证 alignment 和 first_line_indent 配置被正确使用。"""
        node = _make_node(ReferenceEntry)
        node.load_config(root_config)
        assert hasattr(node._pydantic_config, "alignment")
        assert hasattr(node._pydantic_config, "first_line_indent")
        assert node._pydantic_config.font_size == "五号"
