#!/usr/bin/env python
"""
comprehensive tests for style modules:
  check_format, get_some, set_some, style_enum, utils
"""

import pytest
from types import SimpleNamespace
from unittest.mock import MagicMock, PropertyMock, patch as mock_patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from wordformat.style.diff import DIFFResult, CharacterStyle, ParagraphStyle
from wordformat.style.reader import (
    paragraph_get_alignment, paragraph_get_space_before, paragraph_get_space_after,
    paragraph_get_line_spacing, paragraph_get_first_line_indent,
    paragraph_get_builtin_style_name, run_get_font_name, run_get_font_size_pt,
    run_get_font_color, run_get_font_bold, run_get_font_italic,
    run_get_font_underline, GetIndent, _get_style_spacing,
)
from wordformat.style.writer import (
    run_set_font_name, set_paragraph_space_before_by_lines,
    set_paragraph_space_after_by_lines, _paragraph_space_by_lines,
    SetSpacing, SetLineSpacing, SetIndent, SetFirstLineIndent,
)
from wordformat.style.defs import (
    FontName, FontSize, FontColor, Alignment, LineSpacingRule, LineSpacing,
    FirstLineIndent, LeftIndent, RightIndent, BuiltInStyle, SpaceBefore, SpaceAfter,
    Spacing, UnitLabelEnum,
)
from wordformat.style.units import extract_unit_from_string, UnitResult


# ---------------------------------------------------------------------------
# fixtures & helpers
# ---------------------------------------------------------------------------

@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def mock_warning():
    w = MagicMock()
    for attr in ("bold", "italic", "underline", "font_size", "font_color",
                 "font_name", "alignment", "space_before", "space_after",
                 "line_spacing", "line_spacingrule", "first_line_indent",
                 "left_indent", "right_indent", "builtin_style_name"):
        setattr(w, attr, True)
    return w


@pytest.fixture
def mock_warning_off():
    w = MagicMock()
    for attr in ("bold", "italic", "underline", "font_size", "font_color",
                 "font_name", "alignment", "space_before", "space_after",
                 "line_spacing", "line_spacingrule", "first_line_indent",
                 "left_indent", "right_indent", "builtin_style_name"):
        setattr(w, attr, False)
    return w


def _set_warning(w):
    import wordformat.style.diff as m
    m.style_checks_warning = w


def _clear_warning():
    import wordformat.style.diff as m
    m.__dict__.pop("style_checks_warning", None)


# ===========================================================================
# DIFFResult
# ===========================================================================

class TestDIFFResult:
    def test_defaults_and_str(self):
        d = DIFFResult()
        assert d.diff_type is d.expected_value is d.current_value is d.comment is None
        assert d.level == 0
        d2 = DIFFResult(comment="hi")
        assert str(d2) == "hi"

    @pytest.mark.parametrize("lvl", [0, 1, 2, 3])
    def test_level(self, lvl):
        assert DIFFResult(level=lvl).level == lvl



# ===========================================================================
# CharacterStyle
# ===========================================================================

class TestCharacterStyle:
    def test_defaults(self, mock_warning):
        _set_warning(mock_warning)
        cs = CharacterStyle()
        assert cs.font_name_cn.value == "宋体"
        assert cs.font_name_en.value == "Times New Roman"
        assert cs.font_size.value == "小四"
        assert cs.bold is cs.italic is cs.underline is False
        _clear_warning()

    def test_diff_from_run_no_diff_when_matching(self, doc, mock_warning):
        _set_warning(mock_warning)
        cs = CharacterStyle()
        run = doc.add_paragraph().add_run("t")
        run.font.bold = run.font.italic = run.font.underline = False
        run.font.size = Pt(12)
        run_set_font_name(run, "宋体")
        run.font.name = "Times New Roman"
        types = [d.diff_type for d in cs.diff_from_run(run)]
        assert "bold" not in types and "italic" not in types and "underline" not in types
        _clear_warning()

    @pytest.mark.parametrize("prop,expected_val,current_val", [
        ("bold", False, True), ("italic", False, True), ("underline", False, True),
    ])
    def test_diff_boolean_mismatch(self, doc, mock_warning, prop, expected_val, current_val):
        _set_warning(mock_warning)
        cs = CharacterStyle(**{prop: expected_val})
        run = doc.add_paragraph().add_run("t")
        setattr(run.font, prop, current_val)
        types = [d.diff_type for d in cs.diff_from_run(run)]
        assert prop in types
        _clear_warning()

    def test_diff_font_size_and_name_cn(self, doc, mock_warning):
        _set_warning(mock_warning)
        cs = CharacterStyle(font_size="小四", font_name_cn="宋体")
        run = doc.add_paragraph().add_run("测试")
        run.font.size = Pt(14)
        run_set_font_name(run, "黑体")
        types = [d.diff_type for d in cs.diff_from_run(run)]
        assert "font_size" in types
        assert next(d for d in cs.diff_from_run(run) if d.diff_type == "font_name_cn").current_value == "黑体"
        _clear_warning()

    def test_apply_to_run_fixes_bold(self, doc, mock_warning):
        _set_warning(mock_warning)
        run = doc.add_paragraph().add_run("t")
        run.font.bold = False
        result = CharacterStyle(bold=True).apply_to_run(run)
        assert run.font.bold is True
        assert any(d.diff_type == "bold" for d in result)
        _clear_warning()

    def test_to_string_filters_by_warning(self, mock_warning, mock_warning_off):
        diffs = [DIFFResult(diff_type="bold", current_value=True, expected_value=False)]
        _set_warning(mock_warning)
        assert "加粗错误" in CharacterStyle.to_string(diffs)
        _set_warning(mock_warning_off)
        assert CharacterStyle.to_string(diffs) == ""
        _clear_warning()



# ===========================================================================
# ParagraphStyle
# ===========================================================================

class TestParagraphStyle:
    def test_defaults(self, mock_warning):
        _set_warning(mock_warning)
        ps = ParagraphStyle()
        assert ps.alignment.value == "左对齐"
        assert ps.space_before.value == "0.5行"
        assert ps.line_spacing.value == "1.5倍"
        _clear_warning()

    def test_diff_none_returns_empty(self, mock_warning):
        _set_warning(mock_warning)
        assert ParagraphStyle().diff_from_paragraph(None) == []
        _clear_warning()

    def test_diff_detects_alignment(self, doc, mock_warning):
        _set_warning(mock_warning)
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assert "alignment" in [d.diff_type for d in ParagraphStyle(alignment="左对齐").diff_from_paragraph(p)]
        _clear_warning()

    def test_diff_builtin_style_name_match(self, doc, mock_warning):
        """BuiltInStyle('正文').rel_value='Normal' 与 'normal' 应视为一致。"""
        _set_warning(mock_warning)
        p = doc.add_paragraph()
        types = [d.diff_type for d in ParagraphStyle(builtin_style_name="正文").diff_from_paragraph(p)]
        assert "builtin_style_name" not in types  # 英文"normal" = 中文"正文"
        _clear_warning()

    def test_to_string_line_spacing_rule_key_bug(self, doc, mock_warning):
        """BUG: diff_type='line_spacing_rule' but warning dict key='line_spacingrule'."""
        _set_warning(mock_warning)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        diffs = ParagraphStyle(line_spacingrule="单倍行距").diff_from_paragraph(p)
        assert "line_spacing_rule" in [d.diff_type for d in diffs]
        result = ParagraphStyle.to_string(diffs)
        assert "行距选项" not in result  # filtered out due to key mismatch
        _clear_warning()

    def test_from_config(self, mock_warning):
        _set_warning(mock_warning)
        cfg = SimpleNamespace(alignment="居中对齐", space_before="1行")
        ps = ParagraphStyle.from_config(cfg)
        assert ps.alignment.value == "居中对齐"
        assert ps.space_after.value == "0.5行"  # default
        _clear_warning()



# ===========================================================================
# Additional coverage tests for check_format.py
# ===========================================================================


class TestDIFFResultWarningLevel:
    """Cover line 94: DIFFResult with level='warning' (string)"""

    def test_level_string(self):
        d = DIFFResult(level="warning")
        assert d.level == "warning"



class TestParagraphStyleApplyToParagraph:
    """Cover lines 315-358: ParagraphStyle.apply_to_paragraph with various indent types"""

    def test_apply_left_indent(self, doc, mock_warning):
        """apply_to_paragraph with left_indent diff (line 337-338)"""
        _set_warning(mock_warning)
        p = doc.add_paragraph()
        ps = ParagraphStyle(left_indent="2字符")
        diffs = ps.apply_to_paragraph(p)
        assert any(d.diff_type == "left_indent" for d in diffs)
        _clear_warning()

    def test_apply_right_indent(self, doc, mock_warning):
        """apply_to_paragraph with right_indent diff (line 339-341)"""
        _set_warning(mock_warning)
        p = doc.add_paragraph()
        ps = ParagraphStyle(right_indent="2字符")
        diffs = ps.apply_to_paragraph(p)
        assert any(d.diff_type == "right_indent" for d in diffs)
        _clear_warning()

    def test_apply_first_line_indent(self, doc, mock_warning):
        """apply_to_paragraph with first_line_indent diff (line 342-344)"""
        _set_warning(mock_warning)
        p = doc.add_paragraph()
        ps = ParagraphStyle(first_line_indent="2字符")
        diffs = ps.apply_to_paragraph(p)
        assert any(d.diff_type == "first_line_indent" for d in diffs)
        _clear_warning()

    def test_apply_unknown_diff_type_skipped(self, doc, mock_warning):
        """apply_to_paragraph with unknown diff_type -> logged and skipped (line 348-353)"""
        _set_warning(mock_warning)
        p = doc.add_paragraph()
        ps = ParagraphStyle()
        # Force an unknown diff_type by monkey-patching
        from wordformat.style.diff import DIFFResult
        original_diff = ps.diff_from_paragraph
        def fake_diff(paragraph):
            return [DIFFResult(diff_type="unknown_type", comment="test")]
        ps.diff_from_paragraph = fake_diff
        result = ps.apply_to_paragraph(p)
        # unknown_type should be skipped (continue), so result should be empty
        assert len(result) == 0
        ps.diff_from_paragraph = original_diff
        _clear_warning()

    def test_apply_builtin_style_name(self, doc, mock_warning):
        """apply_to_paragraph with builtin_style_name diff (line 345-347)"""
        _set_warning(mock_warning)
        p = doc.add_paragraph()
        ps = ParagraphStyle(builtin_style_name="Heading 1")
        diffs = ps.apply_to_paragraph(p)
        assert any(d.diff_type == "builtin_style_name" for d in diffs)
        _clear_warning()



class TestParagraphStyleFromConfigExtended:
    """Cover line 478: ParagraphStyle.from_config with real config"""

    def test_from_config_all_fields(self, mock_warning):
        """from_config with all fields specified (line 478+)"""
        _set_warning(mock_warning)
        cfg = SimpleNamespace(
            alignment="居中对齐",
            space_before="1行",
            space_after="0.5行",
            line_spacing="2倍",
            line_spacingrule="2倍行距",
            first_line_indent="2字符",
            left_indent="1字符",
            right_indent="1字符",
            builtin_style_name="正文",
        )
        ps = ParagraphStyle.from_config(cfg)
        assert ps.alignment.value == "居中对齐"
        assert ps.space_before.value == "1行"
        assert ps.space_after.value == "0.5行"
        assert ps.line_spacing == 2.0
        assert ps.line_spacingrule.value == "2倍行距"
        assert ps.first_line_indent.value == "2字符"
        assert ps.left_indent.value == "1字符"
        assert ps.right_indent.value == "1字符"
        assert ps.builtin_style_name.rel_value == "Normal"
        _clear_warning()

    def test_from_config_partial_fields(self, mock_warning):
        """from_config with only some fields -> defaults for rest"""
        _set_warning(mock_warning)
        cfg = SimpleNamespace(alignment="右对齐")
        ps = ParagraphStyle.from_config(cfg)
        assert ps.alignment.value == "右对齐"
        assert ps.space_before.value == "0.5行"  # default
        assert ps.builtin_style_name.rel_value == "Normal"  # default
        _clear_warning()



# ===========================================================================
# Additional coverage tests for check_format.py (CharacterStyle / ParagraphStyle)
# ===========================================================================


class TestCharacterStyleInitFromConfig:
    """Cover line 94: CharacterStyle.__init__ with style_checks_warning from get_config()"""

    def test_init_loads_warning_from_config(self, config_path):
        """When style_checks_warning is None, __init__ calls get_config() (line 94)"""
        from wordformat.config.loader import init_config, clear_config
        import wordformat.style.diff as m

        # Ensure style_checks_warning is None so __init__ triggers get_config()
        m.style_checks_warning = None
        init_config(config_path)
        try:
            cs = CharacterStyle()
            # After init, style_checks_warning should have been loaded
            assert m.style_checks_warning is not None
        finally:
            clear_config()
            m.style_checks_warning = None



class TestCharacterStyleDiffFontColor:
    """Cover line 157: CharacterStyle.diff_from_run font_color diff"""

    def test_diff_font_color_mismatch(self, doc, mock_warning):
        """font_color != current_color triggers diff (line 157)"""
        _set_warning(mock_warning)
        cs = CharacterStyle(font_color="红色")
        run = doc.add_paragraph().add_run("t")
        # Default color is black (0,0,0), red is (255,0,0)
        types = [d.diff_type for d in cs.diff_from_run(run)]
        assert "font_color" in types
        _clear_warning()



class TestCharacterStyleApplyToRunItalic:
    """Cover lines 207-208: CharacterStyle.apply_to_run 'italic' case"""

    def test_apply_to_run_fixes_italic(self, doc, mock_warning):
        """italic mismatch: run.italic=False, expected=True (lines 207-208)"""
        _set_warning(mock_warning)
        cs = CharacterStyle(italic=True)
        p = doc.add_paragraph()
        run = p.add_run("test")
        run.italic = False  # Wrong - should be True
        result = cs.apply_to_run(run)
        assert len(result) > 0
        assert run.italic is True
        assert any(d.diff_type == "italic" for d in result)
        _clear_warning()



class TestCharacterStyleApplyToRunUnderline:
    """Cover lines 212-213: CharacterStyle.apply_to_run 'underline' case"""

    def test_apply_to_run_fixes_underline(self, doc, mock_warning):
        """underline mismatch: run.underline=False, expected=True (lines 212-213)"""
        _set_warning(mock_warning)
        cs = CharacterStyle(underline=True)
        p = doc.add_paragraph()
        run = p.add_run("test")
        run.underline = False  # Wrong
        result = cs.apply_to_run(run)
        assert len(result) > 0
        assert run.underline is True
        assert any(d.diff_type == "underline" for d in result)
        _clear_warning()



class TestCharacterStyleApplyToRunFontColor:
    """Cover lines 218-219: CharacterStyle.apply_to_run 'font_color' case"""

    def test_apply_to_run_fixes_font_color(self, doc, mock_warning):
        """font_color mismatch triggers format (lines 218-219)"""
        _set_warning(mock_warning)
        cs = CharacterStyle(font_color="红色")
        p = doc.add_paragraph()
        run = p.add_run("test")
        # Default is black, so red should trigger a fix
        result = cs.apply_to_run(run)
        assert len(result) > 0
        assert any(d.diff_type == "font_color" for d in result)
        _clear_warning()



class TestCharacterStyleApplyToRunFontNameCn:
    """Cover lines 226-227: CharacterStyle.apply_to_run 'font_name_cn' case"""

    def test_apply_to_run_fixes_font_name_cn(self, doc, mock_warning):
        """font_name_cn mismatch triggers format (lines 226-227)"""
        _set_warning(mock_warning)
        cs = CharacterStyle(font_name_cn="黑体")
        p = doc.add_paragraph()
        run = p.add_run("测试")
        # Default CN font is 宋体, so 黑体 should trigger a fix
        run_set_font_name(run, "宋体")
        result = cs.apply_to_run(run)
        assert len(result) > 0
        assert any(d.diff_type == "font_name_cn" for d in result)
        _clear_warning()



class TestCharacterStyleToStringNone:
    """Cover line 244: CharacterStyle.to_string with style_checks_warning is None"""

    def test_to_string_warning_none(self):
        """style_checks_warning=None 时返回所有 diff 的标准格式文本。"""
        import wordformat.style.diff as m
        m.style_checks_warning = None
        diffs = [
            DIFFResult(diff_type="bold", current_value=True, expected_value=False),
            DIFFResult(diff_type="italic", current_value=True, expected_value=False),
        ]
        result = CharacterStyle.to_string(diffs, target="测试")
        assert "加粗错误" in result
        assert "斜体错误" in result
        assert "测试" in result



class TestCharacterStyleToStringBoldFilter:
    """Cover line 250: CharacterStyle.to_string with style_checks_warning.bold = True"""

    def test_to_string_bold_filtered_in(self, mock_warning):
        """warning.bold=True 时 bold diff 被包含。"""
        _set_warning(mock_warning)
        diffs = [DIFFResult(diff_type="bold", current_value=True, expected_value=False)]
        result = CharacterStyle.to_string(diffs, target="测试")
        assert "加粗错误" in result
        _clear_warning()



class TestCharacterStyleToStringVariousFilters:
    """Cover lines 252, 254, 256: CharacterStyle.to_string with italic/underline/font_size/font_color/font_name filters"""

    def test_to_string_italic_filtered(self, mock_warning):
        """warning.italic=True 时 italic diff 被包含。"""
        _set_warning(mock_warning)
        diffs = [DIFFResult(diff_type="italic", current_value=True, expected_value=False)]
        result = CharacterStyle.to_string(diffs)
        assert "斜体错误" in result
        _clear_warning()

    def test_to_string_font_size_filtered(self, mock_warning):
        """warning.font_size=True 时 font_size diff 被包含。"""
        _set_warning(mock_warning)
        diffs = [DIFFResult(diff_type="font_size", current_value=10.0, expected_value=12.0)]
        result = CharacterStyle.to_string(diffs)
        assert "字号错误" in result
        _clear_warning()

    def test_to_string_font_color_filtered(self, mock_warning):
        """warning.font_color=True 时 font_color diff 被包含。"""
        _set_warning(mock_warning)
        diffs = [DIFFResult(diff_type="font_color", current_value="红色", expected_value="黑色")]
        result = CharacterStyle.to_string(diffs)
        assert "字体颜色错误" in result
        _clear_warning()

    def test_to_string_font_name_filtered(self, mock_warning):
        """warning.font_name=True 时 font_name_cn diff 被包含。"""
        _set_warning(mock_warning)
        diffs = [DIFFResult(diff_type="font_name_cn", current_value="宋体", expected_value="黑体")]
        result = CharacterStyle.to_string(diffs)
        assert "中文字体错误" in result
        _clear_warning()



class TestParagraphStyleToStringNone:
    """Cover line 478: ParagraphStyle.to_string with style_checks_warning is None"""

    def test_to_string_warning_none(self):
        """style_checks_warning=None 时返回所有 diff 的标准格式文本。"""
        import wordformat.style.diff as m
        m.style_checks_warning = None
        diffs = [
            DIFFResult(diff_type="alignment", current_value="左对齐", expected_value="居中对齐"),
            DIFFResult(diff_type="space_before", current_value="0行", expected_value="0.5行"),
        ]
        result = ParagraphStyle.to_string(diffs)
        assert "对齐错误" in result
        assert "段前间距错误" in result

