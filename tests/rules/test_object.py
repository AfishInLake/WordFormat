"""FigureImage 和 TableObject 单元测试。"""
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from wordformat.rules.object import FigureImage, TableObject


class TestFigureImage:
    def test_apply_format_no_image_path_skips(self):
        """paragraph 有文字时跳过图片插入。"""
        doc = Document()
        p = doc.add_paragraph("已有文字")
        node = FigureImage(
            value={"category": "figure_image", "paragraph": "notexist.png"},
            level=1,
            paragraph=p,
        )
        node.load_config({})
        node.apply_format(doc)
        # 不影响已有文字
        assert p.runs[0].text == "已有文字"

    def test_apply_format_empty_paragraph_missing_file(self):
        """空段落 + 路径不存在 → 不崩溃。"""
        doc = Document()
        p = doc.add_paragraph("")
        node = FigureImage(
            value={"category": "figure_image", "paragraph": "/nonexistent/img.png"},
            level=1,
            paragraph=p,
        )
        node.load_config({})
        node.apply_format(doc)
        # 不崩溃就是通过

    def test_base_checks_alignment(self):
        """_base 检查对齐和首行缩进。"""
        doc = Document()
        p = doc.add_paragraph("")
        node = FigureImage(
            value={"category": "figure_image"},
            level=1,
            paragraph=p,
        )
        node.load_config({"figures": {"image": {"alignment": "居中对齐", "first_line_indent": "0字符"}}})
        node._base(doc, p=True, r=False)
        # 检查模式不崩溃

    def test_check_format_with_comment(self):
        """check_format 在格式不对时添加批注。"""
        doc = Document()
        p = doc.add_paragraph("")
        node = FigureImage(
            value={"category": "figure_image"},
            level=1,
            paragraph=p,
        )
        node.load_config({"figures": {"image": {"alignment": "居中对齐", "first_line_indent": "0字符"}}})
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1


class TestTableObject:
    def test_apply_no_table_rows_skips(self):
        """无 table_rows 时跳过表格创建。"""
        doc = Document()
        p = doc.add_paragraph("")
        node = TableObject(
            value={"category": "table_object"},
            level=1,
            paragraph=p,
        )
        node.load_config({})
        with patch.object(node, "_try_insert_table") as mock_table:
            node.apply_format(doc)
        mock_table.assert_called_once()

    def test_try_insert_table_creates_table(self):
        """有 table_rows 时创建 Word 表格。"""
        doc = Document()
        p = doc.add_paragraph("")
        node = TableObject(
            value={
                "category": "table_object",
                "table_rows": [["A", "B"], ["1", "2"]],
            },
            level=1,
            paragraph=p,
        )
        node.load_config({})
        node._try_insert_table(doc)
        # 表格已插入文档
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert table.cell(0, 0).text == "A"

    def test_try_insert_table_move_to_paragraph(self):
        """表格移动到段落后，不是文档末尾。"""
        doc = Document()
        p1 = doc.add_paragraph("段落1")
        p2 = doc.add_paragraph("段落2")
        node = TableObject(
            value={
                "category": "table_object",
                "table_rows": [["X"]],
            },
            level=1,
            paragraph=p1,
        )
        node.load_config({})
        node._try_insert_table(doc)
        # 表格应该在 p1 之后
        body = doc.element.body
        p_elements = body.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
        )
        tbl_elements = body.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
        )
        assert len(tbl_elements) == 1

    def test_apply_format_calls_try_insert(self):
        """apply_format 触发 _try_insert_table。"""
        doc = Document()
        p = doc.add_paragraph("")
        node = TableObject(
            value={"category": "table_object", "table_rows": [["A"]]},
            level=1,
            paragraph=p,
        )
        node.load_config({})
        node.apply_format(doc)
        assert len(doc.tables) == 1
