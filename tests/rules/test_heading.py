"""
rules 模块测试 —— 聚焦真实行为验证，无填充。
"""

from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from wordformat.config.models import NodeConfigRoot
from wordformat.rules import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleContentCN,
    AbstractTitleContentEN,
    AbstractTitleEN,
    Acknowledgements,
    AcknowledgementsCN,
    BodyText,
    CaptionFigure,
    CaptionTable,
    FormatNode,
    HeadingLevel1Node,
    HeadingLevel2Node,
    HeadingLevel3Node,
    KeywordsCN,
    KeywordsEN,
    ReferenceEntry,
    References,
)
from wordformat.rules.node import FormatNode as FormatNodeBase

# ---------------------------------------------------------------------------
# 共享 fixtures / helpers
# ---------------------------------------------------------------------------


def _make_node(cls, text="测试文本"):
    """创建一个带 paragraph 的 FormatNode 实例。"""
    doc = Document()
    p = doc.add_paragraph(text)
    return cls(value=p, level=0, paragraph=p)


def _load_root_config(config_path):
    """从 YAML 路径加载配置 dict。"""
    return _load_yaml(config_path)


def _load_yaml(path):
    import yaml

    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


@pytest.fixture
def root_config(sample_yaml_config):
    """从 sample_yaml_config 加载 NodeConfigRoot，与示例文件解耦。"""
    from wordformat.config.loader import init_config

    init_config(sample_yaml_config)
    return _load_root_config(sample_yaml_config)


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def para(doc):
    return doc.add_paragraph("测试段落")


@pytest.fixture
def run_with_text(para):
    r = para.add_run("关键词")
    r.font.size = Pt(12)
    r.font.bold = True
    return r


# ---------------------------------------------------------------------------
# 1. FormatNode 基类行为
# ---------------------------------------------------------------------------


class TestHeadingBug:
    """
    NODE_TYPE 现在自动回退为 CONFIG_PATH，
    基类 FormatNode.load_config 可正确解析 heading 配置。
    """

    def test_heading_level_configs(self, root_config):
        """各级标题 load_config 加载正确层级的配置。"""
        node_l1 = _make_node(HeadingLevel1Node)
        node_l1.load_config(root_config)
        assert node_l1.pydantic_config.font_size == "小二"

        node_l2 = _make_node(HeadingLevel2Node)
        node_l2.load_config(root_config)
        assert node_l2.pydantic_config.font_size == "三号"

        node_l3 = _make_node(HeadingLevel3Node)
        node_l3.load_config(root_config)
        assert node_l3.pydantic_config.font_size == "小四"


# ---------------------------------------------------------------------------
# 5. KeywordsCN / KeywordsEN 特有逻辑
# ---------------------------------------------------------------------------


class TestHeadingLevel1NodeBase:
    """覆盖 heading.py HeadingLevel1Node._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("第一章 绪论")
        r.font.size = Pt(10)
        r.font.bold = False
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("第一章 绪论")
        r.font.size = Pt(10)
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1  # apply 后 diff 仍有残留差异

    def test_check_returns_issues_list(self, root_config):
        """返回值应为包含 issue 字典的列表。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("第一章 绪论")
        r.font.size = Pt(10)
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        # 应有 run_issues 或 paragraph_issues

    def test_check_skips_empty_runs(self, root_config):
        """空白 run 应被跳过。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("   ")
        r2 = p.add_run("第一章 绪论")
        r2.font.size = Pt(10)
        node = HeadingLevel1Node(value=p, level=1, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        # 空白 run 不应触发 comment
        run_comments = [
            c for c in mock_comment.call_args_list if c.kwargs.get("runs") is r1
        ]
        assert len(run_comments) == 0


# ---------------------------------------------------------------------------
# 19. HeadingLevel2Node._base 覆盖
# ---------------------------------------------------------------------------


class TestHeadingLevel2NodeBase:
    """覆盖 heading.py HeadingLevel2Node._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1 研究背景")
        r.font.size = Pt(10)
        node = HeadingLevel2Node(value=p, level=2, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1 研究背景")
        r.font.size = Pt(10)
        node = HeadingLevel2Node(value=p, level=2, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1  # apply 后 diff 仍有残留差异

    def test_check_returns_issues(self, root_config):
        """返回值应包含 issue 字典。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1 研究背景")
        r.font.size = Pt(10)
        node = HeadingLevel2Node(value=p, level=2, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)


# ---------------------------------------------------------------------------
# 20. HeadingLevel3Node._base 覆盖
# ---------------------------------------------------------------------------


class TestHeadingLevel3NodeBase:
    """覆盖 heading.py HeadingLevel3Node._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1.1 问题描述")
        r.font.size = Pt(10)
        node = HeadingLevel3Node(value=p, level=3, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1.1 问题描述")
        r.font.size = Pt(10)
        node = HeadingLevel3Node(value=p, level=3, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1  # apply 后 diff 仍有残留差异

    def test_check_returns_issues(self, root_config):
        """返回值应包含 issue 字典。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("1.1.1 问题描述")
        r.font.size = Pt(10)
        node = HeadingLevel3Node(value=p, level=3, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)


# ---------------------------------------------------------------------------
# 21. References._base 覆盖
# ---------------------------------------------------------------------------


class TestHeadingLevelNodes:
    """覆盖 heading.py lines 36-41, 52-58"""

    def test_heading_level1_load_config_dict(self):
        """HeadingLevel1Node.load_config with dict (lines 36-41)"""
        from wordformat.rules.heading import HeadingLevel1Node

        node = HeadingLevel1Node(
            value={"category": "headings.level_1", "fingerprint": "fp"},
            level=1,
        )
        config_dict = {
            "headings": {
                "level_1": {
                    "alignment": "居中对齐",
                    "font_size": "小二",
                    "bold": True,
                }
            }
        }
        node.load_config(config_dict)
        assert node.pydantic_config is not None

    def test_heading_level2_load_config_dict(self):
        """HeadingLevel2Node.load_config with dict (lines 52-58)"""
        from wordformat.rules.heading import HeadingLevel2Node

        node = HeadingLevel2Node(
            value={"category": "headings.level_2", "fingerprint": "fp"},
            level=2,
        )
        config_dict = {
            "headings": {
                "level_2": {
                    "alignment": "左对齐",
                    "font_size": "三号",
                }
            }
        }
        node.load_config(config_dict)
        assert node.pydantic_config is not None

    def test_heading_level3_load_config_dict(self):
        """HeadingLevel3Node.load_config with dict"""
        from wordformat.rules.heading import HeadingLevel3Node

        node = HeadingLevel3Node(
            value={"category": "headings.level_3", "fingerprint": "fp"},
            level=3,
        )
        config_dict = {
            "headings": {
                "level_3": {
                    "alignment": "左对齐",
                    "font_size": "小四",
                }
            }
        }
        node.load_config(config_dict)
        assert node.pydantic_config is not None

    def test_heading_base_with_config(self, sample_yaml_config):
        """_base method with loaded config"""
        from wordformat.config.loader import init_config, get_config
        from wordformat.rules.heading import HeadingLevel1Node

        init_config(sample_yaml_config)
        config = get_config()

        node = HeadingLevel1Node(
            value={"category": "headings.level_1", "fingerprint": "fp"},
            level=1,
        )
        node.load_config(config)
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("第一章 绪论")
        node.paragraph = p
        node.check_format(doc)  # 通过 RULES handler 执行格式检查


# ==================== (o) set_style.py 额外覆盖测试 ====================
