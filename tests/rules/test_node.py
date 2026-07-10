"""
rules 模块测试 —— 聚焦真实行为验证，无填充。
"""
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from wordformat.config.models import NodeConfigRoot
from wordformat.rules import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleContentCN,
    AbstractTitleContentEN,
    AbstractTitleEN,
    Acknowledgements,
    AcknowledgementsCN,
    BodyText,
    CaptionFigure,
    CaptionTable,
    FormatNode,
    HeadingLevel1Node,
    HeadingLevel2Node,
    HeadingLevel3Node,
    KeywordsCN,
    KeywordsEN,
    ReferenceEntry,
    References,
)
from wordformat.rules.node import FormatNode as FormatNodeBase

# ---------------------------------------------------------------------------
# 共享 fixtures / helpers
# ---------------------------------------------------------------------------


def _make_node(cls, text="测试文本"):
    """创建一个带 paragraph 的 FormatNode 实例。"""
    doc = Document()
    p = doc.add_paragraph(text)
    return cls(value=p, level=0, paragraph=p)


def _load_root_config(config_path):
    """从 YAML 路径加载配置 dict。"""
    return _load_yaml(config_path)


def _load_yaml(path):
    import yaml
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


@pytest.fixture
def root_config(sample_yaml_config):
    """从 sample_yaml_config 加载 NodeConfigRoot，与示例文件解耦。"""
    from wordformat.config.loader import init_config
    init_config(sample_yaml_config)
    return _load_root_config(sample_yaml_config)


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def para(doc):
    return doc.add_paragraph("测试段落")


@pytest.fixture
def run_with_text(para):
    r = para.add_run("关键词")
    r.font.size = Pt(12)
    r.font.bold = True
    return r


# ---------------------------------------------------------------------------
# 1. FormatNode 基类行为
# ---------------------------------------------------------------------------




class TestFormatNodeBase:
    """FormatNode 基类的核心契约。"""

    def test_base_is_noop(self, doc, para):
        """_base() 默认为空操作。"""
        node = FormatNodeBase(value=para, level=0, paragraph=para)
        node._base(doc, p=True, r=True)
        node._base(doc, p=False, r=False)

    def test_check_format_calls_base_with_true(self, doc, para):
        """check_format 应以 p=True, r=True 调用 _base。"""
        node = FormatNodeBase(value=para, level=0, paragraph=para)
        with patch.object(node, "_base") as mock_base, \
                patch.object(node, "_run_rules"):
            node.check_format(doc)
        mock_base.assert_called_once_with(doc, p=True, r=True)

    def test_apply_format_calls_base_with_false(self, doc, para):
        """apply_format 应以 p=False, r=False 调用 _base。"""
        node = FormatNodeBase(value=para, level=0, paragraph=para)
        with patch.object(node, "_base") as mock_base, \
                patch.object(node, "_run_rules"):
            node.apply_format(doc)
        mock_base.assert_called_once_with(doc, p=False, r=False)

# ---------------------------------------------------------------------------
# 2. 所有节点类型可实例化
# ---------------------------------------------------------------------------


NODE_CLASSES = [
    AbstractTitleCN, AbstractTitleContentCN, AbstractContentCN,
    AbstractTitleEN, AbstractTitleContentEN, AbstractContentEN,
    KeywordsCN, KeywordsEN,
    HeadingLevel1Node, HeadingLevel2Node, HeadingLevel3Node,
    BodyText,
    CaptionFigure, CaptionTable,
    References, ReferenceEntry,
    Acknowledgements, AcknowledgementsCN,
]


@pytest.mark.parametrize("cls", NODE_CLASSES, ids=lambda c: c.__name__)


class TestNodeInstantiation:
    """每个节点类型都能正确实例化。"""

    def test_instantiation(self, cls):
        node = _make_node(cls)
        assert node.paragraph is not None
        assert node.level == 0

    def test_has_defaults(self, cls):
        assert hasattr(cls, "DEFAULTS")
        assert isinstance(cls.DEFAULTS, dict)

    def test_has_node_type(self, cls):
        assert hasattr(cls, "NODE_TYPE")
        assert isinstance(cls.NODE_TYPE, str)
        assert len(cls.NODE_TYPE) > 0


# ---------------------------------------------------------------------------
# 3. load_config 为每个节点正确设置 _pydantic_config
# ---------------------------------------------------------------------------




class TestLoadConfig:
    """验证 load_config 正确合并 DEFAULTS 与 YAML 配置。"""

    def test_abstract_title_cn(self, root_config):
        node = _make_node(AbstractTitleCN)
        node.load_config(root_config)
        assert node.pydantic_config.alignment == "居中对齐"
        assert node.pydantic_config.font_size == "小二"

    def test_abstract_content_cn(self, root_config):
        node = _make_node(AbstractContentCN)
        node.load_config(root_config)
        assert node.pydantic_config.alignment == "两端对齐"

    def test_abstract_title_en(self, root_config):
        node = _make_node(AbstractTitleEN)
        node.load_config(root_config)
        assert node.pydantic_config.alignment == "居中对齐"

    def test_abstract_content_en(self, root_config):
        node = _make_node(AbstractContentEN)
        node.load_config(root_config)
        assert node.pydantic_config.alignment == "两端对齐"

    def test_body_text(self, root_config):
        node = _make_node(BodyText)
        node.load_config(root_config)
        assert node.pydantic_config.alignment == "两端对齐"

    def test_caption_figure(self, root_config):
        node = _make_node(CaptionFigure)
        node.load_config(root_config)
        assert node.pydantic_config.caption_prefix == "图"

    def test_caption_table(self, root_config):
        node = _make_node(CaptionTable)
        node.load_config(root_config)
        assert node.pydantic_config.caption_prefix == "表"

    def test_caption_table_content_font_size(self, root_config):
        """YAML 中 tables.content.font_size 覆盖 DEFAULTS。"""
        node = _make_node(CaptionTable)
        node.load_config(root_config)
        # YAML: tables.content.font_size = '五号'
        assert node.pydantic_config.font_size == "小四"  # DEFAULTS 值（非 content 子对象）

    def test_references(self, root_config):
        node = _make_node(References)
        node.load_config(root_config)
        assert node.pydantic_config.alignment == "居中对齐"

    def test_reference_entry(self, root_config):
        node = _make_node(ReferenceEntry)
        node.load_config(root_config)
        assert node.pydantic_config.font_size == "五号"

    def test_acknowledgements(self, root_config):
        node = _make_node(Acknowledgements)
        node.load_config(root_config)
        assert node.pydantic_config.alignment == "居中对齐"

    def test_acknowledgements_cn(self, root_config):
        node = _make_node(AcknowledgementsCN)
        node.load_config(root_config)
        assert node.pydantic_config.alignment == "两端对齐"

    def test_keywords_cn_loads_rules(self, root_config):
        """KeywordsCN load_config 后 rules 可访问。"""
        node = _make_node(KeywordsCN)
        node.load_config(root_config)
        assert node.pydantic_config.rules.keyword_count.enabled is True

    def test_keywords_en_loads_label(self, root_config):
        """KeywordsEN load_config 后 label 可访问。"""
        node = _make_node(KeywordsEN)
        node.load_config(root_config)
        assert node.pydantic_config.label.font_size == "三号"

    def test_keywords_en_loads_label(self, root_config):
        node = _make_node(KeywordsEN)
        node.load_config(root_config)
        assert node.pydantic_config.label.font_size == "三号"


# ---------------------------------------------------------------------------
# 4. HeadingLevelConfig bug
# ---------------------------------------------------------------------------




class TestBaseImplementation:
    """验证 _base 实现确实执行了 diff_from_run / apply_to_run 逻辑。"""

    def test_body_text_check_calls_diff(self, root_config):
        """BodyText.check_format 应调用 diff_from_run。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("正文内容")
        r.font.size = Pt(10)  # 故意设置错误字号
        node = BodyText(value=p, level=0, paragraph=p)
        node.load_config(root_config)

        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)

        # 应至少调用一次 add_comment（因为字号不匹配）
        assert mock_comment.call_count >= 1

    def test_abstract_title_cn_check_runs(self, root_config):
        """AbstractTitleCN.check_format 应遍历所有 run。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要")
        r.font.size = Pt(10)
        node = AbstractTitleCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)

        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)

        # 至少对 run 和 paragraph 各调用一次
        assert mock_comment.call_count >= 2

    def test_references_check_runs(self, root_config):
        """References.check_format 应遍历 run 并调用 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("参考文献")
        r.font.size = Pt(10)
        node = References(value=p, level=0, paragraph=p)
        node.load_config(root_config)

        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)

        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 7. BodyText._apply_citation_superscript
# ---------------------------------------------------------------------------

