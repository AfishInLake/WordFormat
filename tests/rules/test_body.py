"""
rules 模块测试 —— 聚焦真实行为验证，无填充。
"""
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from wordformat.config.models import NodeConfigRoot
from wordformat.rules import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleContentCN,
    AbstractTitleContentEN,
    AbstractTitleEN,
    Acknowledgements,
    AcknowledgementsCN,
    BodyText,
    CaptionFigure,
    CaptionTable,
    FormatNode,
    HeadingLevel1Node,
    HeadingLevel2Node,
    HeadingLevel3Node,
    KeywordsCN,
    KeywordsEN,
    ReferenceEntry,
    References,
)
from wordformat.rules.node import FormatNode as FormatNodeBase

# ---------------------------------------------------------------------------
# 共享 fixtures / helpers
# ---------------------------------------------------------------------------


def _make_node(cls, text="测试文本"):
    """创建一个带 paragraph 的 FormatNode 实例。"""
    doc = Document()
    p = doc.add_paragraph(text)
    return cls(value=p, level=0, paragraph=p)


def _load_root_config(config_path):
    """从 YAML 路径加载配置 dict。"""
    return _load_yaml(config_path)


def _load_yaml(path):
    import yaml
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


@pytest.fixture
def root_config(sample_yaml_config):
    """从 sample_yaml_config 加载 NodeConfigRoot，与示例文件解耦。"""
    from wordformat.config.loader import init_config
    init_config(sample_yaml_config)
    return _load_root_config(sample_yaml_config)


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def para(doc):
    return doc.add_paragraph("测试段落")


@pytest.fixture
def run_with_text(para):
    r = para.add_run("关键词")
    r.font.size = Pt(12)
    r.font.bold = True
    return r


# ---------------------------------------------------------------------------
# 1. FormatNode 基类行为
# ---------------------------------------------------------------------------




class TestBodyTextCitationSuperscript:
    """测试 BodyText.apply_format 中的引用上标自动设置。"""

    @staticmethod
    def _get_vertAlign(run):
        """返回 run 的 w:vertAlign 值，无则返回 None。"""
        rPr = run._element.find(qn("w:rPr"))
        if rPr is None:
            return None
        va = rPr.find(qn("w:vertAlign"))
        return va.get(qn("w:val")) if va is not None else None

    def test_single_citation_gets_superscript(self):
        """单引用 [1] 应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("这是一篇论文[1]的引用。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp1"},
                        level=0, paragraph=p)
        # 直接调用 apply_format（跳过 load_config，不依赖完整配置）
        node._apply_citation_superscript()
        # "[1]" 应在独立的上标 run 中
        runs = p.runs
        superscript_texts = [
            r.text for r in runs if self._get_vertAlign(r) == "superscript"
        ]
        assert "[1]" in superscript_texts

    def test_multiple_citations(self):
        """多个引用 [1] 和 [2,3] 都应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("参见文献[1]和[2,3]的讨论。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp2"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        superscript_texts = [
            r.text for r in p.runs if self._get_vertAlign(r) == "superscript"
        ]
        assert "[1]" in superscript_texts
        assert "[2,3]" in superscript_texts

    def test_range_citation(self):
        """范围引用 [1-3] 应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("文献[1-3]提供了详细分析。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp3"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        superscript_texts = [
            r.text for r in p.runs if self._get_vertAlign(r) == "superscript"
        ]
        assert "[1-3]" in superscript_texts

    def test_chinese_comma_citation(self):
        """中文逗号分隔的引用 [1，2] 应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("见[1，2，3]的研究。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp4"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        superscript_texts = [
            r.text for r in p.runs if self._get_vertAlign(r) == "superscript"
        ]
        assert any("[1，2，3]" in t for t in superscript_texts)

    def test_no_citation_leaves_runs_unchanged(self):
        """无引用的段落应保持原样。"""
        doc = Document()
        p = doc.add_paragraph("这是一段没有引用的正文。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp5"},
                        level=0, paragraph=p)
        original_text = p.text
        node._apply_citation_superscript()
        # 文本不变，且无上标 run
        assert p.text == original_text
        assert all(self._get_vertAlign(r) is None for r in p.runs)

    def test_non_citation_brackets_not_affected(self):
        """非数字方括号如 [注] 不应被设为上标。"""
        doc = Document()
        p = doc.add_paragraph("这是一个[注]释说明。")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp6"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        assert all(self._get_vertAlign(r) is None for r in p.runs)

    def test_citation_split_across_runs(self):
        """引用跨 run 时先分割再设上标。"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("前面文字[1")
        p.add_run("2]后面文字")
        node = BodyText(value={"category": "body_text", "fingerprint": "fp7"},
                        level=0, paragraph=p)
        node._apply_citation_superscript()
        superscript_texts = [
            r.text for r in p.runs if self._get_vertAlign(r) == "superscript"
        ]
        # 分割后 [1 和 2] 分别在两个 run 中，但都是上标
        assert "[1" in superscript_texts
        assert "2]" in superscript_texts


# ---------------------------------------------------------------------------
# 8. AbstractTitleCN._base 完整 diff/apply 覆盖
# ---------------------------------------------------------------------------

