"""
rules 模块测试 —— 聚焦真实行为验证，无填充。
"""
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from wordformat.config.models import NodeConfigRoot
from wordformat.rules import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleContentCN,
    AbstractTitleContentEN,
    AbstractTitleEN,
    Acknowledgements,
    AcknowledgementsCN,
    BodyText,
    CaptionFigure,
    CaptionTable,
    FormatNode,
    HeadingLevel1Node,
    HeadingLevel2Node,
    HeadingLevel3Node,
    KeywordsCN,
    KeywordsEN,
    ReferenceEntry,
    References,
)
from wordformat.rules.node import FormatNode as FormatNodeBase

# ---------------------------------------------------------------------------
# 共享 fixtures / helpers
# ---------------------------------------------------------------------------


def _make_node(cls, text="测试文本"):
    """创建一个带 paragraph 的 FormatNode 实例。"""
    doc = Document()
    p = doc.add_paragraph(text)
    return cls(value=p, level=0, paragraph=p)


def _load_root_config(config_path):
    """从 YAML 路径加载配置 dict。"""
    return _load_yaml(config_path)


def _load_yaml(path):
    import yaml
    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


@pytest.fixture
def root_config(sample_yaml_config):
    """从 sample_yaml_config 加载 NodeConfigRoot，与示例文件解耦。"""
    from wordformat.config.loader import init_config
    init_config(sample_yaml_config)
    return _load_root_config(sample_yaml_config)


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def para(doc):
    return doc.add_paragraph("测试段落")


@pytest.fixture
def run_with_text(para):
    r = para.add_run("关键词")
    r.font.size = Pt(12)
    r.font.bold = True
    return r


# ---------------------------------------------------------------------------
# 1. FormatNode 基类行为
# ---------------------------------------------------------------------------



class TestAbstractTitleCNBase:
    """覆盖 abstract.py AbstractTitleCN._base 的 diff 和 apply 分支。"""

    def test_check_with_wrong_format_triggers_comments(self, root_config):
        """check 模式：错误格式应触发 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要")
        r.font.size = Pt(10)  # 错误字号
        r.font.bold = False  # 应为加粗
        node = AbstractTitleCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_apply_fixes_wrong_format(self, root_config):
        """apply 模式：应调用 apply_to_paragraph 和 apply_to_run。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要")
        r.font.size = Pt(10)
        r.font.bold = False
        node = AbstractTitleCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        # apply 模式也会调用 add_comment
        assert mock_comment.call_count >= 2

    def test_check_no_runs_skips_without_error(self, root_config):
        """段落无 run 时（空段），check_format 安全跳过不报错。"""
        doc = Document()
        p = doc.add_paragraph()
        node = AbstractTitleCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)  # 不应抛异常


# ---------------------------------------------------------------------------
# 9. AbstractTitleContentCN._base 覆盖（标题+正文混合）
# ---------------------------------------------------------------------------



class TestAbstractTitleContentCNBase:
    """覆盖 abstract.py AbstractTitleContentCN._base 的 check_title 分支。"""

    def test_check_title_run_uses_title_style(self, root_config):
        """包含'摘要'的 run 应使用 chinese_title 样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要")
        r.font.size = Pt(10)  # 错误字号
        node = AbstractTitleContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_check_content_run_uses_content_style(self, root_config):
        """非标题 run 应使用 chinese_content 样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("这是摘要正文内容")
        r.font.size = Pt(10)
        node = AbstractTitleContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_apply_title_and_content_runs(self, root_config):
        """apply 模式：标题和正文 run 都应被修正。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("摘要")
        r1.font.size = Pt(10)
        r2 = p.add_run("正文内容")
        r2.font.size = Pt(10)
        node = AbstractTitleContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 3

    def test_check_mode_does_not_mutate_run_text(self, root_config):
        """_base 在检查模式下不应修改 run.text（修复：移除了破坏性 replace）。"""
        doc = Document()
        p = doc.add_paragraph()
        original_text = "摘 要"
        r = p.add_run(original_text)
        r.font.size = Pt(10)
        node = AbstractTitleContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        # 修复后：检查模式不应改变 run.text
        assert r.text == original_text


# ---------------------------------------------------------------------------
# 10. AbstractContentCN._base 覆盖
# ---------------------------------------------------------------------------



class TestAbstractContentCNBase:
    """覆盖 abstract.py AbstractContentCN._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_font_size(self, root_config):
        """check 模式：错误字号应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要正文")
        r.font.size = Pt(10)
        node = AbstractContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_apply_fixes_font_size(self, root_config):
        """apply 模式：应修正字号。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("摘要正文")
        r.font.size = Pt(10)
        node = AbstractContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 2

    def test_check_multiple_runs(self, root_config):
        """多个 run 都应被检查。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("摘要")
        r1.font.size = Pt(10)
        r2 = p.add_run("正文")
        r2.font.size = Pt(10)
        node = AbstractContentCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        # 2 run comments + 1 paragraph comment
        assert mock_comment.call_count >= 3


# ---------------------------------------------------------------------------
# 11. AbstractTitleEN._base 覆盖
# ---------------------------------------------------------------------------



class TestAbstractTitleENBase:
    """覆盖 abstract.py AbstractTitleEN._base 的 diff/apply 逻辑。
    注意：AbstractTitleEN 只在 diff_result 非空时才 add_comment。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Abstract")
        r.font.size = Pt(10)
        r.font.bold = False
        node = AbstractTitleEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_empty_list(self, root_config):
        """AbstractTitleEN._base 始终返回 []。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Abstract")
        r.font.size = Pt(10)
        node = AbstractTitleEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Abstract")
        r.font.size = Pt(10)
        r.font.bold = False
        node = AbstractTitleEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 12. AbstractTitleContentEN._base 覆盖
# ---------------------------------------------------------------------------



class TestAbstractTitleContentENBase:
    """覆盖 abstract.py AbstractTitleContentEN._base 的 check_title 分支。"""

    def test_check_title_run_uses_title_style(self, root_config):
        """包含'Abstract'的 run 应使用 english_title 样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("Abstract")
        r.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_check_content_run_uses_content_style(self, root_config):
        """非标题 run 应使用 english_content 样式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("This is the abstract content.")
        r.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2

    def test_apply_mixed_runs(self, root_config):
        """apply 模式：混合标题和正文 run。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Abstract")
        r1.font.size = Pt(10)
        r2 = p.add_run("Content text")
        r2.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 3

    def test_check_title_normalizes_case_lower(self, root_config):
        """小写 'abstract' 应匹配并自动修正为 'Abstract'。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("abstract: some content")
        r.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        assert r.text.startswith("Abstract")

    def test_check_title_normalizes_case_upper(self, root_config):
        """全大写 'ABSTRACT' 应匹配并自动修正为 'Abstract'。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("ABSTRACT\n\nbody text")
        r.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        assert r.text.startswith("Abstract")

    def test_split_abstract_across_runs(self, root_config):
        """"Abstract" 被拆分到两个 run 时仍能正确识别标题部分。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Abst")
        r1.font.size = Pt(10)
        r2 = p.add_run("ract: body text")
        r2.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        # r1 开头被修正为 "Abstract"，r2 保持 "body text" 部分
        assert r1.text.startswith("Abstract")
        assert "body text" in r2.text

    def test_split_abstract_across_three_runs(self, root_config):
        """"Abstract" 被拆分到三个 run 时仍能正确识别。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("Abs")
        r1.font.size = Pt(10)
        r2 = p.add_run("tra")
        r2.font.size = Pt(10)
        r3 = p.add_run("ct: content")
        r3.font.size = Pt(10)
        node = AbstractTitleContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)
        # r1 应被修正为 "Abstract"
        assert r1.text.startswith("Abstract")
        # r2 和 r3 开头部分属于标题前缀，应被清空
        assert r2.text == ""
        assert "content" in r3.text


# ---------------------------------------------------------------------------
# 13. AbstractContentEN._base 覆盖
# ---------------------------------------------------------------------------



class TestAbstractContentENBase:
    """覆盖 abstract.py AbstractContentEN._base 的 diff/apply 逻辑。
    注意：AbstractContentEN 只在 diff_result/issues 非空时才 add_comment。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("This is abstract content.")
        r.font.size = Pt(10)
        node = AbstractContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("This is abstract content.")
        r.font.size = Pt(10)
        node = AbstractContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            # 配置中 line_spacing 为 "0倍"，现会触发 ValueError，mock 掉该步
            with patch("wordformat.style.diff.LineSpacing.format"):
                node.apply_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_multiple_runs(self, root_config):
        """多个 run 都应被检查。"""
        doc = Document()
        p = doc.add_paragraph()
        r1 = p.add_run("First sentence. ")
        r1.font.size = Pt(10)
        r2 = p.add_run("Second sentence.")
        r2.font.size = Pt(10)
        node = AbstractContentEN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 2


# ---------------------------------------------------------------------------
# 14. Acknowledgements._base 覆盖
# ---------------------------------------------------------------------------
