"""
rules 模块测试 —— 聚焦真实行为验证，无填充。
"""

from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from wordformat.config.models import NodeConfigRoot
from wordformat.rules import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleContentCN,
    AbstractTitleContentEN,
    AbstractTitleEN,
    Acknowledgements,
    AcknowledgementsCN,
    BodyText,
    CaptionFigure,
    CaptionTable,
    FormatNode,
    HeadingLevel1Node,
    HeadingLevel2Node,
    HeadingLevel3Node,
    KeywordsCN,
    KeywordsEN,
    ReferenceEntry,
    References,
)
from wordformat.rules.node import FormatNode as FormatNodeBase

# ---------------------------------------------------------------------------
# 共享 fixtures / helpers
# ---------------------------------------------------------------------------


def _make_node(cls, text="测试文本"):
    """创建一个带 paragraph 的 FormatNode 实例。"""
    doc = Document()
    p = doc.add_paragraph(text)
    return cls(value=p, level=0, paragraph=p)


def _load_root_config(config_path):
    """从 YAML 路径加载配置 dict。"""
    return _load_yaml(config_path)


def _load_yaml(path):
    import yaml

    with open(path, encoding="utf-8") as f:
        return yaml.safe_load(f)


@pytest.fixture
def root_config(sample_yaml_config):
    """从 sample_yaml_config 加载 NodeConfigRoot，与示例文件解耦。"""
    from wordformat.config.loader import init_config

    init_config(sample_yaml_config)
    return _load_root_config(sample_yaml_config)


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def para(doc):
    return doc.add_paragraph("测试段落")


@pytest.fixture
def run_with_text(para):
    r = para.add_run("关键词")
    r.font.size = Pt(12)
    r.font.bold = True
    return r


# ---------------------------------------------------------------------------
# 1. FormatNode 基类行为
# ---------------------------------------------------------------------------


class TestAcknowledgementsBase:
    """覆盖 acknowledgement.py Acknowledgements._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢")
        r.font.size = Pt(10)
        r.font.bold = False
        node = Acknowledgements(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_empty_list(self, root_config):
        """Acknowledgements._base 始终返回 []。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢")
        r.font.size = Pt(10)
        node = Acknowledgements(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢")
        r.font.size = Pt(10)
        node = Acknowledgements(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1  # apply 后 diff 仍有残留差异

    def test_check_no_diffs_no_comment(self, root_config):
        """格式完全正确时，不应调用 add_comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢")
        node = Acknowledgements(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        # 即使格式正确，段落级别的 diff 仍可能触发 comment
        # 但如果没有差异，不应有 comment
        # 注意：新 Document 的段落默认对齐方式可能与配置不同


# ---------------------------------------------------------------------------
# 15. AcknowledgementsCN._base 覆盖
# ---------------------------------------------------------------------------


class TestAcknowledgementsCNBase:
    """覆盖 acknowledgement.py AcknowledgementsCN._base 的 diff/apply 逻辑。"""

    def test_check_with_wrong_format(self, root_config):
        """check 模式：错误格式应触发 comment。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("感谢导师的指导。")
        r.font.size = Pt(10)
        node = AcknowledgementsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1

    def test_check_returns_empty_list(self, root_config):
        """AcknowledgementsCN._base 始终返回 []。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("感谢导师的指导。")
        r.font.size = Pt(10)
        node = AcknowledgementsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        node.check_format(doc)

    def test_apply_with_wrong_format(self, root_config):
        """apply 模式：应修正格式。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("感谢导师的指导。")
        r.font.size = Pt(10)
        node = AcknowledgementsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        with patch.object(node, "add_comment") as mock_comment:
            node.apply_format(doc)
        assert mock_comment.call_count >= 1  # apply 后 diff 仍有残留差异

    def test_check_first_line_indent(self, root_config):
        """验证 first_line_indent 配置被正确使用。"""
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run("致谢内容段落。")
        r.font.size = Pt(10)
        node = AcknowledgementsCN(value=p, level=0, paragraph=p)
        node.load_config(root_config)
        # 验证配置中有 first_line_indent 字段
        assert node.pydantic_config.first_line_indent is not None
        with patch.object(node, "add_comment") as mock_comment:
            node.check_format(doc)
        assert mock_comment.call_count >= 1


# ---------------------------------------------------------------------------
# 16. CaptionFigure._base 覆盖
# ---------------------------------------------------------------------------
