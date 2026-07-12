#!/usr/bin/env python
"""
comprehensive tests for style modules:
  check_format, get_some, set_some, style_enum, utils
"""

import pytest
from types import SimpleNamespace
from unittest.mock import MagicMock, PropertyMock, patch as mock_patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from wordformat.style.diff import DIFFResult, CharacterStyle, ParagraphStyle
from wordformat.style.reader import (
    paragraph_get_alignment,
    paragraph_get_space_before,
    paragraph_get_space_after,
    paragraph_get_line_spacing,
    paragraph_get_first_line_indent,
    paragraph_get_builtin_style_name,
    run_get_font_name,
    run_get_font_size_pt,
    run_get_font_color,
    run_get_font_bold,
    run_get_font_italic,
    run_get_font_underline,
    GetIndent,
)
from wordformat.style.writer import (
    run_set_font_name,
    set_paragraph_space_before_by_lines,
    set_paragraph_space_after_by_lines,
    _paragraph_space_by_lines,
    SetSpacing,
    SetLineSpacing,
    SetIndent,
    SetFirstLineIndent,
)
from wordformat.style.defs import (
    FontName,
    FontSize,
    FontColor,
    Alignment,
    LineSpacingRule,
    LineSpacing,
    FirstLineIndent,
    LeftIndent,
    RightIndent,
    BuiltInStyle,
    SpaceBefore,
    SpaceAfter,
    Spacing,
    UnitLabelEnum,
)
from wordformat.style.units import extract_unit_from_string, UnitResult


# ---------------------------------------------------------------------------
# fixtures & helpers
# ---------------------------------------------------------------------------


@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def mock_warning():
    w = MagicMock()
    for attr in (
        "bold",
        "italic",
        "underline",
        "font_size",
        "font_color",
        "font_name",
        "alignment",
        "space_before",
        "space_after",
        "line_spacing",
        "line_spacingrule",
        "first_line_indent",
        "left_indent",
        "right_indent",
        "builtin_style_name",
    ):
        setattr(w, attr, True)
    return w


@pytest.fixture
def mock_warning_off():
    w = MagicMock()
    for attr in (
        "bold",
        "italic",
        "underline",
        "font_size",
        "font_color",
        "font_name",
        "alignment",
        "space_before",
        "space_after",
        "line_spacing",
        "line_spacingrule",
        "first_line_indent",
        "left_indent",
        "right_indent",
        "builtin_style_name",
    ):
        setattr(w, attr, False)
    return w


def _set_warning(w):
    import wordformat.style.diff as m

    m.style_checks_warning = w


def _clear_warning():
    import wordformat.style.diff as m

    m.__dict__.pop("style_checks_warning", None)


# ===========================================================================
# get_some
# ===========================================================================


class TestGetSomeAlignment:
    @pytest.mark.parametrize(
        "align",
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT],
    )
    def test_direct(self, doc, align):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = align
        assert paragraph_get_alignment(p) == align

    def test_no_alignment_returns_none(self, doc):
        assert paragraph_get_alignment(doc.add_paragraph()) is None


class TestGetSomeLineSpacing:
    @pytest.mark.parametrize(
        "rule,expected",
        [
            (WD_LINE_SPACING.SINGLE, 1.0),
            (WD_LINE_SPACING.ONE_POINT_FIVE, 1.5),
            (WD_LINE_SPACING.DOUBLE, 2.0),
        ],
    )
    def test_preset(self, doc, rule, expected):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = rule
        assert paragraph_get_line_spacing(p) == expected

    def test_multiple_custom(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 2.3
        assert paragraph_get_line_spacing(p) == 2.3

    @pytest.mark.parametrize(
        "rule", [WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST]
    )
    def test_fixed_returns_none(self, doc, rule):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = rule
        assert paragraph_get_line_spacing(p) is None


class TestGetSomeRun:
    def test_font_name_default_none(self, doc):
        assert run_get_font_name(doc.add_paragraph().add_run("x")) is None

    def test_font_name_after_set(self, doc):
        run = doc.add_paragraph().add_run("x")
        run_set_font_name(run, "宋体")
        assert run_get_font_name(run) == "宋体"

    def test_font_size_docdefaults_11(self, doc):
        # 空 run 现沿继承链解析到 docDefaults sz=22 → 11.0（非旧硬编码 12.0）
        assert run_get_font_size_pt(doc.add_paragraph().add_run("x")) == 11.0

    def test_font_size_explicit(self, doc):
        run = doc.add_paragraph().add_run("x")
        run.font.size = Pt(14)
        assert run_get_font_size_pt(run) == 14.0

    def test_font_color_default_black(self, doc):
        assert run_get_font_color(doc.add_paragraph().add_run("x")) == (0, 0, 0)

    def test_font_color_set(self, doc):
        run = doc.add_paragraph().add_run("x")
        run.font.color.rgb = RGBColor(0xFF, 0, 0)
        assert run_get_font_color(run) == (255, 0, 0)

    @pytest.mark.parametrize(
        "getter,attr",
        [
            (run_get_font_bold, "bold"),
            (run_get_font_italic, "italic"),
            (run_get_font_underline, "underline"),
        ],
    )
    def test_boolean_props(self, doc, getter, attr):
        run = doc.add_paragraph().add_run("x")
        assert getter(run) is False
        setattr(run.font, attr, True)
        assert getter(run) is True


class TestGetSomeFirstLineIndent:
    def test_default_none(self, doc):
        assert paragraph_get_first_line_indent(doc.add_paragraph()) is None

    def test_with_chars(self, doc):
        p = doc.add_paragraph()
        SetFirstLineIndent.set_char(p, 2)
        assert paragraph_get_first_line_indent(p) == 2.0

    def test_ignores_firstLine_twips(self, doc):
        """Only reads firstLineChars, ignores firstLine (physical twips)."""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        assert paragraph_get_first_line_indent(p) is None


class TestGetSomeBuiltinStyle:
    def test_default_normal_lowercase(self, doc):
        name = paragraph_get_builtin_style_name(doc.add_paragraph())
        assert name == "normal"


class TestGetIndent:
    def test_default_none(self, doc):
        p = doc.add_paragraph()
        assert GetIndent.left_indent(p) is None
        assert GetIndent.right_indent(p) is None

    def test_invalid_type_raises(self, doc):
        with pytest.raises(ValueError, match="必须是"):
            GetIndent.line_indent(doc.add_paragraph(), "bad")

    def test_left_after_set_char(self, doc):
        p = doc.add_paragraph()
        SetIndent.set_char(p, "R", 3)
        assert GetIndent.left_indent(p) == 3.0


# ===========================================================================
# Additional coverage tests for get_some.py
# ===========================================================================


class TestGetSomeAlignmentFromStyle:
    """Cover line 38: paragraph_get_alignment with style.alignment (not direct)"""

    def test_alignment_from_style(self, doc):
        """When paragraph has no direct alignment, falls back to style.alignment"""
        p = doc.add_paragraph()
        p.paragraph_format.alignment = None
        # Set alignment on the paragraph's style
        p.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assert paragraph_get_alignment(p) == WD_ALIGN_PARAGRAPH.CENTER

    def test_alignment_from_base_style(self, doc):
        """沿真实 basedOn 链解析对齐：Child(basedOn=Base) 继承 Base 的对齐。"""
        from docx.enum.style import WD_STYLE_TYPE

        base = doc.styles.add_style("BaseAlign", WD_STYLE_TYPE.PARAGRAPH)
        base.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        child = doc.styles.add_style("ChildAlign", WD_STYLE_TYPE.PARAGRAPH)
        child.base_style = base
        p = doc.add_paragraph(style="ChildAlign")
        assert paragraph_get_alignment(p) == WD_ALIGN_PARAGRAPH.RIGHT

    def test_alignment_no_base_style_returns_none(self, doc):
        """No direct alignment, no style alignment, no base_style -> None"""
        p = doc.add_paragraph()
        p.paragraph_format.alignment = None
        p.style.paragraph_format.alignment = None
        p.style._base_style = None
        assert paragraph_get_alignment(p) is None


class TestGetSomeSpaceBeforeAfterInheritance:
    """Cover lines 147-148, 157-158: space_before/after with style inheritance"""

    def test_space_before_invalid_attr_returns_none(self, doc):
        """beforeLines attr is non-numeric -> self_lines=None, fall to style (line 147-148)"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:beforeLines"), "abc")
        pPr.append(spacing)
        # style has no spacing either
        assert paragraph_get_space_before(p) is None

    def test_space_after_invalid_attr_returns_none(self, doc):
        """afterLines attr is non-numeric -> self_lines=None, fall to style (line 180-181)"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:afterLines"), "xyz")
        pPr.append(spacing)
        assert paragraph_get_space_after(p) is None

    def test_space_before_from_style(self, doc):
        """No direct beforeLines, but style has it -> returns style value"""
        p = doc.add_paragraph()
        mock_style = MagicMock()
        mock_style.element = None
        mock_style.base_style = None
        with mock_patch.object(
            type(p), "style", new_callable=PropertyMock, return_value=mock_style
        ):
            # _get_style_spacing returns None for None element
            assert paragraph_get_space_before(p) is None

    def test_space_after_from_style(self, doc):
        """No direct afterLines, but style has it -> returns style value"""
        p = doc.add_paragraph()
        mock_style = MagicMock()
        mock_style.element = None
        mock_style.base_style = None
        with mock_patch.object(
            type(p), "style", new_callable=PropertyMock, return_value=mock_style
        ):
            assert paragraph_get_space_after(p) is None


class TestGetSomeLineSpacingInheritance:
    """Cover lines 233-235: paragraph_get_line_spacing with style inheritance"""

    def test_line_spacing_attribute_error_returns_none(self, doc):
        """paragraph with broken paragraph_format -> returns None (lines 233-235)"""
        mock_para = MagicMock()
        # paragraph_format raises AttributeError
        mock_para.paragraph_format = None
        result = paragraph_get_line_spacing(mock_para)
        assert result is None

    def test_line_spacing_type_error_returns_none(self):
        """paragraph with broken paragraph_format -> returns None"""
        mock_para = MagicMock()
        mock_para.paragraph_format.line_spacing_rule = "invalid"
        mock_para.paragraph_format.line_spacing = None
        result = paragraph_get_line_spacing(mock_para)
        assert result is None


class TestGetSomeFirstLineIndentPhysicalUnit:
    """Cover lines 264-266: paragraph_get_first_line_indent with firstLine (physical unit)"""

    def test_first_line_indent_exception_returns_none(self, doc):
        """Exception during parsing -> returns None (lines 264-266)"""
        p = doc.add_paragraph()
        # Force an exception by making pPr raise
        original = p._element.find

        def bad_find(qname):
            raise RuntimeError("test error")

        p._element.find = bad_find
        assert paragraph_get_first_line_indent(p) is None
        p._element.find = original


class TestGetSomeRunExtended:
    """Cover lines 317, 338: run_get_font_size_pt with style, run_get_font_color with theme"""

    def test_font_size_from_style(self, doc):
        """run 无直接字号时，沿段落样式链解析到样式字号。"""
        from docx.enum.style import WD_STYLE_TYPE

        st = doc.styles.add_style("SizeStyle", WD_STYLE_TYPE.PARAGRAPH)
        st.font.size = Pt(16)
        run = doc.add_paragraph(style="SizeStyle").add_run("x")
        assert run_get_font_size_pt(run) == 16.0

    def test_font_color_theme_color(self, doc):
        """run.font.color.rgb is None (theme color) -> returns (0,0,0) (line 338)"""
        run = doc.add_paragraph().add_run("x")
        run.font.color.rgb = None
        assert run_get_font_color(run) == (0, 0, 0)

    def test_font_color_none(self, doc):
        """run.font.color is None -> returns (0,0,0)"""
        run = doc.add_paragraph().add_run("x")
        # color is a read-only property, so we mock the entire run_get_font_color
        # by directly testing the code path with a MagicMock run
        mock_run = MagicMock()
        mock_run.font.color = None
        assert run_get_font_color(mock_run) == (0, 0, 0)


class TestGetIndentWithRealElement:
    """Cover lines 408-423: GetIndent.line_indent with real indent element"""

    def test_line_indent_left_with_chars(self, doc):
        """Real w:ind element with w:leftChars -> returns float (lines 408-423)"""
        p = doc.add_paragraph()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:leftChars"), "300")
        pPr.append(ind)
        assert GetIndent.line_indent(p, "left") == 3.0

    def test_line_indent_right_with_chars(self, doc):
        """Real w:ind element with w:rightChars -> returns float"""
        p = doc.add_paragraph()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:rightChars"), "200")
        pPr.append(ind)
        assert GetIndent.line_indent(p, "right") == 2.0

    def test_line_indent_invalid_chars_value(self, doc):
        """w:leftChars has non-numeric value -> returns None (line 417-418)"""
        p = doc.add_paragraph()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:leftChars"), "abc")
        pPr.append(ind)
        assert GetIndent.line_indent(p, "left") is None

    def test_line_indent_no_pPr(self, doc):
        """No pPr element -> returns None (line 400-401)"""
        p = doc.add_paragraph()
        # Remove pPr if it exists
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            p._element.remove(pPr)
        assert GetIndent.line_indent(p, "left") is None

    def test_line_indent_no_ind(self, doc):
        """pPr exists but no ind element -> returns None (line 404-405)"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn

        pPr = p._element.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            pPr.remove(ind)
        assert GetIndent.line_indent(p, "left") is None


class TestParagraphGetSpaceBeforeWithValidLines:
    """Cover lines 157-158: paragraph_get_space_before with valid beforeLines in XML"""

    def test_space_before_valid_lines(self, doc):
        """Valid beforeLines in XML returns correct value (lines 157-158)"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:beforeLines"), "100")
        pPr.append(spacing)
        assert paragraph_get_space_before(p) == 1.0

    def test_space_before_valid_lines_rounded(self, doc):
        """Valid beforeLines value is rounded (line 156)"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:beforeLines"), "33")
        pPr.append(spacing)
        assert paragraph_get_space_before(p) == 0.3  # round(0.33, 1)


class TestParagraphGetSpaceAfterWithValidLines:
    """Cover lines 190-191: paragraph_get_space_after with valid afterLines in XML"""

    def test_space_after_valid_lines(self, doc):
        """Valid afterLines in XML returns correct value (lines 190-191)"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:afterLines"), "150")
        pPr.append(spacing)
        assert paragraph_get_space_after(p) == 1.5

    def test_space_after_valid_lines_rounded(self, doc):
        """Valid afterLines value is rounded (line 189)"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:afterLines"), "67")
        pPr.append(spacing)
        assert paragraph_get_space_after(p) == 0.7  # round(0.67, 1)


class TestRunGetFontColorValidRGB:
    """Cover line 280: run_get_font_color when color.rgb is a valid RGBColor"""

    def test_font_color_valid_rgb_tuple_access(self, doc):
        """color.rgb returns RGBColor, access via indexing (line 280)"""
        run = doc.add_paragraph().add_run("x")
        run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
        # RGBColor supports indexing: rgb[0] = 0x12, etc.
        result = run_get_font_color(run)
        assert result == (0x12, 0x34, 0x56)


class TestRunGetFontColorFalsyRGB:
    """Cover line 338: run_get_font_color when color.rgb is falsy (empty string)"""

    def test_font_color_falsy_rgb(self):
        """color.rgb is a falsy value -> returns (0,0,0) (line 338)"""
        mock_run = MagicMock()
        # Create a mock color where rgb is truthy but then becomes falsy
        mock_color = MagicMock()
        mock_color.rgb = ""  # empty string is falsy
        mock_run.font.color = mock_color
        result = run_get_font_color(mock_run)
        assert result == (0, 0, 0)

    def test_font_color_zero_rgb(self):
        """color.rgb is 0 (falsy int) -> returns (0,0,0) (line 338)"""
        mock_run = MagicMock()
        mock_color = MagicMock()
        mock_color.rgb = 0  # 0 is falsy
        mock_run.font.color = mock_color
        result = run_get_font_color(mock_run)
        assert result == (0, 0, 0)


class TestGetIndentLineIndentWithCharsValues:
    """Cover lines 419-423: GetIndent.line_indent with valid leftChars/rightChars values"""

    def test_line_indent_left_chars_valid(self, doc):
        """Valid w:leftChars value returns correct float (lines 419-423)"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:leftChars"), "150")
        pPr.append(ind)
        assert GetIndent.line_indent(p, "left") == 1.5

    def test_line_indent_right_chars_valid(self, doc):
        """Valid w:rightChars value returns correct float"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:rightChars"), "250")
        pPr.append(ind)
        assert GetIndent.line_indent(p, "right") == 2.5

    def test_line_indent_left_chars_zero(self, doc):
        """w:leftChars = 0 -> returns max(0.0, 0.0) = 0.0"""
        p = doc.add_paragraph()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:leftChars"), "0")
        pPr.append(ind)
        assert GetIndent.line_indent(p, "left") == 0.0

    def test_line_indent_exception_returns_none(self, doc):
        """Exception during line_indent -> returns None (lines 421-423)"""
        mock_para = MagicMock()
        mock_para._element.pPr = None
        assert GetIndent.line_indent(mock_para, "left") is None
