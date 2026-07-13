#!/usr/bin/env python
"""样式继承链解析测试（style/inheritance.py + reader 委托）。

覆盖完整五层链：直接格式 → 字符样式(rStyle) → 段落样式 basedOn 链
→ docDefaults → 主题字体，以及 Mock 降级与各提取器分支。
"""

import pytest
from unittest.mock import MagicMock

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from wordformat.style.inheritance import (
    StyleResolver,
    ThemeFontTable,
    ThemeRef,
    x_bold,
    x_size_pt,
)
from wordformat.style.reader import (
    GetIndent,
    paragraph_get_alignment,
    paragraph_get_first_line_indent,
    paragraph_get_line_spacing,
    paragraph_get_line_spacing_rule,
    paragraph_get_space_after,
    paragraph_get_space_before,
    run_get_font_bold,
    run_get_font_color,
    run_get_font_italic,
    run_get_font_name,
    run_get_font_name_en,
    run_get_font_size_pt,
    run_get_font_underline,
)


@pytest.fixture
def doc():
    return Document()


def _para_style(doc, name, base=None):
    st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base is not None:
        st.base_style = base
    return st


# ── 1. 字符属性从样式继承（run 无直接格式）─────────────────────────
class TestRunInheritsFromParagraphStyle:
    def test_bold_from_heading_style(self, doc):
        """Heading 1 样式自带加粗，run 无直接 b → True。"""
        p = doc.add_paragraph("标题", style="Heading 1")
        assert run_get_font_bold(p.runs[0]) is True

    def test_italic_from_style(self, doc):
        st = _para_style(doc, "ItalicStyle")
        st.font.italic = True
        r = doc.add_paragraph(style="ItalicStyle").add_run("x")
        assert run_get_font_italic(r) is True

    def test_underline_from_style(self, doc):
        st = _para_style(doc, "ULStyle")
        st.font.underline = True
        r = doc.add_paragraph(style="ULStyle").add_run("x")
        assert run_get_font_underline(r) is True

    def test_size_from_style(self, doc):
        st = _para_style(doc, "SizeS")
        st.font.size = Pt(18)
        r = doc.add_paragraph(style="SizeS").add_run("x")
        assert run_get_font_size_pt(r) == 18.0

    def test_color_from_style(self, doc):
        st = _para_style(doc, "ColorS")
        st.font.color.rgb = RGBColor(0x11, 0x22, 0x33)
        r = doc.add_paragraph(style="ColorS").add_run("x")
        assert run_get_font_color(r) == (0x11, 0x22, 0x33)


# ── 2. docDefaults 兜底 ───────────────────────────────────────────
class TestDocDefaults:
    def test_size_from_docdefaults(self, doc):
        """默认模板 docDefaults sz=22 → 11.0。"""
        assert run_get_font_size_pt(doc.add_paragraph().add_run("x")) == 11.0

    def test_en_font_from_docdefaults_theme(self, doc):
        """docDefaults asciiTheme=minorHAnsi → 主题字体 Cambria。"""
        assert run_get_font_name_en(doc.add_paragraph().add_run("x")) == "Cambria"


# ── 3. basedOn 多级链 ─────────────────────────────────────────────
class TestBasedOnChain:
    def test_three_level_chain(self, doc):
        """孙样式无值，沿 basedOn 上溯到祖样式取加粗。"""
        gp = _para_style(doc, "GP")
        gp.font.bold = True
        parent = _para_style(doc, "Parent", base=gp)
        child = _para_style(doc, "Child", base=parent)
        r = doc.add_paragraph(style="Child").add_run("x")
        assert run_get_font_bold(r) is True

    def test_nearer_style_wins(self, doc):
        """近端样式覆盖远端：Parent 非加粗覆盖 GP 加粗。"""
        gp = _para_style(doc, "GP2")
        gp.font.bold = True
        parent = _para_style(doc, "Parent2", base=gp)
        parent.font.bold = False
        r = doc.add_paragraph(style="Parent2").add_run("x")
        assert run_get_font_bold(r) is False

    def test_alignment_via_chain(self, doc):
        base = _para_style(doc, "BAlign")
        base.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_style(doc, "CAlign", base=base)
        p = doc.add_paragraph(style="CAlign")
        assert paragraph_get_alignment(p) == WD_ALIGN_PARAGRAPH.CENTER


# ── 4. 直接格式优先 ───────────────────────────────────────────────
class TestDirectOverride:
    def test_direct_bold_overrides_style(self, doc):
        st = _para_style(doc, "NoBold")
        st.font.bold = False
        r = doc.add_paragraph(style="NoBold").add_run("x")
        r.bold = True
        assert run_get_font_bold(r) is True

    def test_direct_size_overrides(self, doc):
        r = doc.add_paragraph(style="Heading 1").add_run("x")
        r.font.size = Pt(9)
        assert run_get_font_size_pt(r) == 9.0


# ── 5. 字符样式（rStyle）链 ───────────────────────────────────────
class TestCharacterStyle:
    def test_rstyle_bold(self, doc):
        cst = doc.styles.add_style("CBold", WD_STYLE_TYPE.CHARACTER)
        cst.font.bold = True
        r = doc.add_paragraph().add_run("x")
        r.style = "CBold"
        assert run_get_font_bold(r) is True


# ── 6. 段落属性继承 ───────────────────────────────────────────────
class TestParagraphInheritance:
    def test_space_before_from_style(self, doc):
        st = _para_style(doc, "SB")
        st.element.get_or_add_pPr().append(_spacing("beforeLines", "150"))
        p = doc.add_paragraph(style="SB")
        assert paragraph_get_space_before(p) == 1.5

    def test_space_after_from_style(self, doc):
        st = _para_style(doc, "SA")
        st.element.get_or_add_pPr().append(_spacing("afterLines", "200"))
        p = doc.add_paragraph(style="SA")
        assert paragraph_get_space_after(p) == 2.0

    def test_first_line_indent_from_style(self, doc):
        st = _para_style(doc, "FLI")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLineChars"), "200")
        st.element.get_or_add_pPr().append(ind)
        p = doc.add_paragraph(style="FLI")
        assert paragraph_get_first_line_indent(p) == 2.0

    def test_left_right_indent_from_style(self, doc):
        st = _para_style(doc, "LRI")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:leftChars"), "300")
        ind.set(qn("w:rightChars"), "100")
        st.element.get_or_add_pPr().append(ind)
        p = doc.add_paragraph(style="LRI")
        assert GetIndent.left_indent(p) == 3.0
        assert GetIndent.right_indent(p) == 1.0

    def test_hanging_indent_negative(self, doc):
        p = doc.add_paragraph()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:hangingChars"), "150")
        p._element.get_or_add_pPr().append(ind)
        assert paragraph_get_first_line_indent(p) == -1.5


# ── 行距 ──────────────────────────────────────────────────────────
class TestLineSpacing:
    @pytest.mark.parametrize(
        "rule,value,exp_val,exp_rule",
        [
            (WD_LINE_SPACING.SINGLE, None, 1.0, WD_LINE_SPACING.SINGLE),
            (WD_LINE_SPACING.ONE_POINT_FIVE, None, 1.5, WD_LINE_SPACING.ONE_POINT_FIVE),
            (WD_LINE_SPACING.DOUBLE, None, 2.0, WD_LINE_SPACING.DOUBLE),
        ],
    )
    def test_preset(self, doc, rule, value, exp_val, exp_rule):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = rule
        assert paragraph_get_line_spacing(p) == exp_val
        assert paragraph_get_line_spacing_rule(p) == exp_rule

    def test_multiple_custom(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 2.3
        assert paragraph_get_line_spacing(p) == 2.3
        assert paragraph_get_line_spacing_rule(p) == WD_LINE_SPACING.MULTIPLE

    @pytest.mark.parametrize(
        "rule", [WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST]
    )
    def test_fixed_returns_none(self, doc, rule):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = rule
        assert paragraph_get_line_spacing(p) is None


# ── 主题字体表 ────────────────────────────────────────────────────
class TestThemeFontTable:
    def test_resolve_minor(self, doc):
        t = ThemeFontTable(doc)
        assert t.resolve("minorHAnsi") == "Cambria"
        assert t.resolve("majorHAnsi") == "Calibri"

    def test_resolve_unknown_none(self, doc):
        assert ThemeFontTable(doc).resolve("nope") is None

    def test_empty_table_no_document(self):
        assert ThemeFontTable(None).resolve("minorHAnsi") is None


# ── Mock / 降级 ───────────────────────────────────────────────────
class TestDegradation:
    def test_mock_run_getters_default(self):
        m = MagicMock()
        assert run_get_font_bold(m) is False
        assert run_get_font_italic(m) is False
        assert run_get_font_underline(m) is False
        assert run_get_font_size_pt(m) == 12.0
        assert run_get_font_color(m) == (0, 0, 0)
        assert run_get_font_name(m) is None
        assert run_get_font_name_en(m) is None

    def test_mock_para_getters_default(self):
        m = MagicMock()
        assert paragraph_get_alignment(m) is None
        assert paragraph_get_space_before(m) is None
        assert paragraph_get_line_spacing(m) is None
        assert GetIndent.left_indent(m) is None

    def test_resolver_no_document(self):
        """无文档构建的 resolver 只读直接格式，不抛异常。"""
        res = StyleResolver(None)
        assert res.theme.resolve("minorHAnsi") is None


# ── 提取器直接单测 ────────────────────────────────────────────────
class TestExtractors:
    def test_bold_toggle_off(self, doc):
        r = doc.add_paragraph().add_run("x")
        rPr = r._element.get_or_add_rPr()
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "0")
        rPr.append(b)
        assert x_bold(rPr) is False

    def test_size_invalid(self, doc):
        r = doc.add_paragraph().add_run("x")
        rPr = r._element.get_or_add_rPr()
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "abc")
        rPr.append(sz)
        from wordformat.style.inheritance import _MISS

        assert x_size_pt(rPr) is _MISS

    def test_theme_ref_font(self, doc):
        """rFonts 只给 asciiTheme → 提取器返回 ThemeRef，resolver 兑现字体名。"""
        r = doc.add_paragraph().add_run("x")
        rPr = r._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:asciiTheme"), "majorHAnsi")
        rPr.append(rFonts)
        from wordformat.style.inheritance import x_font_ascii

        assert isinstance(x_font_ascii(rPr), ThemeRef)
        assert run_get_font_name_en(r) == "Calibri"


def _spacing(attr, val):
    sp = OxmlElement("w:spacing")
    sp.set(qn(f"w:{attr}"), val)
    return sp
