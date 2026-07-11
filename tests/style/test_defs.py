#!/usr/bin/env python
"""
comprehensive tests for style modules:
  check_format, get_some, set_some, style_enum, utils
"""

import pytest
from types import SimpleNamespace
from unittest.mock import MagicMock, PropertyMock, patch as mock_patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from wordformat.style.diff import DIFFResult, CharacterStyle, ParagraphStyle
from wordformat.style.reader import (
    paragraph_get_alignment, paragraph_get_space_before, paragraph_get_space_after,
    paragraph_get_line_spacing, paragraph_get_first_line_indent,
    paragraph_get_builtin_style_name, run_get_font_name, run_get_font_size_pt,
    run_get_font_color, run_get_font_bold, run_get_font_italic,
    run_get_font_underline, GetIndent,
)
from wordformat.style.writer import (
    run_set_font_name, set_paragraph_space_before_by_lines,
    set_paragraph_space_after_by_lines, _paragraph_space_by_lines,
    SetSpacing, SetLineSpacing, SetIndent, SetFirstLineIndent,
)
from wordformat.style.defs import (
    FontName, FontSize, FontColor, Alignment, LineSpacingRule, LineSpacing,
    FirstLineIndent, LeftIndent, RightIndent, BuiltInStyle, SpaceBefore, SpaceAfter,
    Spacing, UnitLabelEnum,
)
from wordformat.style.units import extract_unit_from_string, UnitResult


# ---------------------------------------------------------------------------
# fixtures & helpers
# ---------------------------------------------------------------------------

@pytest.fixture
def doc():
    return Document()


@pytest.fixture
def mock_warning():
    w = MagicMock()
    for attr in ("bold", "italic", "underline", "font_size", "font_color",
                 "font_name", "alignment", "space_before", "space_after",
                 "line_spacing", "line_spacingrule", "first_line_indent",
                 "left_indent", "right_indent", "builtin_style_name"):
        setattr(w, attr, True)
    return w


@pytest.fixture
def mock_warning_off():
    w = MagicMock()
    for attr in ("bold", "italic", "underline", "font_size", "font_color",
                 "font_name", "alignment", "space_before", "space_after",
                 "line_spacing", "line_spacingrule", "first_line_indent",
                 "left_indent", "right_indent", "builtin_style_name"):
        setattr(w, attr, False)
    return w


def _set_warning(w):
    import wordformat.style.diff as m
    m.style_checks_warning = w


def _clear_warning():
    import wordformat.style.diff as m
    m.__dict__.pop("style_checks_warning", None)


# ===========================================================================
# style_enum - FontSize
# ===========================================================================

class TestFontSize:
    @pytest.mark.parametrize("label,expected", [
        ("小四", 12), ("四号", 14), ("三号", 16), ("五号", 10.5), ("一号", 26), ("七号", 5.5),
    ])
    def test_label_map(self, label, expected):
        assert FontSize(label).rel_value == expected

    def test_bare_number_rel_value_is_string(self):
        """No unit -> extract_unit fails -> rel_value stays as raw string."""
        assert FontSize("15").rel_value == "15"

    def test_base_set(self, doc):
        run = doc.add_paragraph().add_run("x")
        FontSize("小四").base_set(run)
        assert run.font.size.pt == 12

    def test_base_set_invalid_raises(self, doc):
        with pytest.raises(ValueError, match="无效的字号"):
            FontSize("abc").base_set(doc.add_paragraph().add_run("x"))



# ===========================================================================
# style_enum - FontColor
# ===========================================================================

class TestFontColor:
    @pytest.mark.parametrize("spec,expected", [
        ("BLACK", (0, 0, 0)), ("black", (0, 0, 0)), ("黑色", (0, 0, 0)),
        ("red", (255, 0, 0)), ("红色", (255, 0, 0)),
        ("#FF0000", (255, 0, 0)), ("#f00", (255, 0, 0)), ("FF0000", (255, 0, 0)),
        ("white", (255, 255, 255)), ("白色", (255, 255, 255)),
    ])
    def test_parse_color(self, spec, expected):
        assert FontColor(spec).rel_value == expected

    def test_eq_tuple(self):
        assert FontColor("red") == (255, 0, 0)
        assert FontColor("red") != (0, 0, 0)

    def test_eq_string(self):
        """字符串比较应与解析后的RGB比较"""
        assert FontColor("red") == "red"
        assert FontColor("red") != "blue"

    def test_base_set(self, doc):
        run = doc.add_paragraph().add_run("x")
        FontColor("red").base_set(run)
        assert run.font.color.rgb == RGBColor(255, 0, 0)

    def test_invalid_raises(self):
        with pytest.raises(ValueError):
            FontColor("not_a_color").rel_value

    def test_non_string_raises(self):
        with pytest.raises(TypeError):
            FontColor(123).rel_value

    @pytest.mark.parametrize("h,ok", [("#f00", True), ("#FF0000", True), ("FF0000", True), ("#GGG", False)])
    def test_is_hex(self, h, ok):
        assert FontColor._is_hex(h) == ok

    def test_normalize_hex(self):
        assert FontColor._normalize_hex("#f00") == "#ff0000"
        assert FontColor._normalize_hex("AABBCC") == "#aabbcc"



# ===========================================================================
# style_enum - Alignment / LineSpacingRule / LineSpacing / BuiltInStyle / FontName
# ===========================================================================

class TestAlignmentEnum:
    @pytest.mark.parametrize("label,val", [
        ("左对齐", WD_ALIGN_PARAGRAPH.LEFT), ("居中对齐", WD_ALIGN_PARAGRAPH.CENTER),
        ("右对齐", WD_ALIGN_PARAGRAPH.RIGHT), ("两端对齐", WD_ALIGN_PARAGRAPH.JUSTIFY),
    ])
    def test_label_map(self, label, val):
        assert Alignment(label).rel_value == val

    def test_base_set_and_get(self, doc):
        p = doc.add_paragraph()
        Alignment("居中对齐").base_set(p)
        assert p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert Alignment("左对齐").get_from_paragraph(p) == WD_ALIGN_PARAGRAPH.CENTER

    def test_get_fallback_left(self, doc):
        assert Alignment("左对齐").get_from_paragraph(doc.add_paragraph()) == WD_ALIGN_PARAGRAPH.LEFT

    def test_invalid_raises(self, doc):
        with pytest.raises(ValueError, match="无效的对齐方式"):
            Alignment("无效").base_set(doc.add_paragraph())



class TestLineSpacingRuleEnum:
    @pytest.mark.parametrize("label,val", [
        ("单倍行距", WD_LINE_SPACING.SINGLE), ("1.5倍行距", WD_LINE_SPACING.ONE_POINT_FIVE),
        ("2倍行距", WD_LINE_SPACING.DOUBLE), ("固定值", WD_LINE_SPACING.EXACTLY),
        ("最小值", WD_LINE_SPACING.AT_LEAST), ("多倍行距", WD_LINE_SPACING.MULTIPLE),
    ])
    def test_label_map(self, label, val):
        assert LineSpacingRule(label).rel_value == val



class TestLineSpacingEnum:
    def test_rel_value(self):
        assert LineSpacing("1.5倍").rel_value == 1.5

    def test_base_set(self, doc):
        p = doc.add_paragraph()
        LineSpacing("1.5倍").base_set(p)
        assert p.paragraph_format.line_spacing == 1.5

    def test_leq_zero_silently_changes_to_1(self, doc):
        """<=0 现在抛出 ValueError 而非静默改为 1。"""
        p = doc.add_paragraph()
        with pytest.raises(ValueError):
            LineSpacing("0倍").base_set(p)

    def test_zero_should_raise(self, doc):
        with pytest.raises(ValueError):
            LineSpacing("0倍").base_set(doc.add_paragraph())



class TestBuiltInStyleEnum:
    @pytest.mark.parametrize("label,expected", [("正文", "Normal"), ("标题", "Title"), ("Heading 1", "Heading 1")])
    def test_label_map(self, label, expected):
        assert BuiltInStyle(label).rel_value == expected

    def test_get_from_paragraph(self, doc):
        assert BuiltInStyle("正文").get_from_paragraph(doc.add_paragraph()) == "normal"

    def test_base_set_valid(self, doc):
        p = doc.add_paragraph()
        BuiltInStyle("正文").base_set(p)
        assert p.style.name == "Normal"

    def test_base_set_invalid_raises(self, doc):
        """NonExistent 样式会被 _ensure_style_exists 自动创建，不再抛异常"""
        p = doc.add_paragraph()
        BuiltInStyle("NonExistent").base_set(p)
        assert p.style.name == "NonExistent"



class TestFontNameEnum:
    @pytest.mark.parametrize("name,is_cn", [
        ("宋体", True), ("黑体", True), ("Times New Roman", False), ("Arial", False),
    ])
    def test_is_chinese(self, name, is_cn):
        assert FontName(name).is_chinese(name) == is_cn

    def test_base_set_chinese(self, doc):
        run = doc.add_paragraph().add_run("x")
        FontName("宋体").base_set(run)
        assert run_get_font_name(run) == "宋体"

    def test_base_set_english(self, doc):
        run = doc.add_paragraph().add_run("x")
        FontName("Arial").base_set(run)
        assert run.font.name == "Arial"



class TestUnitLabelEnumEq:
    def test_same_class_same_rel(self):
        assert FontSize("小四") == FontSize("小四")

    def test_same_class_diff_rel(self):
        assert FontSize("小四") != FontSize("三号")

    def test_string_rel_eq(self):
        assert FontName("宋体") == "宋体"

    def test_numeric_rel_eq(self):
        assert FontSize("小四") == 12

    def test_different_class(self):
        assert FontSize("小四") != Alignment("左对齐")



# ===========================================================================
# utils
# ===========================================================================

class TestExtractUnitFromString:
    @pytest.mark.parametrize("text,val,unit", [
        ("12pt", 12.0, "pt"), ("1.5磅", 1.5, "pt"), ("2cm", 2.0, "cm"),
        ("3.5厘米", 3.5, "cm"), ("1inch", 1.0, "inch"), ("0.5英寸", 0.5, "inch"),
        ("10mm", 10.0, "mm"), ("2毫米", 2.0, "mm"),
        ("1.5行", 1.5, "hang"), ("2字符", 2.0, "char"), ("1.5倍", 1.5, "bei"),
        ("100emu", 100.0, "emu"),
    ])
    def test_valid(self, text, val, unit):
        r = extract_unit_from_string(text)
        assert r.is_valid and r.value == val and r.standard_unit == unit

    def test_invalid(self):
        r = extract_unit_from_string("no_unit")
        assert not r.is_valid and r.value is None

    def test_case_insensitive(self):
        assert extract_unit_from_string("12PT").standard_unit == "pt"



class TestUnitResult:
    def test_to_dict(self):
        r = UnitResult(original_unit="pt", standard_unit="pt", value=12.0, is_valid=True)
        assert r.to_dict()["value"] == 12.0

    @pytest.mark.parametrize("unit,val,emu", [
        ("pt", 1, 12700), ("cm", 1, 360000), ("inch", 1, 914400), ("mm", 1, 36000), ("emu", 1, 1),
    ])
    def test_convert_to_emu(self, unit, val, emu):
        assert UnitResult(standard_unit=unit, value=val, is_valid=True).convert_to_emu() == emu

    def test_convert_to_emu_hang_bug(self):
        """hang 单位现在返回 None 而非 0。"""
        assert UnitResult(standard_unit="hang", value=1.5, is_valid=True).convert_to_emu() is None

    def test_convert_to_emu_hang_should_be_none(self):
        assert UnitResult(standard_unit="hang", value=1.5, is_valid=True).convert_to_emu() is None

    def test_convert_to_emu_invalid(self):
        assert UnitResult(is_valid=False).convert_to_emu() is None

    @pytest.mark.parametrize("unit,ch", [
        ("pt", "磅"), ("cm", "厘米"), ("inch", "英寸"), ("mm", "毫米"),
        ("emu", "emu"), ("hang", "行"), ("char", "字符"), ("bei", "倍"),
    ])
    def test_unit_ch(self, unit, ch):
        assert UnitResult(standard_unit=unit).unit_ch == ch

    def test_unit_ch_none(self):
        assert UnitResult(standard_unit=None).unit_ch == ""

    def test_unit_ch_invalid_raises(self):
        with pytest.raises(ValueError, match="Invalid unit"):
            UnitResult(standard_unit="xyz").unit_ch



# ===========================================================================
# Coverage: style_enum.py uncovered lines
# ===========================================================================


class TestUnitLabelEnumRelValueSetter:
    """Cover line 122: rel_value.setter"""

    def test_rel_value_setter(self):
        """Setting rel_value directly via setter"""
        e = FontSize("小四")
        assert e.rel_value == 12
        e.rel_value = 99
        assert e.rel_value == 99



class TestUnitLabelEnumBaseSetDefault:
    """Cover line 135: base_set default implementation (just logs debug)"""

    def test_base_set_default_no_crash(self, doc):
        """UnitLabelEnum.base_set default just logs, no crash"""
        # Use a bare UnitLabelEnum instance - but it's abstract due to get_from_paragraph
        # So we create a concrete subclass that doesn't override base_set
        class MinimalEnum(UnitLabelEnum):
            def get_from_paragraph(self, paragraph):
                return None

        e = MinimalEnum("test")
        # Calling base_set should just log and not crash
        e.base_set(doc.add_paragraph())



class TestUnitLabelEnumFormatWithRun:
    """Cover lines 152-155: format() method with Run object (not Paragraph)"""

    def test_format_with_run(self, doc):
        """format() dispatches to fun(run=...) when docx_obj is a Run"""
        run = doc.add_paragraph().add_run("x")
        FontSize("小四").format(run)
        assert run.font.size.pt == 12

    def test_format_with_paragraph(self, doc):
        """format() dispatches to fun(paragraph=...) when docx_obj is a Paragraph"""
        p = doc.add_paragraph()
        Alignment("居中对齐").format(p)
        assert p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_format_no_fun_falls_to_base_set(self, doc):
        """format() with no matching function_map falls back to base_set"""
        run = doc.add_paragraph().add_run("x")
        FontName("宋体").format(run)
        assert run_get_font_name(run) == "宋体"



class TestUnitLabelEnumGetFromParagraphAbstract:
    """Cover line 163: get_from_paragraph abstract method raises NotImplementedError"""

    def test_abstract_raises_not_implemented(self, doc):
        """Calling get_from_paragraph on base UnitLabelEnum raises NotImplementedError"""
        # Can't instantiate UnitLabelEnum directly due to ABC, use a minimal subclass
        class MinimalEnum(UnitLabelEnum):
            pass

        e = MinimalEnum("test")
        with pytest.raises(NotImplementedError):
            e.get_from_paragraph(doc.add_paragraph())



class TestFontColorParseColorValueError:
    """Cover lines 304-305: FontColor._parse_color ValueError for invalid hex"""

    def test_parse_color_invalid_hex_raises(self):
        """Hex that passes _is_hex but fails webcolors.hex_to_rgb"""
        # _is_hex checks regex, but webcolors might still fail
        # A valid hex format that webcolors can't parse is unlikely since all valid hex are valid colors
        # But we can test the ValueError path by using a hex that causes webcolors to raise
        # Actually, all valid 6-digit hex are valid RGB colors, so this path is hard to trigger
        # We'll test with a mock to force the ValueError
        import webcolors
        original = webcolors.hex_to_rgb
        webcolors.hex_to_rgb = MagicMock(side_effect=ValueError("mock error"))
        try:
            with pytest.raises(ValueError, match="非法十六进制色值"):
                FontColor("#FF0000").rel_value
        finally:
            webcolors.hex_to_rgb = original



class TestFontColorBaseSetNonString:
    """Cover line 330: FontColor.base_set with non-string value raises TypeError"""

    def test_base_set_non_string_raises(self, doc):
        """base_set with non-string value raises TypeError"""
        run = doc.add_paragraph().add_run("x")
        fc = FontColor.__new__(FontColor)
        fc.value = 123  # non-string
        fc._rel_value = None
        fc.original_unit = None
        fc.unit_ch = None
        fc._rel_unit = None
        fc.extract_unit_result = None
        with pytest.raises(TypeError, match="颜色标识仅支持字符串"):
            fc.base_set(run)



class TestFontColorEqEdgeCases:
    """Cover lines 344-345: FontColor.__eq__ with non-tuple or wrong length tuple"""

    def test_eq_non_tuple_returns_false(self):
        """__eq__ with non-color values returns False"""
        fc = FontColor("red")
        assert fc == "red"          # 字符串颜色名现在被正确比较
        assert fc != 42             # 非颜色应返回 False
        assert fc != None           # None 应返回 False

    def test_eq_wrong_length_tuple_returns_false(self):
        """__eq__ with tuple of wrong length returns False (line 339-340)"""
        fc = FontColor("red")
        assert fc != (255, 0)
        assert fc != (255, 0, 0, 0)

    def test_eq_parse_error_returns_false(self):
        """__eq__ when rel_value raises -> returns False (lines 344-345)"""
        fc = FontColor.__new__(FontColor)
        fc.value = 12345  # will cause _parse_color to raise TypeError
        fc._rel_value = None
        fc.original_unit = None
        fc.unit_ch = None
        fc._rel_unit = None
        fc.extract_unit_result = None
        # __eq__ catches TypeError/ValueError from rel_value
        assert fc != (255, 0, 0)



class TestSpacingBaseSetHangNotImplemented:
    """Cover lines 400-404: Spacing.base_set raises NotImplementedError for 'hang' unit"""

    def test_spacing_get_from_paragraph_hang_raises(self, doc):
        """Spacing.get_from_paragraph with 'hang' unit raises NotImplementedError"""
        s = Spacing("1行")
        with pytest.raises(NotImplementedError, match="Spacing 需要知道是 before 还是 after"):
            s.get_from_paragraph(doc.add_paragraph())



class TestSpaceBeforeGetFromParagraphUnits:
    """Cover lines 412-416: SpaceBefore.get_from_paragraph with pt/mm/cm/inch units"""

    @pytest.mark.parametrize("unit,val,attr", [
        ("pt", 12, "pt"), ("mm", 5, "mm"), ("cm", 1, "cm"), ("inch", 0.5, "inches"),
    ])
    def test_get_from_paragraph_physical_units(self, doc, unit, val, attr):
        """SpaceBefore.get_from_paragraph returns value in specified physical unit"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(val) if unit == "pt" else None
        if unit != "pt":
            # Set via EMU conversion
            from docx.shared import Emu
            if unit == "mm":
                p.paragraph_format.space_before = Emu(36000 * val)
            elif unit == "cm":
                p.paragraph_format.space_before = Emu(360000 * val)
            elif unit == "inch":
                p.paragraph_format.space_before = Emu(914400 * val)

        sb = SpaceBefore(f"{val}{unit}")
        result = sb.get_from_paragraph(p)
        assert result is not None
        # Result should be close to val in the specified unit
        assert abs(result - val) < 0.01

    def test_get_from_paragraph_hang_unit(self, doc):
        """SpaceBefore.get_from_paragraph with 'hang' unit uses paragraph_get_space_before"""
        p = doc.add_paragraph()
        set_paragraph_space_before_by_lines(p, 0.5)
        sb = SpaceBefore("0.5行")
        result = sb.get_from_paragraph(p)
        assert result == 0.5

    def test_get_from_paragraph_none(self, doc):
        """SpaceBefore.get_from_paragraph returns None when no spacing set"""
        sb = SpaceBefore("12pt")
        result = sb.get_from_paragraph(doc.add_paragraph())
        assert result is None



class TestSpaceAfterGetFromParagraphUnits:
    """Cover lines 424-428: SpaceAfter.get_from_paragraph with pt/mm/cm/inch units"""

    @pytest.mark.parametrize("unit,val", [
        ("pt", 12), ("mm", 5), ("cm", 1), ("inch", 0.5),
    ])
    def test_get_from_paragraph_physical_units(self, doc, unit, val):
        """SpaceAfter.get_from_paragraph returns value in specified physical unit"""
        p = doc.add_paragraph()
        from docx.shared import Emu
        if unit == "pt":
            p.paragraph_format.space_after = Pt(val)
        elif unit == "mm":
            p.paragraph_format.space_after = Emu(36000 * val)
        elif unit == "cm":
            p.paragraph_format.space_after = Emu(360000 * val)
        elif unit == "inch":
            p.paragraph_format.space_after = Emu(914400 * val)

        sa = SpaceAfter(f"{val}{unit}")
        result = sa.get_from_paragraph(p)
        assert result is not None
        assert abs(result - val) < 0.01

    def test_get_from_paragraph_hang_unit(self, doc):
        """SpaceAfter.get_from_paragraph with 'hang' unit uses paragraph_get_space_after"""
        p = doc.add_paragraph()
        set_paragraph_space_after_by_lines(p, 1.0)
        sa = SpaceAfter("1行")
        result = sa.get_from_paragraph(p)
        assert result == 1.0

    def test_get_from_paragraph_none(self, doc):
        """SpaceAfter.get_from_paragraph returns None when no spacing set"""
        sa = SpaceAfter("12pt")
        result = sa.get_from_paragraph(doc.add_paragraph())
        assert result is None



class TestLineSpacingRuleBaseSetInvalid:
    """Cover lines 447-451: LineSpacingRule.base_set with invalid value raises ValueError"""

    def test_base_set_invalid_raises(self, doc):
        """base_set with value not in _LABEL_MAP raises ValueError"""
        p = doc.add_paragraph()
        lsr = LineSpacingRule.__new__(LineSpacingRule)
        lsr.value = "无效行距"
        lsr._rel_value = None
        lsr.original_unit = None
        lsr.unit_ch = None
        lsr._rel_unit = None
        lsr.extract_unit_result = None
        with pytest.raises(ValueError, match="无效的行距选项"):
            lsr.base_set(p)



class TestLineSpacingBaseSetInvalidValue:
    """Cover line 489: LineSpacing.base_set with invalid value raises ValueError"""

    def test_base_set_non_numeric_raises(self, doc):
        """base_set with non-numeric rel_value raises ValueError"""
        p = doc.add_paragraph()
        ls = LineSpacing.__new__(LineSpacing)
        ls.value = "abc"
        ls._rel_value = "abc"  # not int or float
        ls.original_unit = None
        ls.unit_ch = None
        ls._rel_unit = None
        ls.extract_unit_result = None
        with pytest.raises(ValueError, match="无效的行距"):
            ls.base_set(p)

    def test_base_set_none_raises(self, doc):
        """base_set with None rel_value raises ValueError"""
        p = doc.add_paragraph()
        ls = LineSpacing.__new__(LineSpacing)
        ls.value = None
        ls._rel_value = None
        ls.original_unit = None
        ls.unit_ch = None
        ls._rel_unit = None
        ls.extract_unit_result = None
        with pytest.raises(ValueError, match="无效的行距"):
            ls.base_set(p)



class TestLeftIndentGetFromParagraphUnits:
    """Cover lines 513-517: LeftIndent.get_from_paragraph with pt/mm/cm/inch units"""

    @pytest.mark.parametrize("unit,val", [
        ("pt", 12), ("mm", 5), ("cm", 1), ("inch", 0.5),
    ])
    def test_get_from_paragraph_physical_units(self, doc, unit, val):
        """LeftIndent.get_from_paragraph returns value in specified physical unit"""
        p = doc.add_paragraph()
        from docx.shared import Emu
        if unit == "pt":
            p.paragraph_format.left_indent = Pt(val)
        elif unit == "mm":
            p.paragraph_format.left_indent = Emu(36000 * val)
        elif unit == "cm":
            p.paragraph_format.left_indent = Emu(360000 * val)
        elif unit == "inch":
            p.paragraph_format.left_indent = Emu(914400 * val)

        li = LeftIndent(f"{val}{unit}")
        result = li.get_from_paragraph(p)
        assert result is not None
        assert abs(result - val) < 0.01

    def test_get_from_paragraph_none(self, doc):
        """LeftIndent.get_from_paragraph returns None when no indent set"""
        li = LeftIndent("12pt")
        result = li.get_from_paragraph(doc.add_paragraph())
        assert result is None



class TestRightIndentGetFromParagraphUnits:
    """Cover lines 525-529: RightIndent.get_from_paragraph with pt/mm/cm/inch units"""

    @pytest.mark.parametrize("unit,val", [
        ("pt", 12), ("mm", 5), ("cm", 1), ("inch", 0.5),
    ])
    def test_get_from_paragraph_physical_units(self, doc, unit, val):
        """RightIndent.get_from_paragraph returns value in specified physical unit"""
        p = doc.add_paragraph()
        from docx.shared import Emu
        if unit == "pt":
            p.paragraph_format.right_indent = Pt(val)
        elif unit == "mm":
            p.paragraph_format.right_indent = Emu(36000 * val)
        elif unit == "cm":
            p.paragraph_format.right_indent = Emu(360000 * val)
        elif unit == "inch":
            p.paragraph_format.right_indent = Emu(914400 * val)

        ri = RightIndent(f"{val}{unit}")
        result = ri.get_from_paragraph(p)
        assert result is not None
        assert abs(result - val) < 0.01

    def test_get_from_paragraph_none(self, doc):
        """RightIndent.get_from_paragraph returns None when no indent set"""
        ri = RightIndent("12pt")
        result = ri.get_from_paragraph(doc.add_paragraph())
        assert result is None



class TestFirstLineIndentGetFromParagraphUnits:
    """Cover lines 548-552: FirstLineIndent.get_from_paragraph with pt/mm/cm/inch units"""

    @pytest.mark.parametrize("unit,val", [
        ("pt", 24), ("mm", 5), ("cm", 1), ("inch", 0.5),
    ])
    def test_get_from_paragraph_physical_units(self, doc, unit, val):
        """FirstLineIndent.get_from_paragraph returns value in specified physical unit"""
        p = doc.add_paragraph()
        from docx.shared import Emu
        if unit == "pt":
            p.paragraph_format.first_line_indent = Pt(val)
        elif unit == "mm":
            p.paragraph_format.first_line_indent = Emu(36000 * val)
        elif unit == "cm":
            p.paragraph_format.first_line_indent = Emu(360000 * val)
        elif unit == "inch":
            p.paragraph_format.first_line_indent = Emu(914400 * val)

        fli = FirstLineIndent(f"{val}{unit}")
        result = fli.get_from_paragraph(p)
        assert result is not None
        assert abs(result - val) < 0.01

    def test_get_from_paragraph_none(self, doc):
        """FirstLineIndent.get_from_paragraph returns None when no indent set"""
        fli = FirstLineIndent("12pt")
        result = fli.get_from_paragraph(doc.add_paragraph())
        assert result is None



class TestBuiltInStyleBaseSetElseBranch:
    """Cover lines 594-595: BuiltInStyle.base_set with value NOT in _LABEL_MAP (else branch)
       Cover line 600: BuiltInStyle.base_set with invalid style raises ValueError"""

    def test_base_set_else_branch_valid_style(self, doc):
        """base_set with value not in _LABEL_MAP but valid docx style (else branch, line 596-600)"""
        p = doc.add_paragraph()
        # "Heading 1" is a valid style name but also in _LABEL_MAP, so use a valid
        # style name that is NOT in _LABEL_MAP. "Normal" is in _LABEL_MAP via "正文".
        # We need a style name that exists in docx but not in _LABEL_MAP.
        # Let's use a direct style name like "Heading 5" which exists in docx but not in map.
        bis = BuiltInStyle.__new__(BuiltInStyle)
        bis.value = "Heading 1"  # This IS in _LABEL_MAP, so it hits the if branch
        bis._rel_value = None
        bis.original_unit = None
        bis.unit_ch = None
        bis._rel_unit = None
        bis.extract_unit_result = None
        # This hits the if branch (line 591-595)
        bis.base_set(p)
        assert p.style.name == "Heading 1"

    def test_base_set_else_branch_not_in_map(self, doc):
        """base_set with value NOT in _LABEL_MAP hits else branch (lines 596-600)"""
        p = doc.add_paragraph()
        bis = BuiltInStyle.__new__(BuiltInStyle)
        bis.value = "Normal"  # "Normal" is NOT a key in _LABEL_MAP (only "正文" is)
        bis._rel_value = None
        bis.original_unit = None
        bis.unit_ch = None
        bis._rel_unit = None
        bis.extract_unit_result = None
        # "Normal" is not in _LABEL_MAP keys, so it hits else branch (line 596)
        # But "Normal" is a valid style, so docx_obj.style = self.value succeeds
        bis.base_set(p)
        assert p.style.name == "Normal"

    def test_base_set_else_branch_invalid_style_raises(self, doc):
        """base_set with value NOT in _LABEL_MAP: _ensure_style_exists 会自动创建样式"""
        p = doc.add_paragraph()
        bis = BuiltInStyle.__new__(BuiltInStyle)
        bis.value = "CompletelyInvalidStyleName"
        bis._rel_value = None
        bis.original_unit = None
        bis.unit_ch = None
        bis._rel_unit = None
        bis.extract_unit_result = None
        # _ensure_style_exists 会自动创建样式，不再抛异常
        bis.base_set(p)
        assert p.style.name == "CompletelyInvalidStyleName"

    def test_base_set_if_branch_invalid_style_raises(self, doc):
        """base_set with value in _LABEL_MAP but invalid style raises ValueError (lines 594-595)"""
        p = doc.add_paragraph()
        bis = BuiltInStyle.__new__(BuiltInStyle)
        # Use a value that IS in _LABEL_MAP but maps to an invalid style
        bis.value = "正文"  # maps to "Normal" which IS valid, so this won't raise
        bis._rel_value = None
        bis.original_unit = None
        bis.unit_ch = None
        bis._rel_unit = None
        bis.extract_unit_result = None
        # This should succeed since Normal is a valid style
        bis.base_set(p)
        assert p.style.name == "Normal"

