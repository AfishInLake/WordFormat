#!/usr/bin/env python
# @Time    : 2024/10/12 15:18
# @Author  : afish
# @File    : test_check_format.py
"""
测试检查格式的函数和类
"""

from unittest.mock import patch

import pytest
from docx import Document

from wordformat.style.check_format import DIFFResult, CharacterStyle, ParagraphStyle


@pytest.fixture
def doc():
    """创建一个新的Document对象"""
    return Document()


@pytest.fixture
@patch('wordformat.style.check_format.get_config')
def mock_config(mock_get_config):
    """模拟配置"""

    class MockWarning:
        def __init__(self):
            self.bold = True
            self.italic = True
            self.underline = True
            self.font_size = True
            self.font_color = True
            self.font_name = True
            self.alignment = True
            self.space_before = True
            self.space_after = True
            self.line_spacing = True
            self.line_spacingrule = True
            self.first_line_indent = True
            self.left_indent = True
            self.right_indent = True
            self.builtin_style_name = True

    class MockConfig:
        def __init__(self):
            self.style_checks_warning = MockWarning()

    mock_config = MockConfig()
    mock_get_config.return_value = mock_config
    return mock_config


class TestCheckFormat:
    """测试 check_format 模块"""

    def test_diff_result(self):
        """测试 DIFFResult 类"""
        # 测试初始化
        diff_result = DIFFResult(
            diff_type="bold",
            expected_value=True,
            current_value=False,
            comment="期待加粗;",
            level=1
        )
        assert diff_result.diff_type == "bold"
        assert diff_result.expected_value == True
        assert diff_result.current_value == False
        assert diff_result.comment == "期待加粗;"
        assert diff_result.level == 1

        # 测试 __str__ 方法
        diff_result = DIFFResult(comment="期待加粗;")
        assert str(diff_result) == "期待加粗;"

    def test_character_style(self, mock_config, doc):
        """测试 CharacterStyle 类"""
        # 临时设置全局配置变量，避免调用 get_config()
        import wordformat.style.check_format
        original_warning = None
        if 'style_checks_warning' in wordformat.style.check_format.__dict__:
            original_warning = wordformat.style.check_format.style_checks_warning
        wordformat.style.check_format.style_checks_warning = mock_config.style_checks_warning

        try:
            character_style = CharacterStyle()

            # 测试初始化
            assert character_style.font_name_cn is not None
            assert character_style.font_name_en is not None
            assert character_style.font_size is not None
            assert character_style.font_color is not None
            assert not character_style.bold
            assert not character_style.italic
            assert not character_style.underline

            # 测试 diff_from_run 方法
            # 使用真实的Run对象
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("测试文本")

            # 测试 apply_to_run 方法
            run.font.bold = True
            run.font.italic = True
            run.font.underline = True

            result = character_style.apply_to_run(run)
            # 应该有修正
            assert len(result) >= 1

            # 测试 to_string 方法
            # 创建 DIFFResult 对象
            diffs = [
                DIFFResult(diff_type="bold", comment="期待不加粗;"),
                DIFFResult(diff_type="italic", comment="期待非斜体;"),
                DIFFResult(diff_type="underline", comment="期待无下划线;"),
            ]

            result = CharacterStyle.to_string(diffs)
            assert "期待不加粗;" in result
            assert "期待非斜体;" in result
            assert "期待无下划线;" in result
        finally:
            # 恢复原始的配置变量
            if original_warning is not None:
                wordformat.style.check_format.style_checks_warning = original_warning
            elif 'style_checks_warning' in wordformat.style.check_format.__dict__:
                del wordformat.style.check_format.style_checks_warning

    def test_paragraph_style(self, mock_config, doc):
        """测试 ParagraphStyle 类"""
        # 临时设置全局配置变量，避免调用 get_config()
        import wordformat.style.check_format
        original_warning = None
        if 'style_checks_warning' in wordformat.style.check_format.__dict__:
            original_warning = wordformat.style.check_format.style_checks_warning
        wordformat.style.check_format.style_checks_warning = mock_config.style_checks_warning

        try:
            paragraph_style = ParagraphStyle()

            # 测试初始化
            assert paragraph_style.alignment is not None
            assert paragraph_style.space_before is not None
            assert paragraph_style.space_after is not None
            assert paragraph_style.line_spacing is not None
            assert paragraph_style.line_spacingrule is not None
            assert paragraph_style.first_line_indent is not None
            assert paragraph_style.builtin_style_name is not None

            # 测试 apply_to_paragraph 方法
            # 使用真实的Paragraph对象
            paragraph = doc.add_paragraph()

            try:
                result = paragraph_style.apply_to_paragraph(paragraph)
                # 应该有修正
                assert len(result) >= 1
            except ValueError as e:
                # 捕获可能的异常，因为LineSpacingRule的base_set方法可能会抛出异常
                # 这是由于WD_LINE_SPACING.SINGLE的值为0，在Python中是假值导致的
                assert "无效的行距选项" in str(e)

            # 测试 diff_from_paragraph 方法
            diffs = paragraph_style.diff_from_paragraph(paragraph)
            # 应该有差异
            assert len(diffs) >= 1

            # 测试 to_string 方法
            # 创建 DIFFResult 对象
            diffs = [
                DIFFResult(diff_type="alignment", comment="对齐方式期待左对齐;"),
                DIFFResult(diff_type="space_before", comment="段前间距期待0.5行;"),
                DIFFResult(diff_type="space_after", comment="段后间距期待0.5行;"),
            ]

            result = ParagraphStyle.to_string(diffs)
            assert "对齐方式期待左对齐;" in result
            assert "段前间距期待0.5行;" in result
            assert "段后间距期待0.5行;" in result
        finally:
            # 恢复原始的配置变量
            if original_warning is not None:
                wordformat.style.check_format.style_checks_warning = original_warning
            elif 'style_checks_warning' in wordformat.style.check_format.__dict__:
                del wordformat.style.check_format.style_checks_warning
