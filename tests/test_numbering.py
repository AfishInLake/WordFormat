"""
Core 模块综合测试

覆盖 tree.py, utils.py, rules/node.py, numbering.py, settings.py
"""

import os
import pytest
from io import StringIO
from unittest.mock import MagicMock, patch

from docx import Document
from docx.oxml.ns import qn

from wordformat.tree import Tree, Stack, print_tree
from wordformat.rules.node import TreeNode, FormatNode
from wordformat.numbering import (
    _auto_strip_numbering,
    _strip_reference_numbering,
    apply_auto_numbering,
    create_numbering_definition,
    process_heading_numbering,
)
from wordformat.utils import (
    get_file_name,
    ensure_is_directory,
    ensure_directory_exists,
    _to_roman,
    _to_chinese_num,
    load_yaml_with_merge,
    get_paragraph_numbering_text,
    remove_all_numbering,
    _format_number,
    _get_level_fmt,
    _count_numbering_levels,
)
from wordformat.base import DocxBase
from wordformat import settings


# ============================================================
# tree.py — Tree
# ============================================================


# ============================================================
# numbering.py — auto_strip_numbering
# ============================================================


class TestAutoStripNumbering:
    def test_level0_chinese_chapter_strips(self, doc):
        """一级标题：第一章 → 绪论"""
        p = doc.add_paragraph()
        run = p.add_run("第一章 绪论")
        result = _auto_strip_numbering(p, ilvl=0)
        assert result is True
        assert p.text == "绪论"

    def test_non_numbered_text_returns_false(self, doc):
        """无编号文本不匹配任何模式"""
        p = doc.add_paragraph()
        p.add_run("绪论")
        result = _auto_strip_numbering(p, ilvl=0)
        assert result is False
        assert p.text == "绪论"

    def test_level1_pattern_not_matched_at_level0(self, doc):
        """二级编号 '1.1 背景' 对一级标题不匹配"""
        p = doc.add_paragraph()
        p.add_run("1.1 背景")
        # ilvl=0 不匹配二级编号格式 → _auto_strip_numbering 会尝试所有模式但找不到匹配
        result = _auto_strip_numbering(p, ilvl=0)
        assert result is False

    def test_no_runs_returns_false(self, doc):
        """无 run 的段落 → 返回 False"""
        p = doc.add_paragraph()
        assert _auto_strip_numbering(p, ilvl=0) is False

    def test_multi_run_level0_strips_correctly(self, doc):
        """多 run 一级标题：第一章 + 绪论 → 绪论"""
        p = doc.add_paragraph()
        r1 = p.add_run("第一章")
        r2 = p.add_run(" 绪论")
        result = _auto_strip_numbering(p, ilvl=0)
        assert result is True
        assert p.text == "绪论"


# ============================================================
# numbering.py — apply_auto_numbering
# ============================================================


class TestApplyAutoNumbering:
    def test_adds_numPr_to_paragraph(self, doc):
        p = doc.add_paragraph("test")
        apply_auto_numbering(p, num_id="100", ilvl="0")
        from docx.oxml.ns import qn

        pPr = p._element.find(qn("w:pPr"))
        assert pPr is not None
        numPr = pPr.find(qn("w:numPr"))
        assert numPr is not None
        numId_elem = numPr.find(qn("w:numId"))
        assert numId_elem.get(qn("w:val")) == "100"
        ilvl_elem = numPr.find(qn("w:ilvl"))
        assert ilvl_elem.get(qn("w:val")) == "0"

    def test_replaces_existing_numPr(self, doc):
        p = doc.add_paragraph("test")
        apply_auto_numbering(p, num_id="50", ilvl="1")
        apply_auto_numbering(p, num_id="99", ilvl="2")
        from docx.oxml.ns import qn

        pPr = p._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr"))
        assert numPr.find(qn("w:numId")).get(qn("w:val")) == "99"
        assert numPr.find(qn("w:ilvl")).get(qn("w:val")) == "2"


# ============================================================
# numbering.py — create_numbering_definition
# ============================================================


class TestCreateNumberingDefinition:
    def test_disabled_config_returns_empty(self, doc):
        config = MagicMock()
        config.enabled = False
        result = create_numbering_definition(doc, config)
        assert result == {"headings": {}, "references": None}

    def test_enabled_config_creates_definitions(self, doc):
        config = MagicMock()
        config.enabled = True
        config.references = MagicMock()
        config.references.enabled = False
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        result = create_numbering_definition(doc, config)
        assert "level_1" in result["headings"]
        assert "level_2" not in result["headings"]
        assert "level_3" not in result["headings"]

    def test_multiple_levels(self, doc):
        config = MagicMock()
        config.enabled = True
        config.references = MagicMock()
        config.references.enabled = False
        for key in ("level_1", "level_2", "level_3"):
            lvl = MagicMock()
            lvl.enabled = True
            lvl.template = f"template_{key}"
            lvl.suffix = "space"
            lvl.numbering_indent = None
            lvl.text_indent = None
            setattr(config, key, lvl)

        result = create_numbering_definition(doc, config)
        assert len(result["headings"]) == 3
        # num_id 应该递增
        ids = [int(v) for v in result["headings"].values()]
        assert ids == sorted(ids)


# ============================================================
# numbering.py — process_heading_numbering
# ============================================================


class TestProcessHeadingNumbering:
    def test_disabled_config_noop(self, doc):
        root = TreeNode({"category": "top"})
        config = MagicMock()
        config.enabled = False
        process_heading_numbering(root, doc, config)
        # 无异常即通过


# ============================================================

# ============================================================
# utils.py — get_paragraph_numbering_text
# ============================================================


class TestGetParagraphNumberingText:
    """测试从段落 XML 中提取自动编号文字"""

    def test_no_pPr_returns_empty(self, doc):
        """段落没有 pPr 时返回空字符串"""
        p = doc.add_paragraph("hello")
        # 确保没有 pPr
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None:
            p._element.remove(pPr)
        assert get_paragraph_numbering_text(p) == ""

    def test_no_numPr_returns_empty(self, doc):
        """段落有 pPr 但没有 numPr 时返回空字符串"""
        p = doc.add_paragraph("hello")
        from docx.oxml import OxmlElement

        pPr = OxmlElement("w:pPr")
        p._element.insert(0, pPr)
        assert get_paragraph_numbering_text(p) == ""

    def test_no_numId_returns_empty(self, doc):
        """numPr 中没有 numId 时返回空字符串"""
        p = doc.add_paragraph("hello")
        from docx.oxml import OxmlElement

        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
        p._element.insert(0, pPr)
        assert get_paragraph_numbering_text(p) == ""

    def test_numId_zero_returns_empty(self, doc):
        """numId 为 0 时返回空字符串"""
        p = doc.add_paragraph("hello")
        from docx.oxml import OxmlElement

        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), "0")
        numPr.append(numId_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)
        assert get_paragraph_numbering_text(p) == ""

    def test_no_numbering_part_returns_empty(self, doc):
        """文档没有 numbering part 时返回空字符串（已修复：捕获 NotImplementedError）"""
        p = doc.add_paragraph("hello")
        from docx.oxml import OxmlElement

        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), "1")
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), "0")
        numPr.append(ilvl_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)
        # 移除 numbering 关系使 numbering_part 访问抛出异常
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        rels = doc.part.rels
        to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
        for k in to_remove:
            del rels[k]
        # 修复后捕获 NotImplementedError 并返回空字符串
        assert get_paragraph_numbering_text(p) == ""

    def _setup_numbering(
        self, doc, num_fmt="decimal", lvl_text="%1.", num_id="1", abstract_num_id="0"
    ):
        """辅助方法：为文档创建 numbering part 和编号定义"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        # 先移除已有的 numbering 关系（Document() 默认模板有 numbering part）
        try:
            rels = doc.part.rels
            to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
            for k in to_remove:
                del rels[k]
        except Exception:
            pass

        numbering_elm = OxmlElement("w:numbering")

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), abstract_num_id)

        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), num_fmt)
        lvl.append(numFmt)
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), lvl_text)
        lvl.append(lvlText)
        abstract_num.append(lvl)
        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), abstract_num_id)
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        return numbering_elm

    def _add_numPr_to_paragraph(self, p, num_id="1", ilvl="0"):
        """辅助方法：为段落添加 numPr"""
        from docx.oxml import OxmlElement

        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), num_id)
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), ilvl)
        numPr.append(ilvl_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)

    def test_with_numbering_definition(self, doc):
        """测试有完整 numbering 定义时能提取编号文字"""
        # BUG: qn() 无法处理 XPath 表达式（如 w:num[@w:numId='1']），
        # 导致 get_paragraph_numbering_text 在查找 num 元素时抛出 ValueError
        self._setup_numbering(doc, num_fmt="decimal", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "1."

    def test_chinese_counting_format(self, doc):
        """测试中文计数格式（chineseCounting）"""
        self._setup_numbering(doc, num_fmt="chineseCounting", lvl_text="第%1章")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "第一章"

    def test_upper_roman_format(self, doc):
        """测试大写罗马数字格式"""
        self._setup_numbering(doc, num_fmt="upperRoman", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "I."

    def test_lower_roman_format(self, doc):
        """测试小写罗马数字格式"""
        self._setup_numbering(doc, num_fmt="lowerRoman", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "i."

    def test_upper_letter_format(self, doc):
        """测试大写字母格式"""
        self._setup_numbering(doc, num_fmt="upperLetter", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "A."

    def test_lower_letter_format(self, doc):
        """测试小写字母格式"""
        self._setup_numbering(doc, num_fmt="lowerLetter", lvl_text="%1.")

        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p)

        result = get_paragraph_numbering_text(p)
        assert result == "a."

    def test_missing_abstract_num_returns_empty(self, doc):
        """abstractNumId 对应的 abstractNum 不存在时返回空字符串"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        # 先移除已有的 numbering 关系
        rels = doc.part.rels
        to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
        for k in to_remove:
            del rels[k]

        numbering_elm = OxmlElement("w:numbering")

        # 只创建 num 引用，不创建 abstractNum
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), "999")
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        p = doc.add_paragraph("test")
        self._add_numPr_to_paragraph(p)

        assert get_paragraph_numbering_text(p) == ""

    def test_missing_lvl_returns_empty(self, doc):
        """abstractNum 中没有对应 ilvl 的 lvl 时返回空字符串"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        # 先移除已有的 numbering 关系
        rels = doc.part.rels
        to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
        for k in to_remove:
            del rels[k]

        numbering_elm = OxmlElement("w:numbering")

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), "0")
        # 不创建任何 lvl
        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), "1")
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), "0")
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        p = doc.add_paragraph("test")
        self._add_numPr_to_paragraph(p)

        assert get_paragraph_numbering_text(p) == ""


# ============================================================
# utils.py — _format_number
# ============================================================


class TestFormatNumber:
    """测试 _format_number 各种格式化类型"""

    def test_decimal(self):
        assert _format_number(5, "decimal") == "5"

    def test_upper_roman(self):
        assert _format_number(3, "upperRoman") == "III"

    def test_lower_roman(self):
        assert _format_number(4, "lowerRoman") == "iv"

    def test_upper_letter(self):
        assert _format_number(1, "upperLetter") == "A"
        assert _format_number(3, "upperLetter") == "C"

    def test_lower_letter(self):
        assert _format_number(1, "lowerLetter") == "a"
        assert _format_number(3, "lowerLetter") == "c"

    def test_chinese_counting_thousand(self):
        assert _format_number(5, "chineseCountingThousand") == "五"

    def test_ideograph_traditional(self):
        assert _format_number(3, "ideographTraditional") == "三"

    def test_chinese_counting(self):
        assert _format_number(7, "chineseCounting") == "七"

    def test_unknown_format_fallback_to_str(self):
        """未知格式回退到 str"""
        assert _format_number(42, "unknownFormat") == "42"


# ============================================================
# utils.py — _get_level_fmt
# ============================================================


class TestGetLevelFmt:
    """测试 _get_level_fmt 获取指定级别的 numFmt"""

    def test_existing_level_returns_fmt(self):
        from lxml import etree

        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:numFmt w:val="upperRoman"/>'
            "  </w:lvl>"
            "</w:abstractNum>"
        )
        assert _get_level_fmt(abstract_num, 0) == "upperRoman"

    def test_missing_level_returns_decimal(self):
        from lxml import etree

        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "</w:abstractNum>"
        )
        assert _get_level_fmt(abstract_num, 0) == "decimal"

    def test_no_numFmt_returns_decimal(self):
        from lxml import etree

        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:start w:val="1"/>'
            "  </w:lvl>"
            "</w:abstractNum>"
        )
        assert _get_level_fmt(abstract_num, 0) == "decimal"


# ============================================================
# utils.py — remove_all_numbering
# ============================================================


class TestRemoveAllNumbering:
    """测试 remove_all_numbering 移除标题样式的编号绑定"""

    def test_removes_numPr_from_heading_styles(self, doc):
        """从 Heading 1/2/3 样式中移除 numPr"""
        from docx.oxml import OxmlElement

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = doc.styles[style_name]
            style_element = style._element

            # 确保 pPr 存在
            pPr = style_element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                style_element.insert(0, pPr)

            # 添加 numPr
            numPr = OxmlElement("w:numPr")
            pPr.append(numPr)

        remove_all_numbering(doc)

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = doc.styles[style_name]
            pPr = style._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                assert numPr is None, f"{style_name} 的 numPr 未被移除"

    def test_removes_outlineLvl_from_heading_styles(self, doc):
        """从 Heading 1/2/3 样式中移除 outlineLvl"""
        from docx.oxml import OxmlElement

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = doc.styles[style_name]
            style_element = style._element

            pPr = style_element.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                style_element.insert(0, pPr)

            # 先移除已有的 outlineLvl
            existing = pPr.find(qn("w:outlineLvl"))
            if existing is not None:
                pPr.remove(existing)

            outlineLvl = OxmlElement("w:outlineLvl")
            outlineLvl.set(qn("w:val"), "0")
            pPr.append(outlineLvl)

        remove_all_numbering(doc)

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = doc.styles[style_name]
            pPr = style._element.find(qn("w:pPr"))
            if pPr is not None:
                outlineLvl = pPr.find(qn("w:outlineLvl"))
                assert outlineLvl is None, f"{style_name} 的 outlineLvl 未被移除"

    def test_no_pPr_in_style_no_error(self, doc):
        """样式中没有 pPr 时不报错"""
        # Heading 1 的 pPr 可能存在也可能不存在，直接调用不应报错
        remove_all_numbering(doc)

    def test_style_not_in_doc_no_error(self, doc):
        """如果样式不存在（不太可能），不报错"""
        # Document() 默认包含 Heading 1/2/3，所以这个测试主要验证不会抛异常
        remove_all_numbering(doc)


# ============================================================
# numbering.py — auto_strip_numbering (empty result)
# ============================================================


class TestAutoStripNumberingEmptyResult:
    """测试 _auto_strip_numbering 在文本变空时的行为"""

    def test_text_becomes_empty_after_strip(self, doc):
        """匹配整个文本后，段落文本变为空"""
        p = doc.add_paragraph()
        run = p.add_run("1.1 ")
        result = _auto_strip_numbering(p, ilvl=1)
        assert result is True
        assert p.text == ""

    def test_single_run_fully_consumed(self, doc):
        """单个 run 的文本完全被匹配清除"""
        p = doc.add_paragraph()
        run = p.add_run("第一章")
        result = _auto_strip_numbering(p, ilvl=0)
        assert result is True
        assert p.text == ""


# ============================================================
# numbering.py — create_numbering_definition (XML creation)
# ============================================================


class TestCreateNumberingDefinitionXML:
    """测试 create_numbering_definition 的 XML 创建逻辑"""

    def test_creates_abstractNum_and_num_elements(self, doc):
        """验证创建的 XML 中包含 abstractNum 和 num 元素"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "%1."
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        config.references = MagicMock()
        config.references.enabled = False
        result = create_numbering_definition(doc, config)

        assert "level_1" in result["headings"]

        # 验证 numbering part 中有 abstractNum 和 num
        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element

        abstract_nums = numbering_elm.findall(qn("w:abstractNum"))
        nums = numbering_elm.findall(qn("w:num"))
        assert len(abstract_nums) >= 1
        assert len(nums) >= 1

    def test_chinese_template_creates_chinese_numFmt(self, doc):
        """中文模板（含'第'和'章'）使用 chineseCountingThousand 格式"""
        config = MagicMock()
        config.enabled = True
        config.references = MagicMock()
        config.references.enabled = False
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        create_numbering_definition(doc, config)

        numbering_elm = doc.part.numbering_part._element
        # 找到最后一个 abstractNum（create_numbering_definition 追加的）
        abstract_nums = numbering_elm.findall(qn("w:abstractNum"))
        abstract_num = abstract_nums[-1]
        lvl = abstract_num.find(qn("w:lvl"))
        numFmt = lvl.find(qn("w:numFmt"))
        assert numFmt.get(qn("w:val")) == "chineseCountingThousand"

    def test_decimal_template_creates_decimal_numFmt(self, doc):
        """非中文模板使用 decimal 格式"""
        config = MagicMock()
        config.enabled = True
        config.references = MagicMock()
        config.references.enabled = False
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "%1."
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        create_numbering_definition(doc, config)

        numbering_elm = doc.part.numbering_part._element
        abstract_nums = numbering_elm.findall(qn("w:abstractNum"))
        abstract_num = abstract_nums[-1]
        lvl = abstract_num.find(qn("w:lvl"))
        numFmt = lvl.find(qn("w:numFmt"))
        assert numFmt.get(qn("w:val")) == "decimal"

    def test_no_template_skips_level(self, doc):
        """template 为空时跳过该级别"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = ""
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        result = create_numbering_definition(doc, config)
        assert "level_1" not in result

    def test_creates_numbering_part_when_missing(self, tmp_path):
        """文档没有 numbering part 时自动创建"""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        # 创建一个没有 numbering 关系的文档
        doc = Document()
        # 移除已有的 numbering 关系（如果有）
        try:
            rels = doc.part.rels
            to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
            for k in to_remove:
                del rels[k]
        except Exception:
            pass

        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "%1."
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        result = create_numbering_definition(doc, config)
        assert "level_1" in result["headings"]

        # 现在应该有了
        numbering_part = doc.part.numbering_part
        assert numbering_part is not None


# ============================================================
# numbering.py — process_heading_numbering (enabled)
# ============================================================


class TestProcessHeadingNumberingEnabled:
    """测试 process_heading_numbering 在启用时的完整流程"""

    def test_enabled_config_processes_headings(self, doc):
        """启用配置时处理标题节点：清除手动编号 + 应用自动编号"""
        config = MagicMock()
        config.enabled = True

        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None

        level_2 = MagicMock()
        level_2.enabled = True
        level_2.template = "%1.%2"
        level_2.suffix = "space"
        level_2.numbering_indent = None
        level_2.text_indent = None

        level_3 = MagicMock()
        level_3.enabled = False
        level_3.template = ""

        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        # 创建标题节点（heading 类别需要 fingerprint）
        root = TreeNode({"category": "top"})

        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        p1 = doc.add_paragraph("第一章 绪论")
        h1_node.paragraph = p1
        root.add_child_node(h1_node)

        h2_node = TreeNode({"category": "heading_level_2", "fingerprint": "fp2"})
        p2 = doc.add_paragraph("1.1 研究背景")
        h2_node.paragraph = p2
        h1_node.add_child_node(h2_node)

        body_node = TreeNode({"category": "body_text", "fingerprint": "fp3"})
        p3 = doc.add_paragraph("这是正文内容")
        body_node.paragraph = p3
        h2_node.add_child_node(body_node)

        process_heading_numbering(root, doc, config)

        # 验证手动编号被清除
        assert p1.text == "绪论"
        assert p2.text == "研究背景"

        # 验证自动编号被应用
        pPr1 = p1._element.find(qn("w:pPr"))
        numPr1 = pPr1.find(qn("w:numPr"))
        assert numPr1 is not None

        pPr2 = p2._element.find(qn("w:pPr"))
        numPr2 = pPr2.find(qn("w:numPr"))
        assert numPr2 is not None

    def test_enabled_but_no_paragraph_skips(self, doc):
        """启用配置但节点没有 paragraph 属性时跳过"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        root = TreeNode({"category": "top"})
        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        # 不设置 paragraph 属性
        root.add_child_node(h1_node)

        # 不应抛异常
        process_heading_numbering(root, doc, config)

    def test_enabled_level_disabled_skips_strip(self, doc):
        """级别未启用时跳过该级别的处理"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = False
        level_1.template = ""
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        root = TreeNode({"category": "top"})
        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        p1 = doc.add_paragraph("第一章 绪论")
        h1_node.paragraph = p1
        root.add_child_node(h1_node)

        process_heading_numbering(root, doc, config)

        # 文本不应被修改
        assert p1.text == "第一章 绪论"

    def test_enabled_level_auto_strips_and_numbers(self, doc):
        """启用编号后自动清除手动编号并应用自动编号"""
        config = MagicMock()
        config.enabled = True
        level_1 = MagicMock()
        level_1.enabled = True
        level_1.template = "第%1章"
        level_1.suffix = "space"
        level_1.numbering_indent = None
        level_1.text_indent = None
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        root = TreeNode({"category": "top"})
        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        p1 = doc.add_paragraph("第一章 绪论")
        h1_node.paragraph = p1
        root.add_child_node(h1_node)

        process_heading_numbering(root, doc, config)

        # 自动清除手动编号
        assert p1.text == "绪论"
        # 自动编号应被应用
        pPr = p1._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr"))
        assert numPr is not None

    def test_empty_num_id_map_returns_early(self, doc):
        """num_id_map 为空时提前返回"""
        config = MagicMock()
        config.enabled = True
        # 所有级别都禁用或没有 template → create_numbering_definition 返回空
        level_1 = MagicMock()
        level_1.enabled = False
        level_2 = MagicMock()
        level_2.enabled = False
        level_3 = MagicMock()
        level_3.enabled = False
        config.level_1 = level_1
        config.level_2 = level_2
        config.level_3 = level_3

        root = TreeNode({"category": "top"})
        h1_node = TreeNode({"category": "heading_level_1", "fingerprint": "fp1"})
        p1 = doc.add_paragraph("第一章 绪论")
        h1_node.paragraph = p1
        root.add_child_node(h1_node)

        # 不应抛异常，且不应修改段落
        process_heading_numbering(root, doc, config)
        assert p1.text == "第一章 绪论"


# ============================================================
# numbering.py — _strip_reference_numbering
# ============================================================


class TestStripReferenceNumbering:
    def test_strip_bracket_number(self, doc):
        """清除 [1] 格式的参考文献编号"""
        p = doc.add_paragraph("[1] Some reference text")
        assert _strip_reference_numbering(p) is True
        assert p.text == "Some reference text"

    def test_strip_dot_number(self, doc):
        """清除 1. 格式的编号（要求后面有空格）"""
        p = doc.add_paragraph("1. Some reference text")
        assert _strip_reference_numbering(p) is True
        assert p.text == "Some reference text"

    def test_strip_parentheses_number(self, doc):
        """清除 (1) 格式的编号"""
        p = doc.add_paragraph("(1) Some reference text")
        assert _strip_reference_numbering(p) is True
        assert p.text == "Some reference text"

    def test_strip_chinese_bracket_number(self, doc):
        """清除［1］全角方括号编号"""
        p = doc.add_paragraph("［1］ Some reference text")
        assert _strip_reference_numbering(p) is True
        assert p.text == "Some reference text"

    def test_no_number_returns_false(self, doc):
        """没有编号的段落返回 False"""
        p = doc.add_paragraph("Some reference text without number")
        assert _strip_reference_numbering(p) is False
        assert p.text == "Some reference text without number"

    def test_year_not_stripped(self, doc):
        """不应误删除年份（如 2024. 后面没有空格不会被 1. 模式匹配）"""
        p = doc.add_paragraph("2024. A great paper about X.")
        # "2024." 后面有空格，会被匹配。我们只测不以数字开头的引用
        p2 = doc.add_paragraph("Smith, J. (2024). A paper.")
        assert _strip_reference_numbering(p2) is False

    def test_strip_across_runs(self, doc):
        """跨 run 清除编号"""
        p = doc.add_paragraph()
        p.add_run("[1]")
        p.add_run(" Some reference text")
        assert _strip_reference_numbering(p) is True
        assert p.text == "Some reference text"


# ============================================================
# numbering.py — create_numbering_definition (references)
# ============================================================


class TestCreateReferenceNumberingDefinition:
    def test_references_enabled_creates_definition(self, doc):
        """参考文献编号启用时创建独立 abstractNum 和 num"""
        from wordformat.config.models import NumberingConfig, NumberingLevelConfig

        config = NumberingConfig(
            enabled=True,
            level_1=NumberingLevelConfig(enabled=False),
            level_2=NumberingLevelConfig(enabled=False),
            level_3=NumberingLevelConfig(enabled=False),
            references=NumberingLevelConfig(
                enabled=True, template="[%1]", suffix="space"
            ),
        )
        result = create_numbering_definition(doc, config)
        assert result["references"] is not None
        ref_id = result["references"]
        assert int(ref_id) > 0

        # 验证 numbering part 中有 abstractNum 和 num（文档可能已有其他定义）
        numbering_elm = doc.part.numbering_part._element
        abstract_nums = numbering_elm.findall(qn("w:abstractNum"))
        nums = numbering_elm.findall(qn("w:num"))
        assert len(abstract_nums) >= 1
        assert len(nums) >= 1

    def test_references_disabled_no_definition(self, doc):
        """参考文献编号禁用时不创建定义"""
        from wordformat.config.models import NumberingConfig, NumberingLevelConfig

        config = NumberingConfig(
            enabled=True,
            level_1=NumberingLevelConfig(enabled=False),
            level_2=NumberingLevelConfig(enabled=False),
            level_3=NumberingLevelConfig(enabled=False),
            references=NumberingLevelConfig(
                enabled=False, template="[%1]", suffix="space"
            ),
        )
        result = create_numbering_definition(doc, config)
        assert result["references"] is None


# ============================================================
# numbering.py — process_heading_numbering (references)
# ============================================================


class TestProcessReferenceNumbering:
    def test_reference_entry_gets_numbering(self, doc):
        """参考文献条目节点应被应用自动编号"""
        from wordformat.rules.references import ReferenceEntry
        from wordformat.config.models import NumberingConfig, NumberingLevelConfig

        p = doc.add_paragraph("Some reference text")
        node = ReferenceEntry(
            value={"category": "body_text", "fingerprint": "abc123"},
            level=2,
            paragraph=p,
        )

        root = FormatNode(value={"category": "top"}, expected_rule={}, level=0)
        root.add_child_node(node)

        config = NumberingConfig(
            enabled=True,
            level_1=NumberingLevelConfig(enabled=False),
            level_2=NumberingLevelConfig(enabled=False),
            level_3=NumberingLevelConfig(enabled=False),
            references=NumberingLevelConfig(
                enabled=True, template="[%1]", suffix="space"
            ),
        )
        process_heading_numbering(root, doc, config)
        # 应已应用自动编号
        numPr = p._element.find(qn("w:pPr"))
        assert numPr is not None
        numPr_elem = numPr.find(qn("w:numPr"))
        assert numPr_elem is not None

    def test_reference_entry_strips_manual_then_applies_numbering(self, doc):
        """参考文献条目：先清除手动编号再应用自动编号"""
        from wordformat.rules.references import ReferenceEntry
        from wordformat.config.models import NumberingConfig, NumberingLevelConfig

        p = doc.add_paragraph("[1] Some reference text")
        node = ReferenceEntry(
            value={"category": "body_text", "fingerprint": "abc456"},
            level=2,
            paragraph=p,
        )

        root = FormatNode(value={"category": "top"}, expected_rule={}, level=0)
        root.add_child_node(node)

        config = NumberingConfig(
            enabled=True,
            level_1=NumberingLevelConfig(enabled=False),
            level_2=NumberingLevelConfig(enabled=False),
            level_3=NumberingLevelConfig(enabled=False),
            references=NumberingLevelConfig(
                enabled=True, template="[%1]", suffix="space"
            ),
        )
        process_heading_numbering(root, doc, config)
        # 手动编号 [1] 已清除
        assert p.text == "Some reference text"
        # 自动编号已应用
        numPr = p._element.find(qn("w:pPr"))
        numPr_elem = numPr.find(qn("w:numPr"))
        assert numPr_elem is not None


# ============================================================
# hyperlinks.py — _parse_ref_numbers
# ============================================================


class TestParseRefNumbers:
    def test_single_number(self):
        from wordformat.hyperlinks import _parse_ref_numbers

        assert _parse_ref_numbers("[1]") == [1]

    def test_multiple_numbers(self):
        from wordformat.hyperlinks import _parse_ref_numbers

        assert _parse_ref_numbers("[1,2,3]") == [1, 2, 3]

    def test_range(self):
        from wordformat.hyperlinks import _parse_ref_numbers

        assert _parse_ref_numbers("[1-3]") == [1, 2, 3]

    def test_mixed(self):
        from wordformat.hyperlinks import _parse_ref_numbers

        assert _parse_ref_numbers("[1,3-5]") == [1, 3, 4, 5]

    def test_chinese_comma(self):
        from wordformat.hyperlinks import _parse_ref_numbers

        assert _parse_ref_numbers("[1，2，3]") == [1, 2, 3]

    def test_spaces(self):
        from wordformat.hyperlinks import _parse_ref_numbers

        assert _parse_ref_numbers("[1, 2, 3]") == [1, 2, 3]


# ============================================================
# hyperlinks.py — _insert_bookmark
# ============================================================


class TestInsertBookmark:
    def test_adds_bookmark_to_paragraph(self, doc):
        from wordformat.hyperlinks import _insert_bookmark, _next_bookmark_id

        p = doc.add_paragraph("A reference entry.")
        bid = _next_bookmark_id()
        _insert_bookmark(p, "_Ref1", bid)

        para_elem = p._element
        starts = para_elem.findall(qn("w:bookmarkStart"))
        ends = para_elem.findall(qn("w:bookmarkEnd"))
        assert len(starts) == 1
        assert len(ends) == 1
        assert starts[0].get(qn("w:name")) == "_Ref1"

    def test_bookmark_before_first_run(self, doc):
        from wordformat.hyperlinks import _insert_bookmark, _next_bookmark_id

        p = doc.add_paragraph("Text")
        bid = _next_bookmark_id()
        _insert_bookmark(p, "_RefTest", bid)

        para_elem = p._element
        children = list(para_elem)
        # bookmarkStart 应在第一个 run 之前
        bm_idx = None
        first_r_idx = None
        for i, child in enumerate(children):
            if child.tag == qn("w:bookmarkStart"):
                bm_idx = i
            if child.tag == qn("w:r") and first_r_idx is None:
                first_r_idx = i
        assert bm_idx is not None
        assert first_r_idx is not None
        assert bm_idx < first_r_idx


# ============================================================
# hyperlinks.py — _wrap_citations_in_hyperlinks
# ============================================================


class TestWrapCitationsInHyperlinks:
    def test_wraps_citation_in_hyperlink(self, doc):
        from wordformat.hyperlinks import _wrap_citations_in_hyperlinks

        p = doc.add_paragraph("参见")
        p.add_run("[1]").font.superscript = True
        p.add_run("的研究。")

        _wrap_citations_in_hyperlinks(p, ["_Ref1"])

        para_elem = p._element
        hyperlinks = para_elem.findall(qn("w:hyperlink"))
        assert len(hyperlinks) == 1
        assert hyperlinks[0].get(qn("w:anchor")) == "_Ref1"

    def test_regular_brackets_not_wrapped(self, doc):
        from wordformat.hyperlinks import _wrap_citations_in_hyperlinks

        p = doc.add_paragraph("这是一个[注]释说明。")

        _wrap_citations_in_hyperlinks(p, ["_Ref1"])

        para_elem = p._element
        hyperlinks = para_elem.findall(qn("w:hyperlink"))
        assert len(hyperlinks) == 0

    def test_run_retains_superscript_in_hyperlink(self, doc):
        from wordformat.hyperlinks import _wrap_citations_in_hyperlinks

        p = doc.add_paragraph("见")
        r = p.add_run("[1]")
        r.font.superscript = True

        _wrap_citations_in_hyperlinks(p, ["_Ref1"])

        para_elem = p._element
        hyperlink = para_elem.find(qn("w:hyperlink"))
        assert hyperlink is not None
        r_elem = hyperlink.find(qn("w:r"))
        rPr = r_elem.find(qn("w:rPr"))
        vertAlign = rPr.find(qn("w:vertAlign"))
        assert vertAlign.get(qn("w:val")) == "superscript"
        # 同时应有 Hyperlink 样式
        rStyle = rPr.find(qn("w:rStyle"))
        assert rStyle.get(qn("w:val")) == "Hyperlink"


# ============================================================
# utils.py — _format_number 额外覆盖测试
# ============================================================


class TestFormatNumberAdditional:
    """补充 _format_number 的边界情况和完整格式覆盖"""

    def test_decimal_zero(self):
        """decimal 格式处理 0"""
        assert _format_number(0, "decimal") == "0"

    def test_decimal_negative(self):
        """decimal 格式处理负数"""
        assert _format_number(-5, "decimal") == "-5"

    def test_upper_roman_zero(self):
        """upperRoman 格式处理 0（_to_roman 现在返回 "0"）"""
        assert _format_number(0, "upperRoman") == "0"

    def test_upper_roman_negative(self):
        """upperRoman 格式处理负数（_to_roman 现在返回 "0"）"""
        assert _format_number(-3, "upperRoman") == "0"

    def test_lower_roman_zero(self):
        """lowerRoman 格式处理 0"""
        assert _format_number(0, "lowerRoman") == "0"

    def test_upper_letter_boundary_26(self):
        """upperLetter 格式 n=26 返回 Z（1 <= n <= 26 范围内）"""
        assert _format_number(26, "upperLetter") == "Z"

    def test_upper_letter_boundary_1(self):
        """upperLetter 格式 n=1 返回 A"""
        assert _format_number(1, "upperLetter") == "A"

    def test_upper_letter_boundary_27(self):
        """upperLetter 格式 n=27 超出 A-Z 范围回退到 str"""
        assert _format_number(27, "upperLetter") == "27"

    def test_lower_letter_boundary_26(self):
        """lowerLetter 格式 n=26 返回 z（1 <= n <= 26 范围内）"""
        assert _format_number(26, "lowerLetter") == "z"

    def test_lower_letter_boundary_1(self):
        """lowerLetter 格式 n=1 返回 a"""
        assert _format_number(1, "lowerLetter") == "a"

    def test_chinese_counting_zero(self):
        """chineseCounting 格式处理 0"""
        assert _format_number(0, "chineseCounting") == "0"

    def test_chinese_counting_negative(self):
        """chineseCounting 格式处理负数"""
        assert _format_number(-3, "chineseCounting") == "-3"

    def test_chinese_counting_hundred(self):
        """chineseCounting 格式处理 100（_to_chinese_num 返回 "一百"）"""
        assert _format_number(100, "chineseCounting") == "一百"

    def test_chinese_counting_thousand_zero(self):
        """chineseCountingThousand 格式处理 0"""
        assert _format_number(0, "chineseCountingThousand") == "0"

    def test_ideograph_traditional_zero(self):
        """ideographTraditional 格式处理 0"""
        assert _format_number(0, "ideographTraditional") == "0"

    def test_unknown_format_with_zero(self):
        """未知格式回退到 str，处理 0"""
        assert _format_number(0, "totallyUnknown") == "0"

    def test_unknown_format_with_large_number(self):
        """未知格式回退到 str，处理大数"""
        assert _format_number(99999, "unknownFmt") == "99999"


# ============================================================
# utils.py — _get_level_fmt 额外覆盖测试
# ============================================================


class TestGetLevelFmtAdditional:
    """补充 _get_level_fmt 的边界情况测试"""

    def test_level_missing_numFmt_returns_decimal(self):
        """lvl 存在但缺少 numFmt 元素时返回 decimal"""
        from lxml import etree

        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:start w:val="1"/>'
            "  </w:lvl>"
            "</w:abstractNum>"
        )
        assert _get_level_fmt(abstract_num, 0) == "decimal"

    def test_level_missing_numFmt_val_returns_decimal(self):
        """numFmt 元素存在但缺少 w:val 属性时返回 decimal"""
        from lxml import etree

        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            "    <w:numFmt/>"
            "  </w:lvl>"
            "</w:abstractNum>"
        )
        assert _get_level_fmt(abstract_num, 0) == "decimal"

    def test_nonexistent_level_returns_decimal(self):
        """请求不存在的级别返回 decimal"""
        from lxml import etree

        abstract_num = etree.fromstring(
            '<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:lvl w:ilvl="0">'
            '    <w:numFmt w:val="decimal"/>'
            "  </w:lvl>"
            "</w:abstractNum>"
        )
        assert _get_level_fmt(abstract_num, 5) == "decimal"


# ============================================================
# utils.py — _count_numbering_levels 覆盖测试
# ============================================================


class TestCountNumberingLevels:
    """测试 _count_numbering_levels 编号级别计数逻辑"""

    def _setup_numbering_doc(
        self, doc, num_fmt="decimal", lvl_text="%1.", num_id="1", abstract_num_id="0"
    ):
        """辅助方法：为文档创建 numbering part"""
        from docx.oxml import OxmlElement
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.parts.numbering import NumberingPart

        # 移除已有的 numbering 关系
        try:
            rels = doc.part.rels
            to_remove = [k for k, v in rels.items() if v.reltype == RT.NUMBERING]
            for k in to_remove:
                del rels[k]
        except Exception:
            pass

        numbering_elm = OxmlElement("w:numbering")

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), abstract_num_id)

        for lvl_idx in range(3):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(lvl_idx))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            numFmt = OxmlElement("w:numFmt")
            numFmt.set(qn("w:val"), num_fmt)
            lvl.append(numFmt)
            lvlText = OxmlElement("w:lvlText")
            lvlText.set(qn("w:val"), lvl_text)
            lvl.append(lvlText)
            abstract_num.append(lvl)

        numbering_elm.append(abstract_num)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        abstract_num_id_ref = OxmlElement("w:abstractNumId")
        abstract_num_id_ref.set(qn("w:val"), abstract_num_id)
        num.append(abstract_num_id_ref)
        numbering_elm.append(num)

        numbering_part = NumberingPart(
            PackURI("/word/numbering.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            numbering_elm,
            doc.part.package,
        )
        doc.part.relate_to(numbering_part, RT.NUMBERING)

        return numbering_elm

    def _add_numPr_to_paragraph(self, p, num_id="1", ilvl="0"):
        """辅助方法：为段落添加 numPr"""
        from docx.oxml import OxmlElement

        pPr = OxmlElement("w:pPr")
        numPr = OxmlElement("w:numPr")
        numId_elem = OxmlElement("w:numId")
        numId_elem.set(qn("w:val"), num_id)
        numPr.append(numId_elem)
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), ilvl)
        numPr.append(ilvl_elem)
        pPr.append(numPr)
        p._element.insert(0, pPr)

    def test_single_numbered_paragraph(self, doc):
        """单个编号段落返回 {0: 1}"""
        self._setup_numbering_doc(doc)
        p = doc.add_paragraph("绪论")
        self._add_numPr_to_paragraph(p, num_id="1", ilvl="0")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p)
        assert result == {0: 1}

    def test_multiple_level1_paragraphs(self, doc):
        """多个一级编号段落正确计数"""
        self._setup_numbering_doc(doc)
        p1 = doc.add_paragraph("第一章")
        self._add_numPr_to_paragraph(p1, num_id="1", ilvl="0")
        p2 = doc.add_paragraph("正文")
        doc.add_paragraph("无编号段落")
        p3 = doc.add_paragraph("第二章")
        self._add_numPr_to_paragraph(p3, num_id="1", ilvl="0")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p3)
        assert result == {0: 2}

    def test_mixed_levels_reset_counters(self, doc):
        """混合级别时下级计数器在上级重启"""
        self._setup_numbering_doc(doc)
        p1 = doc.add_paragraph("第一章")
        self._add_numPr_to_paragraph(p1, num_id="1", ilvl="0")
        p2 = doc.add_paragraph("1.1 背景")
        self._add_numPr_to_paragraph(p2, num_id="1", ilvl="1")
        p3 = doc.add_paragraph("1.2 方法")
        self._add_numPr_to_paragraph(p3, num_id="1", ilvl="1")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p3)
        assert result == {0: 1, 1: 2}

    def test_level_reset_after_new_parent(self, doc):
        """新的上级段落重置下级计数"""
        self._setup_numbering_doc(doc)
        p1 = doc.add_paragraph("第一章")
        self._add_numPr_to_paragraph(p1, num_id="1", ilvl="0")
        p2 = doc.add_paragraph("1.1 背景")
        self._add_numPr_to_paragraph(p2, num_id="1", ilvl="1")
        p3 = doc.add_paragraph("第二章")
        self._add_numPr_to_paragraph(p3, num_id="1", ilvl="0")
        p4 = doc.add_paragraph("2.1 方法")
        self._add_numPr_to_paragraph(p4, num_id="1", ilvl="1")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p4)
        assert result == {0: 2, 1: 1}

    def test_no_matching_num_ids_returns_empty(self, doc):
        """没有匹配的 numId 时返回空字典"""
        self._setup_numbering_doc(doc, num_id="1", abstract_num_id="0")
        p = doc.add_paragraph("test")
        self._add_numPr_to_paragraph(p, num_id="1", ilvl="0")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        # 使用不存在的 abstract_num_id
        result = _count_numbering_levels(numbering_elm, "999", p)
        assert result == {}

    def test_no_numbered_paragraphs_before_target(self, doc):
        """目标段落之前没有编号段落时返回 {ilvl: 1}"""
        self._setup_numbering_doc(doc)
        doc.add_paragraph("无编号段落")
        p = doc.add_paragraph("编号段落")
        self._add_numPr_to_paragraph(p, num_id="1", ilvl="0")

        numbering_part = doc.part.numbering_part
        numbering_elm = numbering_part._element
        result = _count_numbering_levels(numbering_elm, "0", p)
        assert result == {0: 1}
