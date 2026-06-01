"""Core 数据层测试 —— Tree, TreeNode, Stack, CategoryRegistry, build_tree"""
import pytest
from docx import Document
from docx.oxml.ns import qn

from wordformat.core.tree import Tree, Stack, TreeNode as CoreTreeNode
from wordformat.rules.node import FormatNode, TreeNode  # rules.TreeNode extends core.TreeNode


class TestTreeCreation:
    """Tree 创建与基本属性"""

    def test_create_tree_with_root_value(self):
        tree = Tree("root")
        assert tree.root.value == "root"

    def test_create_tree_with_dict_value(self):
        tree = Tree({"category": "top"})
        assert tree.root.value == {"category": "top"}

    def test_tree_repr(self):
        tree = Tree("my_root")
        assert repr(tree) == "Tree(root=my_root)"

    def test_is_empty(self):
        """空树（仅 root 无子节点）返回 True，有子节点返回 False"""
        tree = Tree("root")
        assert tree.is_empty() is True
        tree.root.add_child("child")
        assert tree.is_empty() is False


class TestTreeTraversals:
    """三种遍历：前序、后序、层序"""

    def setup_method(self):
        tree = Tree("A")
        b = tree.root.add_child("B")
        c = tree.root.add_child("C")
        b.add_child("D")
        b.add_child("E")
        c.add_child("F")
        self.tree = tree

    def test_preorder(self):
        assert list(self.tree.preorder()) == ["A", "B", "D", "E", "C", "F"]

    def test_postorder(self):
        assert list(self.tree.postorder()) == ["D", "E", "B", "F", "C", "A"]

    def test_level_order(self):
        assert list(self.tree.level_order()) == ["A", "B", "C", "D", "E", "F"]

    def test_preorder_single_node(self):
        tree = Tree("only")
        assert list(tree.preorder()) == ["only"]

    def test_postorder_single_node(self):
        tree = Tree("only")
        assert list(tree.postorder()) == ["only"]

    def test_level_order_single_node(self):
        tree = Tree("only")
        assert list(tree.level_order()) == ["only"]


class TestTreeFindAndMetrics:
    """find_by_condition, height, size"""

    def setup_method(self):
        tree = Tree("A")
        b = tree.root.add_child("B")
        c = tree.root.add_child("C")
        b.add_child("D")
        b.add_child("E")
        c.add_child("F")
        self.tree = tree

    def test_find_by_condition_exists(self):
        node = self.tree.find_by_condition(lambda v: v == "E")
        assert node is not None
        assert node.value == "E"

    def test_find_by_condition_not_exists(self):
        node = self.tree.find_by_condition(lambda v: v == "Z")
        assert node is None

    def test_find_by_condition_first_match(self):
        node = self.tree.find_by_condition(lambda v: isinstance(v, str) and len(v) == 1)
        assert node.value == "A"  # DFS 先遇到根

    def test_height_single_node(self):
        assert Tree("x").height() == 0

    def test_height_deep_tree(self):
        tree = Tree("1")
        tree.root.add_child("2").add_child("3").add_child("4")
        assert tree.height() == 3

    def test_height_balanced_tree(self):
        assert self.tree.height() == 2

    def test_size_single_node(self):
        assert Tree("x").size() == 1

    def test_size(self):
        assert self.tree.size() == 6


# ============================================================
# tree.py — Stack
# ============================================================


class TestStack:
    """Stack 的全部操作"""

    def test_push_and_pop(self):
        s = Stack()
        s.push(10)
        s.push(20)
        assert s.pop() == 20
        assert s.pop() == 10

    def test_peek(self):
        s = Stack()
        s.push("hello")
        assert s.peek() == "hello"
        assert s.size() == 1  # peek 不弹出

    def test_peek_safe_on_empty(self):
        s = Stack()
        assert s.peek_safe() is None

    def test_peek_safe_returns_top(self):
        s = Stack()
        s.push(42)
        assert s.peek_safe() == 42

    def test_is_empty(self):
        s = Stack()
        assert s.is_empty() is True
        s.push(1)
        assert s.is_empty() is False

    def test_size(self):
        s = Stack()
        assert s.size() == 0
        s.push("a")
        s.push("b")
        assert s.size() == 2

    def test_clear(self):
        s = Stack()
        s.push(1)
        s.push(2)
        s.clear()
        assert s.is_empty() is True
        assert s.size() == 0

    def test_pop_empty_raises(self):
        s = Stack()
        with pytest.raises(IndexError, match="pop from empty stack"):
            s.pop()

    def test_peek_empty_raises(self):
        s = Stack()
        with pytest.raises(IndexError, match="peek from empty stack"):
            s.peek()

    def test_bool_truthy(self):
        s = Stack()
        assert bool(s) is False
        s.push(1)
        assert bool(s) is True

    def test_repr(self):
        s = Stack()
        s.push(1)
        s.push(2)
        assert "1" in repr(s)
        assert "2" in repr(s)


# ============================================================
# tree.py — print_tree
# ============================================================


class TestPrintTree:
    """print_tree 输出捕获"""

    def test_print_single_node(self):
        node = TreeNode("hello")
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(node)
        output = buf.getvalue()
        assert "hello" in output
        assert "└──" in output

    def test_print_dict_value_node(self):
        node = TreeNode({
            "category": "body_text",
            "paragraph": "这是一段正文内容",
            "fingerprint": "fp001",
        })
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(node)
        output = buf.getvalue()
        assert "body_text" in output

    def test_print_tree_with_children(self):
        root = TreeNode("root")
        root.add_child("child1")
        root.add_child("child2")
        buf = StringIO()
        with patch("sys.stdout", buf):
            print_tree(root)
        output = buf.getvalue()
        assert "root" in output
        assert "child1" in output
        assert "child2" in output


# ============================================================
# utils.py — check_duplicate_fingerprints
# ============================================================


class TestTreeNode:
    def test_init_with_simple_value(self):
        node = TreeNode("hello")
        assert node.value == "hello"
        assert node.children == []

    def test_init_with_dict_value_and_fingerprint(self):
        node = TreeNode({"category": "body_text", "meta": {"fingerprint": "abc123"}})
        assert node.fingerprint == "abc123"

    def test_init_with_top_level_fingerprint_fallback(self):
        """向后兼容：旧 JSON（fingerprint 在顶层）仍可读取。"""
        node = TreeNode({"category": "body_text", "fingerprint": "old123"})
        assert node.fingerprint == "old123"

    def test_init_with_top_category_skips_fingerprint(self):
        node = TreeNode({"category": "top"})
        assert node.fingerprint is None

    def test_init_missing_fingerprint_no_raise(self):
        """Missing fingerprint no longer raises — it defaults to None."""
        node = TreeNode({"category": "body_text"})
        assert node.fingerprint is None
        assert node.index is None

    def test_config_default_empty(self):
        node = TreeNode("x")
        assert node.config == {}

    def test_load_config_nested_path(self):
        node = TreeNode("x")
        node.NODE_TYPE = "a.b.c"
        full = {"a": {"b": {"c": {"key": "val"}}}}
        node.load_config(full)
        assert node.config == {"key": "val"}

    def test_load_config_missing_path(self):
        node = TreeNode("x")
        node.NODE_TYPE = "x.y.z"
        node.load_config({"a": 1})
        assert node.config == {}

    def test_load_config_non_dict_input(self):
        node = TreeNode("x")
        node.load_config("not a dict")
        assert node.config == {}

    def test_add_child_returns_child(self):
        node = TreeNode("parent")
        child = node.add_child("child_val")
        assert child.value == "child_val"
        assert len(node.children) == 1

    def test_add_child_node(self):
        parent = TreeNode("parent")
        child = TreeNode("child")
        parent.add_child_node(child)
        assert parent.children[0] is child

    def test_repr(self):
        node = TreeNode("test_val")
        assert repr(node) == "TreeNode(test_val)"

    def test_fingerprint_attribute_exists_for_non_dict(self):
        node = TreeNode("simple_string")
        assert hasattr(node, "fingerprint")


# ============================================================
# rules/node.py — FormatNode
# ============================================================


class TestNormalizeText:
    """测试 normalize_text 文本归一化"""

    def test_collapse_whitespace(self):
        from wordformat.utils import normalize_text
        assert normalize_text("  第一章  绪论  ") == "第一章 绪论"

    def test_fullwidth_space(self):
        from wordformat.utils import normalize_text
        assert normalize_text("　摘要　") == "摘要"

    def test_nfc_normalization(self):
        from wordformat.utils import normalize_text
        import unicodedata
        text_nfd = unicodedata.normalize("NFD", "é")
        result = normalize_text(text_nfd)
        assert result == "é"

    def test_empty_string(self):
        from wordformat.utils import normalize_text
        assert normalize_text("") == ""


class TestAlignParagraphs:
    """测试 align_paragraphs 序列对齐"""

    @staticmethod
    def _make_docx_para(text):
        from docx import Document
        doc = Document()
        return doc.add_paragraph(text)

    def test_perfect_match(self):
        from wordformat.utils import align_paragraphs
        json_entries = [{"paragraph": "第一章"}, {"paragraph": "绪论"}]
        docx_paras = [self._make_docx_para("第一章"), self._make_docx_para("绪论")]
        matches, insertions, deletions = align_paragraphs(json_entries, docx_paras)
        assert matches == {0: docx_paras[0], 1: docx_paras[1]}
        assert insertions == set()
        assert deletions == set()

    def test_with_insertion(self):
        from wordformat.utils import align_paragraphs
        json_entries = [{"paragraph": "第一章"}, {"paragraph": "新增"}, {"paragraph": "绪论"}]
        docx_paras = [self._make_docx_para("第一章"), self._make_docx_para("绪论")]
        matches, insertions, deletions = align_paragraphs(json_entries, docx_paras)
        assert matches == {0: docx_paras[0], 2: docx_paras[1]}
        assert insertions == {1}
        assert deletions == set()

    def test_with_deletion(self):
        from wordformat.utils import align_paragraphs
        json_entries = [{"paragraph": "第一章"}, {"paragraph": "绪论"}]
        docx_paras = [
            self._make_docx_para("第一章"),
            self._make_docx_para("多余"),
            self._make_docx_para("绪论"),
        ]
        matches, insertions, deletions = align_paragraphs(json_entries, docx_paras)
        assert matches == {0: docx_paras[0], 1: docx_paras[2]}
        assert insertions == set()
        assert deletions == {1}

    def test_identical_texts(self):
        """两个相同文本的段落：按序列位置区分"""
        from wordformat.utils import align_paragraphs
        json_entries = [{"paragraph": "关键词：A"}, {"paragraph": "关键词：B"}]
        docx_paras = [self._make_docx_para("关键词：A"), self._make_docx_para("关键词：B")]
        matches, insertions, deletions = align_paragraphs(json_entries, docx_paras)
        assert matches == {0: docx_paras[0], 1: docx_paras[1]}
        assert insertions == set()
        assert deletions == set()

    def test_empty_both(self):
        from wordformat.utils import align_paragraphs
        matches, insertions, deletions = align_paragraphs([], [])
        assert matches == {}
        assert insertions == set()
        assert deletions == set()

    def test_replace_with_different_lengths(self):
        from wordformat.utils import align_paragraphs
        json_entries = [{"paragraph": "A"}, {"paragraph": "B"}, {"paragraph": "C"}]
        docx_paras = [
            self._make_docx_para("A"),
            self._make_docx_para("X"),
            self._make_docx_para("C"),
        ]
        matches, insertions, deletions = align_paragraphs(json_entries, docx_paras)
        assert matches == {0: docx_paras[0], 2: docx_paras[2]}
        assert insertions == {1}
        assert deletions == {1}

    def test_whitespace_tolerance(self):
        from wordformat.utils import align_paragraphs
        json_entries = [{"paragraph": "摘要"}]
        docx_paras = [self._make_docx_para("  摘要  ")]
        matches, insertions, deletions = align_paragraphs(json_entries, docx_paras)
        assert len(matches) == 1


# ============================================================
# word_structure/utils.py — collect_nodes_in_order 测试
# ============================================================


class TestCollectNodesInOrder:
    """测试 collect_nodes_in_order 节点收集"""

    def test_dfs_order(self):
        from wordformat.word_structure.utils import collect_nodes_in_order
        root = FormatNode(value={"category": "top"}, level=0)
        child1 = FormatNode(value={"category": "heading_level_1"}, level=1)
        child1a = FormatNode(value={"category": "body_text"}, level=2)
        child2 = FormatNode(value={"category": "heading_level_1"}, level=1)
        child2a = FormatNode(value={"category": "body_text"}, level=2)
        root.add_child_node(child1)
        child1.add_child_node(child1a)
        root.add_child_node(child2)
        child2.add_child_node(child2a)
        result = collect_nodes_in_order(root)
        assert len(result) == 4
        assert result[0] is child1
        assert result[1] is child1a
        assert result[2] is child2
        assert result[3] is child2a

    def test_empty_tree(self):
        from wordformat.word_structure.utils import collect_nodes_in_order
        root = FormatNode(value={"category": "top"}, level=0)
        assert collect_nodes_in_order(root) == []


# ============================================================
# tree.py — 覆盖 print_tree, _build_simple_tree, filter, level_order
# ============================================================


class TestPrintTree:
    """测试 tree.py print_tree 的各种路径"""

    def test_print_tree_from_json_path(self, tmp_path):
        from wordformat.tree import print_tree
        import json
        json_path = tmp_path / "test.json"
        json_path.write_text(json.dumps([
            {"category": "heading_level_1", "paragraph": "绪论"},
            {"category": "body_text", "paragraph": "一些正文"},
        ]))
        print_tree(str(json_path))

    def test_print_tree_with_filter_categories(self):
        from wordformat.tree import print_tree
        from wordformat.rules.node import TreeNode
        root = TreeNode({"category": "top", "paragraph": ""})
        child = TreeNode({"category": "heading_level_1", "paragraph": "第一章"})
        body = TreeNode({"category": "body_text", "paragraph": "正文"})
        root.add_child_node(child)
        child.add_child_node(body)
        print_tree(root, filter_categories=["heading_level_1"])

    def test_print_tree_with_confidence_and_index(self):
        from wordformat.tree import print_tree
        from wordformat.rules.node import TreeNode
        root = TreeNode({"category": "top", "paragraph": ""})
        child = TreeNode({"category": "body_text", "paragraph": "test", "confidence": 0.95})
        root.add_child_node(child)
        print_tree(root, show_confidence=True, show_index=True)

    def test_print_tree_non_dict_value_with_paragraph_attr(self):
        from wordformat.tree import print_tree
        from wordformat.rules.node import TreeNode

        class Wrapper:
            paragraph = {"category": "body_text", "paragraph": "Hello"}
        root = TreeNode(value={"category": "top", "paragraph": ""})
        root.add_child_node(TreeNode(value=Wrapper()))
        print_tree(root)


class TestBuildSimpleTree:
    """测试 _build_simple_tree 树构建"""

    def test_build_mixed_headings_and_body(self):
        from wordformat.tree import _build_simple_tree
        from wordformat.rules.node import TreeNode
        items = [
            {"category": "heading_level_1", "paragraph": "第一章"},
            {"category": "body_text", "paragraph": "正文1"},
            {"category": "heading_level_2", "paragraph": "1.1"},
            {"category": "body_text", "paragraph": "正文2"},
            {"category": "heading_level_1", "paragraph": "第二章"},
            {"category": "body_text", "paragraph": "正文3"},
        ]
        root = TreeNode({"category": "top", "paragraph": ""})
        _build_simple_tree(root, items)
        assert len(root.children) == 2
        h1 = root.children[0]
        assert h1.value["category"] == "heading_level_1"
        assert len(h1.children) == 2
        assert h1.children[1].value["category"] == "heading_level_2"
        h2 = root.children[1]
        assert h2.value["category"] == "heading_level_1"

    def test_build_body_text_attached_to_top(self):
        from wordformat.tree import _build_simple_tree
        from wordformat.rules.node import TreeNode
        items = [
            {"category": "body_text", "paragraph": "独立正文"},
        ]
        root = TreeNode({"category": "top", "paragraph": ""})
        _build_simple_tree(root, items)
        assert len(root.children) == 1
        assert root.children[0].value["category"] == "body_text"


class TestTreeLevelOrder:
    """测试 Tree.level_order 空 root 路径"""

    def test_level_order_empty_root(self):
        from wordformat.tree import Tree
        t = Tree("x")
        t.root = None
        assert list(t.level_order()) == []


# ============================================================
# node.py — 覆盖 load_config, apply_replace, _clean_paragraph_edge_spaces
# ============================================================


class TestTreeNodeLoadConfig:
    """测试 TreeNode.load_config 边缘情况"""

    def test_load_config_non_dict_mid_path(self):
        from wordformat.rules.node import TreeNode
        node = TreeNode("test")
        node.NODE_TYPE = "foo.bar"
        node.load_config({"foo": "not_a_dict"})
        assert node.config == {}


class TestFormatNodeApplyReplace:
    """测试 FormatNode.apply_replace 各种路径"""

    def test_apply_replace_non_dict_value(self):
        node = FormatNode(value="plain_string", level=1)
        assert node.apply_replace() is False

    def test_apply_replace_empty_after_strip(self):
        node = FormatNode(value={"replace": "   "}, level=1)
        assert node.apply_replace() is False

    def test_apply_replace_no_replace_key(self):
        node = FormatNode(value={"category": "body_text"}, level=1)
        assert node.apply_replace() is False

    def test_apply_replace_paragraph_none(self):
        node = FormatNode(value={"replace": "text"}, level=1)
        assert node.apply_replace() is False

    def test_apply_replace_single_run(self):
        doc = Document()
        p = doc.add_paragraph("original text")
        node = FormatNode(value={"replace": "new text"}, level=1, paragraph=p)
        node.apply_replace()
        # replace_text 经过 strip().replace(" ", "") 处理
        assert p.runs[0].text == "newtext"

    def test_apply_replace_multi_run(self):
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("aaa")
        p.add_run("bbb")
        p.add_run("ccc")
        node = FormatNode(value={"replace": "一二三"}, level=1, paragraph=p)
        node.apply_replace()
        combined = "".join(run.text for run in p.runs)
        assert combined == "一二三"


class TestFormatNodeCleanEdgeSpaces:
    """测试 _clean_paragraph_edge_spaces 边缘情况"""

    def test_clean_leading_spaces(self):
        doc = Document()
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run("  hello")
        node = FormatNode(value="test", level=1, paragraph=p)
        node._clean_paragraph_edge_spaces()
        assert run.text == "hello"

    def test_clean_trailing_spaces(self):
        doc = Document()
        p = doc.add_paragraph()
        p.clear()
        run = p.add_run("hello  ")
        node = FormatNode(value="test", level=1, paragraph=p)
        node._clean_paragraph_edge_spaces()
        assert run.text == "hello"


# ============================================================
# set_style.py — 覆盖样式修正、同步插入/删除
# ============================================================


class TestFindAndModifyFirst:
    """测试 find_and_modify_first"""

    def test_finds_matching_node(self):
        from wordformat.word_structure.utils import find_and_modify_first
        root = FormatNode(value={"category": "top"}, level=0)
        target = FormatNode(value={"category": "body_text", "flag": True}, level=2)
        root.add_child_node(target)
        result = find_and_modify_first(root, lambda n: n.value.get("flag"))
        assert result is target

    def test_returns_none_on_no_match(self):
        from wordformat.word_structure.utils import find_and_modify_first
        root = FormatNode(value={"category": "top"}, level=0)
        result = find_and_modify_first(root, lambda n: False)
        assert result is None


# ============================================================
# utils.py — 额外覆盖
# ============================================================


class TestFormatNodeRemaining:
    """测试 node.py 剩余未覆盖路径"""

    def test_apply_replace_empty_runs(self):
        doc = Document()
        p = doc.add_paragraph()
        p.clear()
        node = FormatNode(value={"replace": "text"}, level=1, paragraph=p)
        assert node.apply_replace() is False

    def test_clean_edge_spaces_empty_runs(self):
        doc = Document()
        p = doc.add_paragraph()
        p.clear()
        node = FormatNode(value="test", level=1, paragraph=p)
        node._clean_paragraph_edge_spaces()
        # 不抛异常即可


# ============================================================
# utils.py — 更多覆盖
# ============================================================


class TestCoreTree:
    """测试 core.Tree 的遍历、查找、属性"""

    def test_preorder(self):
        from wordformat.core.tree import Tree
        t = Tree({"category": "top"})
        t.root.add_child({"category": "heading"})
        t.root.add_child({"category": "body"})
        t.root.children[0].add_child({"category": "sub"})
        result = list(t.preorder())
        assert len(result) == 4
        assert result[0]["category"] == "top"

    def test_postorder(self):
        from wordformat.core.tree import Tree
        t = Tree({"category": "top"})
        t.root.add_child({"category": "child1"})
        t.root.add_child({"category": "child2"})
        result = list(t.postorder())
        assert result[0]["category"] == "child1"
        assert result[1]["category"] == "child2"
        assert result[2]["category"] == "top"

    def test_level_order(self):
        from wordformat.core.tree import Tree
        t = Tree({"category": "top"})
        t.root.add_child({"category": "a"})
        t.root.add_child({"category": "b"})
        result = list(t.level_order())
        assert result[0]["category"] == "top"
        assert result[1]["category"] == "a"
        assert result[2]["category"] == "b"

    def test_level_order_empty_root(self):
        from wordformat.core.tree import Tree
        t = Tree("x")
        t.root = None
        assert list(t.level_order()) == []

    def test_find_by_condition_found(self):
        from wordformat.core.tree import Tree
        t = Tree({"category": "top"})
        t.root.add_child({"category": "heading", "flag": True})
        t.root.add_child({"category": "body"})
        result = t.find_by_condition(lambda v: v.get("flag"))
        assert result is not None
        assert result.value["category"] == "heading"

    def test_find_by_condition_not_found(self):
        from wordformat.core.tree import Tree
        t = Tree({"category": "top"})
        assert t.find_by_condition(lambda v: False) is None

    def test_height(self):
        from wordformat.core.tree import Tree
        t = Tree("root")
        t.root.add_child("a")
        t.root.add_child("b")
        t.root.children[1].add_child("c")
        assert t.height() == 2

    def test_height_single_node(self):
        from wordformat.core.tree import Tree
        t = Tree("root")
        assert t.height() == 0

    def test_size(self):
        from wordformat.core.tree import Tree
        t = Tree("root")
        t.root.add_child("a")
        t.root.add_child("b")
        assert t.size() == 3

    def test_is_empty(self):
        from wordformat.core.tree import Tree
        t = Tree("root")
        assert t.is_empty() is True
        t.root.add_child("a")
        assert t.is_empty() is False


class TestCoreTreeNodeTypeContent:
    """测试 TreeNode 的 type/content 序列化"""

    def test_to_dict_with_type(self):
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "table"})
        node.type = "table"
        node.content = {"rows": 2, "cols": 2, "cells": [["A", "B"], ["C", "D"]]}
        d = node.to_dict()
        assert d["type"] == "table"
        assert d["content"]["rows"] == 2

    def test_to_dict_default_paragraph_omitted(self):
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "body_text"})
        d = node.to_dict()
        assert "type" not in d  # paragraph 是默认值，省略

    def test_to_dict_empty_content_omitted(self):
        from wordformat.core.tree import TreeNode
        node = TreeNode({"category": "body_text"})
        d = node.to_dict()
        assert "content" not in d

    def test_from_dict_with_type(self):
        from wordformat.core.tree import TreeNode
        data = {
            "value": {"category": "table"},
            "type": "table",
            "content": {"rows": 3},
            "children": [],
        }
        node = TreeNode.from_dict(data)
        assert node.type == "table"
        assert node.content["rows"] == 3

    def test_from_dict_defaults(self):
        from wordformat.core.tree import TreeNode
        node = TreeNode.from_dict({"value": {"x": 1}, "children": []})
        assert node.type == "paragraph"
        assert node.content == {}


# ============================================================
# core/category.py — CategoryRegistry 边界测试
# ============================================================


class TestCategoryRegistry:
    """测试 CategoryRegistry 扩展和边界"""

    def test_register_duplicate_raises(self):
        from wordformat.core.category import CategoryRegistry
        r = CategoryRegistry()
        with pytest.raises(ValueError, match="已存在"):
            r.register("body_text", level=1, is_heading=True, override=False)

    def test_register_override_succeeds(self):
        from wordformat.core.category import CategoryRegistry
        r = CategoryRegistry()
        r.register("body_text", level=5, is_heading=True, override=True)
        assert r.get_level("body_text") == 5

    def test_update_extend_defaults_false(self):
        from wordformat.core.category import CategoryRegistry
        r = CategoryRegistry()
        r.update({
            "categories": {
                "my_custom": {"level": 1, "is_heading": True},
            },
            "extend_defaults": False,
        })
        # 默认类别被清空，只有新类别
        assert r.get_level("body_text") == 999  # 未注册 → 默认值
        assert r.get_level("my_custom") == 1

    def test_update_extend_defaults_true(self):
        from wordformat.core.category import CategoryRegistry
        r = CategoryRegistry()
        r.update({
            "categories": {
                "my_custom": {"level": 1, "is_heading": True},
            },
            "extend_defaults": True,
        })
        # 默认类别保留
        assert r.get_level("body_text") == 999
        assert r.get_level("my_custom") == 1

    def test_is_heading_false_for_body(self):
        from wordformat.core.category import CategoryRegistry
        r = CategoryRegistry()
        assert r.is_heading("body_text") is False
        assert r.is_heading("heading_level_1") is True


# ============================================================
# core/tree_builder.py — factory 参数测试
# ============================================================


class TestCoreBuildTree:
    """测试 core.tree_builder.build_tree 的 factory 参数"""

    def test_custom_factory(self):
        from wordformat.core.tree_builder import build_tree
        from wordformat.core.tree import TreeNode

        class FlagNode(TreeNode):
            flag: bool = True

        def factory(item, level):
            node = FlagNode(item)
            node.flag = item.get("flag", False)
            return node

        items = [{"category": "body_text", "paragraph": "test", "flag": True}]
        root = build_tree(items, factory=factory)
        assert isinstance(root, TreeNode)
        assert root.children[0].flag is True

    def test_custom_root_value(self):
        from wordformat.core.tree_builder import build_tree
        root = build_tree([], root_value={"custom": "root"})
        assert root.value == {"custom": "root"}


# ============================================================
# orchestration/binding.py — _build_element 测试
# ============================================================


class TestCategoryRegistryReset:
    """测试 reset_registry"""

    def test_reset_restores_defaults(self):
        from wordformat.core.category import reset_registry, get_registry
        r = get_registry()
        r.register("temp_type", level=1, is_heading=True, override=True)
        r2 = reset_registry()
        assert r2.get_level("body_text") == 999
        assert r2.get_level("heading_level_1") == 1


# ============================================================
# tree.py — print_tree JSON path
# ============================================================


class TestPrintTreeJSONPath:
    """测试 print_tree 的 JSON 文件路径分支"""

    def test_print_tree_from_json(self, tmp_path):
        import json
        json_path = tmp_path / "test_tree.json"
        json_path.write_text(json.dumps([
            {"category": "heading_level_1", "paragraph": "绪论"},
            {"category": "body_text", "paragraph": "一些正文"},
        ]))
        from wordformat.tree import print_tree
        print_tree(str(json_path))


# ============================================================
# 防御性代码测试
# ============================================================


class TestTreeNodeEdgeCases:
    """测试 TreeNode 边界"""

    def test_none_value(self):
        from wordformat.core.tree import TreeNode as CoreTreeNode
        node = CoreTreeNode(None)
        assert node.value is None
        d = node.to_dict()
        assert d["value"] is None

    def test_from_dict_missing_children(self):
        from wordformat.core.tree import TreeNode as CoreTreeNode
        node = CoreTreeNode.from_dict({"value": "orphan"})
        assert node.value == "orphan"
        assert node.children == []
        assert node.type == "paragraph"


class TestCategoryRegistryUpdateEdgeCases:
    """CategoryRegistry.update() 边界"""

    def test_update_empty_config(self):
        from wordformat.core.category import CategoryRegistry
        r = CategoryRegistry()
        r.update({})
        assert r.get_level("body_text") == 999
        assert r.is_heading("heading_level_1") is True

    def test_update_extend_false_empty(self):
        from wordformat.core.category import CategoryRegistry
        r = CategoryRegistry()
        r.update({"categories": {}, "extend_defaults": False})
        assert r.get_level("body_text") == 999  # 未注册 → 默认值

    def test_update_missing_keys_defaults(self):
        from wordformat.core.category import CategoryRegistry
        r = CategoryRegistry()
        r.update({
            "categories": {"minimal": {}},
            "extend_defaults": True,
        })
        assert r.get_level("minimal") == 999  # 缺 level → 默认 999
        assert r.is_heading("minimal") is False  # 缺 is_heading → False


