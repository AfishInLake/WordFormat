#! /usr/bin/env python
# @Time    : 2026/7/12
# @Author  : afish
# @File    : stages_md.py
"""Markdown → Docx 转换的 pipeline 阶段。

Markdown 仅提供文本内容和文档结构（标题层级、段落划分），
所有格式（字体、字号、段落样式）由 YAML 配置 + FormatNode 统一处理。
图片插入、表格创建等由对应 FormatNode 子类负责。
"""

from pathlib import Path

from docx import Document

from wordformat.log_config import logger
from wordformat.markdown.parser import parse_markdown
from wordformat.rules.node import FormatNode

from .context import FormatContext


class LoadMarkdownStage:
    """读取 Markdown 文件到内存。"""

    def process(self, ctx: FormatContext) -> FormatContext:
        md_path = Path(ctx.md_path)
        if not md_path.exists():
            raise FileNotFoundError(f"Markdown 文件不存在: {md_path}")
        ctx.md_text = md_path.read_text(encoding="utf-8")
        logger.info(f"已加载 Markdown 文件: {md_path.resolve()}")
        return ctx


class MarkdownParseStage:
    """解析 Markdown 为扁平段落列表（仅提取纯文本，行内标记丢弃）。"""

    def process(self, ctx: FormatContext) -> FormatContext:
        ctx.paragraphs = parse_markdown(ctx.md_text)
        logger.info(f"解析完成，共 {len(ctx.paragraphs)} 个段落")
        return ctx


class DocumentCreationStage:
    """从 FormatNode 树创建新的 .docx Document。

    每个节点创建一个段落，单 run 填充纯文本。
    格式、图片、表格等由后续 FormattingExecutionStage 中各 FormatNode 子类处理。
    """

    def process(self, ctx: FormatContext) -> FormatContext:
        ctx.document = Document()
        nodes = self._flatten_tree_nodes(ctx.root_node)
        logger.info(f"开始创建文档，共 {len(nodes)} 个节点")

        for node in nodes:
            value = node.value
            text = value.get("paragraph", "") if isinstance(value, dict) else str(value)
            category = value.get("category", "") if isinstance(value, dict) else ""
            para = ctx.document.add_paragraph()
            # figure_image: 路径留给 FigureImage._try_insert_image 处理，这里只建空段落
            if category == "figure_image":
                para.add_run("")
            else:
                para.add_run(text.strip())
            node.paragraph = para

        return ctx

    @staticmethod
    def _flatten_tree_nodes(root_node: FormatNode) -> list[FormatNode]:
        result: list[FormatNode] = []

        def dfs(node):
            for child in node.children:
                result.append(child)
                dfs(child)

        dfs(root_node)
        return result
