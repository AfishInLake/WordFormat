#! /usr/bin/env python
# @Time    : 2026/7/7 12:25
# @Author  : afish
# @File    : loadpipline.py
# 加载配置、文档
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.shared import Pt, RGBColor

from wordformat.config.loader import load_config
from wordformat.hyperlinks import create_citation_hyperlinks
from wordformat.log_config import logger
from wordformat.rules.abstract import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleEN,
)
from wordformat.rules.caption import CaptionFigure, CaptionTable
from wordformat.rules.keywords import KeywordsCN, KeywordsEN
from wordformat.rules.node import FormatNode
from wordformat.rules.references import ReferenceEntry, References
from wordformat.settings import VOIDNODELIST
from wordformat.structure.document_builder import DocumentBuilder
from wordformat.structure.utils import promote_bodytext_in_subtrees_of_type
from wordformat.style.defs import (
    Alignment,
    FirstLineIndent,
    FontColor,
    FontSize,
    LeftIndent,
    LineSpacing,
    LineSpacingRule,
    RightIndent,
    SpaceAfter,
    SpaceBefore,
    ensure_style_exists,
)
from wordformat.style.writer import (
    SetFirstLineIndent,
    SetIndent,
    SetSpacing,
)
from wordformat.style.xml_ops import (
    ensure_pPr,
)
from wordformat.utils import (
    count_chinese_chars,
    ensure_directory_exists,
    get_file_name,
    has_chinese,
    parse_caption_text,
)

from .context import FormatContext


class LoadConfigStage:
    """加载配置pipline"""

    def process(self, ctx: FormatContext) -> FormatContext:
        if ctx.config_path:
            try:
                ctx.config_model = load_config(ctx.config_path)
                logger.info("配置文件验证通过")
            except Exception as e:
                logger.error(f"配置加载失败: {str(e)}")
                raise
        else:
            logger.info("未提供配置文件，使用默认配置")
        return ctx


class LoadDocxStage:
    """加载docx的pipline"""

    def process(self, ctx: FormatContext) -> FormatContext:
        """加载docx"""
        ctx.document = Document(ctx.docx_path)
        return ctx


class TreeBuildingStage:
    """构建tree pipline"""

    def process(self, ctx: FormatContext) -> FormatContext:
        ctx.root_node = DocumentBuilder.build_from_json(
            ctx.json_path, config=ctx.config_model
        )
        return ctx


class ParagraphAlignmentStage:
    """展平树并对齐段落 pipline"""

    def process(self, ctx: FormatContext) -> FormatContext:
        nodes = self._flatten_tree_nodes(ctx.root_node)
        for node, para in zip(nodes, ctx.document.paragraphs, strict=False):
            node.paragraph = para
        return ctx

    def _flatten_tree_nodes(self, root_node):
        """DFS 前序遍历展平树中所有节点（排除虚拟根节点），顺序与文档段落一致。"""
        result = []

        def dfs(node):
            for child in node.children:
                result.append(child)
                dfs(child)

        dfs(root_node)
        return result


class TreeNormalizationStage:
    """提升子树（摘要、参考文献）"""

    def process(self, ctx: FormatContext) -> FormatContext:
        mappings = {
            AbstractTitleCN: AbstractContentCN,
            AbstractTitleEN: AbstractContentEN,
            References: ReferenceEntry,
        }
        for parent_cls, target_cls in mappings.items():
            promote_bodytext_in_subtrees_of_type(
                ctx.root_node, parent_type=parent_cls, target_type=target_cls
            )
        return ctx


class StyleDefinitionFixStage:
    """修正样式定义（仅 apply 模式）"""

    def _fix_style_run_properties(self, style, cfg, style_name: str):
        """修正样式定义中的字符格式属性。

        python-docx 的 Style.font API 覆盖 size/color/bold/italic/underline，
        但 eastAsia 字体名需通过 XML 设置（Font 类不支持该属性）。
        """
        # 字体名：西文用 style.font.name，东亚用 XML（python-docx 不支持 eastAsia）
        cn_name = getattr(cfg, "chinese_font_name", None)
        en_name = getattr(cfg, "english_font_name", None)
        if cn_name or en_name:
            from wordformat.style.xml_ops import ensure_rPr, rPr_set_font

            rPr = ensure_rPr(style.element)
            rPr_set_font(rPr, cn_name=cn_name, en_name=en_name)

        font_size = getattr(cfg, "font_size", None)
        if font_size is not None:
            try:
                style.font.size = Pt(FontSize(font_size).rel_value)
            except Exception as e:
                logger.warning(f"设置样式 '{style_name}' 字号失败: {e}")

        font_color = getattr(cfg, "font_color", None)
        if font_color is not None:
            try:
                rgb = FontColor(font_color).rel_value
                style.font.color.rgb = RGBColor(*rgb)
            except Exception as e:
                logger.warning(f"设置样式 '{style_name}' 颜色失败: {e}")

        bold = getattr(cfg, "bold", None)
        if bold is not None:
            style.font.bold = bold

        italic = getattr(cfg, "italic", None)
        if italic is not None:
            style.font.italic = italic

        underline = getattr(cfg, "underline", None)
        if underline is not None:
            style.font.underline = underline

    def _fix_style_paragraph_properties(self, style, cfg, style_name: str):
        """修正样式定义中的段落格式属性。

        优先使用 python-docx 的 style.paragraph_format API（对齐、行距）。
        行单位（段前/段后间距）和字符单位（缩进）因 python-docx 不支持，
        回退到 XML 操作。
        """
        # --- 对齐方式（python-docx API） ---
        alignment = getattr(cfg, "alignment", None)
        if alignment is not None:
            try:
                style.paragraph_format.alignment = Alignment(alignment).rel_value
            except Exception as e:
                logger.warning(f"设置样式 '{style_name}' 对齐方式失败: {e}")

        # --- 行距（python-docx API） ---
        line_spacingrule = getattr(cfg, "line_spacingrule", None)
        if line_spacingrule is not None:
            try:
                lsr = LineSpacingRule(line_spacingrule)
                style.paragraph_format.line_spacing_rule = lsr.rel_value
                line_spacing = getattr(cfg, "line_spacing", None)
                if line_spacing is not None:
                    ls = LineSpacing(line_spacing)
                    if ls.rel_unit == "pt":
                        style.paragraph_format.line_spacing = Pt(ls.rel_value)
                    else:
                        style.paragraph_format.line_spacing = ls.rel_value
                else:
                    logger.warning(
                        f"样式 '{style_name}' 设置了 line_spacingrule 但未设置 line_spacing，已跳过行距"
                    )
            except Exception as e:
                logger.warning(f"设置样式 '{style_name}' 行距失败: {e}")

        # --- 段前/段后间距（仅支持行单位，需 XML） ---
        pPr = ensure_pPr(style.element)
        for attr_name, cls, spacing_type in [
            ("space_before", SpaceBefore, "before"),
            ("space_after", SpaceAfter, "after"),
        ]:
            val = getattr(cfg, attr_name, None)
            if val is None:
                continue
            try:
                inst = cls(val)
                if inst.rel_unit == "hang":
                    SetSpacing._set_hang_on_pPr(pPr, spacing_type, inst.rel_value)
                else:
                    logger.warning(
                        f"样式 '{style_name}' {attr_name} 使用了 '{inst.rel_unit}' 单位，"
                        f"样式定义仅支持'行'单位，已跳过"
                    )
            except Exception as e:
                logger.warning(f"设置样式 '{style_name}' {attr_name} 失败: {e}")

        # --- 缩进（仅支持字符单位，需 XML） ---
        first_line_indent = getattr(cfg, "first_line_indent", None)
        if first_line_indent is not None:
            try:
                inst = FirstLineIndent(first_line_indent)
                if inst.rel_unit == "char":
                    SetFirstLineIndent._clear_ind_on_pPr(pPr)
                    SetFirstLineIndent._set_char_on_pPr(pPr, inst.rel_value)
                else:
                    logger.warning(
                        f"样式 '{style_name}' first_line_indent 使用了 '{inst.rel_unit}' 单位，"
                        f"样式定义仅支持'字符'单位，已跳过"
                    )
            except Exception as e:
                logger.warning(f"设置样式 '{style_name}' first_line_indent 失败: {e}")

        for attr_name, cls, indent_type in [
            ("left_indent", LeftIndent, "R"),
            ("right_indent", RightIndent, "X"),
        ]:
            val = getattr(cfg, attr_name, None)
            if val is None:
                continue
            try:
                inst = cls(val)
                if inst.rel_unit == "char":
                    SetIndent._set_char_on_pPr(pPr, indent_type, inst.rel_value)
                else:
                    logger.warning(
                        f"样式 '{style_name}' {attr_name} 使用了 '{inst.rel_unit}' 单位，"
                        f"样式定义仅支持'字符'单位，已跳过"
                    )
            except Exception as e:
                logger.warning(f"设置样式 '{style_name}' {attr_name} 失败: {e}")

    def _fix_all_style_definitions(self, document: DocumentObject, config_model):
        """在格式化开始前，统一修正文档中所有使用的样式定义。

        遍历配置中所有段（body_text、headings、abstract、figures、tables、
        references、acknowledgements 等），收集唯一的 builtin_style_name，
        然后修正每个样式定义的字符格式和段落格式，使其与配置一致。

        修正内容：
        1. 字符格式：中英文字体、字号、颜色（清除主题色）、加粗、斜体、下划线
        2. 段落格式：对齐方式
        3. 确保样式定义存在（不存在则创建）
        """

        style_configs = config_model.collect_style_configs()

        for eng_style_name, cfg in style_configs.items():
            ensure_style_exists(document, eng_style_name)

            try:
                style = document.styles[eng_style_name]
            except KeyError:
                logger.warning(f"样式 '{eng_style_name}' 创建失败，跳过修正")
                continue

            self._fix_style_run_properties(style, cfg, eng_style_name)
            self._fix_style_paragraph_properties(style, cfg, eng_style_name)

            logger.debug(f"已修正样式定义: {eng_style_name}")

    def process(self, ctx: FormatContext) -> FormatContext:
        if not ctx.check:
            self._fix_all_style_definitions(ctx.document, ctx.config_model)
        return ctx


class FormattingExecutionStage:
    """执行格式化/检查（核心遍历）"""

    def apply_format_check_to_all_nodes(
        self, root_node: FormatNode, document, config, check=True
    ):
        """
        递归遍历文档树中的所有节点，
        对每个具有 check_format 方法的节点执行该方法。

        :param root_node: 树的根节点（FormatNode 或其子类实例）
        :param document: docx文档的实例
        :param config: 配置文件
        :param check: 用来控制是仅检查还是仅修改
        """

        chapter_index: int = 0
        figure_counter: dict[int, int] = {}
        table_counter: dict[int, int] = {}

        def traverse(node, parent_category="", current_chapter: int = 0):
            nonlocal chapter_index

            category = (
                node.value.get("category", "") if isinstance(node.value, dict) else ""
            )

            # 遇到一级标题时递增章节号
            if category == "heading_level_1":
                chapter_index += 1
                current_chapter = chapter_index

            if hasattr(node, "check_format"):
                try:
                    # top 节点直接关联的 body_text 不参与格式化（如封面页、原创性声明等）
                    # 但间接关联的 body_text（作为 heading 子节点）正常格式化
                    is_top_direct_body_text = (
                        parent_category == "top" and category == "body_text"
                    )
                    if category not in VOIDNODELIST and not is_top_direct_body_text:
                        node.load_config(config)

                        # 对题注节点注入章节号和顺序号
                        if isinstance(node, (CaptionFigure, CaptionTable)):
                            # 检查是否为续表/续图：保留原标题注编号，不递增计数器
                            text = node.paragraph.text.strip() if node.paragraph else ""
                            parsed = parse_caption_text(text)
                            if (
                                parsed
                                and parsed.get("is_continued")
                                and parsed.get("chapter_num") is not None
                                and parsed.get("number_num") is not None
                            ):
                                chapter = parsed["chapter_num"]
                                seq = parsed["number_num"]
                            else:
                                chapter = current_chapter if current_chapter > 0 else 0
                                if isinstance(node, CaptionFigure):
                                    counter = figure_counter
                                else:
                                    counter = table_counter
                                counter[chapter] = counter.get(chapter, 0) + 1
                                seq = counter[chapter]
                            node.value["chapter_number"] = chapter
                            node.value["sequence_number"] = seq

                        # 给所有节点注入章节号（BodyText 引用上标需要）
                        if isinstance(node.value, dict):
                            node.value.setdefault("chapter_number", current_chapter)

                        if node.paragraph:
                            # 先执行内容替换（check/format 两种模式均执行）
                            node.apply_replace(document)
                            if check:
                                node.check_format(document)
                            else:
                                node.apply_format(document)
                except Exception as e:
                    logger.warning(f"Node {node} not format, because: {str(e)}")
                    raise e

            # 目录、附录、封面/声明的子节点跳过格式化
            SKIP_CHILDREN_CATEGORIES = {"heading_mulu", "heading_fulu", "other"}
            if category not in SKIP_CHILDREN_CATEGORIES:
                for child in node.children:
                    traverse(
                        child, parent_category=category, current_chapter=current_chapter
                    )

        traverse(root_node)

    def process(self, ctx: FormatContext) -> FormatContext:
        if ctx.check:
            FormatNode.reset_stats()
        self.apply_format_check_to_all_nodes(
            ctx.root_node, ctx.document, ctx.config_model, ctx.check
        )
        return ctx


class SummaryGenerationStage:
    """生成检测报告摘要（仅 check 模式）"""

    def _build_check_summary(self, root_node, document, config_model) -> str:
        """遍历树和错误统计，生成检测报告摘要文本。"""
        stats = FormatNode._error_stats
        total = stats["total"]

        # 遍历树，收集文档级统计

        def _collect_section(node, sections):
            cls_name = type(node).__name__
            para = node.paragraph
            if para and para.text.strip():
                text = para.text.strip()
                if cls_name == "AbstractContentCN":
                    cn_chars = count_chinese_chars(text)
                    if cn_chars:
                        sections["abstract_cn_chars"] = (
                            sections.get("abstract_cn_chars", 0) + cn_chars
                        )
                elif cls_name == "AbstractContentEN":
                    sections["abstract_en_words"] = sections.get(
                        "abstract_en_words", 0
                    ) + len(text.split())
                elif cls_name == "KeywordsCN":
                    kws = KeywordsCN.extract_keywords(text)
                    if kws:
                        sections["keyword_cn_count"] = len(kws)
                elif cls_name == "KeywordsEN":
                    kws = KeywordsEN.extract_keywords(text)
                    if kws:
                        sections["keyword_en_count"] = len(kws)
                elif cls_name == "ReferenceEntry":
                    if has_chinese(text):
                        sections["ref_cn"] = sections.get("ref_cn", 0) + 1
                    else:
                        sections["ref_en"] = sections.get("ref_en", 0) + 1
            # 处理混合节点：AbstractTitleContentCN/EN 的子 BodyText 是摘要正文
            if cls_name == "AbstractTitleContentCN":
                for child in node.children:
                    cp = child.paragraph
                    if cp and cp.text.strip():
                        cnt = count_chinese_chars(cp.text.strip())
                        sections["abstract_cn_chars"] = (
                            sections.get("abstract_cn_chars", 0) + cnt
                        )
            elif cls_name == "AbstractTitleContentEN":
                for child in node.children:
                    cp = child.paragraph
                    if cp and cp.text.strip():
                        cnt = len(cp.text.strip().split())
                        sections["abstract_en_words"] = (
                            sections.get("abstract_en_words", 0) + cnt
                        )
            for child in node.children:
                _collect_section(child, sections)

        sections: dict = {}
        _collect_section(root_node, sections)

        # 计算万字差错率
        total_chars = sum(
            len(p.text) for p in document.paragraphs if p.text and p.text.strip()
        )
        error_rate = (total / max(total_chars, 1)) * 10000 if total else 0

        # 模板名（从 config 读取）
        template_name = getattr(config_model, "template_name", None) or "未知模板"

        lines = [
            "检测结果：",
            f"检测模板：《{template_name}》",
            f"检测错误数：{total}，万字差错率：{error_rate:.1f}",
            f"错误：{stats.get('错误', 0)}，提醒：{stats.get('提醒', 0)}",
        ]

        # 字数问题
        word_issues = []
        if sections.get("abstract_cn_chars"):
            word_issues.append(
                f"中文摘要：规范：300字左右，原文：{sections['abstract_cn_chars']}字"
            )
        if sections.get("abstract_en_words"):
            word_issues.append(
                f"英文摘要：规范：300字左右，原文：{sections['abstract_en_words']}词"
            )
        if sections.get("keyword_cn_count"):
            word_issues.append(
                f"中文关键词：规范：3-5个，原文：{sections['keyword_cn_count']}个"
            )
        if sections.get("keyword_en_count"):
            word_issues.append(
                f"英文关键词：规范：3-5个，原文：{sections['keyword_en_count']}个"
            )
        ref_cn = sections.get("ref_cn", 0)
        ref_en = sections.get("ref_en", 0)
        if ref_cn or ref_en:
            word_issues.append(
                f"参考文献：规范：不少于15条，原文：中文{ref_cn}条;外文{ref_en}条"
            )
        if word_issues:
            lines.append("字数问题：")
            lines.extend(word_issues)

        lines.append("说明：")
        lines.append(
            "1.请确保文档中正确使用换行符，硬回车（Enter）：指换行且生成新段落；软回车（Shift+Enter）：指换行但不生成新段落。"
        )
        lines.append("2.图片请使用“嵌入型”环绕方式，表格为无环绕方式。")
        lines.append("3.提醒不计算错误。")

        return "\n".join(lines)

    def _add_summary_comment(self, document, summary: str) -> None:
        """将检测报告摘要作为批注添加到文档第一段。空段临时塞空 run 做锚点。"""
        para = document.paragraphs[0]
        if not para.runs:
            para.add_run("")
        document.add_comment(
            runs=para.runs, text=summary, author="Wordformat", initials="afish"
        )

    def process(self, ctx: FormatContext) -> FormatContext:
        if ctx.check:
            summary = self._build_check_summary(
                ctx.root_node, ctx.document, ctx.config_model
            )
            if summary:
                self._add_summary_comment(ctx.document, summary)
        return ctx


class PostProcessingStage:
    """后处理（编号 + 超链接，仅 apply 模式）"""

    def process(self, ctx: FormatContext) -> FormatContext:
        if ctx.check:
            return ctx
        config_model = ctx.config_model
        # 标题编号
        if hasattr(config_model, "numbering") and config_model.numbering.enabled:
            from wordformat.numbering import process_heading_numbering

            process_heading_numbering(
                ctx.root_node,
                ctx.document,
                config_model.numbering,
                config_model.headings,
            )

        # 引用超链接
        create_citation_hyperlinks(ctx.root_node, ctx.document)

        return ctx


class DocumentSavingStage:
    """保存文档"""

    def process(self, ctx: FormatContext) -> FormatContext:
        ensure_directory_exists(ctx.save_dir)
        filename = get_file_name(ctx.docx_path)
        suffix = "--标注版.docx" if ctx.check else "--修改版.docx"
        out_path = Path(ctx.save_dir) / f"{filename}{suffix}"

        ctx.document.save(str(out_path))
        logger.info(f"保存文件到 {out_path}")
        ctx.output_path = str(out_path)
        return ctx
