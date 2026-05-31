#! /usr/bin/env python
"""表格格式化 —— 对文档中所有表格单元格内容进行格式校验或应用。"""

from docx import Document

from wordformat.config.datamodel import NodeConfigRoot
from wordformat.style.check_format import CharacterStyle, ParagraphStyle


def format_table_content(  # noqa: C901
    document: Document, config: NodeConfigRoot, check: bool = True
):
    """对文档中所有表格的单元格内容进行格式校验或应用。

    遍历 document.tables → rows → cells → paragraphs，
    根据 config.tables.content 中的格式配置对每个单元格段落进行
    段落样式和字符样式的检查（check=True）或应用（check=False）。

    Args:
        document: docx 文档对象
        config: Pydantic 配置根模型
        check: True 为仅检查（diff），False 为应用格式
    """
    try:
        content_cfg = config.tables.content
    except AttributeError:
        return

    ps = ParagraphStyle.from_config(content_cfg)
    cstyle = CharacterStyle(
        font_name_cn=content_cfg.chinese_font_name,
        font_name_en=content_cfg.english_font_name,
        font_size=content_cfg.font_size,
        font_color=content_cfg.font_color,
        bold=content_cfg.bold,
        italic=content_cfg.italic,
        underline=content_cfg.underline,
    )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue

                    # 段落样式
                    if check:
                        para_issues = ps.diff_from_paragraph(paragraph)
                    else:
                        para_issues = ps.apply_to_paragraph(paragraph)
                    para_text = ParagraphStyle.to_string(para_issues)
                    if para_text.strip():
                        document.add_comment(
                            runs=paragraph.runs,
                            text=para_text,
                            author="论文解析器",
                            initials="afish",
                        )

                    # 字符样式
                    for run in paragraph.runs:
                        if not run.text.strip():
                            continue
                        if check:
                            diff = cstyle.diff_from_run(run)
                        else:
                            diff = cstyle.apply_to_run(run)
                        run_text = CharacterStyle.to_string(diff)
                        if run_text.strip():
                            document.add_comment(
                                runs=run,
                                text=run_text,
                                author="论文解析器",
                                initials="afish",
                            )
