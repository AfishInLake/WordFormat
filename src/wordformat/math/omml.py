"""
OML Renderer — LaTeX math to OMML (Office Math Markup Language) for .docx files.

Uses the pure-Python pipeline:
    LaTeX -> latex2mathml -> MathML -> mathml2omml -> OMML
"""

from __future__ import annotations

import re

import latex2mathml.converter
import mathml2omml
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsmap, qn

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

MATH_FONT = "Cambria Math"
_DEFAULT_SZ = "22"

# mathml2omml 输出使用 m: 前缀但未声明命名空间，解析时需包裹 root 声明
_WRAPPED_MATH_NS = f'<root xmlns:m="{nsmap["m"]}">{{body}}</root>'


# ---------------------------------------------------------------------------
# LaTeX → OMML pipeline
# ---------------------------------------------------------------------------


def _latex_to_omml_xml(latex: str) -> str | None:
    """Convert a LaTeX math string to OMML XML string via MathML."""
    if not latex.strip():
        return None
    try:
        mathml = latex2mathml.converter.convert(latex)
    except Exception:
        return None
    if not mathml:
        return None
    try:
        return mathml2omml.convert(mathml)
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Element builders (python-docx OxmlElement)
# ---------------------------------------------------------------------------


def _make_wrpr():
    """Build a w:rPr element with math font and size."""
    wrpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), MATH_FONT)
    rf.set(qn("w:hAnsi"), MATH_FONT)
    wrpr.append(rf)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), _DEFAULT_SZ)
    wrpr.append(sz)
    return wrpr


def _make_ctrl_pr():
    """Build an m:ctrlPr containing w:rPr."""
    cp = OxmlElement("m:ctrlPr")
    cp.append(_make_wrpr())
    return cp


def _post_process(omath) -> None:
    """Fix up mathml2omml output to match Word/WPS expectations."""
    # Add w:rPr to every m:r that lacks it (WPS needs font/size on each run)
    for mr in om_elements(omath, "m:r"):
        if mr.find(qn("w:rPr")) is None:
            wrpr = _make_wrpr()
            mrpr = mr.find(qn("m:rPr"))
            if mrpr is not None:
                idx = list(mr).index(mrpr)
                mr.insert(idx + 1, wrpr)
            else:
                mr.insert(0, wrpr)

    # Fix rad elements: add radPr, degHide, deg, ctrlPr as needed
    for rad in om_elements(omath, "m:rad"):
        radpr = rad.find(qn("m:radPr"))
        if radpr is None:
            radpr = OxmlElement("m:radPr")
            rad.insert(0, radpr)
        if radpr.find(qn("m:degHide")) is None and rad.find(qn("m:deg")) is None:
            dg_hide = OxmlElement("m:degHide")
            dg_hide.set(qn("m:val"), "1")
            radpr.append(dg_hide)
        if radpr.find(qn("m:ctrlPr")) is None:
            radpr.append(_make_ctrl_pr())
        deg = rad.find(qn("m:deg"))
        if deg is None:
            deg = OxmlElement("m:deg")
            deg.append(_make_ctrl_pr())
            rp_idx = list(rad).index(radpr)
            rad.insert(rp_idx + 1, deg)
        elif deg.find(qn("m:ctrlPr")) is None:
            deg.append(_make_ctrl_pr())


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def latex_to_omath(latex: str) -> OxmlElement | None:
    """Parse a LaTeX math expression into an ``m:oMath`` OMML element.

    Returns ``None`` if the string is empty or parsing fails.
    """
    omml_xml = _latex_to_omml_xml(latex)
    if not omml_xml:
        return None
    try:
        wrapped = _WRAPPED_MATH_NS.format(body=omml_xml)
        tree = parse_xml(wrapped.encode("utf-8"))
        omath = tree.find(qn("m:oMath"))
        if omath is not None:
            _post_process(omath)
        return omath
    except Exception:
        return None


def latex_to_omath_para(latex: str) -> OxmlElement | None:
    """Parse LaTeX into an ``m:oMathPara`` element (for display / block math)."""
    om = latex_to_omath(latex)
    if om is None:
        return None
    omathpara = OxmlElement("m:oMathPara")
    omathpara.append(om)
    return omathpara


def add_display_math(doc, latex: str):
    """Add a centered display-math paragraph to a ``python-docx`` Document.

    The *latex* string should **not** contain ``$$`` delimiters.
    Returns the new paragraph.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ppr = p._p.find(qn("w:pPr"))
    if ppr is not None:
        ppr.append(_make_wrpr())
    om = latex_to_omath(latex)
    if om is not None:
        omathpara = OxmlElement("m:oMathPara")
        omathpara.append(om)
        p._p.append(omathpara)
    else:
        from docx.shared import Pt

        run = p.add_run(latex)
        run.font.size = Pt(11)
    return p


def add_inline_math(paragraph, text: str):
    """Render text containing ``$...$`` inline-math segments into *paragraph*.

    Math segments are rendered as proper OMML ``m:oMath`` elements.
    """
    parts = re.split(r"(\$[^$\n]+\$)", text)
    for part in parts:
        if part.startswith("$") and part.endswith("$"):
            math = part[1:-1].strip()
            om = latex_to_omath(math)
            if om is not None:
                paragraph._p.append(om)
            else:
                paragraph.add_run(math)
        elif part:
            paragraph.add_run(part)


def set_cell_math(cell, text: str):
    """Render text with ``$...$`` / ``$$...$$`` math into a table cell."""
    p = cell.paragraphs[0]
    p.clear()
    stripped = text.strip()
    if stripped.startswith("$$") and stripped.endswith("$$"):
        math = stripped[2:-2].strip()
        om = latex_to_omath(math)
        if om is not None:
            omathpara = OxmlElement("m:oMathPara")
            omathpara.append(om)
            p._p.append(omathpara)
        else:
            p.add_run(math)
    else:
        add_inline_math(p, text)


def om_elements(element, tag: str):
    """Iterate over descendant elements matching *tag* (e.g. ``m:r``, ``m:rad``).

    Uses pre-resolved qn for performance.
    """
    return element.iter(qn(tag))
