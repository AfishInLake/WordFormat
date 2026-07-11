"""批注文本格式化工具。

所有 add_comment 调用最终由此模块生成符合标准格式的批注文本：
    [位置]-[问题类型]：[现状]，规范：[标准]
"""

# diff_type → 中文问题类型
CHAR_DIFF_LABELS: dict[str, str] = {
    "font_size": "字号错误",
    "bold": "加粗错误",
    "italic": "斜体错误",
    "underline": "下划线错误",
    "font_color": "字体颜色错误",
    "font_name_cn": "中文字体错误",
    "font_name_en": "英文字体错误",
}

PARA_DIFF_LABELS: dict[str, str] = {
    "alignment": "对齐错误",
    "first_line_indent": "首行缩进错误",
    "line_spacing": "行距问题",
    "line_spacing_rule": "行距类型问题",
    "space_before": "段前间距错误",
    "space_after": "段后间距错误",
    "left_indent": "左缩进错误",
    "right_indent": "右缩进错误",
    "builtin_style_name": "内置样式错误",
}


# 问题类型 → 严重等级
SEVERITY_MAP: dict[str, str] = {
    "行距问题": "错误",
    "对齐错误": "错误",
    "首行缩进错误": "错误",
    "段前间距错误": "错误",
    "段后间距错误": "错误",
    "行距类型问题": "错误",
    "左缩进错误": "错误",
    "右缩进错误": "错误",
    "内置样式错误": "错误",
    "字号错误": "错误",
    "加粗错误": "错误",
    "数量过少": "错误",
    "数量过多": "错误",
    "编号错误": "错误",
    "章节号错误": "错误",
    "分隔符错误": "提醒",
    "标签错误": "提醒",
    "间距错误": "提醒",
    "格式错误": "提醒",
    "斜体错误": "提醒",
    "下划线错误": "提醒",
    "字体颜色错误": "提醒",
    "中文字体错误": "提醒",
    "英文字体错误": "提醒",
    "标点错误": "提醒",
}

_DEFAULT_SEVERITY = "错误"

# 严重等级排序权重（数值越小越严重）
SEVERITY_ORDER: dict[str, int] = {"错误": 0, "提醒": 1}

# 严重等级 → 文字颜色（十六进制，无 #）
SEVERITY_COLOR: dict[str, str] = {"错误": "FF0000", "提醒": "0000FF"}


def get_severity(comment_text: str) -> str:
    """从批注文本中提取严重等级。格式：位置-问题类型：..."""
    try:
        prop = comment_text.split("-", 1)[1].split("：")[0]
    except (IndexError, AttributeError):
        return _DEFAULT_SEVERITY
    return SEVERITY_MAP.get(prop, _DEFAULT_SEVERITY)


def severity_color(property_name: str) -> str | None:
    """按问题类型返回对应严重等级的颜色，未知类型按默认等级。"""
    return SEVERITY_COLOR.get(SEVERITY_MAP.get(property_name, _DEFAULT_SEVERITY))


def format_comment(target: str, property_name: str, actual: str, standard: str) -> str:
    """生成标准批注文本：位置-问题类型：现状，规范：标准。"""
    return f"{target}-{property_name}：{actual}，规范：{standard}"


# ── 批注富文本接口 ────────────────────────────────────────────────
# 段落 = 若干「片段」，片段 = (文字, 样式)。样式为 None（黑色正文）或
# dict：{"color": "FF0000", "bold": True, "italic": True, "underline": True}。
Style = dict | None
Segment = tuple[str, Style]


def apply_run_style(run, style: Style) -> None:
    """将样式字典应用到批注 run，仅设置字典中出现的属性。"""
    if not style:
        return
    from docx.shared import RGBColor

    color = style.get("color")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for attr in ("bold", "italic", "underline"):
        if attr in style:
            setattr(run, attr, style[attr])


def add_styled_comment(
    doc,
    runs,
    paragraphs: list[list[Segment]],
    *,
    author: str = "Wordformat",
    initials: str = "afish",
):
    """添加批注，并对其中指定文字片段单独设置样式。

    这是给批注「某些字」上色/加粗的统一入口：`paragraphs` 是若干段落，每段是
    一串 `(文字, 样式)` 片段；样式为 None 表示默认黑色正文。返回 Comment 对象。
    """
    comment = doc.add_comment(runs=runs, text="", author=author, initials=initials)
    for i, segments in enumerate(paragraphs):
        para = comment.paragraphs[0] if i == 0 else comment.add_paragraph()
        for r in list(para.runs):  # 清掉 add_comment 建的占位空 run
            r._element.getparent().remove(r._element)
        for text, style in segments:
            apply_run_style(para.add_run(text), style)
    return comment


def split_comment_line(line: str) -> list[Segment]:
    """把一行批注拆成带样式片段：`位置-问题类型：` 按严重度上色，其余黑色。

    输入格式 `位置-问题类型：现状，规范：标准`；无法解析时整行黑色。
    颜色由「问题类型」（首个 `-` 与首个 `：` 之间）决定。
    """
    prefix, sep, tail = line.partition("：")
    if not sep:
        return [(line, None)]
    prop = prefix.split("-", 1)[1] if "-" in prefix else prefix
    color = severity_color(prop)
    segments: list[Segment] = [(prefix + sep, {"color": color} if color else None)]
    if tail:
        segments.append((tail, None))
    return segments
