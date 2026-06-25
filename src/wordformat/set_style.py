#! /usr/bin/env python
# @Time    : 2026/1/11 19:51
# @Author  : afish
# @File    : set_style.py
from pathlib import Path
from typing import Optional

from docx import Document
from loguru import logger

from wordformat.config.config import get_config, init_config
from wordformat.config.datamodel import BaseModel, NodeConfigRoot
from wordformat.rules import (
    AbstractContentCN,
    AbstractContentEN,
    AbstractTitleCN,
    AbstractTitleEN,
    FormatNode,
    ReferenceEntry,
    References,
)
from wordformat.settings import VOIDNODELIST
from wordformat.style.check_format import CharacterStyle, ParagraphStyle
from wordformat.utils import ensure_directory_exists
from wordformat.word_structure.document_builder import DocumentBuilder
from wordformat.word_structure.utils import promote_bodytext_in_subtrees_of_type


def _collect_all_style_configs(config_model) -> dict:
    """遍历 NodeConfigRoot，收集所有唯一的 (英文样式名 → 配置段) 映射。

    对于被多个段引用的同一样式（如 body_text 和 references.content 都用 "Normal"），
    保留先遇到的配置段。
    """
    from wordformat.config.datamodel import GlobalFormatConfig
    from wordformat.style.style_enum import BuiltInStyle

    style_map: dict[str, object] = {}

    def _resolve_style_name(cfg) -> str | None:
        raw = getattr(cfg, "builtin_style_name", None)
        if not raw:
            return None
        try:
            return BuiltInStyle(raw).rel_value  # "正文" → "Normal", "题注" → "Caption"
        except Exception:
            return raw  # 自定义样式名直接使用

    def _walk(obj, path: str = ""):
        if isinstance(obj, GlobalFormatConfig):
            eng_name = _resolve_style_name(obj)
            if eng_name and eng_name not in style_map:
                style_map[eng_name] = obj
        if isinstance(obj, BaseModel):
            for f_name in type(obj).model_fields:
                val = getattr(obj, f_name)
                if isinstance(val, BaseModel):
                    _walk(val, f"{path}.{f_name}")
                elif isinstance(val, dict):
                    for _k, v in val.items():
                        if isinstance(v, BaseModel):
                            _walk(v, f"{path}.{_k}")

    _walk(config_model)
    return style_map


def _fix_style_run_properties(style, cfg, style_name: str):
    """修正样式定义中的字符格式属性（w:rPr）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from wordformat.style.style_enum import FontColor, FontSize

    style_element = style.element

    # 确保 w:rPr 存在
    rPr = style_element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        style_element.insert(0, rPr)

    # --- 字体名称 ---
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    # 中文字体（eastAsia）
    cn_name = getattr(cfg, "chinese_font_name", None)
    if cn_name:
        rFonts.set(qn("w:eastAsia"), str(cn_name))
    # 英文字体（ascii + hAnsi）
    en_name = getattr(cfg, "english_font_name", None)
    if en_name:
        rFonts.set(qn("w:ascii"), str(en_name))
        rFonts.set(qn("w:hAnsi"), str(en_name))

    # --- 字号 ---
    font_size = getattr(cfg, "font_size", None)
    if font_size is not None:
        try:
            fs = FontSize(font_size)
            pt_val = fs.rel_value  # e.g. "小四" → 12.0
            half_pt = str(int(round(pt_val * 2)))
            # 更新或创建 w:sz
            sz = rPr.find(qn("w:sz"))
            if sz is None:
                sz = OxmlElement("w:sz")
                rPr.append(sz)
            sz.set(qn("w:val"), half_pt)
            # w:szCs（复杂脚本字号）
            szCs = rPr.find(qn("w:szCs"))
            if szCs is None:
                szCs = OxmlElement("w:szCs")
                rPr.append(szCs)
            szCs.set(qn("w:val"), half_pt)
        except Exception as e:
            logger.warning(f"设置样式 '{style_name}' 字号失败: {e}")

    # --- 字体颜色（清除主题色） ---
    font_color = getattr(cfg, "font_color", None)
    if font_color is not None:
        try:
            fc = FontColor(font_color)
            rgb = fc.rel_value
            hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
            # 移除旧的 w:color（尤其是有主题色的）
            for old_color in rPr.findall(qn("w:color")):
                rPr.remove(old_color)
            new_color = OxmlElement("w:color")
            new_color.set(qn("w:val"), hex_color)
            rPr.append(new_color)
        except Exception as e:
            logger.warning(f"设置样式 '{style_name}' 颜色失败: {e}")

    # --- 加粗 ---
    bold = getattr(cfg, "bold", None)
    if bold is not None:
        b = rPr.find(qn("w:b"))
        if bold:
            if b is None:
                rPr.append(OxmlElement("w:b"))
        else:
            if b is not None:
                rPr.remove(b)

    # --- 斜体 ---
    italic = getattr(cfg, "italic", None)
    if italic is not None:
        i = rPr.find(qn("w:i"))
        if italic:
            if i is None:
                rPr.append(OxmlElement("w:i"))
        else:
            if i is not None:
                rPr.remove(i)

    # --- 下划线 ---
    underline = getattr(cfg, "underline", None)
    if underline is not None:
        u = rPr.find(qn("w:u"))
        if underline:
            if u is None:
                u = OxmlElement("w:u")
                u.set(qn("w:val"), "single")
                rPr.append(u)
        else:
            if u is not None:
                rPr.remove(u)


def _fix_style_paragraph_properties(style, cfg, style_name: str):
    """修正样式定义中的段落格式属性（w:pPr）。

    设置：对齐方式、段前/段后间距、行距、首行缩进、左右缩进。
    直接在 w:pPr 元素上操作 XML，确保样式定义级别生效。
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from wordformat.style.style_enum import (
        Alignment,
        FirstLineIndent,
        LeftIndent,
        LineSpacing,
        LineSpacingRule,
        RightIndent,
        SpaceAfter,
        SpaceBefore,
    )

    style_element = style.element

    pPr = style_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_element.append(pPr)

    # --- 对齐方式 ---
    alignment = getattr(cfg, "alignment", None)
    if alignment is not None:
        try:
            al = Alignment(alignment)
            jc = pPr.find(qn("w:jc"))
            if jc is None:
                jc = OxmlElement("w:jc")
                pPr.append(jc)
            xml_val_map = {
                0: "left",
                1: "center",
                2: "right",
                3: "both",
                4: "distribute",
            }
            jc.set(qn("w:val"), xml_val_map.get(al.rel_value, "left"))
        except Exception as e:
            logger.warning(f"设置样式 '{style_name}' 对齐方式失败: {e}")

    # --- 段前/段后间距 ---
    from wordformat.style.set_some import _SetSpacing

    for attr_name, cls, spacing_type in [
        ("space_before", SpaceBefore, "before"),
        ("space_after", SpaceAfter, "after"),
    ]:
        val = getattr(cfg, attr_name, None)
        if val is None:
            continue
        try:
            inst = cls(val)
            if inst.rel_unit == "hang":
                _SetSpacing._set_hang_on_pPr(pPr, spacing_type, inst.rel_value)
            else:
                logger.warning(
                    f"样式 '{style_name}' {attr_name} 使用了 '{inst.rel_unit}' 单位，"
                    f"样式定义仅支持'行'单位，已跳过"
                )
        except Exception as e:
            logger.warning(f"设置样式 '{style_name}' {attr_name} 失败: {e}")

    # --- 行距 ---
    from wordformat.style.set_some import _SetLineSpacing

    line_spacingrule = getattr(cfg, "line_spacingrule", None)
    if line_spacingrule is not None:
        try:
            lsr = LineSpacingRule(line_spacingrule)
            rule_map = {
                0: "auto",
                1: "auto",
                2: "auto",
                3: "atLeast",
                4: "exact",
                5: "auto",
            }
            line_rule = rule_map.get(lsr.rel_value, "auto")

            line_spacing = getattr(cfg, "line_spacing", None)
            if line_spacing is not None:
                ls = LineSpacing(line_spacing)
                ls_val = ls.rel_value
                if ls.rel_unit == "倍":
                    line_val = ls_val * 240  # 倍数 → w:line
                elif ls.rel_unit in ("pt",):
                    line_val = ls_val * 20  # pt → twips
                else:
                    line_val = ls_val * 240
                _SetLineSpacing._set_on_pPr(pPr, line_rule, line_val)
            else:
                logger.warning(
                    f"样式 '{style_name}' 设置了 line_spacingrule 但未设置 line_spacing，已跳过行距"
                )
        except Exception as e:
            logger.warning(f"设置样式 '{style_name}' 行距失败: {e}")

    # --- 缩进：先设置首行缩进，再设置左右缩进，避免 _clear_ind_on_pPr 清除 *Chars 属性 ---
    from wordformat.style.set_some import _SetFirstLineIndent, _SetIndent

    first_line_indent = getattr(cfg, "first_line_indent", None)
    if first_line_indent is not None:
        try:
            inst = FirstLineIndent(first_line_indent)
            if inst.rel_unit == "char":
                _SetFirstLineIndent._clear_ind_on_pPr(pPr)
                _SetFirstLineIndent._set_char_on_pPr(pPr, inst.rel_value)
            else:
                logger.warning(
                    f"样式 '{style_name}' first_line_indent 使用了 '{inst.rel_unit}' 单位，"
                    f"样式定义仅支持'字符'单位，已跳过"
                )
        except Exception as e:
            logger.warning(f"设置样式 '{style_name}' first_line_indent 失败: {e}")

    for attr_name, cls, indent_type in [
        ("left_indent", LeftIndent, "R"),
        ("right_indent", RightIndent, "X"),
    ]:
        val = getattr(cfg, attr_name, None)
        if val is None:
            continue
        try:
            inst = cls(val)
            if inst.rel_unit == "char":
                _SetIndent._set_char_on_pPr(pPr, indent_type, inst.rel_value)
            else:
                logger.warning(
                    f"样式 '{style_name}' {attr_name} 使用了 '{inst.rel_unit}' 单位，"
                    f"样式定义仅支持'字符'单位，已跳过"
                )
        except Exception as e:
            logger.warning(f"设置样式 '{style_name}' {attr_name} 失败: {e}")


def _fix_all_style_definitions(document: Document, config_model):
    """在格式化开始前，统一修正文档中所有使用的样式定义。

    遍历配置中所有段（body_text、headings、abstract、figures、tables、
    references、acknowledgements 等），收集唯一的 builtin_style_name，
    然后修正每个样式定义的字符格式和段落格式，使其与配置一致。

    修正内容：
    1. 字符格式：中英文字体、字号、颜色（清除主题色）、加粗、斜体、下划线
    2. 段落格式：对齐方式
    3. 确保样式定义存在（不存在则创建）
    """
    from wordformat.style.style_enum import _ensure_style_exists

    style_configs = _collect_all_style_configs(config_model)

    for eng_style_name, cfg in style_configs.items():
        _ensure_style_exists(document, eng_style_name)

        try:
            style = document.styles[eng_style_name]
        except KeyError:
            logger.warning(f"样式 '{eng_style_name}' 创建失败，跳过修正")
            continue

        _fix_style_run_properties(style, cfg, eng_style_name)
        _fix_style_paragraph_properties(style, cfg, eng_style_name)

        logger.debug(f"已修正样式定义: {eng_style_name}")


def apply_format_check_to_all_nodes(
    root_node: FormatNode, document, config, check=True
):
    """
    递归遍历文档树中的所有节点，
    对每个具有 check_format 方法的节点执行该方法。

    :param root_node: 树的根节点（FormatNode 或其子类实例）
    :param document: docx文档的实例
    :param config: 配置文件
    :param check: 用来控制是仅检查还是仅修改
    """
    from wordformat.rules.caption import CaptionFigure, CaptionTable
    from wordformat.utils import parse_caption_text

    chapter_index: int = 0
    figure_counter: dict[int, int] = {}
    table_counter: dict[int, int] = {}

    def traverse(node, parent_category="", current_chapter: int = 0):
        nonlocal chapter_index

        category = (
            node.value.get("category", "") if isinstance(node.value, dict) else ""
        )

        # 遇到一级标题时递增章节号
        if category == "heading_level_1":
            chapter_index += 1
            current_chapter = chapter_index

        if hasattr(node, "check_format"):
            try:
                # top 节点直接关联的 body_text 不参与格式化（如封面页、原创性声明等）
                # 但间接关联的 body_text（作为 heading 子节点）正常格式化
                is_top_direct_body_text = (
                    parent_category == "top" and category == "body_text"
                )
                if category not in VOIDNODELIST and not is_top_direct_body_text:
                    node.load_config(config)

                    # 对题注节点注入章节号和顺序号
                    if isinstance(node, (CaptionFigure, CaptionTable)):
                        # 检查是否为续表/续图：保留原标题注编号，不递增计数器
                        text = node.paragraph.text.strip() if node.paragraph else ""
                        parsed = parse_caption_text(text)
                        if (
                            parsed
                            and parsed.get("is_continued")
                            and parsed.get("chapter_num") is not None
                            and parsed.get("number_num") is not None
                        ):
                            chapter = parsed["chapter_num"]
                            seq = parsed["number_num"]
                        else:
                            chapter = current_chapter if current_chapter > 0 else 0
                            if isinstance(node, CaptionFigure):
                                counter = figure_counter
                            else:
                                counter = table_counter
                            counter[chapter] = counter.get(chapter, 0) + 1
                            seq = counter[chapter]
                        node.value["chapter_number"] = chapter
                        node.value["sequence_number"] = seq

                    # 给所有节点注入章节号（BodyText 引用上标需要）
                    if isinstance(node.value, dict):
                        node.value.setdefault("chapter_number", current_chapter)

                    if node.paragraph:
                        # 先执行内容替换（check/format 两种模式均执行）
                        node.apply_replace(document)
                        if check:
                            node.check_format(document)
                        else:
                            node.apply_format(document)
            except Exception as e:
                logger.warning(f"Node {node} not format, because: {str(e)}")
                raise e

        # 目录、附录、封面/声明的子节点跳过格式化
        SKIP_CHILDREN_CATEGORIES = {"heading_mulu", "heading_fulu", "other"}
        if category not in SKIP_CHILDREN_CATEGORIES:
            for child in node.children:
                traverse(
                    child, parent_category=category, current_chapter=current_chapter
                )

    traverse(root_node)


def _flatten_tree_nodes(root_node):
    """DFS 前序遍历展平树中所有节点（排除虚拟根节点），顺序与文档段落一致。"""
    result = []

    def dfs(node):
        for child in node.children:
            result.append(child)
            dfs(child)

    dfs(root_node)
    return result


def format_table_content(
    document: Document, config: NodeConfigRoot, check: bool = True
):
    """对文档中所有表格的单元格内容进行格式校验或应用。

    遍历 document.tables → rows → cells → paragraphs，
    根据 config.tables.content 中的格式配置对每个单元格段落进行
    段落样式和字符样式的检查（check=True）或应用（check=False）。

    Args:
        document: docx 文档对象
        config: Pydantic 配置根模型
        check: True 为仅检查（diff），False 为应用格式
    """
    try:
        content_cfg = config.tables.content
    except AttributeError:
        return

    ps = ParagraphStyle.from_config(content_cfg)
    cstyle = CharacterStyle(
        font_name_cn=content_cfg.chinese_font_name,
        font_name_en=content_cfg.english_font_name,
        font_size=content_cfg.font_size,
        font_color=content_cfg.font_color,
        bold=content_cfg.bold,
        italic=content_cfg.italic,
        underline=content_cfg.underline,
    )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        continue

                    # 段落样式
                    if check:
                        para_issues = ps.diff_from_paragraph(paragraph)
                    else:
                        para_issues = ps.apply_to_paragraph(paragraph)
                    para_text = ParagraphStyle.to_string(para_issues, target="表格内容")
                    if para_text.strip():
                        document.add_comment(
                            runs=paragraph.runs,
                            text=para_text,
                            author="Wordformat",
                            initials="afish",
                        )

                    # 字符样式
                    for run in paragraph.runs:
                        if not run.text.strip():
                            continue
                        if check:
                            diff = cstyle.diff_from_run(run)
                        else:
                            diff = cstyle.apply_to_run(run)
                        run_text = CharacterStyle.to_string(diff, target="表格内容")
                        if run_text.strip():
                            document.add_comment(
                                runs=run,
                                text=run_text,
                                author="Wordformat",
                                initials="afish",
                            )


def _build_check_summary(root_node, document, config_model) -> str:
    """遍历树和错误统计，生成检测报告摘要文本。"""
    from wordformat.rules.node import FormatNode

    stats = FormatNode._error_stats
    total = stats["total"]

    # 遍历树，收集文档级统计
    import re

    def _collect_section(node, sections):
        cls_name = type(node).__name__
        para = node.paragraph
        if para and para.text.strip():
            text = para.text.strip()
            if cls_name == "AbstractContentCN":
                chinese_only = "".join(ch for ch in text if "一" <= ch <= "鿿")
                if chinese_only:
                    sections["abstract_cn_chars"] = sections.get(
                        "abstract_cn_chars", 0
                    ) + len(chinese_only)
            elif cls_name == "AbstractContentEN":
                sections["abstract_en_words"] = sections.get(
                    "abstract_en_words", 0
                ) + len(text.split())
            elif cls_name == "KeywordsCN":
                kw_text = re.sub(r"关键词[:：]?\s*", "", text)
                kws = [k.strip() for k in re.split(r"[；;]", kw_text) if k.strip()]
                if kws:
                    sections["keyword_cn_count"] = len(kws)
            elif cls_name == "KeywordsEN":
                kw_text = re.sub(r"Keywords?:?\s*", "", text, flags=re.IGNORECASE)
                kws = [k.strip() for k in re.split(r"[,;]", kw_text) if k.strip()]
                if kws:
                    sections["keyword_en_count"] = len(kws)
            elif cls_name == "ReferenceEntry":
                has_chinese = any("一" <= ch <= "鿿" for ch in text)
                if has_chinese:
                    sections["ref_cn"] = sections.get("ref_cn", 0) + 1
                else:
                    sections["ref_en"] = sections.get("ref_en", 0) + 1
        # 处理混合节点：AbstractTitleContentCN/EN 的子 BodyText 是摘要正文
        if cls_name == "AbstractTitleContentCN":
            for child in node.children:
                cp = child.paragraph
                if cp and cp.text.strip():
                    cn = "".join(ch for ch in cp.text.strip() if "一" <= ch <= "鿿")
                    cnt = len(cn)
                    cur = sections.get("abstract_cn_chars", 0)
                    sections["abstract_cn_chars"] = cur + cnt
        elif cls_name == "AbstractTitleContentEN":
            for child in node.children:
                cp = child.paragraph
                if cp and cp.text.strip():
                    cnt = len(cp.text.strip().split())
                    sections["abstract_en_words"] = (
                        sections.get("abstract_en_words", 0) + cnt
                    )
        for child in node.children:
            _collect_section(child, sections)

    sections: dict = {}
    _collect_section(root_node, sections)

    # 计算万字差错率
    total_chars = sum(
        len(p.text) for p in document.paragraphs if p.text and p.text.strip()
    )
    error_rate = (total / max(total_chars, 1)) * 10000 if total else 0

    # 模板名（从 config 读取）
    template_name = getattr(config_model, "template_name", None) or "未知模板"

    lines = [
        "检测结果：",
        f"检测模板：《{template_name}》",
        f"检测错误数：{total}，万字差错率：{error_rate:.1f}",
        f"严重错误：{stats.get('严重', 0)}，一般错误：{stats.get('一般', 0)}，提醒：{stats.get('提醒', 0)}",
    ]

    # 字数问题
    word_issues = []
    if sections.get("abstract_cn_chars"):
        word_issues.append(
            f"中文摘要：规范：300字左右，原文：{sections['abstract_cn_chars']}字"
        )
    if sections.get("abstract_en_words"):
        word_issues.append(
            f"英文摘要：规范：300字左右，原文：{sections['abstract_en_words']}词"
        )
    if sections.get("keyword_cn_count"):
        word_issues.append(
            f"中文关键词：规范：3-5个，原文：{sections['keyword_cn_count']}个"
        )
    if sections.get("keyword_en_count"):
        word_issues.append(
            f"英文关键词：规范：3-5个，原文：{sections['keyword_en_count']}个"
        )
    ref_cn = sections.get("ref_cn", 0)
    ref_en = sections.get("ref_en", 0)
    if ref_cn or ref_en:
        word_issues.append(
            f"参考文献：规范：不少于15条，原文：中文{ref_cn}条;外文{ref_en}条"
        )
    if word_issues:
        lines.append("字数问题：")
        lines.extend(word_issues)

    lines.append("说明：")
    lines.append(
        "1.请确保文档中正确使用换行符，硬回车（Enter）：指换行且生成新段落；软回车（Shift+Enter）：指换行但不生成新段落。"
    )
    lines.append("2.图片请使用“嵌入型”环绕方式，表格为无环绕方式。")
    lines.append("3.提醒不计算错误。")

    return "\n".join(lines)


def _add_summary_comment(document, summary: str) -> None:
    """将检测报告摘要作为批注添加到文档第一段。空段临时塞空 run 做锚点。"""
    para = document.paragraphs[0]
    if not para.runs:
        para.add_run("")
    document.add_comment(
        runs=para.runs, text=summary, author="Wordformat", initials="afish"
    )


def auto_format_thesis_document(
    jsonpath: str | list,
    docxpath: str,
    configpath: Optional[str] = None,
    savepath: str = "output/",
    check=True,
):
    """自动对学位论文文档进行格式校验与批注。

    该函数根据结构化 JSON 描述和 YAML 格式配置，对指定的 Word 文档进行格式合规性检查，
    并在不符合规范的位置插入批注（comments）。主要用于学术论文（如本科/硕士/博士论文）
    的自动化格式审查。

    流程说明：
        1. 从 JSON 文件加载文档逻辑结构树；
        2. 加载 Word 文档，并将每个非空段落匹配到对应的结构节点；
        4. 对特定子树（如中英文摘要、参考文献）执行节点提升操作，确保内容节点正确挂载；
        5. 遍历所有结构节点，依据配置文件中的格式规则进行校验，并在文档中添加批注；
        6. 保存带批注的文档到指定路径。

    Args:
        check (bool): 用来控制是仅检查还是仅修改
        jsonpath (str): 文档逻辑结构的 JSON 文件路径 或 json 数据，描述各章节/段落的语义类型。
        docxpath (str): 待处理的原始 Word (.docx) 文档路径。
        savepath (str): 处理完成后带批注的文档保存路径。
        configpath (Optional[str]): 格式规范配置文件（YAML）路径，支持继承与合并。
                                 为 None 时使用内置默认配置。

    Side Effects:
        - 读取 jsonpath、docxpath 和 configpath 指定的文件；
        - 在 docx 文档中插入批注（不修改原文内容，仅添加审阅意见）；
        - 将结果文档写入 savepath。

    Example:
        >>> auto_format_thesis_document(
        ...     "thesis_structure.json",
        ...     "draft.docx",
        ...     "formatted_with_comments.docx",
        ...     "format_rules.yaml"
        ... )
    """
    from wordformat.utils import get_file_name

    if configpath:
        init_config(configpath)
        try:
            config_model = get_config()
            logger.info("配置文件验证通过")
        except Exception as e:
            logger.error(f"配置加载失败: {str(e)}")
            raise
    else:
        config_model = NodeConfigRoot()
        logger.info("未提供配置文件，使用默认配置")

    ensure_directory_exists(savepath)

    filename_without_ext = get_file_name(docxpath)
    root_node = DocumentBuilder.build_from_json(jsonpath, config=config_model)
    # 注意：不再过滤 body_text 节点，body_text 也需要格式化（首行缩进、字体等）
    document = Document(docxpath)

    if not check:
        style_list = []
        for style in document.styles:
            style_list.append(style.name)
        logger.info(f"可用的样式有：{style_list}")

    nodes = _flatten_tree_nodes(root_node)
    for node, paragraph in zip(nodes, document.paragraphs, strict=False):
        node.paragraph = paragraph

    # 替换摘要节点
    subtress_dict = {
        AbstractTitleCN: AbstractContentCN,
        AbstractTitleEN: AbstractContentEN,
        References: ReferenceEntry,
    }
    for key, value in subtress_dict.items():
        promote_bodytext_in_subtrees_of_type(
            root_node, parent_type=key, target_type=value
        )
    # 执行格式化前，先统一修正样式定义（清除主题色、设置字体等）
    if not check:
        _fix_all_style_definitions(document, config_model)

    # 执行格式化
    if check:
        FormatNode.reset_stats()
    apply_format_check_to_all_nodes(root_node, document, config_model, check)

    # 表格内容格式化（已移除：表格内格式由 Word 表格样式控制，不做段落/字符级检查）
    # format_table_content(document, config_model, check)

    # 检测报告摘要（仅 check 模式）
    if check:
        summary = _build_check_summary(root_node, document, config_model)
        if summary:
            _add_summary_comment(document, summary)

    # 处理标题自动编号（仅在格式化模式下执行，检查模式不修改编号）
    if (
        not check
        and hasattr(config_model, "numbering")
        and config_model.numbering.enabled
    ):
        from wordformat.numbering import process_heading_numbering

        process_heading_numbering(
            root_node, document, config_model.numbering, config_model.headings
        )

    # 创建引用超链接（仅在格式化模式下执行）
    if not check:
        from wordformat.hyperlinks import create_citation_hyperlinks

        create_citation_hyperlinks(root_node, document)

    savepath = Path(savepath)
    if check:
        docx_path = str(savepath / f"{filename_without_ext}--标注版.docx")
    else:
        docx_path = str(savepath / f"{filename_without_ext}--修改版.docx")
    logger.info(f"保存文件到 {docx_path}")
    document.save(docx_path)
    return docx_path
