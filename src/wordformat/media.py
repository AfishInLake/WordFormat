"""媒体层 —— 图片 blob 提取、注册、插入、文档合并。"""

from __future__ import annotations

import hashlib
import io
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def insert_element(
    elem: OxmlElement,
    document: Document,
    *,
    after: Paragraph | Table = None,
    before: Paragraph | Table = None,
    index: int | None = None,
) -> Paragraph | None:
    """将 OOXML 元素插入到文档中指定位置。

    支持段落（Paragraph）和表格（Table）作为锚点。
    四种定位方式（优先级从高到低）：
      - after:   插入到该段落/表格之后
      - before:  插入到该段落/表格之前
      - index:   插入到 document.paragraphs[index] 之后
      - 均未指定：追加到 body 末尾

    w:p 元素会包装为 Paragraph 返回；w:tbl 等元素返回 None。
    """
    from docx.text.paragraph import Paragraph as _P

    anchor = None
    use_after = True
    if after is not None:
        anchor = after._element if hasattr(after, "_element") else after
        use_after = True
    elif before is not None:
        anchor = before._element if hasattr(before, "_element") else before
        use_after = False
    elif index is not None:
        paras = document.paragraphs
        if 0 <= index < len(paras):
            anchor = paras[index]._element
            use_after = True

    if anchor is not None:
        if use_after:
            anchor.addnext(elem)
        else:
            anchor.addprevious(elem)
    else:
        try:
            document.element.body.append(elem)
        except Exception:
            pass

    if elem.tag.endswith("}p"):
        return _P(elem, document._body)
    return None


def merge_docx(cover_path: str, content_path: str, output_path: str) -> None:
    """将封面拼接到正文之前，合并两份文档的样式定义。

    以正文文档为底，迁移封面中独有的样式定义到正文，
    再将封面元素插入到正文最前面。图片 rId 自动重映射。
    """
    from docx.oxml.ns import qn as _qn

    cover = Document(cover_path)
    content = Document(content_path)

    # 1. 合并样式定义：封面独有样式 → 正文
    _merge_styles(cover, content)

    # 2. 合并编号定义：封面独有编号 → 正文
    _merge_numbering(cover, content)

    # 3. 封面元素插入正文最前面
    body = content.element.body
    cover_children = [c for c in cover.element.body if not c.tag.endswith("}sectPr")]
    for child in reversed(cover_children):
        cloned = deepcopy(child)
        for blip in cloned.iter(_qn("a:blip")):
            old_rId = blip.get(_qn("r:embed"))
            if old_rId is None:
                continue
            try:
                src_part = cover.part.related_parts.get(old_rId)
            except Exception:
                src_part = None
            if src_part is not None:
                new_rId = ImageRegistry.register_blob(content, src_part.blob)
                blip.set(_qn("r:embed"), new_rId)
        body.insert(0, cloned)

    content.save(output_path)


def _merge_styles(source: Document, target: Document) -> None:
    """将 source 中目标不存在的样式定义复制到 target。"""
    try:
        src_styles_elm = source.styles.element
    except Exception:
        return
    try:
        tgt_styles_elm = target.styles.element
    except Exception:
        return

    existing_ids = {
        s.get(qn("w:styleId"))
        for s in tgt_styles_elm.findall(qn("w:style"))
        if s.get(qn("w:styleId"))
    }

    for style_elm in src_styles_elm.findall(qn("w:style")):
        style_id = style_elm.get(qn("w:styleId"))
        if style_id and style_id not in existing_ids:
            tgt_styles_elm.append(deepcopy(style_elm))
            existing_ids.add(style_id)


def _merge_numbering(source: Document, target: Document) -> None:
    """将 source 的编号定义合并到 target。"""
    try:
        src_num = source.part.numbering_part
    except Exception:
        return
    try:
        tgt_num = target.part.numbering_part
    except Exception:
        return

    # 合并 abstractNum 和 num 定义
    existing_abstract = {
        e.get(qn("w:abstractNumId"))
        for e in tgt_num._element.findall(qn("w:abstractNum"))
        if e.get(qn("w:abstractNumId"))
    }
    existing_num = {
        e.get(qn("w:numId"))
        for e in tgt_num._element.findall(qn("w:num"))
        if e.get(qn("w:numId"))
    }

    for elem in src_num._element:
        tag = elem.tag.split("}")[-1]
        if tag == "abstractNum":
            nid = elem.get(qn("w:abstractNumId"))
            if nid and nid not in existing_abstract:
                tgt_num._element.append(deepcopy(elem))
                existing_abstract.add(nid)
        elif tag == "num":
            nid = elem.get(qn("w:numId"))
            if nid and nid not in existing_num:
                tgt_num._element.append(deepcopy(elem))
                existing_num.add(nid)


class ImageRegistry:
    """管理图片 blob 的提取、缓存和跨文档注册。"""

    @staticmethod
    def extract_from_paragraph(paragraph: Paragraph) -> dict:
        """从含 w:drawing 的段落提取图片元数据和二进制。"""
        from docx.oxml.ns import qn as _qn

        for r_elem in paragraph._element:
            if r_elem.tag.split("}")[-1] != "r":
                continue
            drawing = r_elem.find(_qn("w:drawing"))
            if drawing is None:
                continue
            blip = drawing.find(".//" + _qn("a:blip"))
            if blip is None:
                continue
            rId = blip.get(_qn("r:embed"))
            if rId is None:
                continue
            related_part = paragraph.part.related_parts.get(rId)
            if related_part is None:
                continue
            blob = related_part.blob
            extent = drawing.find(".//" + _qn("wp:extent"))
            cx = int(extent.get("cx", 0)) if extent is not None else 0
            cy = int(extent.get("cy", 0)) if extent is not None else 0
            # 对齐方式
            pPr = paragraph._element.find(_qn("w:pPr"))
            alignment = "center"
            if pPr is not None:
                jc = pPr.find(_qn("w:jc"))
                if jc is not None:
                    val = jc.get(_qn("w:val"), "center")
                    jc_map = {"left": "left", "right": "right", "both": "justify"}
                    alignment = jc_map.get(val, "center")
            return {
                "blob": blob,
                "rId": rId,
                "source": related_part.partname,
                "width_emu": cx,
                "height_emu": cy,
                "sha256": hashlib.sha256(blob).hexdigest(),
                "alignment": alignment,
            }
        return {
            "blob": None,
            "rId": None,
            "source": "",
            "width_emu": 0,
            "height_emu": 0,
            "sha256": "",
            "alignment": "center",
        }

    @staticmethod
    def register_blob(document: Document, blob: bytes) -> str:
        """将图片 blob 注册到目标文档，返回 rId。重复 blob 复用已有关系。"""
        stream = io.BytesIO(blob)
        rId, _ = document.part.get_or_add_image(stream)
        return rId

    @staticmethod
    def transfer_image(source_para: Paragraph, target_doc: Document) -> str | None:
        """将图片从源段落传输到目标文档，返回目标文档中的 rId。"""
        info = ImageRegistry.extract_from_paragraph(source_para)
        if info["blob"] is None:
            return None
        return ImageRegistry.register_blob(target_doc, info["blob"])
