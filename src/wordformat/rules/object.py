"""图片段落和表格对象节点。"""

from pathlib import Path

from docx.shared import Inches

from wordformat.rules.node import FormatNode
from wordformat.structure.registry import register
from wordformat.style.comments import format_comment


@register("figure_image")
class FigureImage(FormatNode):
    """图片段落节点（包含内联图片的段落，非题注）。

    只检查对齐和首行缩进，不检查行距、字体等。
    md→docx 时会尝试从 value 中的路径插入实际图片。
    """

    NODE_TYPE = "figures.image"
    NODE_LABEL = "图片段落"
    DEFAULTS = {"alignment": "居中对齐", "first_line_indent": "0字符"}
    DEFAULT_RULES = {}

    def _base(self, doc, p: bool, r: bool):
        """仅检查对齐和首行缩进。"""
        from wordformat.style.defs import Alignment, FirstLineIndent
        from wordformat.style.diff import _format_para_value

        self._try_insert_image()

        cfg = self.pydantic_config
        expected_align = Alignment(str(cfg.alignment or "居中对齐"))
        actual_align = expected_align.get_from_paragraph(self.paragraph)
        if expected_align != actual_align:
            self.add_comment(
                doc=doc,
                runs=self.paragraph.runs,
                text=format_comment(
                    self.NODE_LABEL,
                    "对齐错误",
                    _format_para_value("alignment", actual_align),
                    _format_para_value("alignment", expected_align),
                ),
            )

        expected_indent = FirstLineIndent(str(cfg.first_line_indent or "0字符"))
        actual_indent = expected_indent.get_from_paragraph(self.paragraph)
        if expected_indent != actual_indent:
            self.add_comment(
                doc=doc,
                runs=self.paragraph.runs,
                text=format_comment(
                    self.NODE_LABEL,
                    "首行缩进错误",
                    _format_para_value("first_line_indent", actual_indent),
                    _format_para_value("first_line_indent", expected_indent),
                ),
            )

    def _try_insert_image(self) -> None:
        """如果段落为空且 value 中有图片路径，尝试插入图片。"""
        if self.paragraph is None:
            return
        # 已有非空 run 时跳过（图片已存在或来自 docx 而非 md）
        for run in self.paragraph.runs:
            if run.text and run.text.strip():
                return
        value = self.value
        if not isinstance(value, dict):
            return
        path = value.get("paragraph", "").strip()
        if not path:
            return
        img_path = Path(path)
        if not img_path.exists():
            return
        try:
            # 清空占位 run，插入图片
            for run in self.paragraph.runs:
                run.text = ""
            if self.paragraph.runs:
                self.paragraph.runs[0].add_picture(str(img_path), width=Inches(5.5))
            else:
                run = self.paragraph.add_run()
                run.add_picture(str(img_path), width=Inches(5.5))
        except Exception:
            pass  # 图片格式不支持等，保留占位文本


@register("table_object")
class TableObject(FormatNode):
    """表格对象节点（表格整体格式，非题注）。

    md→docx 时会从 value.table_rows 创建 Word 表格。
    """

    NODE_TYPE = "tables.object"
    NODE_LABEL = "表格对象"
    DEFAULTS = {}
    DEFAULT_RULES = {}

    def _base(self, doc, p: bool, r: bool):
        self._try_insert_table(doc)

    def _try_insert_table(self, doc) -> None:
        """如果 value 中有 table_rows，在段落位置后插入 Word 表格。"""
        value = self.value
        if not isinstance(value, dict):
            return
        rows = value.get("table_rows", [])
        if not rows:
            return
        num_cols = max(len(r) for r in rows) if rows else 1
        table = doc.add_table(rows=len(rows), cols=num_cols, style="Table Grid")
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text
        # 将表格从文档末尾移到当前段落后
        if self.paragraph is not None:
            self.paragraph._element.addnext(table._element)
