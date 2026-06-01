"""DOCX 生成 —— Markdown / 图片 / 表格 → Word 文档。

Markdown 文本转 Word 文档，支持标题、段落、行内公式 ($...$)、
块级公式 ($$...$$)、HTML 表格、图片、Markdown pipe 表格。
"""

import io
import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from wordformat.math import add_display_math, add_inline_math, set_cell_math

# ── HTML table parser ────────────────────────────────────────────────


class _TableParser(HTMLParser):
    """将 HTML <table> 解析为二维 list。"""

    def __init__(self):
        super().__init__()
        self.rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell: str = ""
        self._in_td = False

    def handle_starttag(self, tag, attrs):
        if tag in ("td", "th"):
            self._in_td = True
            self._current_cell = ""

    def handle_endtag(self, tag):
        if tag in ("td", "th"):
            self._in_td = False
            self._current_row.append(self._current_cell.strip())
        elif tag == "tr":
            if self._current_row:
                self.rows.append(self._current_row)
            self._current_row = []

    def handle_data(self, data):
        if self._in_td:
            self._current_cell += data


# ── Paragraph rendering ──────────────────────────────────────────────


def _add_formatted_paragraph(doc: Document, text: str):
    """将 Markdown 段落路由到展示数学或内联数学渲染。"""
    text = text.strip()
    if text.startswith("$$") and text.endswith("$$"):
        math = text[2:-2].strip()
        if math:
            return add_display_math(doc, math)
    p = doc.add_paragraph()
    add_inline_math(p, text)
    return p


def _render_html_table(doc: Document, html_str: str):
    """将 HTML 表格字符串渲染为 python-docx 表格。"""
    parser = _TableParser()
    parser.feed(html_str)
    parser.close()
    if parser.rows:
        ncols = max(len(row) for row in parser.rows)
        table = doc.add_table(rows=len(parser.rows), cols=ncols, style="Table Grid")
        for i, row in enumerate(parser.rows):
            for j, cell_text in enumerate(row):
                if j < ncols:
                    set_cell_math(table.cell(i, j), cell_text)


def _insert_image(doc: Document, img_buf: io.BytesIO, width):
    """向文档插入居中图片。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_buf, width=width)


def _extract_html_block(lines: list[str], start: int) -> str:
    """从行列表中提取多行 HTML 块。"""
    block = lines[start]
    stripped = block.strip()
    tag_match = re.match(r"</?(\w+)", stripped.lstrip("<"))
    tag = tag_match.group(1) if tag_match else ""
    close_tag = f"</{tag}>" if tag else ""
    if close_tag and close_tag in stripped:
        return block
    if stripped.endswith(">") and not tag:
        return block
    i = start + 1
    while i < len(lines):
        block += "\n" + lines[i]
        if close_tag and close_tag in lines[i]:
            return block
        i += 1
    return block


# ── Block content ────────────────────────────────────────────────────


def _render_block_content(doc: Document, content: str):
    """渲染混合内容（文本 + HTML 表格块）。"""
    parts = re.split(r"(<table>.*?</table>)", content, flags=re.DOTALL | re.IGNORECASE)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if re.match(r"<table>.*?</table>", part, re.DOTALL | re.IGNORECASE):
            _render_html_table(doc, part)
        else:
            _add_formatted_paragraph(doc, part)


# ── Markdown renderer ────────────────────────────────────────────────


def _render_markdown(  # noqa: C901
    doc: Document,
    lines: list[str],
    image_map: Optional[dict[str, io.BytesIO]] = None,
    image_width=None,
):
    """将 Markdown 行列表渲染到 python-docx Document。"""
    if image_width is None:
        image_width = Cm(15.5)
    image_map = image_map or {}

    i = 0
    text_buf: list[str] = []
    math_buf: list[str] = []
    in_display_math = False

    def flush_text():
        nonlocal text_buf
        if text_buf:
            content = "\n".join(text_buf).strip()
            if content:
                _add_formatted_paragraph(doc, content)
            text_buf = []

    def flush_math():
        nonlocal math_buf
        if math_buf:
            math = " ".join(math_buf).strip()
            if math:
                add_display_math(doc, math)
            math_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── 展示数学块 $$...$$ ──
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            # 单行 $$...$$
            flush_text()
            math = stripped[2:-2].strip()
            if math:
                add_display_math(doc, math)
            i += 1
            continue

        if stripped.startswith("$$"):
            flush_text()
            in_display_math = True
            math_buf.append(stripped[2:].strip())
            i += 1
            continue

        if in_display_math:
            if stripped.endswith("$$"):
                math_buf.append(stripped[:-2].strip())
                flush_math()
                in_display_math = False
            else:
                math_buf.append(stripped)
            i += 1
            continue

        # ── HTML 块 ──
        if stripped.startswith("<"):
            flush_text()
            html_block = _extract_html_block(lines, i)
            i += html_block.count("\n") + 1
            s = html_block.strip()

            img_match = re.search(r'<img\s+src="([^"]+)"', s)
            if img_match:
                img_name = Path(img_match.group(1)).name
                if img_name in image_map:
                    _insert_image(doc, image_map[img_name], image_width)
                continue

            if re.match(r"<table", s, re.IGNORECASE):
                _render_html_table(doc, s)
                continue

            div_match = re.match(r"<div[^>]*>(.*?)</div>", s, re.DOTALL)
            if div_match:
                inner = div_match.group(1).strip()
                if not re.match(r"<img", inner, re.IGNORECASE):
                    p = _add_formatted_paragraph(doc, inner)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue
            continue

        # ── 标题 ──
        heading_match = re.match(r"^(#{1,5})\s+(.+)", line)
        if heading_match:
            flush_text()
            heading_text = heading_match.group(2)
            level = min(len(heading_match.group(1)), 9)
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # ── 空行：段落分隔 ──
        if not stripped:
            flush_text()
            i += 1
            continue

        # ── 表格行（Markdown pipe table） ──
        if stripped.startswith("|") and stripped.endswith("|"):
            # 检测连续的 pipe 表格行，收集后渲染为 HTML table 再处理
            table_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j].strip())
                j += 1
            # 只处理有分隔行的表格（|---|---|）
            if any(re.match(r"^\|[\s\-:|]+\|$", tl) for tl in table_lines):
                flush_text()
                _render_markdown_table(doc, table_lines)
                i = j
                continue

        text_buf.append(line)
        i += 1

    flush_text()
    flush_math()


def _render_markdown_table(doc: Document, pipe_lines: list[str]):
    """将 Markdown pipe 表格行渲染为 python-docx 表格。"""
    rows: list[list[str]] = []
    for line in pipe_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        # 跳过分隔行（|---|---|）
        if all(re.match(r"^[\-:]+$", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=ncols, style="Table Grid")
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            if cell_text:
                set_cell_math(table.cell(i, j), cell_text)


# ── Public API ───────────────────────────────────────────────────────


def markdown_to_docx(
    markdown_text: str,
    image_map: Optional[dict[str, io.BytesIO]] = None,
    image_width=None,
    page_width=None,
    page_height=None,
    top_margin=None,
    bottom_margin=None,
    left_margin=None,
    right_margin=None,
) -> io.BytesIO:
    """将 Markdown 文本渲染为 DOCX 文件（返回 BytesIO 缓冲区）。

    Args:
        markdown_text:  Markdown 文本
        image_map:      图片名 → BytesIO 映射（Markdown 中 <img src="name"> 引用）
        image_width:    图片宽度（默认 15.5cm）
        page_width:     页面宽度
        page_height:    页面高度
        top_margin:     上边距
        bottom_margin:  下边距
        left_margin:    左边距
        right_margin:   右边距

    Returns:
        包含 DOCX 数据的 BytesIO 缓冲区
    """
    if image_width is None:
        image_width = Cm(15.5)
    if page_width is None:
        page_width = Cm(21)
    if page_height is None:
        page_height = Cm(29.7)
    if top_margin is None:
        top_margin = Cm(3)
    if bottom_margin is None:
        bottom_margin = Cm(2)
    if left_margin is None:
        left_margin = Cm(3)
    if right_margin is None:
        right_margin = Cm(2.5)

    doc = Document()
    for section in doc.sections:
        section.page_width = page_width
        section.page_height = page_height
        section.top_margin = top_margin
        section.bottom_margin = bottom_margin
        section.left_margin = left_margin
        section.right_margin = right_margin

    lines = markdown_text.splitlines()
    _render_markdown(doc, lines, image_map, image_width)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
