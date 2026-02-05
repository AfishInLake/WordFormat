#! /usr/bin/env python
# @Time    : 2025/12/22 21:47
# @Author  : afish
# @File    : DocxBase.py
import re
from typing import Tuple

from loguru import logger

from config.config import get_config, init_config
from src.agent.onnx_single_infer import onnx_single_infer
from src.utils import get_paragraph_xml_fingerprint


class DocxBase:
    def __init__(self, docx_file, system_prompt, configpath):
        from docx import Document

        self.re_dict = {}
        self.docx_file = docx_file
        self.document = Document(docx_file)
        init_config(configpath)
        try:
            self.config_model = get_config()  # 首次调用：触发load()
            self.config = self.config_model.model_dump()
            logger.info("配置文件验证通过")
        except Exception as e:
            logger.error(f"配置加载失败: {str(e)}")
            raise
        self.regex_init()

    def parse(self) -> list[dict]:
        result = []
        for paragraph in self.document.paragraphs:
            # 跳过空段落
            text = paragraph.text.strip()
            if not paragraph.text.strip():
                continue
            try:
                onnx_result = onnx_single_infer(text)
                tag, score = onnx_result["label"], onnx_result["score"]
                response = {
                    "category": tag,
                    "comment": f"置信度：{score}",
                    "paragraph": text,
                }
                if score < 0.6:  # 置信度过低
                    response["category"] = "body_text"
            except Exception as e:
                response = {
                    "category": "body_text",
                    "comment": str(e),
                    "paragraph": text,
                }

            response["fingerprint"] = get_paragraph_xml_fingerprint(paragraph)
            logger.info(response)
            if response["category"] == "heading_fulu":
                break
            result.append(response)
        return result

    def get_tag_by_regex(self, paragraph: str) -> Tuple[str, str]:
        """
        检查段落是否符合正则字典中的指定规则（regex_str为正则字符串）
        Args:
            paragraph: 待匹配的段落纯文本
        Returns:
            匹配成功：(标签名, 匹配成功的正则字符串)
            匹配失败：('', '')
        """
        # 1. 空值/空白文本直接返回，避免无效匹配
        if not isinstance(paragraph, str) or not paragraph.strip():
            return "", ""

        # 2. 文本预处理：去除首尾空白/制表符，标准化匹配文本（适配Word文档冗余空白）
        norm_paragraph = paragraph.strip()

        # 3. 遍历正则字典，逐一生成正则对象并匹配
        for tag, regex_str in self.re_dict.items():
            # 防护：正则字符串为空则跳过
            if not isinstance(regex_str, str) or not regex_str.strip():
                continue
            try:
                # 编译正则字符串为正则对象（核心修复点）
                regex = re.compile(regex_str)
                # 前缀匹配：从段落开头开始匹配（符合标题/关键词等场景的排版习惯）
                if regex.match(norm_paragraph):
                    return tag, regex_str  # 匹配成功返回标签+原正则字符串
            except re.error as e:
                # 防护：无效正则字符串不崩溃，仅打印警告
                logger.warning(
                    f"警告：正则字符串编译失败（tag={tag}），错误：{e}，正则内容：{regex_str}"
                )
                continue
        # 无匹配项返回空元组
        return "", ""

    def regex_init(self):
        """初始化正则表达式字典 + 统一验证所有正则（非空+语法合法）"""
        self.re_dict = {}
        # 定义「标签-正则字符串」映射：集中管理所有需要加载的正则，便于统一校验
        # 格式：(字典key, 正则字符串, 配置项描述) → 描述用于异常信息，快速定位配置问题
        regex_mappings = [
            # 摘要相关
            (
                "abstract_chinese_title",
                self.config_model.abstract.chinese.chinese_title.section_title_re,
                "abstract.chinese.chinese_title.section_title_re",
            ),
            (
                "abstract_english_title",
                self.config_model.abstract.english.english_title.section_title_re,
                "abstract.english.english_title.section_title_re",
            ),
            (
                "keywords_chinese",
                self.config_model.abstract.keywords.get("chinese").section_title_re,
                "abstract.keywords.chinese.section_title_re",
            ),
            (
                "keywords_english",
                self.config_model.abstract.keywords.get("english").section_title_re,
                "abstract.keywords.english.section_title_re",
            ),
            # 致谢相关
            (
                "acknowledgements_title",
                self.config_model.acknowledgements.title.section_title_re,
                "acknowledgements.title.section_title_re",
            ),
            # 图题/表题
            (
                "caption_figure",
                self.config_model.figures.section_title_re,
                "figures.title.section_title_re",
            ),
            (
                "caption_table",
                self.config_model.tables.section_title_re,
                "tables.title.section_title_re",
            ),
            # 各级标题
            (
                "heading_level_1",
                self.config_model.headings.level_1.section_title_re,
                "headings.level_1.section_title_re",
            ),
            (
                "heading_level_2",
                self.config_model.headings.level_2.section_title_re,
                "headings.level_2.section_title_re",
            ),
            (
                "heading_level_3",
                self.config_model.headings.level_3.section_title_re,
                "headings.level_3.section_title_re",
            ),
            # 参考文献
            (
                "references_title",
                self.config_model.references.title.section_title_re,
                "references.title.section_title_re",
            ),
        ]

        # 统一遍历+加载+验证所有正则
        for tag, regex_str, config_desc in regex_mappings:
            # 校验1：正则字符串非空（必传，因section_title_re已设为必选字段，双重保障）
            if not isinstance(regex_str, str) or not regex_str.strip():
                raise ValueError(
                    f"正则初始化失败！标签[{tag}]对应的配置项[{config_desc}]值为空，"
                    f"请检查配置文件，该字段为必传项！"
                )
            # 校验2：正则字符串语法合法
            try:
                re.compile(
                    regex_str
                )  # 编译校验语法，无需保留对象（匹配时直接用字符串）
            except re.error as e:
                raise ValueError(
                    f"正则初始化失败！标签[{tag}]对应的配置项[{config_desc}]语法错误，"
                    f"业务无法继续运行！错误原因：{str(e)}，错误正则：{regex_str}"
                ) from e
            # 校验通过，加入正则字典
            self.re_dict[tag] = regex_str

        logger.info(f"正则表达式初始化成功，共加载{len(self.re_dict)}条有效正则规则")
